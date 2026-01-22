import os
import json
import csv
import io
import smtplib
import secrets
import uuid
import hmac
import hashlib
from email.message import EmailMessage
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.parser import BytesParser
from datetime import datetime, timedelta
from pathlib import Path
import psycopg2
from psycopg2.extras import DictCursor
from flask import Flask, g, render_template, render_template_string, request, redirect, url_for, session, flash, abort, send_from_directory, send_file, make_response, jsonify
from werkzeug.security import check_password_hash, generate_password_hash
from werkzeug.utils import secure_filename
import openpyxl
from openpyxl.utils.datetime import from_excel
import re

from shipstation_sync import (
    sync_units_and_grouping,
    normalize_company_key,
    deep_rescan_since_2024,
    set_normalize_fields_json_fn,
    set_customer_helper_fns,
)
from sync_progress import run_sync_in_background, get_sync_progress, is_sync_running, request_sync_cancel, reset_if_stale, is_cancel_requested

try:
    from hospital_targets import get_targets, build_facility_cache, zip_to_latlon
except ImportError:
    # hospital_targets module not available (missing dependencies)
    get_targets = None
    build_facility_cache = None

try:
    from doctor_targets import get_doctors_for_facility
except ImportError:
    # doctor_targets module not available (missing dependencies)
    get_doctors_for_facility = None

from threading import Thread

BASE_DIR = Path(__file__).resolve().parent
EXCLUDED_DIST_SOURCES = {"test", "cancelled", "canceled"}


def normalize_rep_slug(slug: str | None) -> str:
    """Normalize rep slug for filesystem paths (lowercase, trimmed)."""
    return (slug or "").strip().lower()


def normalize_year_month(year_month: str | None) -> str:
    """Normalize YYYY-MM strings; returns '' if invalid."""
    ym = (year_month or "").strip()
    if not re.match(r"^\d{4}-\d{2}$", ym):
        return ""
    return ym


def get_tracing_report_path(tracing_dir: Path, rep_slug: str, year_month: str, ext: str = "csv") -> Path:
    """Canonical tracing report file path."""
    rep_slug_n = normalize_rep_slug(rep_slug)
    ym = normalize_year_month(year_month)
    if not rep_slug_n or not ym:
        # Return a deterministic path, but caller should treat invalid inputs as error
        return tracing_dir / rep_slug_n / ym / f"Tracing_Report_{ym}.{ext.lstrip('.')}"
    return tracing_dir / rep_slug_n / ym / f"Tracing_Report_{ym}.{ext.lstrip('.')}"


# ---------------------------------------------------------------------------
# ShipStation sync helpers (module-scope fallbacks)
#
# Production has previously failed to boot when these helpers were missing.
# We keep module-scope implementations so `create_app()` can always wire
# `shipstation_sync.set_customer_helper_fns(...)` without crashing.


def _customer_helpers_get_database_url() -> str | None:
    return os.environ.get("DATABASE_URL")


def _customer_helpers_query_db(query: str, args: tuple = (), *, one: bool = False):
    url = _customer_helpers_get_database_url()
    if not url:
        return None if one else []
    conn = psycopg2.connect(url, sslmode="require")
    try:
        cur = conn.cursor(cursor_factory=DictCursor)
        cur.execute(query, args)
        rows = cur.fetchall()
        cur.close()
        return (rows[0] if rows else None) if one else rows
    finally:
        conn.close()


def _customer_helpers_execute_db(query: str, args: tuple = (), *, returning_id: bool = False):
    url = _customer_helpers_get_database_url()
    if not url:
        return None
    conn = psycopg2.connect(url, sslmode="require")
    try:
        cur = conn.cursor()
        cur.execute(query, args)
        new_id = None
        if returning_id:
            row = cur.fetchone()
            new_id = row[0] if row else None
        conn.commit()
        cur.close()
        return new_id
    finally:
        conn.close()


def _dict_row(row):
    if not row:
        return None
    try:
        return dict(row)
    except Exception:
        # DictCursor rows should be dict-coercible; fall back gracefully.
        return {k: row[k] for k in row.keys()}


def canonical_customer_key(name: str | None) -> str:
    return normalize_company_key(name or "")


def find_or_create_customer(
    *,
    facility_name: str | None,
    city: str | None = None,
    state: str | None = None,
    address1: str | None = None,
    address2: str | None = None,
    zip_code: str | None = None,
    contact_name: str | None = None,
    contact_phone: str | None = None,
    contact_email: str | None = None,
    primary_rep_id: int | None = None,
):
    company_key = canonical_customer_key(facility_name)
    if not company_key:
        return None

    existing = _dict_row(_customer_helpers_query_db("SELECT * FROM customers WHERE company_key = %s", (company_key,), one=True))
    if existing:
        sets: list[str] = []
        vals: list[object] = []

        def _maybe_set(col: str, value: str | None):
            v = (value or "").strip()
            if not v:
                return
            if (existing.get(col) or "").strip() == v:
                return
            sets.append(f"{col} = %s")
            vals.append(v)

        _maybe_set("facility_name", facility_name)
        _maybe_set("city", city)
        _maybe_set("state", state)
        _maybe_set("address1", address1)
        _maybe_set("address2", address2)
        _maybe_set("zip", zip_code)
        _maybe_set("contact_name", contact_name)
        _maybe_set("contact_phone", contact_phone)
        _maybe_set("contact_email", contact_email)

        if primary_rep_id and not existing.get("primary_rep_id"):
            sets.append("primary_rep_id = %s")
            vals.append(int(primary_rep_id))

        if sets:
            sets.append("updated_at = CURRENT_TIMESTAMP")
            vals.append(existing.get("id"))
            _customer_helpers_execute_db(f"UPDATE customers SET {', '.join(sets)} WHERE id = %s", tuple(vals))
            existing = _dict_row(_customer_helpers_query_db("SELECT * FROM customers WHERE id = %s", (existing.get("id"),), one=True))
        return existing

    cust_id = _customer_helpers_execute_db(
        """
        INSERT INTO customers (company_key, facility_name, address1, address2, city, state, zip, contact_name, contact_phone, contact_email, primary_rep_id)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        RETURNING id
        """,
        (
            company_key,
            (facility_name or "Unknown").strip() or "Unknown",
            (address1 or "").strip() or None,
            (address2 or "").strip() or None,
            (city or "").strip() or None,
            (state or "").strip() or None,
            (zip_code or "").strip() or None,
            (contact_name or "").strip() or None,
            (contact_phone or "").strip() or None,
            (contact_email or "").strip() or None,
            int(primary_rep_id) if primary_rep_id else None,
        ),
        returning_id=True,
    )
    if not cust_id:
        return None
    return _dict_row(_customer_helpers_query_db("SELECT * FROM customers WHERE id = %s", (cust_id,), one=True))


def ensure_rep_assignment(customer_id: int | None, rep_id: int | None, *, make_primary_if_none: bool = False) -> None:
    if not customer_id or not rep_id:
        return

    _customer_helpers_execute_db(
        """
        INSERT INTO customer_rep_assignments (customer_id, rep_id, is_primary)
        VALUES (%s, %s, FALSE)
        ON CONFLICT (customer_id, rep_id) DO NOTHING
        """,
        (int(customer_id), int(rep_id)),
    )

    if make_primary_if_none:
        cust = _dict_row(_customer_helpers_query_db("SELECT id, primary_rep_id FROM customers WHERE id = %s", (int(customer_id),), one=True))
        if cust and not cust.get("primary_rep_id"):
            _customer_helpers_execute_db(
                "UPDATE customers SET primary_rep_id = %s, updated_at=CURRENT_TIMESTAMP WHERE id = %s",
                (int(rep_id), int(customer_id)),
            )

    cust = _dict_row(_customer_helpers_query_db("SELECT primary_rep_id FROM customers WHERE id = %s", (int(customer_id),), one=True))
    primary_rep_id = cust.get("primary_rep_id") if cust else None
    if primary_rep_id:
        _customer_helpers_execute_db(
            "UPDATE customer_rep_assignments SET is_primary = (rep_id = %s) WHERE customer_id = %s",
            (int(primary_rep_id), int(customer_id)),
        )


def pick_rep_for_customer(customer_id: int | None, *, fallback_rep_id: int | None = None) -> int | None:
    if not customer_id:
        return fallback_rep_id
    cust = _dict_row(_customer_helpers_query_db("SELECT primary_rep_id FROM customers WHERE id = %s", (int(customer_id),), one=True))
    if cust and cust.get("primary_rep_id"):
        return int(cust.get("primary_rep_id"))
    row = _dict_row(
        _customer_helpers_query_db(
            """
            SELECT rep_id
            FROM customer_rep_assignments
            WHERE customer_id = %s
            ORDER BY is_primary DESC, rep_id ASC
            LIMIT 1
            """,
            (int(customer_id),),
            one=True,
        )
    )
    if row and row.get("rep_id"):
        return int(row.get("rep_id"))
    return fallback_rep_id


def create_app():
    app = Flask(__name__)
    # SECURITY: Secret key must come from environment variable (never hardcode)
    secret_key = os.environ.get("EQMS_SECRET_KEY")
    if not secret_key:
        print("[WARNING] EQMS_SECRET_KEY environment variable not set. App may not function correctly in production.")
        secret_key = "dev-change-me-in-production"  # Dev fallback only
    app.config["SECRET_KEY"] = secret_key
    
    app.config["DATABASE_URL"] = os.environ.get("DATABASE_URL")
    
    # SECURITY: Admin password must come from environment variable (never hardcode)
    admin_password = os.environ.get("EQMS_ADMIN_PASSWORD") or os.environ.get("ADMIN_PASSWORD")
    if not admin_password:
        print("[WARNING] EQMS_ADMIN_PASSWORD or ADMIN_PASSWORD environment variable not set. Admin login may not work.")
        admin_password = None  # No default - admin login will fail if not set
    app.config["ADMIN_PASSWORD"] = admin_password
    app.config["DIST_RECORDS_IMPORT_TOKEN"] = (os.environ.get("DIST_RECORDS_IMPORT_TOKEN") or "").strip()
    app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=7)
    
    # Jinja2 template filters
    @app.template_filter('format_date')
    def format_date_filter(value):
        """Format a YYYY-MM-DD string or datetime object as 'Month Day, Year'"""
        if not value:
            return "Unknown"
        try:
            # If it's already a string in YYYY-MM-DD format
            if isinstance(value, str) and len(value) >= 10:
                date_str = value[:10]
                from datetime import datetime
                dt = datetime.strptime(date_str, "%Y-%m-%d")
                return dt.strftime("%B %d, %Y")
            # If it's a datetime object
            elif hasattr(value, 'strftime'):
                return value.strftime("%B %d, %Y")
        except Exception as e:
            import traceback
            print(f"[ERROR] format_date_filter: {e}")
            traceback.print_exc()
            return str(value)[:10] if value else "Unknown"
        return str(value)[:10] if value else "Unknown"
    
    # ShipStation credentials must come from environment variables (never hardcode secrets).
    # If missing, ShipStation sync should be disabled but app boot must not crash.
    app.config["SS_API_KEY"] = os.environ.get("SHIPSTATION_API_KEY")
    app.config["SS_API_SECRET"] = os.environ.get("SHIPSTATION_API_SECRET")
    app.config["TRACING_DIR"] = str(BASE_DIR / "tracing_reports")
    app.config["DIST_LOG_APPROVALS_DIR"] = str(BASE_DIR / "distribution_log_approvals")
    app.config["UPLOAD_ROOT"] = str(BASE_DIR / "uploads")
    app.config["DIST_RECORDS_DIR"] = os.environ.get(
        "DIST_RECORDS_DIR", str(Path(app.config["UPLOAD_ROOT"]) / "distribution_records")
    )
    # Optional: store distribution record PDFs in S3 instead of the dyno filesystem.
    # This is strongly recommended on Heroku (dyno filesystem is ephemeral across deploys/restarts).
    app.config["DIST_RECORDS_S3_BUCKET"] = (
        os.environ.get("DIST_RECORDS_S3_BUCKET")
        or os.environ.get("BUCKETEER_BUCKET_NAME")
        or ""
    ).strip()
    app.config["DIST_RECORDS_S3_PREFIX"] = (os.environ.get("DIST_RECORDS_S3_PREFIX") or "distribution_records").strip().strip("/")
    app.config["DIST_RECORDS_S3_REGION"] = (
        os.environ.get("AWS_REGION")
        or os.environ.get("AWS_DEFAULT_REGION")
        or os.environ.get("BUCKETEER_AWS_REGION")
        or ""
    ).strip()
    # For S3-compatible providers (e.g. DigitalOcean Spaces), set a custom endpoint URL.
    # Example: https://nyc3.digitaloceanspaces.com
    app.config["DIST_RECORDS_S3_ENDPOINT_URL"] = (
        os.environ.get("DIST_RECORDS_S3_ENDPOINT_URL")
        or os.environ.get("AWS_ENDPOINT_URL")
        or os.environ.get("S3_ENDPOINT_URL")
        or ""
    ).strip()
    app.config["FORM_TEMPLATES_DIR"] = os.environ.get("FORM_TEMPLATES_DIR", str(BASE_DIR / "form_templates"))
    app.config["QUALITY_DOCS_DIR"] = os.environ.get("QUALITY_DOCS_DIR", str(BASE_DIR / "quality_docs"))
    app.config["TRAINING_DOCS_DIR"] = os.environ.get("TRAINING_DOCS_DIR", str(BASE_DIR / "training_docs"))
    app.config["PRODUCT_COMPLAINTS_DIR"] = os.environ.get("PRODUCT_COMPLAINTS_DIR", str(BASE_DIR / "product_complaints"))
    app.config["CUSTOMER_INFO_DIR"] = os.environ.get("CUSTOMER_INFO_DIR", str(BASE_DIR / "customer_info"))
    
    # Security: File upload size limits (default: 50MB)
    max_upload_size = int(os.environ.get("MAX_UPLOAD_SIZE_MB", "50")) * 1024 * 1024
    app.config["MAX_CONTENT_LENGTH"] = max_upload_size
    
    # Allowed file extensions for uploads (admin-only)
    app.config["ALLOWED_EXTENSIONS"] = {
        ".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".eml"
    }
    
    # Allowed extensions for specific upload types
    app.config["ALLOWED_TRACING_EXTENSIONS"] = {".eml", ".pdf", ".docx", ".doc"}
    app.config["ALLOWED_DISTRIBUTION_EXTENSIONS"] = {".pdf", ".docx", ".doc", ".xlsx", ".xls"}
    app.config["ALLOWED_CSV_EXTENSIONS"] = {".csv"}
    app.config["ALLOWED_APPROVAL_EXTENSIONS"] = {".eml", ".pdf", ".docx", ".doc"}

    # Ensure critical directories exist (non-fatal: create if missing, log if creation fails)
    try:
        Path(app.config["DIST_RECORDS_DIR"]).mkdir(parents=True, exist_ok=True)
        Path(app.config["DIST_LOG_APPROVALS_DIR"]).mkdir(parents=True, exist_ok=True)
        Path(app.config["TRACING_DIR"]).mkdir(parents=True, exist_ok=True)
        Path(app.config["UPLOAD_ROOT"]).mkdir(parents=True, exist_ok=True)
    except Exception as e:
        print(f"[STARTUP] WARNING: Failed to create critical directories: {e}")

    # Email / SMTP config (set via environment variables in production)
    app.config["SMTP_SERVER"] = os.environ.get("SMTP_SERVER")
    app.config["SMTP_PORT"] = os.environ.get("SMTP_PORT", "587")
    app.config["SMTP_USE_TLS"] = os.environ.get("SMTP_USE_TLS", "1") == "1"
    app.config["SMTP_USERNAME"] = os.environ.get("SMTP_USERNAME")
    app.config["SMTP_PASSWORD"] = os.environ.get("SMTP_PASSWORD")
    email_from = (os.environ.get("EMAIL_FROM") or os.environ.get("SMTP_USERNAME") or "").strip()
    app.config["EMAIL_FROM"] = email_from or "no-reply@example.com"

    # region agent log
    def _agent_log(hypothesisId: str, location: str, message: str, data: dict | None = None, *, runId: str = "run1"):
        """Write one NDJSON line to .cursor/debug.log (no secrets / no PII)."""
        try:
            import json as _json
            from pathlib import Path as _Path
            payload = {
                "sessionId": "debug-session",
                "runId": runId,
                "hypothesisId": hypothesisId,
                "location": location,
                "message": message,
                "data": data or {},
                "timestamp": int(datetime.utcnow().timestamp() * 1000),
            }
            log_path = _Path(r"C:\Users\Ethan\OneDrive\Desktop\UI\RepsQMS\.cursor\debug.log")
            log_path.parent.mkdir(parents=True, exist_ok=True)
            with open(log_path, "a", encoding="utf-8") as f:
                f.write(_json.dumps(payload, ensure_ascii=False) + "\n")
        except Exception:
            # Never break runtime due to logging
            pass
    # endregion

    def validate_upload_file(file_storage, allowed_extensions=None):
        """
        Validate an uploaded file for security.
        
        Returns: (is_valid: bool, error_msg: str | None, secure_filename: str | None)
        """
        if not file_storage or not getattr(file_storage, "filename", None):
            return False, "No file provided", None
        
        filename = file_storage.filename
        secure_name = secure_filename(filename)
        
        if not secure_name:
            return False, "Invalid filename", None
        
        # Check extension
        if allowed_extensions:
            ext = Path(secure_name).suffix.lower()
            if ext not in allowed_extensions:
                return False, f"File type not allowed. Allowed: {', '.join(allowed_extensions)}", None
        
        # Check file size (Flask's MAX_CONTENT_LENGTH handles this, but we can double-check)
        try:
            file_storage.seek(0, 2)  # Seek to end
            size = file_storage.tell()
            file_storage.seek(0)  # Reset
            max_size = app.config.get("MAX_CONTENT_LENGTH", 50 * 1024 * 1024)
            if size > max_size:
                max_mb = max_size / (1024 * 1024)
                return False, f"File too large. Maximum size: {max_mb:.0f} MB", None
        except Exception:
            pass  # If we can't check size, let Flask handle it
        
        return True, None, secure_name

    def get_db():
        if "db" not in g:
            url = app.config["DATABASE_URL"]
            if not url:
                return None
            g.db = psycopg2.connect(url, sslmode='require')
        return g.db

    @app.teardown_appcontext
    def close_db(exception=None):
        db = g.pop("db", None)
        if db is not None:
            db.close()
    
    @app.after_request
    def add_security_headers(response):
        """Add security headers to all responses"""
        # X-Frame-Options: Prevent clickjacking
        response.headers['X-Frame-Options'] = 'SAMEORIGIN'
        # X-Content-Type-Options: Prevent MIME type sniffing
        response.headers['X-Content-Type-Options'] = 'nosniff'
        # Basic CSP (conservative, allow inline styles/scripts for Bootstrap/jQuery compatibility)
        # Updated to allow OpenStreetMap tiles, Leaflet marker icons, and reverse geocoding API
        response.headers['Content-Security-Policy'] = (
            "default-src 'self'; "
            "script-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://unpkg.com; "
            "style-src 'self' 'unsafe-inline' https://cdn.jsdelivr.net https://fonts.googleapis.com; "
            "font-src 'self' https://fonts.gstatic.com; "
            "img-src 'self' data: https://*.tile.openstreetmap.org https://unpkg.com; "
            "connect-src 'self' https://nominatim.openstreetmap.org"
        )
        # region agent log
        try:
            p = request.path if request else ""
            if p.startswith("/rep/") or p.startswith("/api/rep-targets") or p.startswith("/api/facility-doctors"):
                _agent_log(
                    "MAP_CSP",
                    "Proto1.py:add_security_headers",
                    "CSP applied",
                    {
                        "path": p,
                        "has_csp": bool(response.headers.get("Content-Security-Policy")),
                        "csp_len": len(response.headers.get("Content-Security-Policy") or ""),
                    },
                )
        except Exception:
            pass
        # endregion
        return response

    def send_email(to: str, subject: str, body: str, *, html: str | None = None, attachments: list[tuple[str, bytes, str]] | None = None) -> tuple[bool, str]:
        """
        Send an email using SMTP configuration from app.config.
        
        Args:
            to: Recipient email address
            subject: Email subject
            body: Plain text email body
            html: Optional HTML email body (if provided, email will be multipart)
            attachments: Optional list of (filename, bytes, content_type) tuples
            
        Returns:
            Tuple of (success: bool, error_message: str)
        """
        import smtplib
        from email.mime.text import MIMEText
        from email.mime.multipart import MIMEMultipart
        from email.mime.base import MIMEBase
        from email import encoders
        import traceback
        
        try:
            smtp_server = app.config.get("SMTP_SERVER", "").strip()
            smtp_port = app.config.get("SMTP_PORT")
            smtp_use_tls = app.config.get("SMTP_USE_TLS", True)
            smtp_username = app.config.get("SMTP_USERNAME", "").strip()
            smtp_password = app.config.get("SMTP_PASSWORD", "").strip()
            email_from = app.config.get("EMAIL_FROM", "").strip()
            
            if not smtp_server:
                error_msg = "SMTP server not configured (SMTP_SERVER environment variable missing)"
                print(f"[EMAIL_ERROR] {error_msg}")
                return False, error_msg
            
            if not email_from:
                error_msg = "Email from address not configured (EMAIL_FROM environment variable missing)"
                print(f"[EMAIL_ERROR] {error_msg}")
                return False, error_msg
            
            # Create message
            if html:
                msg = MIMEMultipart('alternative')
                msg.attach(MIMEText(body, 'plain'))
                msg.attach(MIMEText(html, 'html'))
            else:
                msg = MIMEText(body, 'plain')
            
            msg['Subject'] = subject
            msg['From'] = email_from
            msg['To'] = to
            
            # Add attachments if provided
            if attachments:
                if not isinstance(msg, MIMEMultipart):
                    # Convert to multipart if we have attachments
                    multipart_msg = MIMEMultipart()
                    multipart_msg['Subject'] = msg['Subject']
                    multipart_msg['From'] = msg['From']
                    multipart_msg['To'] = msg['To']
                    multipart_msg.attach(msg)
                    msg = multipart_msg
                
                for filename, file_bytes, content_type in attachments:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(file_bytes)
                    encoders.encode_base64(part)
                    part.add_header(
                        'Content-Disposition',
                        f'attachment; filename= {filename}'
                    )
                    msg.attach(part)
            
            # Connect to SMTP server
            try:
                if smtp_port:
                    server = smtplib.SMTP(smtp_server, int(smtp_port))
                else:
                    server = smtplib.SMTP(smtp_server)
                
                if smtp_use_tls:
                    server.starttls()
                
                if smtp_username and smtp_password:
                    server.login(smtp_username, smtp_password)
                
                server.send_message(msg)
                server.quit()
                
                print(f"[EMAIL_SUCCESS] Sent email to {to} with subject: {subject}")
                return True, "sent"
                
            except smtplib.SMTPAuthenticationError as e:
                error_msg = f"SMTP authentication failed: {str(e)}"
                print(f"[EMAIL_ERROR] {error_msg}")
                return False, error_msg
            except smtplib.SMTPException as e:
                error_msg = f"SMTP error: {str(e)}"
                print(f"[EMAIL_ERROR] {error_msg}")
                return False, error_msg
            except Exception as e:
                error_msg = f"SMTP connection error: {str(e)}"
                print(f"[EMAIL_ERROR] {error_msg}")
                traceback.print_exc()
                return False, error_msg
                
        except Exception as e:
            error_msg = f"Email send failed: {str(e)}"
            print(f"[EMAIL_ERROR] {error_msg}")
            traceback.print_exc()
            return False, error_msg

    @app.errorhandler(500)
    def handle_internal_error(e):
        """Global error handler - show friendly error with correlation ID, no stack traces in production"""
        error_id = secrets.token_hex(8)
        is_debug = app.config.get("DEBUG", False) or os.environ.get("FLASK_ENV") == "development"
        
        # Log error (always log to server console)
        print(f"ERROR {error_id}: {e}")
        
        # Only print full traceback in debug/development mode
        if is_debug:
            import traceback
            traceback.print_exc()
        
        # Check if this is an admin route
        if request.path.startswith("/admin"):
            return f"""<html><head><title>Error</title>
                <style>
                    body {{ font-family: Arial, sans-serif; max-width: 600px; margin: 50px auto; padding: 20px; }}
                    .error-box {{ border: 1px solid #dc3545; background: #f8d7da; padding: 20px; border-radius: 5px; }}
                    .error-id {{ font-family: monospace; background: #fff; padding: 5px; border-radius: 3px; }}
                </style>
            </head><body>
                <div class="error-box">
                    <h1>Error</h1>
                    <p>An error occurred while processing your request.</p>
                    <p><strong>Error ID:</strong> <span class="error-id">{error_id}</span></p>
                    <p>Please check server logs for details. If this problem persists, contact support with the Error ID above.</p>
                    <p><a href="/admin/dashboard">Return to Dashboard</a> | <a href="javascript:history.back()">Go Back</a></p>
                </div>
            </body></html>""", 500
        # For non-admin routes, return generic error
        return f"<h1>Error</h1><p>Error ID: {error_id}</p>", 500

    def query_db(query, args=(), one=False):
        conn = get_db()
        if not conn:
            return [] if not one else None
        cur = conn.cursor(cursor_factory=DictCursor)
        cur.execute(query, args)
        rv = cur.fetchall()
        cur.close()
        return (rv[0] if rv else None) if one else rv

    def execute_db(query, args=(), returning_id=False):
        conn = get_db()
        if not conn:
            return None
        cur = conn.cursor()
        cur.execute(query, args)
        if returning_id:
            new_id_row = cur.fetchone()
            conn.commit()
            return new_id_row[0] if new_id_row else None
        conn.commit()
        return None

    # Cache for reps.active column type detection (module-level cache)
    _reps_active_type_cache = None

    def convert_active_to_db_type(active_value: bool | int) -> int | bool:
        """
        Convert active value (boolean or int) to the correct database type.
        
        Detects the actual column type for reps.active at runtime and converts:
        - If column is INTEGER → returns 1/0
        - If column is BOOLEAN → returns True/False
        
        Caches the result to avoid repeated queries.
        """
        nonlocal _reps_active_type_cache
        
        # Check cache first
        if _reps_active_type_cache is not None:
            is_integer = _reps_active_type_cache
        else:
            # Detect column type using information_schema
            try:
                conn = get_db()
                if conn:
                    cur = conn.cursor()
                    try:
                        cur.execute("""
                            SELECT data_type 
                            FROM information_schema.columns 
                            WHERE table_name = 'reps' AND column_name = 'active'
                        """)
                        row = cur.fetchone()
                        if row:
                            data_type = row[0].upper()
                            is_integer = data_type in ('INTEGER', 'INT', 'SMALLINT', 'BIGINT')
                        else:
                            # Column doesn't exist yet (shouldn't happen, but fallback to integer)
                            is_integer = True
                        _reps_active_type_cache = is_integer
                    finally:
                        cur.close()
                else:
                    # Can't query, default to integer (safer for compatibility)
                    is_integer = True
                    _reps_active_type_cache = is_integer
            except Exception as e:
                # On error, default to integer (safer for compatibility)
                print(f"[WARNING] Could not detect reps.active column type: {e}. Defaulting to integer.")
                is_integer = True
                _reps_active_type_cache = is_integer
        
        # Convert value based on detected type
        if isinstance(active_value, bool):
            if is_integer:
                return 1 if active_value else 0
            # BOOLEAN column: must never return 1/0
            return True if active_value else False
        elif isinstance(active_value, int):
            # Already an int, but ensure it's 0/1 if column is integer
            if is_integer:
                return 1 if active_value else 0
            else:
                return bool(active_value)
        else:
            # Fallback: treat truthy as 1/True, falsy as 0/False
            if is_integer:
                return 1 if active_value else 0
            return bool(active_value)

    def init_db():
        conn = get_db()
        if not conn:
            return
        with conn.cursor() as cur:
            cur.execute("""
                CREATE TABLE IF NOT EXISTS reps (
                    id SERIAL PRIMARY KEY,
                    name TEXT NOT NULL,
                    slug TEXT NOT NULL UNIQUE,
                    password_hash TEXT NOT NULL,
                    active INTEGER NOT NULL DEFAULT 1,
                    email TEXT
                )
            """)
            
            # Schema compatibility: Migrate reps.active from boolean to integer if needed
            try:
                cur.execute("""
                    SELECT data_type 
                    FROM information_schema.columns 
                    WHERE table_name = 'reps' AND column_name = 'active'
                """)
                row = cur.fetchone()
                if row:
                    data_type = row[0].upper()
                    if data_type in ('BOOLEAN', 'BOOL'):
                        print("[INIT_DB] Detected reps.active is BOOLEAN, migrating to INTEGER...")
                        cur.execute("""
                            ALTER TABLE reps ALTER COLUMN active TYPE INTEGER USING (CASE WHEN active THEN 1 ELSE 0 END)
                        """)
                        conn.commit()
                        print("[INIT_DB] Migration complete: reps.active is now INTEGER")
                    else:
                        print(f"[INIT_DB] reps.active column type: {data_type} (no migration needed)")
                else:
                    # Column doesn't exist yet (shouldn't happen after CREATE TABLE, but safe fallback)
                    print("[INIT_DB] reps.active column not found in information_schema (may be new table)")
            except Exception as e:
                # Fail softly - log warning but don't crash
                print(f"[INIT_DB] WARNING: Could not check/migrate reps.active column type: {e}")
                print("[INIT_DB] Continuing with default behavior (assuming INTEGER)")
            
            # Add address columns if they don't exist (safe migration)
            try:
                cur.execute("""
                    DO $$ 
                    BEGIN
                        IF NOT EXISTS (SELECT 1 FROM information_schema.columns WHERE table_name='reps' AND column_name='address_line1') THEN
                            ALTER TABLE reps ADD COLUMN address_line1 TEXT;
                        END IF;
                        IF NOT EXISTS (SELECT 1 FROM information_schema.columns WHERE table_name='reps' AND column_name='address_line2') THEN
                            ALTER TABLE reps ADD COLUMN address_line2 TEXT;
                        END IF;
                        IF NOT EXISTS (SELECT 1 FROM information_schema.columns WHERE table_name='reps' AND column_name='city') THEN
                            ALTER TABLE reps ADD COLUMN city TEXT;
                        END IF;
                        IF NOT EXISTS (SELECT 1 FROM information_schema.columns WHERE table_name='reps' AND column_name='state') THEN
                            ALTER TABLE reps ADD COLUMN state TEXT;
                        END IF;
                        IF NOT EXISTS (SELECT 1 FROM information_schema.columns WHERE table_name='reps' AND column_name='zip') THEN
                            ALTER TABLE reps ADD COLUMN zip TEXT;
                        END IF;
                    END $$;
                """)
            except Exception as e:
                # Columns may already exist, or database doesn't support DO blocks - try individual ALTERs
                for col in ['address_line1', 'address_line2', 'city', 'state', 'zip']:
                    try:
                        cur.execute(f"ALTER TABLE reps ADD COLUMN IF NOT EXISTS {col} TEXT")
                    except Exception:
                        pass  # Column likely already exists
            cur.execute("""
                CREATE TABLE IF NOT EXISTS devices_distributed (
                    id SERIAL PRIMARY KEY,
                    rep_id INTEGER NOT NULL REFERENCES reps(id),
                    shipment_id INTEGER REFERENCES devices_received(id),
                    created_at TEXT NOT NULL,
                    order_number TEXT,
                    ship_date TEXT,
                    tracking_number TEXT,
                    source TEXT DEFAULT 'manual'
                )
            """)
            # Persistent distribution numbering, e.g. D-2025-000123
            cur.execute("ALTER TABLE devices_distributed ADD COLUMN IF NOT EXISTS distribution_number TEXT")
            cur.execute("ALTER TABLE devices_distributed ADD COLUMN IF NOT EXISTS ss_shipment_id TEXT")
            cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_devices_distributed_ss_shipment ON devices_distributed(ss_shipment_id) WHERE ss_shipment_id IS NOT NULL")
            cur.execute(
                "CREATE UNIQUE INDEX IF NOT EXISTS idx_devices_distributed_distribution_number ON devices_distributed(distribution_number) WHERE distribution_number IS NOT NULL AND distribution_number <> ''"
            )

            # Per-year distribution numbering counters.
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS distribution_number_counters (
                    year INTEGER PRIMARY KEY,
                    last_seq INTEGER NOT NULL DEFAULT 0,
                    updated_at TIMESTAMPTZ NOT NULL DEFAULT now()
                )
                """
            )

            # Atomic increment for the per-year sequence.
            cur.execute(
                """
                CREATE OR REPLACE FUNCTION next_distribution_seq(p_year INTEGER)
                RETURNS INTEGER
                LANGUAGE plpgsql
                AS $$
                DECLARE
                    v_seq INTEGER;
                BEGIN
                    INSERT INTO distribution_number_counters(year, last_seq)
                    VALUES (p_year, 1)
                    ON CONFLICT (year) DO UPDATE
                      SET last_seq = distribution_number_counters.last_seq + 1,
                          updated_at = now()
                    RETURNING last_seq INTO v_seq;
                    RETURN v_seq;
                END;
                $$;
                """
            )

            # Trigger to auto-assign a distribution_number on insert if missing.
            cur.execute(
                """
                CREATE OR REPLACE FUNCTION assign_distribution_number_trigger()
                RETURNS TRIGGER
                LANGUAGE plpgsql
                AS $$
                DECLARE
                    v_year INTEGER;
                    v_seq INTEGER;
                    ship_date_text TEXT;
                    created_at_text TEXT;
                BEGIN
                    IF NEW.distribution_number IS NOT NULL AND btrim(NEW.distribution_number) <> '' THEN
                        RETURN NEW;
                    END IF;

                    ship_date_text := COALESCE(NEW.ship_date, '');
                    created_at_text := COALESCE(NEW.created_at, '');

                    IF ship_date_text ~ '^[0-9]{4}-' THEN
                        v_year := substring(ship_date_text from 1 for 4)::INTEGER;
                    ELSIF created_at_text ~ '^[0-9]{4}-' THEN
                        v_year := substring(created_at_text from 1 for 4)::INTEGER;
                    ELSE
                        v_year := EXTRACT(YEAR FROM now())::INTEGER;
                    END IF;

                    v_seq := next_distribution_seq(v_year);
                    NEW.distribution_number := 'D-' || v_year::TEXT || '-' || lpad(v_seq::TEXT, 6, '0');
                    RETURN NEW;
                END;
                $$;
                """
            )
            cur.execute("DROP TRIGGER IF EXISTS trig_devices_distributed_assign_distribution_number ON devices_distributed")
            cur.execute(
                """
                CREATE TRIGGER trig_devices_distributed_assign_distribution_number
                BEFORE INSERT ON devices_distributed
                FOR EACH ROW
                EXECUTE FUNCTION assign_distribution_number_trigger();
                """
            )
            cur.execute("""
                CREATE TABLE IF NOT EXISTS device_distribution_records (
                    id SERIAL PRIMARY KEY,
                    rep_id INTEGER NOT NULL REFERENCES reps(id),
                    dist_id INTEGER NOT NULL REFERENCES devices_distributed(id),
                    stored_filename TEXT NOT NULL,
                    original_filename TEXT NOT NULL,
                    uploaded_at TEXT NOT NULL,
                    fields_json TEXT NOT NULL
                )
            """)
            cur.execute("ALTER TABLE device_distribution_records ADD COLUMN IF NOT EXISTS missing_required_json TEXT")
            cur.execute("ALTER TABLE device_distribution_records ADD COLUMN IF NOT EXISTS unexpected_fields_json TEXT")
            cur.execute("ALTER TABLE device_distribution_records ADD COLUMN IF NOT EXISTS file_type TEXT")
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS master_import_jobs (
                    job_key TEXT PRIMARY KEY,
                    status TEXT NOT NULL,
                    started_at TEXT,
                    updated_at TEXT,
                    cancel_requested INTEGER NOT NULL DEFAULT 0,
                    error TEXT,
                    progress_json TEXT NOT NULL
                )
                """
            )
            # Seed the single master Sales Order import job row.
            cur.execute(
                """
                INSERT INTO master_import_jobs (job_key, status, started_at, updated_at, cancel_requested, error, progress_json)
                VALUES (%s, %s, %s, %s, 0, NULL, %s)
                ON CONFLICT (job_key) DO NOTHING
                """,
                (
                    "master_salesorder_pdf",
                    "idle",
                    None,
                    datetime.utcnow().isoformat(timespec="seconds") + "Z",
                    json.dumps({"status": "idle", "message": "", "started_at": None, "updated_at": None}),
                ),
            )

            # Seed the Label_Bulk import job row.
            cur.execute(
                """
                INSERT INTO master_import_jobs (job_key, status, started_at, updated_at, cancel_requested, error, progress_json)
                VALUES (%s, %s, %s, %s, 0, NULL, %s)
                ON CONFLICT (job_key) DO NOTHING
                """,
                (
                    "label_bulk_pdf",
                    "idle",
                    None,
                    datetime.utcnow().isoformat(timespec="seconds") + "Z",
                    json.dumps({"status": "idle", "message": "", "started_at": None, "updated_at": None}),
                ),
            )

            # Backfill distribution numbers for existing rows that predate the column/trigger.
            # This is idempotent (only fills blanks) and also updates counters to the max used.
            try:
                cur.execute(
                    """
                    WITH dist AS (
                        SELECT
                            id,
                            CASE
                                WHEN COALESCE(ship_date, '') ~ '^[0-9]{4}-' THEN substring(ship_date from 1 for 4)::INTEGER
                                WHEN COALESCE(created_at, '') ~ '^[0-9]{4}-' THEN substring(created_at from 1 for 4)::INTEGER
                                ELSE EXTRACT(YEAR FROM now())::INTEGER
                            END AS yr,
                            row_number() OVER (
                                PARTITION BY CASE
                                    WHEN COALESCE(ship_date, '') ~ '^[0-9]{4}-' THEN substring(ship_date from 1 for 4)::INTEGER
                                    WHEN COALESCE(created_at, '') ~ '^[0-9]{4}-' THEN substring(created_at from 1 for 4)::INTEGER
                                    ELSE EXTRACT(YEAR FROM now())::INTEGER
                                END
                                ORDER BY NULLIF(ship_date, '') NULLS LAST, id ASC
                            ) AS rn
                        FROM devices_distributed
                        WHERE distribution_number IS NULL OR btrim(distribution_number) = ''
                    ),
                    upd AS (
                        UPDATE devices_distributed dd
                        SET distribution_number = 'D-' || dist.yr::TEXT || '-' || lpad(dist.rn::TEXT, 6, '0')
                        FROM dist
                        WHERE dd.id = dist.id
                        RETURNING dist.yr AS yr, dist.rn AS rn
                    )
                    SELECT yr, MAX(rn) AS max_rn
                    FROM upd
                    GROUP BY yr
                    """
                )
                backfill_rows = cur.fetchall() or []
                for yr, max_rn in backfill_rows:
                    try:
                        cur.execute(
                            """
                            INSERT INTO distribution_number_counters(year, last_seq)
                            VALUES (%s, %s)
                            ON CONFLICT (year) DO UPDATE
                              SET last_seq = GREATEST(distribution_number_counters.last_seq, EXCLUDED.last_seq),
                                  updated_at = now()
                            """,
                            (int(yr), int(max_rn)),
                        )
                    except Exception:
                        pass
            except Exception:
                # Don't block app startup if backfill fails; the trigger will cover new rows.
                pass
            cur.execute("""
                CREATE TABLE IF NOT EXISTS new_customer_records (
                    id SERIAL PRIMARY KEY,
                    rep_id INTEGER REFERENCES reps(id),
                    dist_id INTEGER REFERENCES devices_distributed(id),
                    stored_filename TEXT NOT NULL,
                    original_filename TEXT NOT NULL,
                    uploaded_at TEXT NOT NULL,
                    fields_json TEXT NOT NULL,
                    company_key TEXT
                )
            """)
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS customers (
                    id SERIAL PRIMARY KEY,
                    company_key TEXT NOT NULL UNIQUE,
                    facility_name TEXT NOT NULL,
                    address1 TEXT,
                    address2 TEXT,
                    city TEXT,
                    state TEXT,
                    zip TEXT,
                    contact_name TEXT,
                    contact_phone TEXT,
                    contact_email TEXT,
                    primary_rep_id INTEGER REFERENCES reps(id),
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS customer_rep_assignments (
                    id SERIAL PRIMARY KEY,
                    customer_id INTEGER NOT NULL REFERENCES customers(id) ON DELETE CASCADE,
                    rep_id INTEGER NOT NULL REFERENCES reps(id),
                    is_primary BOOLEAN DEFAULT FALSE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    UNIQUE(customer_id, rep_id)
                )
                """
            )
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS customer_notes (
                    id SERIAL PRIMARY KEY,
                    customer_id INTEGER NOT NULL REFERENCES customers(id) ON DELETE CASCADE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    note_date DATE DEFAULT CURRENT_DATE,
                    author TEXT,
                    note_text TEXT NOT NULL
                )
                """
            )
            cur.execute("ALTER TABLE customer_notes ADD COLUMN IF NOT EXISTS note_date DATE DEFAULT CURRENT_DATE")
            cur.execute("ALTER TABLE customer_notes ADD COLUMN IF NOT EXISTS updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP")
            cur.execute(
                """
                CREATE TABLE IF NOT EXISTS rep_password_resets (
                    id SERIAL PRIMARY KEY,
                    rep_id INTEGER NOT NULL REFERENCES reps(id),
                    token TEXT NOT NULL UNIQUE,
                    expires_at TIMESTAMP NOT NULL,
                    used BOOLEAN DEFAULT FALSE,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    used_at TIMESTAMP
                )
                """
            )
            cur.execute("ALTER TABLE devices_distributed ADD COLUMN IF NOT EXISTS customer_id INTEGER REFERENCES customers(id)")
            cur.execute("ALTER TABLE device_distribution_records ADD COLUMN IF NOT EXISTS customer_id INTEGER REFERENCES customers(id)")
            cur.execute("ALTER TABLE new_customer_records ADD COLUMN IF NOT EXISTS customer_id INTEGER REFERENCES customers(id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_customers_company_state ON customers(company_key, state)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_customer_notes_customer ON customer_notes(customer_id, created_at DESC)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_customer_rep_assignments ON customer_rep_assignments(customer_id, rep_id)")
            cur.execute("""
                CREATE TABLE IF NOT EXISTS tracing_reports (
                    id SERIAL PRIMARY KEY,
                    rep_id INTEGER NOT NULL REFERENCES reps(id),
                    month TEXT NOT NULL,
                    generated_at TIMESTAMP NOT NULL,
                    status TEXT NOT NULL DEFAULT 'draft',
                    report_path TEXT NOT NULL,
                    report_csv_content TEXT,
                    email_to TEXT,
                    email_sent_at TIMESTAMP,
                    approval_uploaded_at TIMESTAMP,
                    approval_file_path TEXT,
                    UNIQUE(rep_id, month)
                )
            """)
            conn.commit()
            
            # Add edited_at column to tracing_reports table
            try:
                cur.execute("ALTER TABLE tracing_reports ADD COLUMN IF NOT EXISTS edited_at TIMESTAMP")
                conn.commit()
            except Exception:
                pass  # Column may already exist
            
            # Add report_csv_content column for persistent storage (replaces filesystem dependency)
            try:
                cur.execute("ALTER TABLE tracing_reports ADD COLUMN IF NOT EXISTS report_csv_content TEXT")
                conn.commit()
            except Exception:
                pass  # Column may already exist
            
            # Ensure tracing_approval_tokens table exists
            ensure_tracing_approval_tokens_table()
            
            # Ensure distribution_log_approvals table exists
            ensure_distribution_log_approvals_table()
            
            # Create devices_received table if it doesn't exist
            cur.execute("""
                CREATE TABLE IF NOT EXISTS devices_received (
                    id SERIAL PRIMARY KEY,
                    rep_id INTEGER NOT NULL REFERENCES reps(id),
                    title TEXT,
                    description TEXT,
                    created_at TEXT NOT NULL,
                    packing_list_filename TEXT,
                    packing_list_original TEXT,
                    recv_inspection_filename TEXT,
                    recv_inspection_original TEXT
                )
            """)
            # Add columns if they don't exist (for backward compatibility)
            cur.execute("ALTER TABLE devices_received ADD COLUMN IF NOT EXISTS packing_list_filename TEXT")
            cur.execute("ALTER TABLE devices_received ADD COLUMN IF NOT EXISTS packing_list_original TEXT")
            cur.execute("ALTER TABLE devices_received ADD COLUMN IF NOT EXISTS recv_inspection_filename TEXT")
            cur.execute("ALTER TABLE devices_received ADD COLUMN IF NOT EXISTS recv_inspection_original TEXT")
            
            # Create shipment_line_items table for lot + units tracking
            cur.execute("""
                CREATE TABLE IF NOT EXISTS shipment_line_items (
                    id SERIAL PRIMARY KEY,
                    shipment_id INTEGER NOT NULL REFERENCES devices_received(id) ON DELETE CASCADE,
                    lot_number TEXT NOT NULL,
                    units_received INTEGER NOT NULL CHECK (units_received > 0),
                    created_at TEXT NOT NULL
                )
            """)
            cur.execute("CREATE INDEX IF NOT EXISTS idx_shipment_line_items_shipment ON shipment_line_items(shipment_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_shipment_line_items_lot ON shipment_line_items(lot_number)")
            
            # Create lot_log table if it doesn't exist (for lot selection)
            cur.execute("""
                CREATE TABLE IF NOT EXISTS lot_log (
                    id SERIAL PRIMARY KEY,
                    lot_number TEXT NOT NULL UNIQUE,
                    sku TEXT NOT NULL,
                    correct_lot TEXT,
                    mfg_date TEXT,
                    exp_date TEXT,
                    total_units INTEGER,
                    updated_at TEXT
                )
            """)
            conn.commit()

        # Seed default rep for imports
        if not query_db("SELECT id FROM reps WHERE slug='ethan'", one=True):
            from werkzeug.security import generate_password_hash
            pw_hash = generate_password_hash("password")
            # Use convert_active_to_db_type to handle integer/boolean column type compatibility
            active_value = convert_active_to_db_type(True)  # Default rep is active
            execute_db(
                "INSERT INTO reps (name, slug, password_hash, email, active) VALUES (%s, %s, %s, %s, %s) ON CONFLICT (slug) DO NOTHING",
                ("Ethan (Test)", "ethan", pw_hash, "ethan@example.com", active_value),
            )
        
        # Ensure 'shipstation' rep exists for ShipStation-imported distributions
        if not query_db("SELECT id FROM reps WHERE slug='shipstation'", one=True):
            from werkzeug.security import generate_password_hash
            pw_hash = generate_password_hash(secrets.token_urlsafe(32))  # Random password (not used for login)
            # Use convert_active_to_db_type to handle integer/boolean column type compatibility
            active_value = convert_active_to_db_type(False)  # ShipStation rep is inactive (system use)
            execute_db(
                "INSERT INTO reps (name, slug, password_hash, email, active) VALUES (%s, %s, %s, %s, %s) ON CONFLICT (slug) DO NOTHING",
                ("ShipStation Imports", "shipstation", pw_hash, "shipstation@silq.tech", active_value),
            )

    def is_admin():
        return session.get("is_admin") is True

    def has_dist_records_import_token() -> bool:
        """Allow a headless importer to upload distribution-record attachments.

        This is intended for one-off migrations from a local folder into the production server's
        DIST_RECORDS_DIR, without requiring an interactive admin session.

        Configure via env var: DIST_RECORDS_IMPORT_TOKEN
        Use via request: /admin/distribution-records/upload?token=... (POST)
        """
        expected = (app.config.get("DIST_RECORDS_IMPORT_TOKEN") or "").strip()
        if not expected:
            return False
        provided = (request.args.get("token") or request.headers.get("X-Dist-Records-Token") or "").strip()
        if not provided:
            return False
        try:
            return hmac.compare_digest(provided, expected)
        except Exception:
            return provided == expected

    def reconcile_distribution_records_from_lotlog(*, only_unknown: bool = True, force: bool = False, limit: int | None = None) -> dict:
        """Re-apply lot_log mappings to existing distribution line-items.

        ShipStation sync (and other imports) may have created distribution records with missing/unknown SKU.
        Updating the lot_log does not automatically rewrite historical device_distribution_records rows.

        This function updates device_distribution_records.fields_json in-place.
        """

        def _row_value(row, key: str, index: int):
            try:
                if hasattr(row, "get"):
                    return row.get(key)
            except Exception:
                pass
            try:
                if hasattr(row, "keys"):
                    return row[key]
            except Exception:
                pass
            try:
                return row[index]
            except Exception:
                return None

        # Load the latest lot_log mapping into memory
        lot_rows = query_db(
            "SELECT lot_number, sku, correct_lot, mfg_date, exp_date FROM lot_log",
        ) or []
        lot_map = {}
        for r in lot_rows:
            ln = _row_value(r, "lot_number", 0)
            sku = _row_value(r, "sku", 1)
            correct = _row_value(r, "correct_lot", 2)
            mfg = _row_value(r, "mfg_date", 3)
            exp = _row_value(r, "exp_date", 4)

            key = (str(ln or "").strip().upper())
            if not key:
                continue
            lot_map[key] = {
                "sku": (str(sku or "").strip() if sku is not None else ""),
                "correct_lot": (str(correct or "").strip() if correct is not None else ""),
                "mfg_date": (str(mfg or "").strip() if mfg is not None else ""),
                "exp_date": (str(exp or "").strip() if exp is not None else ""),
            }

        updated = 0
        skipped = 0
        no_lot = 0
        no_mapping = 0
        errors = 0

        # Candidate scan: only rows that look like line-items (not shipment_record attachments)
        rows = query_db(
            """
            SELECT id, fields_json
            FROM device_distribution_records
            WHERE fields_json ILIKE '%Lot%'
            ORDER BY id ASC
            """
        ) or []

        for r in rows:
            if limit is not None and updated >= limit:
                break

            record_id = _row_value(r, "id", 0)
            fields_raw = _row_value(r, "fields_json", 1)
            try:
                f = json.loads(fields_raw or "{}")
            except Exception:
                errors += 1
                continue

            # Skip attachment records
            rt = (f.get("Record Type") or f.get("record_type") or f.get("type") or "").strip().lower()
            src = (f.get("Source") or f.get("source") or "").strip().lower()
            if rt in {"shipment_record", "distribution_record", "document", "attachment"} or src in {"shipment_record", "distribution_record"}:
                skipped += 1
                continue

            lot = (f.get("Lot") or f.get("lot") or "").strip()
            if not lot:
                no_lot += 1
                continue

            lot_key = lot.strip().upper()
            mapping = lot_map.get(lot_key)
            if not mapping:
                no_mapping += 1
                continue

            existing_sku = (f.get("SKU") or f.get("sku") or "").strip()
            existing_sku_norm = existing_sku.strip().lower()
            is_unknown = (not existing_sku) or existing_sku_norm in {"unknown", "-", "n/a", "na"}

            # Conservative by default: only fill in when unknown/missing
            if only_unknown and not is_unknown and not force:
                skipped += 1
                continue

            new_sku = (mapping.get("sku") or "").strip()
            if not new_sku:
                no_mapping += 1
                continue

            if not force and not is_unknown and existing_sku.strip().upper() == new_sku.strip().upper():
                skipped += 1
                continue

            f["SKU"] = new_sku

            # Optional helpful metadata (doesn't affect rollups)
            if mapping.get("correct_lot"):
                """Enqueue a master Sales Order PDF import.

                Dry run has been removed: every import is a real import.

                Auth: admin session OR DIST_RECORDS_IMPORT_TOKEN.
                Request: multipart/form-data
                  - file: PDF
                  - attach_all: optional (default 1)
                  - limit: optional int (max pages to process)
                """
                if not is_admin() and not has_dist_records_import_token():
                    abort(403)

                file = request.files.get("file")
                if not file or not getattr(file, "filename", None):
                    return make_response(json.dumps({"ok": False, "error": "No file"}), 400, {"Content-Type": "application/json"})

                attach_all = (request.form.get("attach_all") or request.args.get("attach_all") or "1").strip().lower() not in {"0", "false", "no", "n"}
                limit_raw = (request.form.get("limit") or request.args.get("limit") or "").strip()
                limit = None
                if limit_raw:
                    try:
                        limit = max(0, int(limit_raw))
                    except Exception:
                        limit = None

                ok, msg, http_status = _start_master_salesorder_import_job_from_filestorage(file, attach_all=attach_all, limit=limit)
                return make_response(json.dumps({"ok": ok, "message": msg}), http_status, {"Content-Type": "application/json"})
        if not customer_id or not note_text:
            return None
        note_date = datetime.utcnow().date()
        return execute_db(
            "INSERT INTO customer_notes (customer_id, author, note_text, note_date, created_at, updated_at) VALUES (%s, %s, %s, %s, %s, %s) RETURNING id",
            (customer_id, author or "admin", note_text.strip(), note_date, datetime.utcnow(), datetime.utcnow()),
            returning_id=True,
        )

    def get_sync_freshness():
        """Return persistent sync timestamps for dashboards (no in-memory globals)."""
        def _extract_scalar(row, key=None, idx=0):
            """
            Safely extract a scalar from a DB row that may be:
              - dict-like (dict, psycopg2 DictRow via dict(row), sqlite Row-ish)
              - tuple/list
              - scalar (datetime/string/None)
            """
            if row is None:
                return None

            # Prefer key lookup when provided
            if key is not None:
                try:
                    if hasattr(row, "get"):
                        v = row.get(key)
                        if v is not None:
                            return v
                except Exception:
                    pass
                try:
                    d = dict(row)
                    if key in d:
                        return d.get(key)
                except Exception:
                    pass
                if isinstance(row, (list, tuple)):
                    return row[idx] if len(row) > idx else None
                # Scalar row with a key requested: cannot extract reliably
                return None

            # No key requested: return first element from tuple/list, else return scalar as-is
            if isinstance(row, (list, tuple)):
                return row[idx] if len(row) > idx else None
            return row

        def _as_dt(val):
            if val is None:
                return None
            try:
                from datetime import datetime, timezone
                if isinstance(val, str):
                    dt = datetime.fromisoformat(val)
                else:
                    dt = val
                if dt and dt.tzinfo is None:
                    dt = dt.replace(tzinfo=timezone.utc)
                return dt
            except Exception:
                return None

        last_all_row = query_db("SELECT MAX(uploaded_at) AS ts FROM device_distribution_records", one=True)
        last_ship_row = query_db(
            """
            SELECT MAX(ddr.uploaded_at) AS ts
            FROM device_distribution_records ddr
            JOIN devices_distributed dd ON dd.id = ddr.dist_id
            WHERE dd.source = 'shipstation'
            """,
            one=True,
        )
        last_run_row = query_db(
            """
            SELECT ran_at, message FROM shipstation_sync_runs
            ORDER BY ran_at DESC
            LIMIT 1
            """,
            one=True,
        )

        last_all_updated = _as_dt(_extract_scalar(last_all_row, key="ts", idx=0))
        last_shipstation_updated = _as_dt(_extract_scalar(last_ship_row, key="ts", idx=0))
        last_sync_ran_at = _as_dt(_extract_scalar(last_run_row, key="ran_at", idx=0))
        last_sync_message = _extract_scalar(last_run_row, key="message", idx=1)

        dashboard_reflects_last_sync = bool(last_shipstation_updated and last_sync_ran_at and last_shipstation_updated >= last_sync_ran_at)

        return {
            "last_all_updated": last_all_updated,
            "last_shipstation_updated": last_shipstation_updated,
            "last_sync_ran_at": last_sync_ran_at,
            "last_sync_message": last_sync_message,
            "dashboard_reflects_last_sync": dashboard_reflects_last_sync,
        }

    def _normalize_ship_date_ymd(val):
        """
        Canonicalize a ship_date into a YYYY-MM-DD string for templates/sorting.
        Returns "" when missing/unparseable (callers can render "Unknown").
        """
        if val is None:
            return ""
        try:
            from datetime import datetime, date
            if isinstance(val, datetime):
                return val.strftime("%Y-%m-%d")
            if isinstance(val, date):
                return val.strftime("%Y-%m-%d")
        except Exception:
            pass
        try:
            if isinstance(val, str):
                s = val.strip()
                return s[:10] if len(s) >= 10 else s
        except Exception:
            return ""
        try:
            s = str(val).strip()
            return s[:10] if len(s) >= 10 else s
        except Exception:
            return ""
    
    def canonical_fields_schema():
        """
        Returns the canonical field names for fields_json.
        All data sources (CSV, DOCX, ShipStation, manual) 
        should map to these standard keys.
        """
        return {
            # Customer identity
            "Facility Name": "",
            "Company Key": "",
            
            # Address
            "Address1": "",
            "Address2": "",
            "City": "",
            "State": "",
            "Zip": "",
            "Country": "",
            
            # Contact
            "Contact Name": "",
            "Contact Phone": "",
            "Contact Email": "",
            
            # Distribution details
            "SKU": "",
            "Lot": "",
            "Quantity": "",
            "Distribution Date": "",
            "Order Number": "",
            "Tracking Number": "",
            
            # Source tracking
            "Source": ""  # 'shipstation', 'csv_import', 'manual', 'rep_upload'
        }
    
    def normalize_fields_json(fields_dict, source="manual"):
        """
        Maps various field name formats to canonical schema.
        Handles legacy field names from DOCX forms, CSV imports, etc.
        Returns dict with standardized keys.
        """
        canonical = canonical_fields_schema()
        
        # Common field mappings from legacy formats
        field_mappings = {
            # Legacy DOCX form fields → canonical
            "shipping Address Line 1": "Address1",
            "Shipping Address Line 1": "Address1",
            "Shipping address line 2 (if Applicable)": "Address2",
            "Address Line 2": "Address2",
            "Zip Code": "Zip",
            "Contact Phone Number": "Contact Phone",
            "contact email": "Contact Email",
            "Contact Email": "Contact Email",
            
            # CSV import fields → canonical
            "BillTo Name": "Contact Name",
            "BillTo Phone": "Contact Phone",
            "ShipToAddr1": "Address1",
            "ShipToAddr2": "Address2",
            "ShipToCity": "City",
            "ShipToState": "State",
            "ShipToPostal": "Zip",
            "ShipToCountry": "Country",
            
            # Already canonical (passthrough)
            "Facility Name": "Facility Name",
            "Address1": "Address1",
            "City": "City",
            "State": "State",
            "SKU": "SKU",
            "Lot": "Lot",
            "Quantity": "Quantity",
            "Order Number": "Order Number"
        }
        
        # Apply mappings
        for old_key, canonical_key in field_mappings.items():
            if old_key in fields_dict:
                canonical[canonical_key] = fields_dict[old_key]
        
        # Copy any unmapped fields (for debugging)
        for key, value in fields_dict.items():
            if key not in field_mappings and key not in canonical:
                canonical[key] = value  # Preserve unknown fields
        
        # Set source
        canonical["Source"] = source
        
        # Auto-compute Company Key if Facility Name present
        if canonical["Facility Name"]:
            canonical["Company Key"] = canonical_customer_key(canonical["Facility Name"])
        
        return canonical

    def fetch_distribution_records(
        exclude_sources=None,
        *,
        customer_ids: list[int] | None = None,
        date_from: str | None = None,
        date_to: str | None = None,
    ):
        """Fetch distribution rows and normalize customer identity consistently.

        This is the canonical loader used by dashboards/logs/CRM so order counts and units stay aligned.

        Args:
            exclude_sources: iterable of source strings (lowercased compare).
            customer_ids: optional list of customer IDs to restrict results.
            date_from/date_to: optional YYYY-MM-DD bounds (filters on dd.ship_date).
        """
        where = []
        params = []
        if customer_ids:
            where.append("COALESCE(dd.customer_id, ddr.customer_id) = ANY(%s)")
            params.append(customer_ids)
        if date_from:
            where.append("dd.ship_date IS NOT NULL AND dd.ship_date >= %s")
            params.append(date_from)
        if date_to:
            where.append("dd.ship_date IS NOT NULL AND dd.ship_date <= %s")
            params.append(date_to)
        where_sql = ("WHERE " + " AND ".join(where)) if where else ""

        rows = query_db(
            f"""
            SELECT dd.id,
                   dd.distribution_number,
                   dd.ship_date,
                   dd.order_number,
                   dd.source,
                   dd.rep_id,
                   dd.ss_shipment_id,
                   dd.tracking_number,
                   dd.customer_id AS dist_customer_id,
                   ddr.customer_id AS record_customer_id,
                   ddr.id AS record_id,
                   r.name AS rep_name,
                   ddr.fields_json,
                   c.id AS customer_id,
                   c.company_key AS company_key,
                   c.facility_name AS customer_name,
                   c.address1 AS customer_addr1,
                   c.address2 AS customer_addr2,
                   c.city AS customer_city,
                   c.state AS customer_state,
                   c.zip AS customer_zip
            FROM devices_distributed dd
            JOIN device_distribution_records ddr ON ddr.dist_id = dd.id
            LEFT JOIN reps r ON r.id = dd.rep_id
            LEFT JOIN customers c ON c.id = COALESCE(dd.customer_id, ddr.customer_id)
            {where_sql}
            """,
            tuple(params),
        ) or []

        excluded = {s.lower() for s in (exclude_sources or set())}
        normalized = []

        def _is_shipment_record_fields(fields: dict) -> bool:
            if not isinstance(fields, dict):
                return False
            rt = (fields.get("Record Type") or fields.get("record_type") or fields.get("type") or "").strip().lower()
            src = (fields.get("Source") or fields.get("source") or "").strip().lower()
            if rt in {"shipment_record", "distribution_record", "document", "attachment"}:
                return True
            if src in {"shipment_record", "distribution_record"}:
                return True
            return False

        for r in rows:
            source = (r.get("source") or "").lower()
            if excluded and source in excluded:
                continue
            try:
                f = json.loads(r["fields_json"])
            except Exception:
                f = {}

            # Attachment / evidence records live in device_distribution_records as well.
            # Skip these so dashboards/logs don't treat them as SKU/lot line items.
            if _is_shipment_record_fields(f):
                continue

            customer_id = r.get("customer_id") or r.get("dist_customer_id") or r.get("record_customer_id")
            facility_name = (r.get("customer_name") or f.get("Facility Name") or "Unknown").strip()
            city = (r.get("customer_city") or f.get("City") or "").strip()
            state = (r.get("customer_state") or f.get("State") or "").strip()
            addr1 = (r.get("customer_addr1") or f.get("Address1") or "").strip()
            addr2 = (r.get("customer_addr2") or f.get("Address2") or "").strip()
            postal = (r.get("customer_zip") or f.get("Zip") or "").strip()

            if not customer_id and facility_name:
                cust_row = normalize_row(find_or_create_customer(
                    facility_name=facility_name,
                    city=city,
                    state=state,
                    address1=addr1,
                    address2=addr2,
                    zip_code=postal,
                    contact_name=f.get("Contact Name", ""),
                    contact_phone=f.get("Contact Phone", ""),
                    contact_email=f.get("Contact Email", ""),
                ))
                if cust_row:
                    customer_id = cust_row.get("id") if isinstance(cust_row, dict) else cust_row
                    try:
                        execute_db(
                            "UPDATE devices_distributed SET customer_id = %s WHERE id = %s",
                            (customer_id, r.get("id")),
                        )
                        execute_db(
                            "UPDATE device_distribution_records SET customer_id = %s WHERE id = %s",
                            (customer_id, r.get("record_id")),
                        )
                        execute_db(
                            "UPDATE new_customer_records SET customer_id = %s WHERE dist_id = %s",
                            (customer_id, r.get("id")),
                        )
                    except Exception:
                        pass

            company_key_val = r.get("company_key") or canonical_customer_key(facility_name)
            ship_date = r.get("ship_date") or ""
            dist_num = (r.get("distribution_number") or "").strip()
            logical_order_number = r.get("order_number") or dist_num or f"DIST-{r.get('id')}"
            order_key = canonical_order_key(logical_order_number, ship_date)
            order_group = canonical_order_number(logical_order_number)
            shipment_id = r.get("ss_shipment_id")
            shipment_key = shipment_id or f"{order_group}::{ship_date or ''}::{r.get('id')}"
            qty = int(float(f.get("Quantity", 0))) if f.get("Quantity") else 0

            normalized.append(
                {
                    "dist_id": r.get("id"),
                    "distribution_number": dist_num,
                    "customer_id": customer_id,
                    "rep_id": r.get("rep_id"),
                    "rep_name": r.get("rep_name") or "Unassigned",
                    "order_key": order_key,
                    "order_group": order_group,
                    "order_number": logical_order_number,
                    "ship_date": ship_date,
                    "shipment_id": shipment_id,
                    "shipment_key": shipment_key,
                    "source": source or "manual",
                    "facility_key": company_key_val,
                    "facility_label": facility_name or "Unknown",
                    "addr1": addr1.upper(),
                    "city": city.upper(),
                    "state": state.upper(),
                    "postal": postal,
                    "qty": qty,
                    "sku": f.get("SKU", "Unknown"),
                    "lot": f.get("Lot", "-"),
                    "tracking_number": r.get("tracking_number"),
                    "fields": f,
                }
            )

        return normalized

    def _normalize_ship_date(value: str | None) -> str:
        if not value:
            return ""
        s = str(value)
        return s[:10]

    def _dist_records_root() -> Path:
        return Path(app.config.get("DIST_RECORDS_DIR") or (Path(app.config["UPLOAD_ROOT"]) / "distribution_records"))

    def _dist_records_s3_enabled() -> bool:
        return bool((app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip())

    def _normalize_s3_endpoint_url(raw: str | None) -> str | None:
        """Normalize S3 endpoint URL.

        Users sometimes paste a Spaces URL including the bucket (e.g.
        https://sfo3.digitaloceanspaces.com/<bucket>) or with a common typo
        in the domain (digitaloceanspaces.co). Boto3 expects endpoint_url to
        be the service root (no bucket/path).
        """
        if not raw:
            return None
        s = str(raw).strip()
        if not s:
            return None
        if not re.match(r"^https?://", s, flags=re.IGNORECASE):
            s = "https://" + s

        try:
            from urllib.parse import urlparse, urlunparse

            u = urlparse(s)
            netloc = (u.netloc or "").strip()
            scheme = u.scheme or "https"

            # Fix common Spaces domain typo.
            netloc = re.sub(r"digitaloceanspaces\.co$", "digitaloceanspaces.com", netloc, flags=re.IGNORECASE)

            # Boto3 endpoint_url should not include any path (bucket/prefix).
            return urlunparse((scheme, netloc, "", "", "", ""))
        except Exception:
            # Fall back to the raw string; boto3 will raise a clear error.
            return s

    def _dist_records_s3_client():
        # Lazy import so local/dev doesn't require boto3 unless S3 is enabled.
        import boto3
        from botocore.config import Config

        region = (app.config.get("DIST_RECORDS_S3_REGION") or "").strip() or None
        endpoint_url = _normalize_s3_endpoint_url(app.config.get("DIST_RECORDS_S3_ENDPOINT_URL"))
        access_key = os.environ.get("AWS_ACCESS_KEY_ID") or os.environ.get("BUCKETEER_AWS_ACCESS_KEY_ID")
        secret_key = os.environ.get("AWS_SECRET_ACCESS_KEY") or os.environ.get("BUCKETEER_AWS_SECRET_ACCESS_KEY")
        session_token = os.environ.get("AWS_SESSION_TOKEN") or os.environ.get("BUCKETEER_AWS_SESSION_TOKEN")

        kwargs = {}
        if region:
            kwargs["region_name"] = region
        if endpoint_url:
            kwargs["endpoint_url"] = endpoint_url
        if access_key and secret_key:
            kwargs["aws_access_key_id"] = access_key
            kwargs["aws_secret_access_key"] = secret_key
            if session_token:
                kwargs["aws_session_token"] = session_token

        # Keep retries/timeouts bounded so a single slow S3 call doesn't hang the request
        # long enough to trigger a Gunicorn worker timeout.
        kwargs["config"] = Config(
            connect_timeout=5,
            read_timeout=60,
            retries={"max_attempts": 3, "mode": "standard"},
        )

        return boto3.client("s3", **kwargs)

    def _dist_records_s3_key(stored_rel: str) -> str:
        prefix = (app.config.get("DIST_RECORDS_S3_PREFIX") or "").strip().strip("/")
        rel = (stored_rel or "").lstrip("/")
        return f"{prefix}/{rel}" if prefix else rel

    def dist_record_blob_exists(stored_filename: str) -> bool:
        if not stored_filename:
            return False
        if _dist_records_s3_enabled():
            try:
                bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
                if not bucket:
                    return False
                key = _dist_records_s3_key(stored_filename)
                client = _dist_records_s3_client()
                client.head_object(Bucket=bucket, Key=key)
                return True
            except Exception:
                return False

        try:
            p = _safe_dist_record_path(stored_filename)
            return p.exists() and p.is_file()
        except Exception:
            return False

    def save_distribution_record_file_to(stored_rel: str, file_storage) -> None:
        """Write the uploaded file to an explicit stored filename (used for repairs).

        When S3/Spaces is enabled, this overwrites the object at the computed key.
        When using filesystem storage, this writes into DIST_RECORDS_DIR.
        """
        if not file_storage or not getattr(file_storage, "filename", None):
            raise ValueError("No file provided")
        if not stored_rel:
            raise ValueError("Missing stored filename")

        if _dist_records_s3_enabled():
            bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
            if not bucket:
                raise RuntimeError("DIST_RECORDS_S3_BUCKET is not set")
            key = _dist_records_s3_key(stored_rel)
            client = _dist_records_s3_client()
            try:
                try:
                    file_storage.stream.seek(0)
                except Exception:
                    pass
                extra = {}
                ct = getattr(file_storage, "mimetype", None) or ""
                if ct:
                    extra["ContentType"] = ct
                if extra:
                    client.upload_fileobj(file_storage.stream, bucket, key, ExtraArgs=extra)
                else:
                    client.upload_fileobj(file_storage.stream, bucket, key)
            except Exception as e:
                raise RuntimeError(f"S3 upload failed: {e}")
            return

        # Filesystem
        root = _dist_records_root()
        root.mkdir(parents=True, exist_ok=True)
        dest_path = root / stored_rel
        # Enforce sandboxing
        _ = _safe_dist_record_path(stored_rel)
        dest_path.parent.mkdir(parents=True, exist_ok=True)
        file_storage.save(dest_path)

    def save_distribution_record_bytes_to(
        stored_rel: str,
        data: bytes,
        *,
        content_type: str = "application/pdf",
    ) -> None:
        """Write bytes to an explicit stored filename.

        Mirrors save_distribution_record_file_to but for callers that generate PDFs in-memory.
        """
        if not stored_rel:
            raise ValueError("Missing stored filename")
        if data is None:
            raise ValueError("No data provided")

        if _dist_records_s3_enabled():
            bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
            if not bucket:
                raise RuntimeError("DIST_RECORDS_S3_BUCKET is not set")
            key = _dist_records_s3_key(stored_rel)
            client = _dist_records_s3_client()
            try:
                import io

                bio = io.BytesIO(data or b"")
                try:
                    bio.seek(0)
                except Exception:
                    pass
                extra = {}
                if content_type:
                    extra["ContentType"] = content_type
                if extra:
                    client.upload_fileobj(bio, bucket, key, ExtraArgs=extra)
                else:
                    client.upload_fileobj(bio, bucket, key)
            except Exception as e:
                raise RuntimeError(f"S3 upload failed: {e}")
            return

        root = _dist_records_root()
        root.mkdir(parents=True, exist_ok=True)
        dest_path = root / stored_rel
        _ = _safe_dist_record_path(stored_rel)
        dest_path.parent.mkdir(parents=True, exist_ok=True)
        with dest_path.open("wb") as fp:
            fp.write(data)

    def delete_distribution_record_blob(stored_filename: str) -> None:
        if not stored_filename:
            return
        if _dist_records_s3_enabled():
            try:
                bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
                if not bucket:
                    return
                key = _dist_records_s3_key(stored_filename)
                client = _dist_records_s3_client()
                client.delete_object(Bucket=bucket, Key=key)
            except Exception:
                pass
            return

        try:
            p = _safe_dist_record_path(stored_filename)
            if p.exists() and p.is_file():
                p.unlink()
        except Exception:
            pass

    def save_distribution_record_file(dist_id: int, file_storage) -> dict:
        """Save a distribution record attachment into DIST_RECORDS_DIR.

        Returns: {stored_filename, original_filename}
        where stored_filename is a path relative to DIST_RECORDS_DIR.
        """
        if not file_storage or not getattr(file_storage, "filename", None):
            raise ValueError("No file provided")

        original = secure_filename(file_storage.filename)
        if not original:
            original = "upload"

        token = uuid.uuid4().hex[:10]
        stored_rel = f"{int(dist_id)}/{token}_{original}"

        # If configured, store in S3 (recommended on Heroku).
        if _dist_records_s3_enabled():
            bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
            if not bucket:
                raise RuntimeError("DIST_RECORDS_S3_BUCKET is not set")
            key = _dist_records_s3_key(stored_rel)
            client = _dist_records_s3_client()
            try:
                try:
                    file_storage.stream.seek(0)
                except Exception:
                    pass
                extra = {}
                ct = getattr(file_storage, "mimetype", None) or ""
                if ct:
                    extra["ContentType"] = ct
                if extra:
                    client.upload_fileobj(file_storage.stream, bucket, key, ExtraArgs=extra)
                else:
                    client.upload_fileobj(file_storage.stream, bucket, key)
            except Exception as e:
                raise RuntimeError(f"S3 upload failed: {e}")
        else:
            root = _dist_records_root()
            root.mkdir(parents=True, exist_ok=True)
            dest_path = (root / stored_rel)
            dest_path.parent.mkdir(parents=True, exist_ok=True)
            file_storage.save(dest_path)

        return {"stored_filename": stored_rel, "original_filename": original}

    def save_distribution_record_bytes(
        dist_id: int,
        original_filename: str,
        data: bytes,
        *,
        content_type: str = "application/pdf",
    ) -> dict:
        """Save a distribution record PDF generated in-memory.

        Returns: {stored_filename, original_filename}
        where stored_filename is a path relative to DIST_RECORDS_DIR.
        """
        if data is None:
            raise ValueError("No data provided")

        original = secure_filename(original_filename or "")
        if not original:
            original = "upload.pdf"
        if not original.lower().endswith(".pdf"):
            original = f"{original}.pdf"

        token = uuid.uuid4().hex[:10]
        stored_rel = f"{int(dist_id)}/{token}_{original}"
        save_distribution_record_bytes_to(stored_rel, data, content_type=content_type)
        return {"stored_filename": stored_rel, "original_filename": original}

    def _safe_dist_record_path(stored_filename: str) -> Path:
        root = _dist_records_root().resolve()
        candidate = (root / (stored_filename or "")).resolve()
        if root != candidate and root not in candidate.parents:
            raise ValueError("Invalid stored filename")
        return candidate

    def ensure_safe_path(base_dir: Path, filename: str) -> Path:
        """
        Ensure a file path is safe for saving (prevents path traversal).
        
        Args:
            base_dir: Base directory where file should be saved
            filename: Filename (will be sanitized with secure_filename)
        
        Returns:
            Safe Path object within base_dir
        
        Raises:
            ValueError: If path would escape base_dir
        """
        base_dir = Path(base_dir).resolve()
        safe_filename = secure_filename(filename)
        if not safe_filename:
            safe_filename = "upload"
        
        candidate = (base_dir / safe_filename).resolve()
        
        # Ensure candidate is within base_dir (prevent path traversal)
        try:
            candidate.relative_to(base_dir)
        except ValueError:
            raise ValueError(f"Invalid filename: {filename} (would escape base directory)")
        
        return candidate

    def _source_label(sources: set[str]):
        # Keep display stable and human-friendly
        cleaned = sorted({(s or "").strip().lower() for s in (sources or set()) if s})
        if not cleaned:
            return ""
        mapping = {
            "shipstation": "ShipStation",
            "manual": "Manual",
            "rep_upload": "Rep Upload",
            "csv_import": "CSV Import",
        }
        return ", ".join(mapping.get(s, s) for s in cleaned)

    def build_order_group_summaries(rows: list[dict]):
        """Aggregate canonical rows into order_group-level summaries.

        Groups by (customer_id, order_group) to safely handle identical order numbers across customers.
        Returns: dict[(customer_id, order_group)] -> summary dict.
        """
        out = {}
        for r in rows or []:
            if not isinstance(r, dict):
                continue
            customer_id = r.get("customer_id")
            order_group = r.get("order_group") or canonical_order_number(r.get("order_number") or "")
            if not customer_id or not order_group:
                continue

            key = (customer_id, order_group)
            rec = out.get(key)
            if not rec:
                rec = {
                    "customer_id": customer_id,
                    "order_group": order_group,
                    "order_number": r.get("order_number") or order_group,
                    "dist_ids": set(),
                    "sources": set(),
                    "ship_dates": [],
                    "sku_units": {},
                    "total_units": 0,
                }
                out[key] = rec

            if r.get("dist_id"):
                rec["dist_ids"].add(r.get("dist_id"))
            rec["sources"].add(r.get("source") or "")

            sd = _normalize_ship_date(r.get("ship_date"))
            if sd:
                rec["ship_dates"].append(sd)

            sku = (r.get("sku") or "Unknown").strip() or "Unknown"
            qty = r.get("qty") or 0
            try:
                qty_int = int(qty)
            except Exception:
                qty_int = 0

            rec["sku_units"][sku] = int(rec["sku_units"].get(sku, 0)) + qty_int
            rec["total_units"] += qty_int

        # finalize computed fields
        for rec in out.values():
            dates = [d for d in rec.get("ship_dates") or [] if d]
            rec["first_ship_date"] = min(dates) if dates else ""
            rec["last_ship_date"] = max(dates) if dates else ""
            sku_breakdown = [
                {"sku": sku, "units": units}
                for sku, units in sorted((rec.get("sku_units") or {}).items(), key=lambda kv: kv[1], reverse=True)
            ]
            rec["sku_breakdown"] = sku_breakdown
            rec["sku_count"] = len(sku_breakdown)
            rec["source_label"] = _source_label(rec.get("sources") or set())
            rec["sku_summary"] = ", ".join(
                f"{x['sku']} ({x['units']})" for x in (sku_breakdown[:4] if sku_breakdown else [])
            )
            if sku_breakdown and len(sku_breakdown) > 4:
                rec["sku_summary"] += ", …"

            # Convert set for safe template iteration
            rec["dist_ids"] = sorted([d for d in rec.get("dist_ids") or set() if d])

        return out

    def clean_company_name(name):
        """Legacy wrapper - use canonical_customer_key instead"""
        return canonical_customer_key(name)

    def canonical_customer_key(name: str | None) -> str:
        """Canonical customer key.

        This intentionally matches the ShipStation sync's notion of a company key
        so imports/CRM grouping stay consistent.
        """
        return normalize_company_key(name or "")

    def canonical_order_number(order_number: str | None) -> str:
        if not order_number:
            return ""
        return re.sub(r"[^A-Z0-9]", "", str(order_number).upper())

    def canonical_order_key(order_number: str | None, ship_date: str | None) -> str:
        return f"{canonical_order_number(order_number)}::{_normalize_ship_date(ship_date)}"

    def _maybe_set_helper(existing: dict, sets: list[str], vals: list[object], col: str, value: str | None) -> None:
        """Helper function to conditionally add a column update to sets/vals if value differs from existing."""
        v = (value or "").strip()
        if not v:
            return
        if (existing.get(col) or "").strip() == v:
            return
        sets.append(f"{col} = %s")
        vals.append(v)

    def find_or_create_customer(
        *,
        facility_name: str | None,
        city: str | None = None,
        state: str | None = None,
        address1: str | None = None,
        address2: str | None = None,
        zip_code: str | None = None,
        contact_name: str | None = None,
        contact_phone: str | None = None,
        contact_email: str | None = None,
        primary_rep_id: int | None = None,
    ):
        """Create or update a customer row keyed by canonical_customer_key(facility_name)."""
        company_key = canonical_customer_key(facility_name)
        if not company_key:
            return None

        # Look up existing
        existing = normalize_row(query_db("SELECT * FROM customers WHERE company_key = %s", (company_key,), one=True))
        if existing:
            sets = []
            vals = []

            _maybe_set_helper(existing, sets, vals, "facility_name", facility_name)
            _maybe_set_helper(existing, sets, vals, "city", city)
            _maybe_set_helper(existing, sets, vals, "state", state)
            _maybe_set_helper(existing, sets, vals, "address1", address1)
            _maybe_set_helper(existing, sets, vals, "address2", address2)
            _maybe_set_helper(existing, sets, vals, "zip", zip_code)
            _maybe_set_helper(existing, sets, vals, "contact_name", contact_name)
            _maybe_set_helper(existing, sets, vals, "contact_phone", contact_phone)
            _maybe_set_helper(existing, sets, vals, "contact_email", contact_email)

            if primary_rep_id and not existing.get("primary_rep_id"):
                sets.append("primary_rep_id = %s")
                vals.append(int(primary_rep_id))

            if sets:
                sets.append("updated_at = CURRENT_TIMESTAMP")
                vals.append(existing.get("id"))
                execute_db(f"UPDATE customers SET {', '.join(sets)} WHERE id = %s", tuple(vals))
                existing = normalize_row(query_db("SELECT * FROM customers WHERE id = %s", (existing.get("id"),), one=True))
            return existing

        # Insert new
        cust_id = execute_db(
            """
            INSERT INTO customers (company_key, facility_name, address1, address2, city, state, zip, contact_name, contact_phone, contact_email, primary_rep_id)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
            """,
            (
                company_key,
                (facility_name or "Unknown").strip() or "Unknown",
                (address1 or "").strip() or None,
                (address2 or "").strip() or None,
                (city or "").strip() or None,
                (state or "").strip() or None,
                (zip_code or "").strip() or None,
                (contact_name or "").strip() or None,
                (contact_phone or "").strip() or None,
                (contact_email or "").strip() or None,
                int(primary_rep_id) if primary_rep_id else None,
            ),
            returning_id=True,
        )
        if not cust_id:
            return None
        return normalize_row(query_db("SELECT * FROM customers WHERE id = %s", (cust_id,), one=True))

    def ensure_rep_assignment(customer_id: int | None, rep_id: int | None, *, make_primary_if_none: bool = False) -> None:
        if not customer_id or not rep_id:
            return

        # Ensure assignment row exists
        execute_db(
            """
            INSERT INTO customer_rep_assignments (customer_id, rep_id, is_primary)
            VALUES (%s, %s, FALSE)
            ON CONFLICT (customer_id, rep_id) DO NOTHING
            """,
            (int(customer_id), int(rep_id)),
        )

        if make_primary_if_none:
            cust = normalize_row(query_db("SELECT id, primary_rep_id FROM customers WHERE id = %s", (int(customer_id),), one=True))
            if cust and not cust.get("primary_rep_id"):
                execute_db("UPDATE customers SET primary_rep_id = %s, updated_at=CURRENT_TIMESTAMP WHERE id = %s", (int(rep_id), int(customer_id)))

        # Sync is_primary flag to match customers.primary_rep_id
        cust = normalize_row(query_db("SELECT primary_rep_id FROM customers WHERE id = %s", (int(customer_id),), one=True))
        primary_rep_id = cust.get("primary_rep_id") if cust else None
        if primary_rep_id:
            execute_db("UPDATE customer_rep_assignments SET is_primary = (rep_id = %s) WHERE customer_id = %s", (int(primary_rep_id), int(customer_id)))

    def pick_rep_for_customer(customer_id: int | None, *, fallback_rep_id: int | None = None) -> int | None:
        if not customer_id:
            return fallback_rep_id
        cust = normalize_row(query_db("SELECT primary_rep_id FROM customers WHERE id = %s", (int(customer_id),), one=True))
        if cust and cust.get("primary_rep_id"):
            return int(cust.get("primary_rep_id"))
        row = normalize_row(
            query_db(
                """
                SELECT rep_id
                FROM customer_rep_assignments
                WHERE customer_id = %s
                ORDER BY is_primary DESC, rep_id ASC
                LIMIT 1
                """,
                (int(customer_id),),
                one=True,
            )
        )
        if row and row.get("rep_id"):
            return int(row.get("rep_id"))
        return fallback_rep_id

    def ensure_rep_documents_table():
        """Ensure the rep_documents table exists for document uploads"""
        execute_db(
            """
            CREATE TABLE IF NOT EXISTS rep_documents (
                id SERIAL PRIMARY KEY,
                rep_id INTEGER NOT NULL REFERENCES reps(id),
                doc_type TEXT NOT NULL,
                stored_filename TEXT NOT NULL,
                original_filename TEXT NOT NULL,
                uploaded_at TIMESTAMP NOT NULL,
                uploaded_by TEXT NOT NULL,
                order_number TEXT,
                lot_number TEXT,
                distribution_id INTEGER REFERENCES devices_distributed(id),
                notes TEXT,
                match_reason TEXT
            )
            """
        )
        # Add match_reason column if it doesn't exist (for existing tables)
        try:
            execute_db("ALTER TABLE rep_documents ADD COLUMN IF NOT EXISTS match_reason TEXT")
        except Exception:
            pass  # Column may already exist or table may not exist yet
    
    def link_document_to_distribution(order_number: str, rep_id: int) -> tuple[int | None, str | None]:
        """
        Attempt to link a document to a distribution record by matching order_number.
        
        Args:
            order_number: Order number to match (will be normalized)
            rep_id: Rep ID to filter matches
            
        Returns:
            Tuple of (distribution_id, match_reason) or (None, None) if no match found
        """
        if not order_number or not rep_id:
            return None, None
        
        # Normalize order_number: strip whitespace, uppercase for case-insensitive matching
        normalized_order = (order_number or "").strip().upper()
        if not normalized_order:
            return None, None
        
        # Query for matching distributions
        # Use ILIKE for case-insensitive matching (PostgreSQL) or LIKE with UPPER
        matches = query_db(
            """
            SELECT id, order_number, ship_date, source
            FROM devices_distributed
            WHERE rep_id = %s 
              AND UPPER(TRIM(COALESCE(order_number, ''))) = %s
            ORDER BY ship_date DESC NULLS LAST, id DESC
            LIMIT 1
            """,
            (rep_id, normalized_order)
        ) or []
        
        if matches:
            match = normalize_row(matches[0])
            dist_id = match["id"]
            source = match.get("source", "unknown")
            match_reason = f"Matched by order_number '{order_number}' (normalized: '{normalized_order}') from {source} source"
            return dist_id, match_reason
        
        return None, None
    
    def get_rep_by_slug(slug: str):
        return query_db("SELECT * FROM reps WHERE slug = %s AND active = 1", (slug.lower(),), one=True)
    
    def get_logged_in_rep_slug():
        return session.get("rep_slug")

    def find_rep_by_identifier(identifier: str | None):
        if not identifier:
            return None
        ident = identifier.strip().lower()
        return query_db(
            "SELECT * FROM reps WHERE active = 1 AND (slug = %s OR LOWER(email) = %s)",
            (ident, ident),
            one=True,
        )

    def generate_simple_rep_password(name: str | None, slug: str | None) -> str:
        """Create a predictable, easy-to-recall password based on the rep identity."""
        base = (slug or name or "rep").lower()
        base = "".join(ch for ch in base if ch.isalnum()) or "rep"
        # Keep it short and human friendly while ensuring a minimum length
        base = base[:8] if len(base) >= 6 else (base + "portal")[:6]
        word = base.capitalize()
        suffix = "123!"
        candidate = f"{word}{suffix}"
        if len(candidate) < 8:
            candidate = f"{word}2024!"
        return candidate
    
    def get_current_rep():
        """Get the currently logged-in rep object from session using rep_id when available"""
        rep_slug = session.get("rep_slug")
        rep_id = session.get("rep_id")
        if not rep_slug:
            return None
        rep = get_rep_by_slug(rep_slug)
        rep = normalize_row(rep)
        if rep and rep_id and rep.get("id") != rep_id:
            # Session mismatch; clear stale session to force re-auth
            session.pop("rep_slug", None)
            session.pop("rep_id", None)
            return None
        return rep

    def normalize_row(row):
        """Normalize psycopg2 rows to plain dicts.

        - DictCursor returns DictRow, which is list-like (so naive list unwrapping breaks it).
        - Some callers may accidentally pass a one-element list of rows.
        """
        if row is None:
            return None

        # If this is a list of rows (common for query_db(one=False)), keep it as-is.
        # If it's a one-element list, unwrap.
        if isinstance(row, list) and not hasattr(row, "keys"):
            if len(row) == 1:
                return normalize_row(row[0])
            return row

        # DictRow (and similar) provides .keys() but isn't a dict.
        if hasattr(row, "keys"):
            try:
                return {k: row[k] for k in row.keys()}
            except Exception:
                pass

        return row
    
    def check_rep_auth(slug):
        """Check if current session has access to the given rep slug and cache rep in g"""
        slug = slug.lower()
        rep = normalize_row(get_rep_by_slug(slug))
        if rep is None:
            abort(404)
        if is_admin():
            g.current_rep = rep
            return None  # Admins can view any rep page
        current_slug = session.get("rep_slug")
        current_id = session.get("rep_id")
        if current_slug and current_slug.lower() == slug and current_id == rep.get("id"):
            g.current_rep = rep
            return None  # Rep can access their own page
        flash("Please log in to access this portal.", "warning")
        return redirect(url_for("rep_login", slug=slug, next=request.url))

    from functools import wraps

    def require_rep(view_fn):
        """Decorator enforcing rep authentication (or admin override) for rep routes"""
        @wraps(view_fn)
        def wrapper(slug, *args, **kwargs):
            auth = check_rep_auth(slug)
            if auth:
                return auth
            return view_fn(slug.lower(), *args, **kwargs)
        return wrapper
    
    def normalize_phone(value: str) -> str:
        if not value:
            return ""
        digits = "".join(ch for ch in value if ch.isdigit())
        if len(digits) == 10:
            return f"({digits[0:3]}) {digits[3:6]}-{digits[6:10]}"
        if len(digits) == 11 and digits[0] == "1":
            return f"({digits[1:4]}) {digits[4:7]}-{digits[7:11]}"
        return value.strip()

    def docx_label_value_pairs(doc_path: Path):
        from docx import Document
        fields = {}
        doc = Document(str(doc_path))
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if len(cells) >= 2:
                    label, value = cells[0], cells[1]
                    if label:
                        fields[label] = value
        for para in doc.paragraphs:
            text = para.text.strip()
            if ":" in text:
                label, value = text.split(":", 1)
                label = label.strip()
                if label:
                    fields[label] = value.strip()
        return fields

    def parse_new_customer_docx(doc_path: Path):
        raw_fields = docx_label_value_pairs(doc_path)
        out_fields = {}
        for label, value in raw_fields.items():
            if "phone" in label.lower():
                out_fields[label] = normalize_phone(value)
            else:
                out_fields[label] = value
        expected = ["Facility Name", "shipping Address Line 1", "City", "State", "Zip Code", "Contact name", "Contact Phone Number"]
        missing_required = [e for e in expected if e not in out_fields or not str(out_fields[e]).strip()]
        unexpected = [k for k in out_fields.keys() if k not in expected]
        return out_fields, missing_required, unexpected

    def parse_device_distribution_docx(doc_path: Path):
        raw_pairs = docx_label_value_pairs(doc_path)
        
        def canon_key(label: str) -> str:
            ll = label.lower()
            if "facility" in ll and "name" in ll:
                return "Facility Name"
            if "facility" in ll and "city" in ll:
                return "City"
            if "facility" in ll and "state" in ll:
                return "State"
            if "device" in ll:
                return "Device"
            if "sku" in ll:
                return "SKU"
            if "lot" in ll:
                return "Lot"
            if "qty" in ll or "quantity" in ll:
                return "Quantity"
            if "distribution" in ll and "date" in ll:
                return "Distribution Date"
            if "recipient" in ll and "name" in ll:
                return "Contact name"
            if "recipient" in ll and "email" in ll:
                return "contact email"
            if "address" in ll:
                return "Address1"
            if "zip" in ll or "postal" in ll:
                return "Zip"
            return label
            
        out = {}
        for k, v in raw_pairs.items():
            key = canon_key(k)
            if "phone" in key.lower():
                v = normalize_phone(v)
            out[key] = v
            
        EXPECTED = ["Facility Name", "Device", "Lot", "Quantity", "Distribution Date"]
        missing = [f for f in EXPECTED if f not in out or not str(out[f]).strip()]
        unexpected = [k for k in out.keys() if k not in EXPECTED and k not in ["City", "State", "Contact name", "contact email", "Address1", "Zip"]]
        
        return out, missing, unexpected

    def parse_tracing_approval_docx(doc_path: Path):
        fields = docx_label_value_pairs(doc_path)
        expected = ["Tracing Report number", "Tracing report approved (Yes/no)", "Distribution Rep Name", "Distribution Rep signature", "Date"]
        missing = [e for e in expected if e not in fields or not str(fields[e]).strip()]
        unexpected = [k for k in fields.keys() if k not in expected]
        return fields, missing, unexpected

    def append_distribution_logs(rep_row, shipment_row, dist_id: int):
        import shutil
        rep_slug = rep_row["slug"]
        master_dir = Path(app.config["UPLOAD_ROOT"]) / "distributed" / rep_slug / "master"
        master_dir.mkdir(parents=True, exist_ok=True)
        master_path = master_dir / "Distribution Log.xlsx"
        shipment_dir = Path(app.config["UPLOAD_ROOT"]) / "received" / rep_slug / str(shipment_row["id"])
        shipment_dir.mkdir(parents=True, exist_ok=True)
        shipment_path = shipment_dir / "Distribution Log.xlsx"
        today_str = datetime.utcnow().strftime("%Y-%m-%d")
        order_no = f"AUTO-{dist_id}"
        row = ["SEE FORMS", order_no, today_str, None, "", "", "", "", "", "", "", "", "", "", "", "", "", "", "", ""]

        def ensure_and_append(path: Path):
            if not path.exists():
                template = Path(app.config.get("FORM_TEMPLATES_DIR", BASE_DIR / "form_templates")) / "Distribution Log.xlsx"
                if template.exists():
                    path.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copyfile(template, path)
                else:
                    wb = openpyxl.Workbook()
                    ws = getattr(wb, "active", None)
                    if ws:
                        ws.append(["Customer", "Order #", "Date Units Distributed", "Total Units", "Physician Name", "Phone", "ShipToAddr1", "ShipToAddr2", "ShipToCity", "ShipToState", "ShipToPostal", "SKU1", "SKU1 Qty", "Lot1", "SKU2", "SKU2 Qty", "Lot2", "SKU3", "SKU3 Qty", "Lot3"])
                    path.parent.mkdir(parents=True, exist_ok=True)
                    wb.save(path)
            wb = openpyxl.load_workbook(path)
            ws = getattr(wb, "active", None)
            if ws:
                ws.append(row)
                wb.save(path)

        ensure_and_append(master_path)
        ensure_and_append(shipment_path)

    # ----------------------------------------------------------------------
    # Tracing helpers & routes
    # NOTE: kept together here so they can be moved into a dedicated
    # tracing module/blueprint in a future refactor.
    # ----------------------------------------------------------------------
    
    def generate_tracing_report_for_rep(rep, month: str):
        """Generate a CSV tracing report for a specific rep and month from actual distributions"""
        try:
            rep = normalize_row(rep)
            year = int(month[:4])
            mon = int(month[5:7])
            start_date = f"{year}-{mon:02d}-01"
            end_date = f"{year + 1}-01-01" if mon == 12 else f"{year}-{mon + 1:02d}-01"

            rep_id = rep["id"]
            rep_slug = rep["slug"]

            # Join with customers table to get city/state when not in fields_json
            rows = query_db(
                """
                SELECT dd.id, dd.order_number, dd.ship_date, dd.customer_id,
                       ddr.fields_json,
                       c.city AS customer_city, c.state AS customer_state,
                       c.facility_name AS customer_facility_name,
                       c.address1 AS customer_address1
                FROM devices_distributed dd
                JOIN device_distribution_records ddr ON ddr.dist_id = dd.id
                LEFT JOIN customers c ON c.id = COALESCE(dd.customer_id, ddr.customer_id)
                WHERE dd.rep_id = %s 
                  AND dd.ship_date >= %s 
                  AND dd.ship_date < %s
                ORDER BY dd.ship_date, dd.id
                """,
                (rep_id, start_date, end_date),
            ) or []
            if not rows:
                return None

            report_rows = []
            for r in rows:
                try:
                    fields = json.loads(r["fields_json"])
                except Exception:
                    fields = {}
                
                # Extract values
                sku = fields.get("SKU", "").strip() if fields.get("SKU") else ""
                lot = str(fields.get("Lot", "") or "").strip()
                quantity_str = fields.get("Quantity", 0) or 0
                try:
                    quantity = int(float(quantity_str))
                except (ValueError, TypeError):
                    quantity = 0
                
                # Filter rule: Exclude rows where:
                # 1. SKU is missing/empty AND Quantity is 0 or less, OR
                # 2. ALL of these are blank: SKU, Lot, Quantity (even if order # exists)
                # This removes "order header" rows and placeholder records that don't represent actual shipments
                if (not sku and quantity <= 0) or (not sku and not lot and quantity <= 0):
                    # Skip this row - it's a header/placeholder, not a real shipment line item
                    continue
                
                # Include this row - it's a valid shipment line item
                # Format Ship Date consistently as YYYY-MM-DD (ensure 10 chars)
                ship_date_str = ""
                if r.get("ship_date"):
                    ship_date_val = r["ship_date"]
                    if isinstance(ship_date_val, str):
                        ship_date_str = ship_date_val[:10] if len(ship_date_val) >= 10 else ship_date_val
                    else:
                        # datetime object
                        ship_date_str = ship_date_val.strftime("%Y-%m-%d") if hasattr(ship_date_val, 'strftime') else str(ship_date_val)[:10]
                
                # Get city/state - prefer fields_json, fallback to customers table
                city = fields.get("City", "") or r.get("customer_city") or ""
                state = fields.get("State", "") or r.get("customer_state") or ""
                facility = fields.get("Facility Name", "") or r.get("customer_facility_name") or ""
                
                # Normalize field values (handle commas/quotes safely via CSV writer)
                report_rows.append(
                    {
                        "Ship Date": ship_date_str,
                        "Order #": str(r.get("order_number") or f"DIST-{r['id']}"),
                        "Facility": str(facility),
                        "City": str(city),
                        "State": str(state),
                        "SKU": str(sku),
                        "Lot": str(lot),
                        "Quantity": quantity,  # Integer - CSV writer handles this
                    }
                )

            tracing_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports"))
            rep_slug_normalized = normalize_rep_slug(rep_slug)
            month_normalized = normalize_year_month(month)
            if not month_normalized:
                print(f"ERROR: Invalid month format: {month} (expected YYYY-MM)")
                return None

            csv_path = get_tracing_report_path(tracing_dir, rep_slug_normalized, month_normalized, ext="csv")
            csv_path.parent.mkdir(parents=True, exist_ok=True)

            fieldnames = ["Ship Date", "Order #", "Facility", "City", "State", "SKU", "Lot", "Quantity"]
            
            # Regression check: Verify no invalid rows before writing
            # Check for rows with empty SKU AND Quantity <= 0, OR all blank (SKU, Lot, Quantity)
            invalid_rows = []
            for row in report_rows:
                sku_val = (row.get("SKU", "") or "").strip()
                lot_val = (row.get("Lot", "") or "").strip()
                qty_val = row.get("Quantity", 0)
                try:
                    qty = int(qty_val) if qty_val else 0
                except (ValueError, TypeError):
                    qty = 0
                
                if (not sku_val and qty <= 0) or (not sku_val and not lot_val and qty <= 0):
                    invalid_rows.append(row)
            
            if invalid_rows:
                print(f"WARNING: Found {len(invalid_rows)} invalid row(s) with empty SKU/Lot and Quantity <= 0. This should not happen after filtering.")
                print(f"Invalid rows (first 3): {invalid_rows[:3]}")  # Show first 3 for debugging
                # Remove invalid rows as a safety measure
                report_rows = [row for row in report_rows if row not in invalid_rows]
            
            # Generate CSV content as string (for DB storage)
            import io
            csv_buffer = io.StringIO()
            writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(report_rows)
            csv_content = csv_buffer.getvalue()
            csv_buffer.close()
            
            # Still write to filesystem for backward compatibility (if directory exists)
            try:
                csv_path.parent.mkdir(parents=True, exist_ok=True)
                with open(csv_path, "w", newline="", encoding="utf-8") as f:
                    f.write(csv_content)
            except Exception as e:
                print(f"WARNING: Could not write CSV file to filesystem: {e}. Report will be stored in DB only.")

            rel_path = str(csv_path.relative_to(BASE_DIR))
            email_to = rep.get("email") or ""
            existing = normalize_row(
                query_db("SELECT id FROM tracing_reports WHERE rep_id = %s AND month = %s", (rep_id, month), one=True)
            )

            generated_at = datetime.now()
            report_id_for_history = None
            
            if existing:
                execute_db(
                    "UPDATE tracing_reports SET generated_at = %s, report_path = %s, report_csv_content = %s, email_to = %s WHERE id = %s",
                    (generated_at, rel_path, csv_content, email_to, existing["id"]),
                )
                report_id_for_history = existing["id"]
                # Record history for regeneration
                try:
                    execute_db(
                        """INSERT INTO tracing_report_history (report_id, event_type, event_at, performed_by, file_path, note)
                           VALUES (%s, 'generated', %s, 'system', %s, 'Report regenerated')""",
                        (report_id_for_history, generated_at, rel_path)
                    )
                except Exception:
                    pass  # History recording is optional
            else:
                execute_db(
                    """INSERT INTO tracing_reports (rep_id, month, generated_at, status, report_path, report_csv_content, email_to)
                       VALUES (%s, %s, %s, 'draft', %s, %s, %s)""",
                    (rep_id, month, generated_at, rel_path, csv_content, email_to),
                )
                # Get the new report_id for history
                try:
                    new_report = query_db(
                        "SELECT id FROM tracing_reports WHERE rep_id = %s AND month = %s",
                        (rep_id, month),
                        one=True
                    )
                    if new_report:
                        report_id_for_history = new_report.get("id")
                        # Record history for generation
                        execute_db(
                            """INSERT INTO tracing_report_history (report_id, event_type, event_at, performed_by, file_path, note)
                               VALUES (%s, 'generated', %s, 'system', %s, 'Report generated')""",
                            (report_id_for_history, generated_at, rel_path)
                        )
                except Exception:
                    pass  # History recording is optional

            return csv_path
        except Exception as e:
            print(f"Error generating tracing report for {rep.get('name') if rep else 'unknown'}: {e}")
            return None
    
    def generate_tracing_reports_for_month(month: str):
        """Generate tracing reports for all active reps for a given month"""
        reps = query_db("SELECT * FROM reps WHERE active = 1 ORDER BY name") or []
        results = []
        for rep in reps:
            path = generate_tracing_report_for_rep(rep, month)
            results.append((rep["slug"], path is not None))
        return results
    
    def generate_global_tracing_report(month: str):
        """Generate a single CSV tracing report for all reps (including shipstation) for a given month"""
        try:
            year = int(month[:4])
            mon = int(month[5:7])
            start_date = f"{year}-{mon:02d}-01"
            end_date = f"{year + 1}-01-01" if mon == 12 else f"{year}-{mon + 1:02d}-01"
            
            # Query all distributions for all active reps (including shipstation) for the month
            rows = query_db(
                """
                SELECT dd.id, dd.order_number, dd.ship_date, dd.customer_id, dd.rep_id,
                       ddr.fields_json,
                       c.city AS customer_city, c.state AS customer_state,
                       c.facility_name AS customer_facility_name,
                       c.address1 AS customer_address1,
                       r.name AS rep_name, r.slug AS rep_slug
                FROM devices_distributed dd
                JOIN device_distribution_records ddr ON ddr.dist_id = dd.id
                LEFT JOIN customers c ON c.id = COALESCE(dd.customer_id, ddr.customer_id)
                LEFT JOIN reps r ON r.id = dd.rep_id
                WHERE dd.ship_date >= %s 
                  AND dd.ship_date < %s
                  AND (dd.rep_id IS NOT NULL OR (r.id IS NOT NULL AND r.active = 1))
                ORDER BY dd.ship_date, dd.id, r.name
                """,
                (start_date, end_date),
            ) or []
            
            if not rows:
                return None
            
            report_rows = []
            for r in rows:
                try:
                    fields = json.loads(r["fields_json"])
                except Exception:
                    fields = {}
                
                # Extract values
                sku = fields.get("SKU", "").strip() if fields.get("SKU") else ""
                lot = str(fields.get("Lot", "") or "").strip()
                quantity_str = fields.get("Quantity", 0) or 0
                try:
                    quantity = int(float(quantity_str))
                except (ValueError, TypeError):
                    quantity = 0
                
                # Filter rule: Exclude rows where SKU is missing/empty AND Quantity is 0 or less
                if (not sku and quantity <= 0) or (not sku and not lot and quantity <= 0):
                    continue
                
                # Format Ship Date
                ship_date_str = ""
                if r.get("ship_date"):
                    ship_date_val = r["ship_date"]
                    if isinstance(ship_date_val, str):
                        ship_date_str = ship_date_val[:10] if len(ship_date_val) >= 10 else ship_date_val
                    else:
                        ship_date_str = ship_date_val.strftime("%Y-%m-%d") if hasattr(ship_date_val, 'strftime') else str(ship_date_val)[:10]
                
                # Get city/state - prefer fields_json, fallback to customers table
                city = fields.get("City", "") or r.get("customer_city") or ""
                state = fields.get("State", "") or r.get("customer_state") or ""
                facility = fields.get("Facility Name", "") or r.get("customer_facility_name") or ""
                rep_name = r.get("rep_name") or "Unknown"
                
                report_rows.append({
                    "Ship Date": ship_date_str,
                    "Order #": str(r.get("order_number") or f"DIST-{r['id']}"),
                    "Rep": str(rep_name),
                    "Facility": str(facility),
                    "City": str(city),
                    "State": str(state),
                    "SKU": str(sku),
                    "Lot": str(lot),
                    "Quantity": quantity,
                })
            
            # Save to global tracing reports directory
            global_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")) / "global" / month
            global_dir.mkdir(parents=True, exist_ok=True)
            csv_path = global_dir / f"Global_Tracing_Report_{month}.csv"
            
            fieldnames = ["Ship Date", "Order #", "Rep", "Facility", "City", "State", "SKU", "Lot", "Quantity"]
            
            try:
                with open(csv_path, "w", newline="", encoding="utf-8") as f:
                    writer = csv.DictWriter(f, fieldnames=fieldnames)
                    writer.writeheader()
                    writer.writerows(report_rows)
            except Exception as e:
                print(f"ERROR writing global tracing CSV file: {e}")
                import traceback
                traceback.print_exc()
                return None
            
            return csv_path
        except Exception as e:
            print(f"Error generating global tracing report: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    # Email functionality removed - use CSV downloads instead
    # def send_global_tracing_email(month: str, recipient_email: str = None):
    #     """Send global monthly tracing report CSV to specified email (default: ethanr@silq.tech)"""
    #     if not recipient_email:
    #         recipient_email = os.environ.get("GLOBAL_TRACING_EMAIL", "ethanr@silq.tech")
    #     
    #     # Generate global report
    #     csv_path = generate_global_tracing_report(month)
    #     if csv_path is None or not csv_path.exists():
    #         return False, "Failed to generate global tracing report"
    #     
    #     server = app.config.get("SMTP_SERVER")
    #     if not server:
    #         # In dev, log instead of failing
    #         print(f"[DEV] Global tracing email send requested for {month} to {recipient_email}")
    #         print(f"[DEV] CSV generated at: {csv_path}")
    #         return False, "SMTP not configured (dev mode - email not sent, CSV generated)"
    #     
    #     from_addr = app.config.get("EMAIL_FROM") or app.config.get("SMTP_USERNAME") or "no-reply@example.com"
    #     
    #     msg = MIMEMultipart()
    #     msg["From"] = from_addr
    #     msg["To"] = recipient_email
    #     msg["Subject"] = f"Silq – Global Monthly Tracing Report – {month}"
    #     
    #     html_body = f"""
    #     <html>
    #       <body>
    #         <h2>Global Monthly Tracing Report – {month}</h2>
    #         <p>Attached is the global monthly tracing report for {month}, including all rep distributions and ShipStation distributions.</p>
    #         <p>This report contains all device distributions for the month across all reps.</p>
    #       </body>
    #     </html>"""
    #     msg.attach(MIMEText(html_body, 'html'))
    #     
    #     # Attach CSV file
    #     from email.mime.base import MIMEBase
    #     from email import encoders
    #     with open(csv_path, "rb") as f:
    #         attachment = MIMEBase("text", "csv")
    #         attachment.set_payload(f.read())
    #         encoders.encode_base64(attachment)
    #         attachment.add_header("Content-Disposition", f"attachment; filename={csv_path.name}")
    #         msg.attach(attachment)
    #     
    #     try:
    #         username = app.config.get("SMTP_USERNAME") or ""
    #         password = app.config.get("SMTP_PASSWORD") or ""
    #         with smtplib.SMTP(server, int(app.config.get("SMTP_PORT", 587))) as s:
    #             if app.config.get("SMTP_USE_TLS"):
    #                 s.starttls()
    #             if username:
    #                 s.login(username, password)
    #             s.send_message(msg)
    #         return True, f"Email sent to {recipient_email}"
    #     except Exception as e:
    #         return False, f"Failed to send email: {str(e)}"
    
    # Email functionality removed - use CSV downloads instead
    # Entire function body commented out - use CSV downloads instead
    # def send_tracing_email_for_rep(rep, month: str):
    #     """Send tracing report email to a rep with CSV attachment and approval link"""
    #     pass
    
    def generate_distribution_log_for_month(month: str) -> Path | None:
        """Generate monthly distribution log CSV for all distributions in the month"""
        # Validate month format (YYYY-MM)
        month_normalized = month.strip()
        if len(month_normalized) != 7 or month_normalized[4] != '-':
            print(f"ERROR: Invalid month format '{month_normalized}', expected YYYY-MM")
            return None
        
        try:
            year, month_num = month_normalized.split('-')
            year_int = int(year)
            month_int = int(month_num)
            if month_int < 1 or month_int > 12:
                print(f"ERROR: Invalid month number {month_int}")
                return None
        except ValueError:
            print(f"ERROR: Invalid month format '{month_normalized}'")
            return None
        
        # Calculate date range for the month
        date_from = f"{year}-{month_num:>02s}-01"
        if month_int == 12:
            date_to = f"{int(year)+1}-01-01"
        else:
            date_to = f"{year}-{month_int+1:>02d}-01"
        
        # Fetch distribution records for the month
        normalized_rows = fetch_distribution_records(exclude_sources=EXCLUDED_DIST_SOURCES)
        filtered_rows = []
        for r in normalized_rows or []:
            ship_date = r.get("ship_date") or ""
            if ship_date and date_from and ship_date < date_from:
                continue
            if ship_date and date_to and ship_date >= date_to:
                continue
            if not ship_date and date_from:
                continue
            filtered_rows.append(r)
        
        # Group by shipment
        shipments = {}
        for r in filtered_rows:
            ship_date = (r.get("ship_date") or "Unknown")
            shipment_key = r.get("shipment_key") or f"{r.get('order_group')}::{ship_date}::{r.get('dist_id')}"
            order_num = r.get("order_number") or f"DIST-{r.get('dist_id')}"
            if shipment_key not in shipments:
                shipments[shipment_key] = {
                    "dist_id": r.get("dist_id"),
                    "order_number": order_num,
                    "ship_date": ship_date,
                    "facility": r.get("facility_label") or r.get("facility_key"),
                    "rep_name": r.get("rep_name") or "Unassigned",
                    "source": r.get("source") or "manual",
                    "tracking_number": r.get("tracking_number") or "",
                    "total_qty": 0,
                    "item_count": 0,
                    "items": [],
                }
            shipments[shipment_key]["items"].append({
                "sku": r.get("sku", "Unknown"),
                "lot": r.get("lot", "UNKNOWN"),
                "quantity": r.get("qty", 0),
            })
            shipments[shipment_key]["total_qty"] += r.get("qty", 0)
            shipments[shipment_key]["item_count"] = len(shipments[shipment_key]["items"])
        
        distributions = list(shipments.values())
        distributions.sort(key=lambda x: x.get("ship_date") or "", reverse=True)
        
        # Create output directory
        month_dir = Path(app.config.get("DIST_LOG_APPROVALS_DIR", BASE_DIR / "distribution_log_approvals")) / month_normalized
        month_dir.mkdir(parents=True, exist_ok=True)
        csv_path = month_dir / f"Distribution_Log_{month_normalized}.csv"
        
        # Write CSV
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow([
                "Ship Date", "Order #", "Facility", "Rep", "Items", "Total Qty", 
                "Source", "Tracking", "Dist ID"
            ])
            for d in distributions:
                writer.writerow([
                    (d.get("ship_date") or "")[:10],
                    d.get("order_number") or "",
                    d.get("facility") or "",
                    d.get("rep_name") or "",
                    d.get("item_count") or 0,
                    d.get("total_qty") or 0,
                    d.get("source") or "",
                    d.get("tracking_number") or "",
                    d.get("dist_id") or "",
                ])
        
        # Create or update database record
        report_path = str(csv_path.relative_to(BASE_DIR))
        generated_at = datetime.now()
        
        try:
            existing = query_db(
                "SELECT id FROM distribution_log_approvals WHERE month = %s",
                (month_normalized,),
                one=True
            )
            if existing:
                execute_db(
                    """UPDATE distribution_log_approvals 
                       SET generated_at = %s, report_path = %s, status = 'draft'
                       WHERE month = %s""",
                    (generated_at, report_path, month_normalized)
                )
            else:
                execute_db(
                    """INSERT INTO distribution_log_approvals 
                       (month, generated_at, status, report_path)
                       VALUES (%s, %s, 'draft', %s)""",
                    (month_normalized, generated_at, report_path)
                )
        except Exception as e:
            print(f"ERROR: Failed to update distribution_log_approvals: {e}")
            import traceback
            traceback.print_exc()
        
        return csv_path
    
    def send_distribution_log_email(month: str) -> tuple[bool, str]:
        """Send distribution log email to approver (Ethanr@silq.tech)"""
        # Get approval email (default to Ethanr@silq.tech, allow override via env)
        approval_email = os.environ.get("DIST_LOG_APPROVAL_EMAIL", "Ethanr@silq.tech").strip()
        if not approval_email:
            return False, "approval email not configured"
        
        # Generate/get report path
        path = generate_distribution_log_for_month(month)
        if path is None or not path.exists():
            return False, "failed to generate distribution log"
        
        # Get or create distribution_log_approvals record
        month_normalized = month.strip()
        if len(month_normalized) != 7 or month_normalized[4] != '-':
            return False, f"invalid month format: {month_normalized}"
        
        try:
            approval_record = normalize_row(
                query_db(
                    "SELECT id FROM distribution_log_approvals WHERE month = %s",
                    (month_normalized,),
                    one=True
                )
            )
            if not approval_record:
                return False, "distribution log record not found"
        except Exception as e:
            return False, f"database error: {str(e)}"
        
        server = app.config.get("SMTP_SERVER")
        if not server:
            return False, "SMTP not configured"
        
        from_addr = app.config.get("EMAIL_FROM") or app.config.get("SMTP_USERNAME") or "no-reply@example.com"
        
        # Build base URL for report link (if applicable)
        try:
            base_url = request.url_root.rstrip('/')
        except RuntimeError:
            base_url = app.config.get("BASE_URL") or "http://localhost:5000"
        
        # Create email
        msg = MIMEMultipart('alternative')
        msg["Subject"] = f"Silq – Monthly Distribution Log Approval – {month}"
        msg["From"] = from_addr
        msg["To"] = approval_email
        
        # Plain text version
        text_body = f"""ACTION REQUIRED: Monthly Distribution Log Approval

Dear Approver,

Please review the monthly distribution log for {month} (attached CSV file).

If the distribution log is accurate and complete, please reply to this email with "I approve".

Your approval will be stored in Silq's quality management system for audit purposes.

Best regards,
Silq Quality Management System
{from_addr}
"""
        
        # HTML version
        html_body = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
</head>
<body style="font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 20px;">
<div style="background-color: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
<h1 style="color: #d32f2f; border-bottom: 3px solid #4CAF50; padding-bottom: 10px;">ACTION REQUIRED: Monthly Distribution Log Approval</h1>

<div style="background-color: #e8f5e9; padding: 15px; border-left: 4px solid #4CAF50; margin: 20px 0;">
<p style="margin: 0;"><strong>Month:</strong> {month}</p>
</div>

<p>Please review the monthly distribution log for <strong>{month}</strong> (attached CSV file).</p>

<div style="background-color: #fff3cd; padding: 20px; border-left: 4px solid #ffc107; margin: 30px 0;">
<h2 style="color: #856404; margin-top: 0;">APPROVAL REQUIRED</h2>
<p>If the distribution log is accurate and complete, please <strong>reply to this email with "I approve"</strong>.</p>
<p style="font-size: 13px; color: #666; margin-top: 15px;">
Note: Your admin can upload your approval email (.eml file) to the system if you prefer.
</p>
</div>

<p style="margin-top: 30px; padding-top: 20px; border-top: 1px solid #ddd; color: #666; font-size: 13px;">
Your approval will be stored in Silq's quality management system for audit purposes.<br><br>
Best regards,<br>
<strong>Silq Quality Management System</strong><br>
{from_addr}
</p>
</div>
</body>
</html>"""
        
        msg.attach(MIMEText(text_body, 'plain'))
        msg.attach(MIMEText(html_body, 'html'))
        
        # Attach CSV file
        from email.mime.base import MIMEBase
        from email import encoders
        with open(path, "rb") as f:
            attachment = MIMEBase("text", "csv")
            attachment.set_payload(f.read())
            encoders.encode_base64(attachment)
            attachment.add_header("Content-Disposition", f"attachment; filename={path.name}")
            msg.attach(attachment)
        
        try:
            username = app.config.get("SMTP_USERNAME") or ""
            password = app.config.get("SMTP_PASSWORD") or ""
            with smtplib.SMTP(server, int(app.config.get("SMTP_PORT", 587))) as s:
                if app.config.get("SMTP_USE_TLS"):
                    s.starttls()
                if username:
                    s.login(username, password)
                s.send_message(msg)
            
            # Update distribution_log_approvals record
            email_sent_at = datetime.now()
            execute_db(
                """UPDATE distribution_log_approvals 
                   SET status = 'sent', email_sent_at = %s, email_to = %s
                   WHERE month = %s""",
                (email_sent_at, approval_email, month_normalized)
            )
            
            return True, "sent"
        except Exception as e:
            return False, f"error: {e}"
    
    def parse_approval_eml(eml_path: Path) -> dict:
        """Parse .eml file and extract metadata"""
        try:
            with open(eml_path, 'rb') as f:
                parser = BytesParser()
                msg = parser.parse(f)
            
            # Extract from
            from_addr = msg.get("From", "")
            if from_addr:
                # Extract email address from "Name <email>" format
                import re
                email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', from_addr)
                if email_match:
                    from_addr = email_match.group(0)
            
            # Extract subject
            subject = msg.get("Subject", "")
            
            # Extract date
            date_str = msg.get("Date", "")
            date_obj = None
            if date_str:
                try:
                    from email.utils import parsedate_to_datetime
                    date_obj = parsedate_to_datetime(date_str)
                except Exception:
                    pass
            
            # Extract body (try to get text/plain first, then text/html)
            body = ""
            if msg.is_multipart():
                for part in msg.walk():
                    content_type = part.get_content_type()
                    if content_type == "text/plain":
                        try:
                            body = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                            break
                        except Exception:
                            pass
                    elif content_type == "text/html" and not body:
                        try:
                            body = part.get_payload(decode=True).decode('utf-8', errors='ignore')
                            # Strip HTML tags for validation
                            import re
                            body = re.sub(r'<[^>]+>', '', body)
                        except Exception:
                            pass
            else:
                try:
                    body = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
                except Exception:
                    body = msg.get_payload()
            
            return {
                "from": from_addr.lower() if from_addr else "",
                "subject": subject,
                "date": date_obj,
                "body": body.lower() if body else "",
            }
        except Exception as e:
            print(f"ERROR: Failed to parse .eml file: {e}")
            import traceback
            traceback.print_exc()
            return {}
    
    def validate_approval_eml(eml_data: dict) -> tuple[bool, str]:
        """Validate approval .eml file (sender and body content)"""
        # Get approved sender email (default to Ethanr@silq.tech, allow whitelist via env)
        approved_senders_str = os.environ.get("DIST_LOG_APPROVAL_SENDERS", "Ethanr@silq.tech").strip()
        approved_senders = [s.strip().lower() for s in approved_senders_str.split(',') if s.strip()]
        if not approved_senders:
            approved_senders = ["ethanr@silq.tech"]
        
        sender_email = eml_data.get("from", "").lower()
        if not sender_email:
            return False, "Could not extract sender email from .eml file"
        
        if sender_email not in approved_senders:
            return False, f"Sender email '{sender_email}' is not in the approved sender list"
        
        body = eml_data.get("body", "").lower()
        if "i approve" not in body:
            return False, "Email body does not contain 'I approve' (case-insensitive)"
        
        return True, "valid"

    def sync_shipstation_data(progress_callback=None, force_rescan=False):
        try:
            api_key = app.config.get("SS_API_KEY")
            api_secret = app.config.get("SS_API_SECRET")
            if not api_key or not api_secret:
                return False, "ShipStation API credentials not configured"

            # Use 'shipstation' rep for ShipStation-imported distributions
            shipstation_rep = normalize_row(query_db("SELECT id FROM reps WHERE slug='shipstation'", one=True))
            if not shipstation_rep:
                return False, "ShipStation rep not found. Please ensure 'shipstation' rep exists."
            shipstation_rep_id = shipstation_rep["id"]

            if force_rescan:
                ok, msg = deep_rescan_since_2024(api_key, api_secret, shipstation_rep_id, query_db, execute_db,
                                                 progress_callback=progress_callback, should_cancel_fn=is_cancel_requested)
            else:
                ok, msg = sync_units_and_grouping(
                    api_key, api_secret, shipstation_rep_id, query_db, execute_db,
                    days=90, start_date_override=None, max_orders=2000,
                    force_rescan=False,
                    progress_callback=progress_callback,
                    should_cancel_fn=is_cancel_requested
                )
            return ok, msg
        except Exception as e:
            import traceback
            traceback.print_exc()
            return False, f"Sync Failed: {str(e)}"

    @app.route("/")
    def index():
        return render_template("index.html")

    @app.route("/admin/login", methods=["GET", "POST"])
    def admin_login():
        if request.method == "POST":
            admin_password = app.config.get("ADMIN_PASSWORD")
            if not admin_password:
                flash("Admin password not configured. Please contact system administrator.", "danger")
                return render_template("admin_login.html")
            if request.form.get("password", "") == admin_password:
                session.clear()
                session["is_admin"] = True
                session.permanent = True
                flash("Logged in as admin.", "success")
                return redirect(url_for("admin_dashboard"))
            else:
                flash("Invalid admin password.", "danger")
        return render_template("admin_login.html")

    @app.route("/admin/logout")
    def admin_logout():
        session.pop("is_admin", None)
        flash("Logged out.", "info")
        return redirect(url_for("index"))

    @app.route("/admin")
    def admin_dashboard():
        if not is_admin():
            return redirect(url_for("admin_login"))
        # One-time warning per admin session if ShipStation credentials are missing.
        if (
            not session.get("_warned_shipstation_missing")
            and (not app.config.get("SS_API_KEY") or not app.config.get("SS_API_SECRET"))
        ):
            flash(
                "ShipStation API credentials are not configured (SHIPSTATION_API_KEY / SHIPSTATION_API_SECRET). "
                "ShipStation sync is disabled until these are set.",
                "warning",
            )
            session["_warned_shipstation_missing"] = True
        try:
            reset_if_stale()
            progress = get_sync_progress()
            sync_running = is_sync_running()
        except Exception:
            progress = None
            sync_running = False

        try:
            ensure_system_tool_runs_table()
            last_deep_rebuild = get_last_system_tool_run("deep_rebuild")
            last_deep_rebuild_display = None
            if last_deep_rebuild:
                ts = last_deep_rebuild.get("ran_at")
                try:
                    from datetime import datetime as _dt, date as _date

                    if isinstance(ts, (_dt, _date)):
                        last_deep_rebuild_display = ts.strftime("%Y-%m-%d %H:%M")
                    else:
                        last_deep_rebuild_display = str(ts or "")
                except Exception:
                    last_deep_rebuild_display = str(ts or "")
        except Exception:
            last_deep_rebuild = None
            last_deep_rebuild_display = None
        try:
            last_sync_row = query_db(
                """
                SELECT ran_at, synced, skipped, message, orders_seen, shipments_seen, duration_seconds
                FROM shipstation_sync_runs
                ORDER BY ran_at DESC
                LIMIT 1
                """,
                one=True,
            )
            last_sync = dict(last_sync_row) if last_sync_row else None
            last_sync_display = None
            if last_sync:
                ts = last_sync.get('ran_at')
                try:
                    from datetime import datetime as _dt, date as _date
                    if isinstance(ts, (_dt, _date)):
                        last_sync_display = ts.strftime('%Y-%m-%d %H:%M')
                    else:
                        last_sync_display = str(ts or "")
                except Exception:
                    last_sync_display = str(ts or "")
        except Exception:
            last_sync = None
            last_sync_display = None

        try:
            skip_breakdown = query_db(
                "SELECT reason, COUNT(*) as c FROM shipstation_skipped_orders GROUP BY reason ORDER BY c DESC"
            ) or []
        except Exception:
            skip_breakdown = []

        # Get cache status for admin diagnostic (using same priority as hospital_targets.py)
        cache_status = {}
        try:
            from pathlib import Path
            from datetime import datetime
            from data_bootstrap import get_cache_metadata
            from hospital_targets import FACILITY_TARGETS_CACHE_PARQUET, FACILITY_TARGETS_CACHE_CSV, FACILITY_CACHE_CSV
            
            BASE_DIR = Path(__file__).resolve().parent
            CACHE_DIR = BASE_DIR / "cache"
            
            # Check hospital cache in priority order (same as load_facility_cache)
            hospital_cache = None
            hospital_cache_type = None
            if FACILITY_TARGETS_CACHE_PARQUET.exists():
                hospital_cache = FACILITY_TARGETS_CACHE_PARQUET
                hospital_cache_type = "parquet"
            elif FACILITY_TARGETS_CACHE_CSV.exists():
                hospital_cache = FACILITY_TARGETS_CACHE_CSV
                hospital_cache_type = "csv"
            elif FACILITY_CACHE_CSV.exists():
                hospital_cache = FACILITY_CACHE_CSV
                hospital_cache_type = "legacy"
            
            # Check doctor cache (standardized name)
            doctor_cache = CACHE_DIR / "facility_doctors_cache.json"
            doctor_cache_legacy = CACHE_DIR / "facility_doctors.json"
            if not doctor_cache.exists() and doctor_cache_legacy.exists():
                doctor_cache = doctor_cache_legacy  # Use legacy if new doesn't exist
            
            metadata = get_cache_metadata()
            
            cache_status = {
                "hospital_cache": {
                    "exists": hospital_cache is not None and hospital_cache.exists(),
                    "filename": hospital_cache.name if hospital_cache else None,
                    "type": hospital_cache_type,
                    "size_mb": round(hospital_cache.stat().st_size / (1024 * 1024), 2) if hospital_cache and hospital_cache.exists() else 0,
                    "last_modified": datetime.fromtimestamp(hospital_cache.stat().st_mtime).strftime("%Y-%m-%d %H:%M") if hospital_cache and hospital_cache.exists() else None,
                    "build_timestamp": metadata.get("facility_catheter_days", {}).get("built_at") or metadata.get("facility_targets_cache", {}).get("built_at"),
                    # Check all possible filenames for debugging
                    "files_checked": {
                        "facility_targets_cache.parquet": FACILITY_TARGETS_CACHE_PARQUET.exists(),
                        "facility_targets_cache.csv": FACILITY_TARGETS_CACHE_CSV.exists(),
                        "facility_catheter_days.csv": FACILITY_CACHE_CSV.exists()
                    }
                },
                "doctor_cache": {
                    "exists": doctor_cache.exists(),
                    "filename": doctor_cache.name if doctor_cache.exists() else None,
                    "size_mb": round(doctor_cache.stat().st_size / (1024 * 1024), 2) if doctor_cache.exists() else 0,
                    "last_modified": datetime.fromtimestamp(doctor_cache.stat().st_mtime).strftime("%Y-%m-%d %H:%M") if doctor_cache.exists() else None,
                    "build_timestamp": metadata.get("facility_doctors", {}).get("built_at"),
                    "legacy_exists": doctor_cache_legacy.exists() if doctor_cache.name == "facility_doctors_cache.json" else False
                }
            }
        except Exception as e:
            # Non-fatal - just log and continue
            print(f"[Admin Dashboard] Could not get cache status: {e}")
            import traceback
            traceback.print_exc()
            cache_status = {
                "hospital_cache": {"exists": False, "error": str(e)},
                "doctor_cache": {"exists": False, "error": str(e)}
            }
        
        return render_template(
            "admin_dashboard.html",
            last_sync=last_sync,
            last_sync_display=last_sync_display,
            last_deep_rebuild=last_deep_rebuild,
            last_deep_rebuild_display=last_deep_rebuild_display,
            skip_breakdown=skip_breakdown,
            sync_running=sync_running,
            sync_progress=progress,
            cache_status=cache_status,
        )

    @app.route("/init-db")
    def init_db_route():
        init_db()
        return "DB initialized"

    @app.route("/admin/sync/shipstation", methods=["GET", "POST"])
    def admin_sync_shipstation():
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        # Check if sync is already running
        if is_sync_running():
            flash("Sync is already in progress. Check the status banner.", "info")
            return redirect(url_for("admin_dashboard"))
        
        mode = request.args.get("mode") or request.form.get("mode") or "deep"
        try:
            print(f"[ShipStation] Starting {mode} sync at {datetime.utcnow().isoformat()}Z")
            ok, msg = run_sync_in_background(
                (lambda progress_callback=None: sync_shipstation_data(progress_callback=progress_callback, force_rescan=(mode=="deep"))),
                app_context=app
            )
            print(f"[ShipStation] Background launch result: ok={ok}, msg={msg}, running={is_sync_running()}")
            flash(msg, "success" if ok else "danger")
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f"Failed to start sync: {str(e)}", "danger")
        return redirect(url_for("admin_dashboard"))

    @app.route("/admin/sync/progress")
    def admin_sync_progress():
        """API endpoint to get current sync progress (JSON, deterministic, safe)"""
        # Read-only progress is safe to expose; used by global banner
        # Mark stale runs (no heartbeat) as completed/stale
        reset_if_stale()
        
        try:
            progress = get_sync_progress()
            running = is_sync_running()
            
            # Ensure progress dict is always valid (defensive)
            if not isinstance(progress, dict):
                progress = {}
            
            # Return deterministic JSON response
            response = {
                "is_running": bool(running),
                "progress": {
                    "synced": int(progress.get("synced", 0)),
                    "skipped": int(progress.get("skipped", 0)),
                    "total_checked": int(progress.get("total_checked", 0)),
                    "current_page": int(progress.get("current_page", 0)),
                    "status": str(progress.get("status", "idle")),
                    "message": str(progress.get("message", "")),
                    "started_at": progress.get("started_at"),
                    "updated_at": progress.get("updated_at"),
                    "error": progress.get("error"),
                },
                "pid": os.getpid(),
            }
            return jsonify(response)
        except Exception as e:
            # Never fail the progress endpoint; return safe defaults
            import traceback
            traceback.print_exc()
            return jsonify({
                "is_running": False,
                "progress": {
                    "synced": 0,
                    "skipped": 0,
                    "total_checked": 0,
                    "current_page": 0,
                    "status": "error",
                    "message": f"Error reading progress: {str(e)}",
                    "started_at": None,
                    "updated_at": None,
                    "error": str(e),
                },
                "pid": os.getpid(),
            })

    @app.route("/admin/sync/stop", methods=["POST"])
    def admin_sync_stop():
        """Cancel a running sync (best-effort)"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            request_sync_cancel()
            flash("Sync cancel requested. It will stop shortly.", "warning")
        except Exception as e:
            flash(f"Failed to request cancel: {str(e)}", "danger")
        return redirect(url_for("admin_dashboard"))
    
    @app.route("/admin/sync/clear-progress", methods=["POST"])
    def admin_sync_clear_progress():
        """Manually clear stuck sync progress state"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            from sync_progress import clear_sync_progress
            clear_sync_progress()
            flash("✅ Sync progress cleared. Polling will stop.", "success")
        except Exception as e:
            flash(f"Failed to clear progress: {str(e)}", "danger")
        return redirect(url_for("admin_dashboard"))

    @app.route("/admin/documents/reconcile")
    def admin_documents_reconcile():
        """Admin page to view and link unlinked document uploads"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        ensure_rep_documents_table()
        
        # Get unlinked documents
        unlinked = query_db(
            """
            SELECT rd.*, r.name as rep_name, r.slug as rep_slug
            FROM rep_documents rd
            JOIN reps r ON rd.rep_id = r.id
            WHERE rd.distribution_id IS NULL
            ORDER BY rd.uploaded_at DESC
            """,
        ) or []
        
        # For each unlinked doc, try to find suggested matches
        for doc in unlinked:
            doc["suggested_matches"] = []
            if doc.get("order_number"):
                matches = query_db(
                    """
                    SELECT id, order_number, ship_date, source, rep_id
                    FROM devices_distributed
                    WHERE regexp_replace(upper(coalesce(order_number, '')), '[^A-Z0-9]', '', 'g') = 
                          regexp_replace(upper(coalesce(%s, '')), '[^A-Z0-9]', '', 'g')
                    ORDER BY 
                        CASE WHEN source = 'shipstation' THEN 0 ELSE 1 END,
                        ship_date NULLS LAST,
                        id ASC
                    LIMIT 5
                    """,
                    (doc.get("order_number"),),
                ) or []
                doc["suggested_matches"] = matches
        
        return render_template("admin_documents_reconcile.html", unlinked_docs=unlinked)

    @app.route("/admin/documents/reconcile/link", methods=["POST"])
    def admin_documents_reconcile_link():
        """Link a document to a distribution"""
        if not is_admin():
            abort(403)
        ensure_rep_documents_table()
        
        doc_id = request.form.get("doc_id")
        dist_id = request.form.get("dist_id")
        match_reason = request.form.get("match_reason", "manual_admin_link")
        
        if not doc_id or not dist_id:
            flash("Missing doc_id or dist_id.", "danger")
            return redirect(url_for("admin_documents_reconcile"))
        
        try:
            # Verify document exists and is unlinked
            doc = query_db(
                "SELECT * FROM rep_documents WHERE id = %s",
                (int(doc_id),),
                one=True,
            )
            if not doc:
                flash("Document not found.", "danger")
                return redirect(url_for("admin_documents_reconcile"))
            
            if doc.get("distribution_id"):
                flash("Document is already linked.", "warning")
                return redirect(url_for("admin_documents_reconcile"))
            
            # Verify distribution exists
            dist = query_db(
                "SELECT * FROM devices_distributed WHERE id = %s",
                (int(dist_id),),
                one=True,
            )
            if not dist:
                flash("Distribution not found.", "danger")
                return redirect(url_for("admin_documents_reconcile"))
            
            # Link them
            execute_db(
                "UPDATE rep_documents SET distribution_id = %s, match_reason = %s WHERE id = %s",
                (int(dist_id), match_reason, int(doc_id)),
            )
            
            flash(f"Document #{doc_id} linked to distribution #{dist_id}.", "success")
        except Exception as e:
            flash(f"Error linking document: {str(e)}", "danger")
        
        return redirect(url_for("admin_documents_reconcile"))

    @app.route("/admin/documents/reconcile/auto", methods=["POST"])
    def admin_documents_reconcile_auto():
        """Attempt to auto-link all unlinked documents"""
        if not is_admin():
            abort(403)
        ensure_rep_documents_table()
        
        unlinked = query_db(
            "SELECT * FROM rep_documents WHERE distribution_id IS NULL AND order_number IS NOT NULL",
        ) or []
        
        linked_count = 0
        for doc in unlinked:
            try:
                uploaded_at = doc.get("uploaded_at")
                if isinstance(uploaded_at, str):
                    uploaded_at = datetime.fromisoformat(uploaded_at.replace('Z', '+00:00'))
                elif not isinstance(uploaded_at, datetime):
                    uploaded_at = datetime.utcnow()
                
                dist_id, match_reason = link_document_to_distribution(
                    order_number=doc.get("order_number"),
                    rep_id=doc.get("rep_id")
                )
                
                if dist_id:
                    execute_db(
                        "UPDATE rep_documents SET distribution_id = %s, match_reason = %s WHERE id = %s",
                        (dist_id, match_reason, doc.get("id")),
                    )
                    linked_count += 1
            except Exception as e:
                continue  # Skip errors, continue with next
        
        flash(f"Auto-linked {linked_count} of {len(unlinked)} unlinked documents.", "success")
        return redirect(url_for("admin_documents_reconcile"))

    @app.route("/admin/documents/reconcile/backfill", methods=["POST"])
    def admin_documents_reconcile_backfill():
        """One-time backfill: attempt to link all historical unlinked documents"""
        if not is_admin():
            abort(403)
        ensure_rep_documents_table()
        
        # Get all unlinked documents (including those that might have been linked before)
        all_docs = query_db(
            """
            SELECT * FROM rep_documents 
            WHERE distribution_id IS NULL 
            ORDER BY uploaded_at ASC
            """,
        ) or []
        
        linked_count = 0
        errors = []
        
        for doc in all_docs:
            try:
                uploaded_at = doc.get("uploaded_at")
                if isinstance(uploaded_at, str):
                    try:
                        uploaded_at = datetime.fromisoformat(uploaded_at.replace('Z', '+00:00'))
                    except Exception as e:
                        import traceback
                        print(f"[ERROR] admin_tracing_upload (datetime parse): {e}")
                        traceback.print_exc()
                        uploaded_at = datetime.utcnow()
                elif not isinstance(uploaded_at, datetime):
                    uploaded_at = datetime.utcnow()
                
                dist_id, match_reason = link_document_to_distribution(
                    order_number=doc.get("order_number"),
                    rep_id=doc.get("rep_id")
                )
                
                if dist_id:
                    execute_db(
                        "UPDATE rep_documents SET distribution_id = %s, match_reason = %s WHERE id = %s",
                        (dist_id, f"backfill:{match_reason}" if match_reason else "backfill:auto", doc.get("id")),
                    )
                    linked_count += 1
            except Exception as e:
                errors.append(f"Doc #{doc.get('id')}: {str(e)}")
                continue
        
        message = f"Backfill complete: Linked {linked_count} of {len(all_docs)} documents."
        if errors:
            message += f" {len(errors)} errors occurred."
        flash(message, "success" if linked_count > 0 else "info")
        
        if errors and len(errors) <= 10:
            for err in errors:
                flash(err, "warning")
        
        return redirect(url_for("admin_documents_reconcile"))

    @app.route("/admin/shipstation/skipped")
    def admin_shipstation_skipped():
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            reason = request.args.get("reason") or None
            hide_cancelled = request.args.get("hide_cancelled", "1") == "1"
            params = []
            where = []
            if reason:
                where.append("reason = %s")
                params.append(reason)
            if hide_cancelled:
                where.append("reason != 'cancelled'")
            where_sql = ("WHERE " + " AND ".join(where)) if where else ""

            skipped_rows = query_db(
                f"""
                SELECT order_id, order_number, order_date, reason, details, updated_at
                FROM shipstation_skipped_orders
                {where_sql}
                ORDER BY updated_at DESC
                LIMIT 500
                """,
                tuple(params) if params else None,
            ) or []

            reasons = query_db(
                "SELECT reason, COUNT(*) as c FROM shipstation_skipped_orders GROUP BY reason ORDER BY c DESC"
            ) or []

            return render_template(
                "admin_shipstation_skipped.html",
                skipped=skipped_rows,
                reasons=reasons,
                selected_reason=reason or "",
                hide_cancelled=hide_cancelled,
            )
        except Exception as e:
            import traceback
            traceback.print_exc()
            return f"Error loading skipped orders: {e}", 500

    @app.route("/admin/tools/validate/<kind>")
    def admin_tools_validate(kind):
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            if kind == "distributions":
                from dev_scripts.validate_distributions_against_reference import validate as validate_dist

                report = validate_dist(query_db)
                title = "DB vs Distributions CSV"
            elif kind == "orders":
                from dev_scripts.validate_orders_against_reference import validate as validate_orders

                report = validate_orders(query_db)
                title = "DB vs Orders CSV"
            elif kind == "shipstation":
                from dev_scripts.compare_shipstation_to_orders_csv import compare as compare_ss

                report = compare_ss(query_db)
                title = "ShipStation vs Orders CSV"
            else:
                abort(404)
            return render_template("admin_validation_report.html", title=title, report=report, kind=kind)
        except FileNotFoundError as fnf:
            flash(str(fnf), "warning")
            return redirect(url_for("admin_dashboard"))
        except Exception as e:
            import traceback
            traceback.print_exc()
            return f"Validation error: {e}", 500

    def ensure_tracing_approval_tokens_table():
        """Create table for secure token-based tracing report approvals"""
        execute_db(
            """
            CREATE TABLE IF NOT EXISTS tracing_approval_tokens (
                id SERIAL PRIMARY KEY,
                report_id INTEGER NOT NULL REFERENCES tracing_reports(id) ON DELETE CASCADE,
                token TEXT NOT NULL UNIQUE,
                expires_at TIMESTAMP NOT NULL,
                used BOOLEAN DEFAULT FALSE,
                used_at TIMESTAMP,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                ip_address TEXT,
                user_agent TEXT
            )
            """
        )
        execute_db(
            """
            CREATE INDEX IF NOT EXISTS idx_tracing_approval_tokens_token 
            ON tracing_approval_tokens(token) WHERE used = FALSE
            """
        )
        execute_db(
            """
            CREATE INDEX IF NOT EXISTS idx_tracing_approval_tokens_report 
            ON tracing_approval_tokens(report_id)
            """
        )

    def ensure_distribution_log_approvals_table():
        """Create table for monthly distribution log approvals"""
        execute_db(
            """
            CREATE TABLE IF NOT EXISTS distribution_log_approvals (
                id SERIAL PRIMARY KEY,
                month TEXT NOT NULL UNIQUE,
                generated_at TIMESTAMP NOT NULL,
                status TEXT NOT NULL DEFAULT 'draft',
                report_path TEXT NOT NULL,
                email_to TEXT,
                email_sent_at TIMESTAMP,
                approval_eml_path TEXT,
                approval_sender_email TEXT,
                approval_subject TEXT,
                approval_date TIMESTAMP,
                approval_uploaded_at TIMESTAMP,
                approved_by TEXT
            )
            """
        )
        execute_db(
            """
            CREATE INDEX IF NOT EXISTS idx_distribution_log_approvals_month 
            ON distribution_log_approvals(month)
            """
        )
        execute_db(
            """
            CREATE INDEX IF NOT EXISTS idx_distribution_log_approvals_status 
            ON distribution_log_approvals(status)
            """
        )

    def ensure_system_tool_runs_table():
        execute_db(
            """
            CREATE TABLE IF NOT EXISTS system_tool_runs (
                id SERIAL PRIMARY KEY,
                tool_name TEXT NOT NULL,
                ran_at TIMESTAMP NOT NULL,
                status TEXT NOT NULL,
                orders INTEGER,
                units INTEGER,
                summary_json TEXT,
                output_text TEXT
            )
            """
        )
        execute_db(
            """
            CREATE INDEX IF NOT EXISTS idx_system_tool_runs_tool_ran_at
            ON system_tool_runs (tool_name, ran_at DESC)
            """
        )

    def get_last_system_tool_run(tool_name: str):
        try:
            row = query_db(
                """
                SELECT *
                FROM system_tool_runs
                WHERE tool_name = %s
                ORDER BY ran_at DESC
                LIMIT 1
                """,
                (tool_name,),
                one=True,
            )
            return normalize_row(row)
        except Exception:
            return None

    def load_reference_merged_csv(path: Path):
        """Load per-company totals from 'Distribution Log (merged).csv' (Customers sheet export)."""
        import csv

        out = {}
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f)
            for row in reader:
                ck = (row.get("CompanyKey") or "").strip().upper()
                if not ck:
                    continue
                try:
                    units = int(float(row.get("TotalUnits") or 0))
                except Exception:
                    units = 0
                try:
                    shipments = int(float(row.get("ShipmentsCount") or 0))
                except Exception:
                    shipments = 0
                first_ship = (row.get("FirstShipDate") or "")[:10]
                last_ship = (row.get("LastShipDate") or "")[:10]
                out[ck] = {
                    "company_key": ck,
                    "company": row.get("Company") or "",
                    "shipments": shipments,
                    "units": units,
                    "first_ship": first_ship,
                    "last_ship": last_ship,
                }
        return out

    def run_deep_rebuild_diagnostics(date_from: str, date_to: str):
        """Compute canonical totals over DB and compare vs reference CSV if available."""
        rows = fetch_distribution_records(
            exclude_sources=EXCLUDED_DIST_SOURCES,
            date_from=date_from,
            date_to=date_to,
        )

        # Order totals (order_group rollups)
        order_summaries = build_order_group_summaries(rows)
        total_orders = len(order_summaries)
        total_units = sum(int(v.get("total_units") or 0) for v in order_summaries.values())

        # SKU totals
        sku_totals = {}
        for r in rows or []:
            if not isinstance(r, dict):
                continue
            sku = (r.get("sku") or "Unknown").strip() or "Unknown"
            qty = r.get("qty") or 0
            try:
                qty_int = int(qty)
            except Exception:
                qty_int = 0
            sku_totals[sku] = int(sku_totals.get(sku, 0)) + qty_int

        sku_breakdown = [
            {"sku": sku, "units": units}
            for sku, units in sorted(sku_totals.items(), key=lambda kv: kv[1], reverse=True)
        ]

        # Customer (company_key) totals to compare vs merged CSV
        by_company = {}
        for r in rows or []:
            ck = (r.get("facility_key") or "").strip().upper()
            if not ck:
                continue
            rec = by_company.get(ck)
            if not rec:
                rec = {
                    "company_key": ck,
                    "company": r.get("facility_label") or ck,
                    "shipments": set(),
                    "units": 0,
                    "first_ship": "",
                    "last_ship": "",
                }
                by_company[ck] = rec
            rec["shipments"].add(r.get("shipment_key") or f"{r.get('order_group')}::{r.get('ship_date')}::{r.get('dist_id')}")
            rec["units"] += int(r.get("qty") or 0)
            sd = _normalize_ship_date(r.get("ship_date"))
            if sd:
                rec["first_ship"] = sd if not rec["first_ship"] else min(rec["first_ship"], sd)
                rec["last_ship"] = sd if not rec["last_ship"] else max(rec["last_ship"], sd)
        for rec in by_company.values():
            rec["shipments"] = len(rec["shipments"])  # convert set -> count

        # Reference CSV (optional)
        ref_path_candidates = [
            BASE_DIR / "dev_scripts" / "Distribution Log (merged).csv",
            BASE_DIR / "Distribution Log (merged).csv",
        ]
        ref_path = next((p for p in ref_path_candidates if p.exists()), None)
        ref = load_reference_merged_csv(ref_path) if ref_path else None

        mismatches = []
        if ref:
            for ck, ref_rec in ref.items():
                db_rec = by_company.get(ck)
                db_units = int(db_rec.get("units") or 0) if db_rec else 0
                db_ship = int(db_rec.get("shipments") or 0) if db_rec else 0
                du = db_units - int(ref_rec.get("units") or 0)
                ds = db_ship - int(ref_rec.get("shipments") or 0)
                if du != 0 or ds != 0:
                    mismatches.append(
                        {
                            "company_key": ck,
                            "company": ref_rec.get("company") or (db_rec.get("company") if db_rec else ck),
                            "db_units": db_units,
                            "ref_units": int(ref_rec.get("units") or 0),
                            "delta_units": du,
                            "db_shipments": db_ship,
                            "ref_shipments": int(ref_rec.get("shipments") or 0),
                            "delta_shipments": ds,
                        }
                    )
            mismatches.sort(key=lambda m: (abs(int(m.get("delta_units") or 0)), abs(int(m.get("delta_shipments") or 0))), reverse=True)

        summary = {
            "date_from": date_from,
            "date_to": date_to,
            "excluded_sources": sorted(EXCLUDED_DIST_SOURCES) if isinstance(EXCLUDED_DIST_SOURCES, (set, list, tuple)) else [],
            "orders": total_orders,
            "units": total_units,
            "sku_breakdown": sku_breakdown,
            "reference_csv": str(ref_path) if ref_path else "",
            "mismatch_count": len(mismatches) if mismatches else 0,
        }

        # Human-readable output for copy/paste
        lines = []
        lines.append("=== Deep Rebuild / Diagnostics (Non-destructive) ===")
        lines.append(f"Window: {date_from} to {date_to}")
        lines.append(f"Orders (order_group): {total_orders}")
        lines.append(f"Units: {total_units}")
        lines.append(f"Excluded sources: {', '.join(summary['excluded_sources']) or 'none'}")
        if ref_path:
            lines.append(f"Reference CSV: {ref_path}")
            lines.append(f"Customer mismatches vs reference: {len(mismatches)}")
        lines.append("")
        lines.append("Top SKUs:")
        for row in sku_breakdown[:15]:
            lines.append(f"  {row['sku']}: {row['units']}")
        if mismatches:
            lines.append("")
            lines.append("Top customer mismatches (delta units / delta shipments):")
            for m in mismatches[:25]:
                lines.append(
                    f"  {m['company_key']}: units {m['db_units']} vs {m['ref_units']} ({m['delta_units']:+}), shipments {m['db_shipments']} vs {m['ref_shipments']} ({m['delta_shipments']:+})"
                )

        return {
            "summary": summary,
            "sku_breakdown": sku_breakdown,
            "mismatches": mismatches[:200] if mismatches else [],
            "output_text": "\n".join(lines),
        }

    @app.route("/admin/system/deep-rebuild", methods=["GET", "POST"])
    def admin_system_deep_rebuild():
        if not is_admin():
            return redirect(url_for("admin_login"))

        ensure_system_tool_runs_table()

        today = datetime.utcnow().strftime("%Y-%m-%d")
        date_from = (request.values.get("date_from") or "2025-01-01").strip()
        date_to = (request.values.get("date_to") or today).strip()

        last_run = get_last_system_tool_run("deep_rebuild")
        result = None
        if request.method == "POST":
            ran_at = datetime.utcnow()
            status = "ok"
            try:
                result = run_deep_rebuild_diagnostics(date_from=date_from, date_to=date_to)
            except Exception as e:
                status = "failed"
                result = {
                    "summary": {
                        "date_from": date_from,
                        "date_to": date_to,
                        "orders": 0,
                        "units": 0,
                        "error": str(e),
                    },
                    "sku_breakdown": [],
                    "mismatches": [],
                    "output_text": f"Deep rebuild failed: {e}",
                }

            try:
                summary = result.get("summary") if isinstance(result, dict) else {}
                execute_db(
                    """
                    INSERT INTO system_tool_runs (tool_name, ran_at, status, orders, units, summary_json, output_text)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                    """,
                    (
                        "deep_rebuild",
                        ran_at,
                        status,
                        int((summary or {}).get("orders") or 0),
                        int((summary or {}).get("units") or 0),
                        json.dumps(summary or {}),
                        result.get("output_text") if isinstance(result, dict) else "",
                    ),
                )
                last_run = get_last_system_tool_run("deep_rebuild")
            except Exception:
                pass

        return render_template(
            "admin_deep_rebuild.html",
            date_from=date_from,
            date_to=date_to,
            result=result,
            last_run=last_run,
        )

    @app.route("/admin/diagnose-duplicates")
    def admin_diagnose_duplicates():
        """Diagnostic page to check for duplicate records"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            import json
            from collections import defaultdict
            
            # Check specific problematic order
            test_order = request.args.get('order', 'SO 0000263')
            
            output = []
            output.append(f"<h2>Diagnostic for order: {test_order}</h2>")
            
            # Get distribution headers for this order
            headers = query_db("""
                SELECT id, ship_date, source, order_number, created_at
                FROM devices_distributed
                WHERE order_number = %s
            """, (test_order,))
            
            if not headers:
                output.append(f"<p>❌ No distribution header found for {test_order}</p>")
            else:
                output.append(f"<h3>Distribution Headers ({len(headers)} found):</h3>")
                output.append("<ul>")
                for h in headers:
                    output.append(f"<li>dist_id={h['id']}, ship_date={h['ship_date']}, source={h['source']}, created={h['created_at']}</li>")
                output.append("</ul>")
                
                # Get all device_distribution_records
                dist_ids = [h['id'] for h in headers]
                records = query_db("""
                    SELECT dist_id, fields_json, id
                    FROM device_distribution_records
                    WHERE dist_id = ANY(%s)
                    ORDER BY dist_id, id
                """, (dist_ids,))
                records = records or []
                output.append(f"<h3>Device Distribution Records ({len(records)} found):</h3>")
                
                by_dist = defaultdict(list)
                for r in records:
                    by_dist[r['dist_id']].append(r)
                
                for dist_id, recs in by_dist.items():
                    output.append(f"<h4>dist_id {dist_id}: {len(recs)} records</h4>")
                    output.append("<ul>")
                    for r in recs:
                        try:
                            fields = json.loads(r['fields_json'])
                            sku = fields.get('SKU', 'N/A')
                            lot = fields.get('Lot', 'N/A')
                            qty = fields.get('Quantity', 'N/A')
                            output.append(f"<li>record_id={r['id']}: SKU={sku}, Lot={lot}, Qty={qty}</li>")
                        except Exception as e:
                            import traceback
                            print(f"[ERROR] admin_diagnostics (JSON parse): {e}")
                            traceback.print_exc()
                            output.append(f"<li>record_id={r['id']}: [Error parsing JSON]</li>")
                    output.append("</ul>")
                    
                    # Check for exact duplicates
                    seen = {}
                    for r in recs:
                        try:
                            fields = json.loads(r['fields_json'])
                            key = (fields.get('SKU'), fields.get('Lot'), fields.get('Quantity'))
                            if key in seen:
                                output.append(f"<p style='color:red'>⚠️ DUPLICATE: {key} appears in record_id {seen[key]} AND {r['id']}</p>")
                            else:
                                seen[key] = r['id']
                        except Exception as e:
                            import traceback
                            print(f"[ERROR] admin_diagnostics (duplicate check): {e}")
                            traceback.print_exc()
            
            # Check for orders with high record counts
            output.append("<hr><h3>Orders with high record counts (potential issues):</h3>")
            high_counts = query_db("""
                SELECT dd.order_number, dd.id as dist_id, COUNT(ddr.id) as record_count
                FROM devices_distributed dd
                LEFT JOIN device_distribution_records ddr ON ddr.dist_id = dd.id
                WHERE dd.ship_date >= '2025-09-01'
                GROUP BY dd.order_number, dd.id
                HAVING COUNT(ddr.id) > 3
                ORDER BY record_count DESC
                LIMIT 20
            """)
            
            if high_counts:
                output.append("<ul>")
                for row in high_counts:
                    output.append(f"<li><a href='/admin/diagnose-duplicates?order={row['order_number']}'>{row['order_number']}</a>: {row['record_count']} records</li>")
                output.append("</ul>")
            else:
                output.append("<p>✅ No orders with unusually high record counts</p>")
            
            # Check for duplicate headers
            output.append("<hr><h3>Orders with multiple distribution headers:</h3>")
            dupe_headers = query_db("""
                SELECT order_number, COUNT(*) as count
                FROM devices_distributed
                WHERE ship_date >= '2025-09-01'
                GROUP BY order_number
                HAVING COUNT(*) > 1
                ORDER BY count DESC
                LIMIT 20
            """)
            
            if dupe_headers:
                output.append("<ul>")
                for row in dupe_headers:
                    output.append(f"<li><a href='/admin/diagnose-duplicates?order={row['order_number']}'>{row['order_number']}</a>: {row['count']} headers</li>")
                output.append("</ul>")
            else:
                output.append("<p>✅ No duplicate distribution headers found</p>")
            
            output.append("<p><a href='/admin'>Back to Dashboard</a></p>")
            
            return "<html><body>" + "\n".join(output) + "</body></html>"
        
        except Exception as e:
            import traceback
            return f"<html><body><h2>Error</h2><pre>{traceback.format_exc()}</pre></body></html>", 500
    
    @app.route("/admin/delete-csv-imports", methods=["POST"])
    def admin_delete_csv_imports():
        """Delete ALL csv_import distribution records to eliminate duplicates"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            # Count CSV imports
            csv_dists = query_db("SELECT id FROM devices_distributed WHERE source = 'csv_import'") or []
            count = len(csv_dists) if csv_dists else 0
            
            if count == 0:
                flash("No CSV imports found.", "info")
                return redirect(url_for("admin_dashboard"))
            
            dist_ids = [d['id'] for d in csv_dists]
            
            # Delete device_distribution_records
            execute_db("DELETE FROM device_distribution_records WHERE dist_id = ANY(%s)", (dist_ids,))
            
            # Delete new_customer_records
            execute_db("DELETE FROM new_customer_records WHERE dist_id = ANY(%s)", (dist_ids,))
            
            # Delete distribution headers
            execute_db("DELETE FROM devices_distributed WHERE source = 'csv_import'")
            
            flash(f"✅ Deleted {count} CSV import distributions. Dashboard will now show correct data.", "success")
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f"Delete error: {str(e)}", "danger")
        
        return redirect(url_for("admin_dashboard"))
    
    @app.route("/admin/cleanup-unknown-lots", methods=["POST"])
    def admin_cleanup_unknown_lots():
        """Remove device distributions with UNKNOWN/MIXED/NONE lot data so they can be re-synced"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            # Count distributions to delete
            dists_to_delete = query_db("""
                SELECT dd.id
                FROM devices_distributed dd
                WHERE dd.source = 'shipstation'
                AND EXISTS (
                    SELECT 1 FROM device_distribution_records ddr2
                    WHERE ddr2.dist_id = dd.id
                    AND (ddr2.fields_json::text LIKE '%%"Lot": "UNKNOWN"%%'
                         OR ddr2.fields_json::text LIKE '%%"SKU": "MIXED"%%'
                         OR ddr2.fields_json::text LIKE '%%"SKU": "NONE"%%')
                )
            """)
            
            count = len(dists_to_delete) if dists_to_delete else 0
            
            if count > 0:
                # Delete device_distribution_records
                execute_db("""
                    DELETE FROM device_distribution_records
                    WHERE dist_id IN (
                        SELECT dd.id
                        FROM devices_distributed dd
                        WHERE dd.source = 'shipstation'
                        AND EXISTS (
                            SELECT 1 FROM device_distribution_records ddr2
                            WHERE ddr2.dist_id = dd.id
                            AND (ddr2.fields_json::text LIKE '%%"Lot": "UNKNOWN"%%'
                                 OR ddr2.fields_json::text LIKE '%%"SKU": "MIXED"%%'
                                 OR ddr2.fields_json::text LIKE '%%"SKU": "NONE"%%')
                        )
                    )
                """)
                
                # Delete new_customer_records
                execute_db("""
                    DELETE FROM new_customer_records
                    WHERE dist_id IN (
                        SELECT dd.id
                        FROM devices_distributed dd
                        WHERE dd.source = 'shipstation'
                        AND NOT EXISTS (
                            SELECT 1 FROM device_distribution_records ddr
                            WHERE ddr.dist_id = dd.id
                        )
                    )
                """)
                
                # Delete distribution headers
                execute_db("""
                    DELETE FROM devices_distributed
                    WHERE source = 'shipstation'
                    AND NOT EXISTS (
                        SELECT 1 FROM device_distribution_records ddr
                        WHERE ddr.dist_id = devices_distributed.id
                    )
                """)
                
                flash(f"✅ Cleaned up {count} distributions with UNKNOWN lots. Run sync again to re-import with correct lot data.", "success")
            else:
                flash("No records with UNKNOWN lots found.", "info")
                
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f"Cleanup error: {str(e)}", "danger")
        
        return redirect(url_for("admin_dashboard"))

    @app.route("/admin/import-csv", methods=["GET", "POST"])
    def admin_import_csv():
        """Import distributions from CSV file (bypass ShipStation API)"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        if request.method == "GET":
            return """
            <html><body>
            <h2>Import Distribution Log CSV</h2>
            <p>Upload your "Distribution Log (distributions).csv" file from SS4.py output:</p>
            <form method="post" enctype="multipart/form-data">
                <input type="file" name="csv_file" accept=".csv" required>
                <button type="submit">Import CSV</button>
            </form>
            <p><a href="/admin">Back to Dashboard</a></p>
            </body></html>
            """
        
        # POST - handle upload
        import csv
        from io import StringIO
        from werkzeug.utils import secure_filename
        
        file = request.files.get('csv_file')
        if not file:
            flash("No file uploaded", "danger")
            return redirect(url_for("admin_import_csv"))
        
        # Validate file upload (security hardening)
        is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions=app.config["ALLOWED_CSV_EXTENSIONS"])
        if not is_valid:
            flash(error_msg or "Invalid file. Only CSV files allowed.", "danger")
            return redirect(url_for("admin_import_csv"))
        
        try:
            # Read CSV and strip BOM if present
            content = file.read().decode('utf-8-sig')  # utf-8-sig removes BOM automatically
            reader = csv.DictReader(StringIO(content))
            
            # Get default rep
            default_rep = normalize_row(query_db("SELECT id FROM reps WHERE slug='ethan'", one=True))
            if not default_rep:
                flash("No default rep found", "danger")
                return redirect(url_for("admin_dashboard"))
            default_rep_id = default_rep["id"]
            
            imported = 0
            skipped = 0
            
            for row in reader:
                company = row.get('Company', '').strip()
                order_number = row.get('Order #', '').strip()
                ship_date_str = row.get('Ship Date', '').strip()
                
                if not order_number or not company:
                    continue
                
                # Check if exists
                existing = query_db(
                    "SELECT id FROM devices_distributed WHERE order_number = %s",
                    (order_number,), one=True
                )
                if existing:
                    skipped += 1
                    continue
                
                # Parse date
                from datetime import datetime as dt
                if ship_date_str:
                    try:
                        if ' ' in ship_date_str:
                            ship_date = dt.strptime(ship_date_str, '%Y-%m-%d %H:%M').date()
                        else:
                            ship_date = dt.strptime(ship_date_str, '%Y-%m-%d').date()
                    except Exception as e:
                        import traceback
                        print(f"[ERROR] admin_import_from_csv (date parse): {e}")
                        traceback.print_exc()
                        ship_date = None
                else:
                    ship_date = None
                
                # Get SKUs
                skus = []
                for i in range(1, 10):
                    sku = row.get(f'SKU{i}', '').strip()
                    qty = row.get(f'SKU{i} Qty', '').strip()
                    lot = row.get(f'Lot{i}', '').strip()
                    
                    if sku:
                        try:
                            qty_int = int(float(qty)) if qty else 0
                        except Exception as e:
                            import traceback
                            print(f"[ERROR] admin_import_from_csv (qty parse): {e}")
                            traceback.print_exc()
                            qty_int = 0
                        skus.append({'sku': sku, 'qty': qty_int, 'lot': lot or 'UNKNOWN'})
                
                # Insert distribution
                dist_id = execute_db(
                    """
                    INSERT INTO devices_distributed 
                    (rep_id, order_number, ship_date, source, created_at) 
                    VALUES (%s, %s, %s, %s, %s) 
                    RETURNING id
                    """,
                    (default_rep_id, order_number, ship_date, 'csv_import', datetime.now()),
                    returning_id=True
                )
                
                # Insert records
                if skus:
                    for sku_info in skus:
                        fields = {
                            "Facility Name": company,
                            "Company Key": clean_company_name(company),
                            "SKU": sku_info['sku'],
                            "Lot": sku_info['lot'],
                            "Quantity": sku_info['qty'],
                            "Distribution Date": ship_date.strftime('%Y-%m-%d') if ship_date else "",
                            "Order Number": order_number,
                        }
                        execute_db(
                            """
                            INSERT INTO device_distribution_records 
                            (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, missing_required_json, unexpected_fields_json)
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                            """,
                            (default_rep_id, dist_id, f"CSV_{order_number}", "Distribution Log CSV", datetime.now(), json.dumps(fields), json.dumps([]), json.dumps([]))
                        )
                    
                    # Also create new_customer_records entry for customer grouping
                    customer_fields = {
                        "Facility Name": company,
                        "Order Number": order_number,
                        "Ship Date": ship_date.strftime('%Y-%m-%d') if ship_date else "",
                        "Total SKUs": len(skus),
                        "Total Units": sum(s['qty'] for s in skus),
                    }
                    execute_db(
                        """
                        INSERT INTO new_customer_records 
                        (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, company_key)
                        VALUES (%s, %s, %s, %s, %s, %s, %s)
                        """,
                        (default_rep_id, dist_id, f"CSV_{order_number}", "CSV Import", datetime.now(), json.dumps(customer_fields), clean_company_name(company))
                    )
                else:
                    fields = {
                        "Facility Name": company,
                        "Company Key": clean_company_name(company),
                        "Order Number": order_number,
                        "Distribution Date": ship_date.strftime('%Y-%m-%d') if ship_date else "",
                        "SKU": "MIXED",
                        "Lot": "UNKNOWN",
                    }
                    execute_db(
                        """
                        INSERT INTO device_distribution_records 
                        (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, missing_required_json, unexpected_fields_json)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        """,
                        (default_rep_id, dist_id, f"CSV_{order_number}", "Distribution Log CSV", datetime.now(), json.dumps(fields), json.dumps([]), json.dumps([]))
                    )
                
                imported += 1
            
            flash(f"✅ Imported {imported} orders ({skipped} already existed)", "success")
            return redirect(url_for("admin_dashboard"))
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f"Import failed: {str(e)}", "danger")
            return redirect(url_for("admin_import_csv"))

    @app.route("/admin/sales-dashboard")
    def admin_sales_dashboard():
        # Allow iframe access with a special token or from specific domains
        # Check if accessed from iframe with auth token
        auth_token = request.args.get('token')
        iframe_auth_token = "silq2412"  # Simple token for iframe access
        
        if not is_admin() and auth_token != iframe_auth_token:
            return redirect(url_for("admin_login"))
        
        # Check if 'shipstation' rep exists (guardrail warning)
        shipstation_rep = query_db("SELECT id, name FROM reps WHERE slug='shipstation'", one=True)
        if not shipstation_rep:
            flash("⚠️ Warning: 'shipstation' rep not found. ShipStation imports may fail. The rep will be auto-created on next app startup.", "warning")
        
        try:
            import json
            from datetime import datetime, timedelta
            import re
            from zoneinfo import ZoneInfo
            from collections import defaultdict
            
            # Dashboard freshness metrics
            progress = get_sync_progress()
            _fresh = get_sync_freshness() or {}
            last_all_updated = _fresh.get("last_all_updated")
            last_shipstation_updated = _fresh.get("last_shipstation_updated")

            # Determine if dashboard reflects last sync completion
            dashboard_reflects_last_sync = False
            try:
                from datetime import timezone
                def _parse_any(ts):
                    if not ts:
                        return None
                    try:
                        dt = datetime.fromisoformat(ts) if isinstance(ts, str) else ts
                        if dt and dt.tzinfo is None:
                            dt = dt.replace(tzinfo=timezone.utc)
                        return dt
                    except Exception:
                        return None
                prog_dt = _parse_any(progress.get('updated_at')) if (progress and progress.get('status') == 'completed') else None
                ship_dt = _parse_any(last_shipstation_updated)
                if prog_dt and ship_dt and ship_dt >= prog_dt:
                    dashboard_reflects_last_sync = True
            except Exception:
                dashboard_reflects_last_sync = False

            # Canonical customer window for windowed stats (orders/units only)
            start_date = '2026-01-01'

            normalized_rows = fetch_distribution_records(exclude_sources=EXCLUDED_DIST_SOURCES)

            if not normalized_rows:
                return render_template(
                    "admin_device_distributions.html",
                    first_time=[],
                    reorders=[],
                    stats={
                        'total_customers': 0,
                        'first_time_customers': 0,
                        'repeat_customers': 0,
                        'total_orders': 0,
                        'new_this_quarter': 0,
                        'total_units': 0,
                        'sku_breakdown': [],
                        'total_customers_lifetime': 0,
                        'first_time_customers_lifetime': 0,
                        'repeat_customers_lifetime': 0,
                    },
                )

            # Group shipments into orders so multi-shipment orders are aggregated correctly
            orders = {}
            facility_orders = defaultdict(list)
            facility_first_seen = {}
            sku_totals_window = defaultdict(int)
            lot_totals_window = defaultdict(int)  # Track lot consumption
            lot_details = defaultdict(lambda: {"sku": set(), "first_used": None, "last_used": None, "units": 0})  # Track lot details
            total_units_window = 0
            shipments_in_window = set()

            for r in normalized_rows:
                facility_key = r["facility_key"]
                ship_date = _normalize_ship_date_ymd(r.get("ship_date"))
                order_id = r["order_group"]  # Canonical order number only
                shipment_key = r.get("shipment_key") or f"{order_id}::{ship_date}::{r.get('dist_id')}"

                if order_id not in orders:
                    orders[order_id] = {
                        "dist_id": r["dist_id"],
                        "customer_id": r.get("customer_id"),
                        "ship_dates": [],
                        "order_number": r["order_number"],
                        "source": r["source"],
                        "facility": r.get("facility_label") or facility_key,
                        "facility_key": facility_key,
                        "addr1": r["addr1"],
                        "city": r["city"],
                        "state": r["state"],
                        "postal": r["postal"],
                        "shipments": {},
                        "total_qty": 0,
                    }

                order_rec = orders[order_id]
                shipment_rec = order_rec["shipments"].setdefault(
                    shipment_key,
                    {
                        "shipment_id": r.get("shipment_id"),
                        "ship_date": ship_date or "Unknown",
                        "items": [],
                        "total_qty": 0,
                        "source": r["source"],
                    },
                )

                shipment_rec["items"].append({
                    "sku": r["sku"],
                    "lot": r.get("lot", "-"),
                    "qty": r["qty"],
                })
                shipment_rec["total_qty"] += r["qty"]

                order_rec["ship_dates"].append(ship_date)
                order_rec["total_qty"] += r["qty"]

                first_date = ship_date or "9999-12-31"
                if facility_key not in facility_first_seen or first_date < facility_first_seen[facility_key]:
                    facility_first_seen[facility_key] = first_date

                if ship_date and ship_date >= start_date:
                    sku_totals_window[r["sku"]] += r["qty"]
                    total_units_window += r["qty"]
                    shipments_in_window.add(shipment_key)
                
                # Track lot consumption (all-time, not windowed by start_date)
                lot = r.get("lot", "").strip()
                if lot and lot != "UNKNOWN" and lot != "-" and ship_date:
                    lot_totals_window[lot] += r["qty"]
                    lot_details[lot]["sku"].add(r["sku"])
                    lot_details[lot]["units"] += r["qty"]
                    # Track first and last used dates
                    if lot_details[lot]["first_used"] is None or ship_date < lot_details[lot]["first_used"]:
                        lot_details[lot]["first_used"] = ship_date
                    if lot_details[lot]["last_used"] is None or ship_date > lot_details[lot]["last_used"]:
                        lot_details[lot]["last_used"] = ship_date

            # Build facility -> orders mapping after shipments are folded into orders
            total_orders_window = 0
            for order in orders.values():
                dates_clean = [d for d in (order.get("ship_dates") or []) if d]
                if dates_clean:
                    max_date = max(dates_clean)
                    order["ship_date"] = _normalize_ship_date_ymd(max_date) or "Unknown"
                else:
                    order["ship_date"] = "Unknown"
                order["shipment_count"] = len(order["shipments"])

                # Aggregate items across shipments by SKU/lot
                sku_lot_totals = defaultdict(int)
                for sh in order["shipments"].values():
                    for item in sh["items"]:
                        sku_lot_totals[(item.get("sku", "Unknown"), item.get("lot", "-"))] += item.get("qty", 0)
                order["items"] = [
                    {"sku": sku, "lot": lot, "qty": qty}
                    for (sku, lot), qty in sku_lot_totals.items()
                ]

                if order["ship_date"] and order["ship_date"] >= start_date:
                    total_orders_window += 1

                facility_orders[order['facility_key']].append(order)

            # Sort each facility's orders by date (most recent first)
            for facility, orders_list in facility_orders.items():
                orders_list.sort(key=lambda x: x['ship_date'], reverse=True)

            # Lifetime classification (no windowing): customers with exactly one order ever are first-time
            first_time_customers = []
            repeat_customers = []
            for facility, orders_list in facility_orders.items():
                most_recent_order = orders_list[0]
                order_count = len(orders_list)
                most_recent_order['order_count'] = order_count
                if order_count == 1:
                    first_time_customers.append(most_recent_order)
                else:
                    repeat_customers.append(most_recent_order)

            # Sort by date (most recent first)
            first_time_customers.sort(key=lambda x: x['ship_date'], reverse=True)
            repeat_customers.sort(key=lambda x: x['ship_date'], reverse=True)

            facility_name_by_key = {}
            for o in orders.values():
                if o.get("facility_key") and o.get("facility"):
                    facility_name_by_key[o.get("facility_key")] = o.get("facility")

            # Facility suggestions for manual entry search/autocomplete
            facility_suggestions = sorted(set(facility_name_by_key.values()))

            # Notes for first-time customers (up to 3 most recent)
            ft_customer_ids = sorted({o.get("customer_id") for o in first_time_customers if o.get("customer_id")})
            recent_notes_by_customer = {}
            if ft_customer_ids:
                note_rows = query_db(
                    """
                    SELECT id, customer_id, note_text, note_date, created_at, author
                    FROM customer_notes
                    WHERE customer_id = ANY(%s)
                    ORDER BY customer_id, created_at DESC
                    """,
                    (ft_customer_ids,),
                ) or []
                for nr in note_rows:
                    # Normalize row -> dict so we can safely add/overwrite fields (DictRow is not safely mutable)
                    try:
                        nr = dict(nr)
                    except Exception:
                        nr = nr if isinstance(nr, dict) else {}

                    # Normalize dates to YYYY-MM-DD strings for template slicing
                    if nr.get("created_at"):
                        nr["created_at"] = _normalize_ship_date_ymd(nr.get("created_at")) or ""
                    if nr.get("note_date"):
                        nr["note_date"] = _normalize_ship_date_ymd(nr.get("note_date")) or ""

                    # Defensive: ensure expected keys exist
                    if "author" not in nr:
                        nr["author"] = None
                    if "id" not in nr:
                        nr["id"] = None

                    cid = nr.get("customer_id")
                    if cid not in recent_notes_by_customer:
                        recent_notes_by_customer[cid] = []
                    if len(recent_notes_by_customer[cid]) < 3:
                        recent_notes_by_customer[cid].append(nr)

            pac = ZoneInfo('America/Los_Angeles')
            now = datetime.now(pac)
            current_quarter_start_dt = now.replace(month=((now.month-1)//3)*3+1, day=1)
            current_quarter_start = current_quarter_start_dt.strftime('%Y-%m-%d')

            # Lifetime stats
            first_time_count_lifetime = len(first_time_customers)
            repeat_count_lifetime = len(repeat_customers)
            total_customers_lifetime = first_time_count_lifetime + repeat_count_lifetime

            # Windowed new-this-quarter (still helpful for ops reporting)
            new_this_quarter = sum(
                1 for _, first_date in facility_first_seen.items()
                if first_date and first_date >= current_quarter_start
            )

            # SKU breakdown (windowed)
            sku_breakdown = []
            for sku, units in sorted(sku_totals_window.items(), key=lambda x: x[1], reverse=True):
                if sku != "Unknown" and total_units_window > 0:
                    percentage = round((units / total_units_window) * 100, 1)
                    sku_breakdown.append({
                        'sku': sku,
                        'units': units,
                        'percentage': percentage
                    })

            # Lot consumption breakdown (all-time, not windowed)
            # Calculate total all-time units for percentage calculation
            total_units_alltime = sum(lot_totals_window.values())
            lot_breakdown = []
            for lot, units in sorted(lot_totals_window.items(), key=lambda x: x[1], reverse=True):
                if lot and lot != "UNKNOWN" and lot != "-":
                    details = lot_details[lot]
                    # Get primary SKU (most common SKU for this lot, or first if multiple)
                    primary_sku = sorted(details["sku"])[0] if details["sku"] else "Unknown"
                    lot_breakdown.append({
                        'lot': lot,
                        'units': units,
                        'sku': primary_sku,
                        'skus': sorted(list(details["sku"])),  # All SKUs for this lot
                        'first_used': details["first_used"],
                        'last_used': details["last_used"],
                        'percentage': round((units / total_units_alltime) * 100, 1) if total_units_alltime > 0 else 0
                    })

            stats = {
                'total_orders': total_orders_window,
                'total_customers': total_customers_lifetime,
                'first_time_customers': first_time_count_lifetime,
                'repeat_customers': repeat_count_lifetime,
                'new_this_quarter': new_this_quarter,
                'total_units': total_units_window,
                'sku_breakdown': sku_breakdown[:5],
                'lot_breakdown': lot_breakdown[:20],  # Top 20 lots by consumption
                'total_customers_lifetime': total_customers_lifetime,
                'first_time_customers_lifetime': first_time_count_lifetime,
                'repeat_customers_lifetime': repeat_count_lifetime,
            }
            
            # Format PT display for last-updated timestamps
            def _fmt_pt(ts):
                if not ts:
                    return None
                try:
                    if isinstance(ts, str):
                        dt = datetime.fromisoformat(ts)
                    else:
                        dt = ts
                    if dt.tzinfo is None:
                        from datetime import timezone
                        dt = dt.replace(tzinfo=timezone.utc)
                    dt = dt.astimezone(pac)
                    return dt.strftime('%Y-%m-%d %H:%M PT')
                except Exception:
                    return str(ts)

            last_all_updated_display = _fmt_pt(last_all_updated)
            last_shipstation_updated_display = _fmt_pt(last_shipstation_updated)

            response = make_response(render_template("admin_device_distributions.html", 
                                 first_time=first_time_customers,
                                 reorders=repeat_customers,
                                 stats=stats,
                                 start_date=start_date,
                                 last_all_updated=last_all_updated_display,
                                 last_shipstation_updated=last_shipstation_updated_display,
                                 sync_progress=progress,
                                 dashboard_reflects_last_sync=dashboard_reflects_last_sync,
                                 all_reps=query_db("SELECT * FROM reps WHERE active = 1 ORDER BY name") or [],
                                 today=now.strftime('%Y-%m-%d'),
                                 facility_suggestions=facility_suggestions,
                                 recent_notes_by_customer=recent_notes_by_customer))
            
            # Allow iframe embedding from silq.tech
            response.headers['X-Frame-Options'] = 'ALLOW-FROM https://www.silq.tech'
            response.headers['Content-Security-Policy'] = "frame-ancestors 'self' https://www.silq.tech https://silq.tech"
            return response
        
        except Exception as e:
            import traceback
            traceback.print_exc()
            return f"Error loading sales dashboard: {e}", 500
    
    @app.route("/admin/device-distributions")
    def admin_device_distributions():
        """Redirect old URL to new sales dashboard"""
        return redirect(url_for("admin_sales_dashboard"))
    
    @app.route("/admin/sales-dashboard/export")
    def admin_sales_dashboard_export():
        """Export Sales Dashboard current view as CSV"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            import io
            import csv as _csv
            from datetime import datetime
            from collections import defaultdict
            
            # Use same logic as dashboard to get current view
            start_date = '2026-01-01'
            normalized_rows = fetch_distribution_records(exclude_sources=EXCLUDED_DIST_SOURCES)
            
            if not normalized_rows:
                flash("No data to export.", "warning")
                return redirect(url_for("admin_sales_dashboard"))
            
            # Group shipments into orders (same logic as dashboard)
            orders = {}
            for r in normalized_rows:
                facility_key = r["facility_key"]
                ship_date = _normalize_ship_date_ymd(r.get("ship_date")) or "Unknown"
                order_id = r["order_group"]
                
                if order_id not in orders:
                    orders[order_id] = {
                        "dist_id": r["dist_id"],
                        "customer_id": r.get("customer_id"),
                        "ship_dates": [],
                        "order_number": r["order_number"],
                        "source": r["source"],
                        "facility": r.get("facility_label") or facility_key,
                        "facility_key": facility_key,
                        "addr1": r["addr1"],
                        "city": r["city"],
                        "state": r["state"],
                        "postal": r["postal"],
                        "items": [],
                        "total_qty": 0,
                        "order_count": 0,  # Will be set later
                    }
                
                order_rec = orders[order_id]
                order_rec["ship_dates"].append(ship_date)
                order_rec["items"].append({
                    "sku": r["sku"],
                    "lot": r.get("lot", "-"),
                    "qty": r["qty"],
                })
                order_rec["total_qty"] += r["qty"]
            
            # Build facility -> orders mapping
            facility_orders = defaultdict(list)
            facility_first_seen = {}
            for order in orders.values():
                dates_clean = [d for d in (order.get("ship_dates") or []) if d and d != "Unknown"]
                order["ship_date"] = max(dates_clean) if dates_clean else "Unknown"
                
                # Aggregate items by SKU/lot
                sku_lot_totals = defaultdict(int)
                for item in order["items"]:
                    key = (item.get("sku", "Unknown"), item.get("lot", "-"))
                    sku_lot_totals[key] += item.get("qty", 0)
                order["items"] = [
                    {"sku": sku, "lot": lot, "qty": qty}
                    for (sku, lot), qty in sku_lot_totals.items()
                ]
                
                first_date = order["ship_date"] or "9999-12-31"
                facility_key = order["facility_key"]
                if facility_key not in facility_first_seen or first_date < facility_first_seen[facility_key]:
                    facility_first_seen[facility_key] = first_date
                
                facility_orders[facility_key].append(order)
            
            # Classify as first-time or repeat
            first_time_customers = []
            repeat_customers = []
            for facility, orders_list in facility_orders.items():
                orders_list.sort(key=lambda x: x['ship_date'], reverse=True)
                most_recent_order = orders_list[0]
                order_count = len(orders_list)
                most_recent_order['order_count'] = order_count
                if order_count == 1:
                    first_time_customers.append(most_recent_order)
                else:
                    repeat_customers.append(most_recent_order)
            
            # Sort by date (most recent first)
            first_time_customers.sort(key=lambda x: x['ship_date'], reverse=True)
            repeat_customers.sort(key=lambda x: x['ship_date'], reverse=True)
            
            # Build CSV
            buf = io.StringIO()
            w = _csv.writer(buf)
            w.writerow([
                "Type",
                "Ship Date",
                "Order Number",
                "Facility Name",
                "Address",
                "City",
                "State",
                "Zip",
                "Total Units",
                "Item Count",
                "SKU",
                "Lot",
                "Quantity",
                "Source",
                "Customer ID",
                "Dist ID",
            ])
            
            # Export first-time customers
            for order in first_time_customers:
                base_row = [
                    "First-Time",
                    (order.get("ship_date") or "")[:10],
                    order.get("order_number") or "",
                    order.get("facility") or "",
                    order.get("addr1") or "",
                    order.get("city") or "",
                    order.get("state") or "",
                    order.get("postal") or "",
                    order.get("total_qty") or 0,
                    len(order.get("items", [])),
                ]
                
                # One row per SKU/Lot combination
                items = order.get("items", [])
                if items:
                    for item in items:
                        w.writerow(base_row + [
                            item.get("sku", ""),
                            item.get("lot", ""),
                            item.get("qty", 0),
                            order.get("source") or "",
                            order.get("customer_id") or "",
                            order.get("dist_id") or "",
                        ])
                else:
                    # No items, still write one row
                    w.writerow(base_row + ["", "", 0, order.get("source") or "", order.get("customer_id") or "", order.get("dist_id") or ""])
            
            # Export repeat customers
            for order in repeat_customers:
                base_row = [
                    "Repeat",
                    (order.get("ship_date") or "")[:10],
                    order.get("order_number") or "",
                    order.get("facility") or "",
                    order.get("addr1") or "",
                    order.get("city") or "",
                    order.get("state") or "",
                    order.get("postal") or "",
                    order.get("total_qty") or 0,
                    len(order.get("items", [])),
                ]
                
                # One row per SKU/Lot combination
                items = order.get("items", [])
                if items:
                    for item in items:
                        w.writerow(base_row + [
                            item.get("sku", ""),
                            item.get("lot", ""),
                            item.get("qty", 0),
                            order.get("source") or "",
                            order.get("customer_id") or "",
                            order.get("dist_id") or "",
                        ])
                else:
                    # No items, still write one row
                    w.writerow(base_row + ["", "", 0, order.get("source") or "", order.get("customer_id") or "", order.get("dist_id") or ""])
            
            # Generate filename with timestamp
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"sales_dashboard_{timestamp}.csv"
            
            out = make_response(buf.getvalue())
            out.headers["Content-Type"] = "text/csv; charset=utf-8"
            out.headers["Content-Disposition"] = f"attachment; filename={filename}"
            return out
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f"Error exporting CSV: {str(e)}", "danger")
            return redirect(url_for("admin_sales_dashboard"))
    
    @app.route("/admin/customer/note/add", methods=["POST"])
    def admin_customer_note_add():
        """Add a note to a customer - supports both AJAX (JSON) and form POST (redirect)"""
        if not is_admin():
            if request.headers.get("Content-Type", "").startswith("application/json") or request.args.get("ajax") == "1":
                return jsonify({"ok": False, "error": "Unauthorized"}), 403
            return redirect(url_for("admin_login"))
        try:
            customer_id = request.form.get("customer_id")
            note_text = (request.form.get("note_text") or "").strip()
            note_date_str = request.form.get("note_date") or datetime.utcnow().date().isoformat()
            
            # Check if this is an AJAX request (via header or query param)
            is_ajax = (
                request.headers.get("X-Requested-With") == "XMLHttpRequest" or
                request.args.get("ajax") == "1" or
                request.headers.get("Accept", "").find("application/json") >= 0
            )
            
            if not customer_id or not note_text:
                if is_ajax:
                    return jsonify({"ok": False, "error": "Customer ID and note text are required."}), 400
                flash("Customer ID and note text are required.", "danger")
                return redirect(url_for("admin_sales_dashboard"))
            
            try:
                customer_id_int = int(customer_id)
            except Exception:
                if is_ajax:
                    return jsonify({"ok": False, "error": "Invalid customer ID."}), 400
                flash("Invalid customer ID.", "danger")
                return redirect(url_for("admin_sales_dashboard"))
            
            # Verify customer exists
            customer = query_db("SELECT id FROM customers WHERE id = %s", (customer_id_int,), one=True)
            if not customer:
                if is_ajax:
                    return jsonify({"ok": False, "error": "Customer not found."}), 404
                flash("Customer not found.", "danger")
                return redirect(url_for("admin_sales_dashboard"))
            
            # Parse note date
            try:
                note_date = datetime.fromisoformat(note_date_str).date() if note_date_str else datetime.utcnow().date()
            except Exception:
                note_date = datetime.utcnow().date()
            
            execute_db(
                "INSERT INTO customer_notes (customer_id, note_text, note_date, created_at, updated_at, author) VALUES (%s, %s, %s, %s, %s, %s)",
                (customer_id_int, note_text, note_date, datetime.utcnow(), datetime.utcnow(), "admin"),
            )
            
            if is_ajax:
                return jsonify({"ok": True, "message": "Note added successfully."}), 200
            
            flash("Note added successfully.", "success")
            return redirect(url_for("admin_sales_dashboard"))
        except Exception as e:
            import traceback
            traceback.print_exc()
            if request.headers.get("X-Requested-With") == "XMLHttpRequest" or request.args.get("ajax") == "1":
                return jsonify({"ok": False, "error": f"Error adding note: {str(e)}"}), 500
            flash(f"Error adding note: {str(e)}", "danger")
            return redirect(url_for("admin_sales_dashboard"))
    
    @app.route("/admin/customer/<int:customer_id>/note/<int:note_id>/edit", methods=["POST"])
    def admin_customer_note_edit(customer_id, note_id):
        """Edit an existing customer note"""
        if not is_admin():
            is_ajax = (
                request.headers.get("X-Requested-With") == "XMLHttpRequest"
                or request.args.get("ajax") == "1"
                or ("application/json" in (request.headers.get("Accept") or ""))
            )
            if is_ajax:
                return jsonify({"ok": False, "error": "Unauthorized"}), 403
            return redirect(url_for("admin_login"))
        try:
            note_text = (request.form.get("note_text") or "").strip()
            note_date_str = request.form.get("note_date") or ""
            
            is_ajax = (
                request.headers.get("X-Requested-With") == "XMLHttpRequest"
                or request.args.get("ajax") == "1"
                or ("application/json" in (request.headers.get("Accept") or ""))
            )
            
            if not note_text:
                if is_ajax:
                    return jsonify({"ok": False, "error": "Note text is required."}), 400
                flash("Note text is required.", "danger")
                return redirect(url_for("admin_customer_crm_profile", customer_id=customer_id))

            # Verify note exists and belongs to customer (avoid silent no-ops)
            note_row = query_db(
                "SELECT id, customer_id, note_text, note_date, created_at, updated_at, author FROM customer_notes WHERE id = %s AND customer_id = %s",
                (note_id, customer_id),
                one=True,
            )
            if not note_row:
                if is_ajax:
                    return jsonify({"ok": False, "error": "Note not found."}), 404
                flash("Note not found.", "danger")
                return redirect(url_for("admin_customer_crm_profile", customer_id=customer_id))
            
            # Parse note date
            try:
                note_date = datetime.fromisoformat(note_date_str).date() if note_date_str else None
            except Exception:
                note_date = None

            updated_at = datetime.utcnow()
            
            if note_date:
                execute_db(
                    "UPDATE customer_notes SET note_text = %s, note_date = %s, updated_at = %s WHERE id = %s AND customer_id = %s",
                    (note_text, note_date, updated_at, note_id, customer_id),
                )
            else:
                execute_db(
                    "UPDATE customer_notes SET note_text = %s, updated_at = %s WHERE id = %s AND customer_id = %s",
                    (note_text, updated_at, note_id, customer_id),
                )
            
            if is_ajax:
                return jsonify(
                    {
                        "ok": True,
                        "note": {
                            "id": int(note_id),
                            "customer_id": int(customer_id),
                            "note_text": note_text,
                            "note_date": note_date.isoformat() if note_date else None,
                            "updated_at": updated_at.isoformat(),
                            "author": (note_row.get("author") if isinstance(note_row, dict) else None) or "admin",
                        },
                    }
                ), 200
            
            flash("Note updated successfully.", "success")
            return redirect(url_for("admin_customer_crm_profile", customer_id=customer_id))
        except Exception as e:
            import traceback
            traceback.print_exc()
            is_ajax = (
                request.headers.get("X-Requested-With") == "XMLHttpRequest"
                or request.args.get("ajax") == "1"
                or ("application/json" in (request.headers.get("Accept") or ""))
            )
            if is_ajax:
                return jsonify({"ok": False, "error": "Error updating note."}), 500
            flash(f"Error updating note: {str(e)}", "danger")
            return redirect(url_for("admin_customer_crm_profile", customer_id=customer_id))
    
    @app.route("/admin/customer/<int:customer_id>/note/<int:note_id>/delete", methods=["POST", "DELETE"])
    def admin_customer_note_delete(customer_id, note_id):
        """Delete a customer note - supports both AJAX (JSON) and form POST (redirect)"""
        if not is_admin():
            is_ajax = (
                request.headers.get("X-Requested-With") == "XMLHttpRequest"
                or request.args.get("ajax") == "1"
                or ("application/json" in (request.headers.get("Accept") or ""))
            )
            if is_ajax:
                return jsonify({"ok": False, "error": "Unauthorized"}), 403
            return redirect(url_for("admin_login"))
        
        try:
            # Verify note exists and belongs to customer
            note = query_db(
                "SELECT id FROM customer_notes WHERE id = %s AND customer_id = %s",
                (note_id, customer_id),
                one=True
            )
            
            if not note:
                is_ajax = (
                    request.headers.get("X-Requested-With") == "XMLHttpRequest"
                    or request.args.get("ajax") == "1"
                    or ("application/json" in (request.headers.get("Accept") or ""))
                )
                if is_ajax:
                    return jsonify({"ok": False, "error": "Note not found."}), 404
                flash("Note not found.", "danger")
                return redirect(url_for("admin_customer_crm_profile", customer_id=customer_id))
            
            execute_db(
                "DELETE FROM customer_notes WHERE id = %s AND customer_id = %s",
                (note_id, customer_id)
            )
            
            if is_ajax:
                return jsonify({"ok": True, "deleted_id": int(note_id), "customer_id": int(customer_id)}), 200
            
            flash("Note deleted successfully.", "success")
            return redirect(url_for("admin_customer_crm_profile", customer_id=customer_id))
        except Exception as e:
            import traceback
            traceback.print_exc()
            is_ajax = (
                request.headers.get("X-Requested-With") == "XMLHttpRequest"
                or request.args.get("ajax") == "1"
                or ("application/json" in (request.headers.get("Accept") or ""))
            )
            if is_ajax:
                return jsonify({"ok": False, "error": "Error deleting note."}), 500
            flash(f"Error deleting note: {str(e)}", "danger")
            return redirect(url_for("admin_customer_crm_profile", customer_id=customer_id))
    
    @app.route("/admin/distribution-log")
    def admin_distribution_log():
        """Simple chronological list of all distributions with filters"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            import json
            from datetime import datetime, timedelta
            from collections import defaultdict

            # Date range defaults - show all 2025 data
            today = datetime.now()
            # Support date filtering from query params (e.g., ?date_from=2025-01-01&date_to=2025-12-31)
            date_from = request.args.get('date_from', '2025-01-01')  # Show all 2025 orders by default
            date_to = request.args.get('date_to', today.strftime('%Y-%m-%d'))

            # Validate and swap if date_from > date_to
            if date_from and date_to and date_from > date_to:
                date_from, date_to = date_to, date_from

            # Optional: filter to a specific shipment/distribution id
            dist_id_filter = None
            dist_id_raw = (request.args.get("dist_id") or "").strip()
            if dist_id_raw:
                try:
                    dist_id_filter = int(dist_id_raw)
                except Exception:
                    dist_id_filter = None

            # Optional: focus on a specific customer (used when navigating from the customer database)
            customer_id_filter = None
            customer_id_raw = (request.args.get("customer_id") or "").strip()
            if customer_id_raw:
                try:
                    customer_id_filter = int(customer_id_raw)
                except Exception:
                    customer_id_filter = None

            # Optional: filter by canonical order group (used by CRM "recent orders" links)
            order_group_filter = (request.args.get("order_group") or "").strip()
            if order_group_filter:
                order_group_filter = canonical_order_number(order_group_filter)

            # Optional: highlight a particular distribution row without filtering the list
            focus_dist_id = None
            focus_dist_id_raw = (request.args.get("focus_dist_id") or "").strip()
            if focus_dist_id_raw:
                try:
                    focus_dist_id = int(focus_dist_id_raw)
                except Exception:
                    focus_dist_id = None

            def _fallback_distribution_number(dist_id_val, ship_date: str | None) -> str:
                try:
                    n = int(dist_id_val)
                except Exception:
                    return ""
                year = ""
                sd = (ship_date or "").strip()
                if re.match(r"^[0-9]{4}-", sd):
                    year = sd[:4]
                else:
                    year = str(datetime.now().year)
                return f"D-{year}-{n:06d}"

            def _build_distributions(rows: list[dict]) -> list[dict]:
                shipments = {}
                for r in rows:
                    ship_date = r.get("ship_date") or "Unknown"
                    shipment_key = r.get("shipment_key") or f"{r.get('order_group')}::{ship_date}::{r.get('dist_id')}"
                    dist_id_val = r.get("dist_id")
                    dist_number = (r.get("distribution_number") or "").strip() or _fallback_distribution_number(dist_id_val, ship_date)
                    sales_order_number = r.get("order_number")
                    display_order = sales_order_number or dist_number or (f"DIST-{dist_id_val}" if dist_id_val else "")

                    if shipment_key not in shipments:
                        shipments[shipment_key] = {
                            "dist_id": dist_id_val,
                            "distribution_number": dist_number,
                            "sales_order_number": sales_order_number,
                            # Back-compat for templates/JS search
                            "order_number": display_order,
                            "ship_date": ship_date,
                            "facility": r.get("facility_label") or r.get("facility_key"),
                            "city": r.get("city"),
                            "state": r.get("state"),
                            "rep_name": r.get("rep_name") or "Unassigned",
                            "source": r.get("source") or "manual",
                            "total_qty": 0,
                            "items": [],
                            "item_count": 0,
                            "files_count": 0,
                        }

                    shipments[shipment_key]["items"].append(
                        {
                            "sku": r.get("sku", "Unknown"),
                            "lot": r.get("lot", "UNKNOWN"),
                            "quantity": r.get("qty", 0),
                        }
                    )
                    shipments[shipment_key]["total_qty"] += r.get("qty", 0)
                    shipments[shipment_key]["item_count"] = len(shipments[shipment_key]["items"])

                distributions_out = list(shipments.values())
                distributions_out.sort(key=lambda x: x.get("ship_date") or "", reverse=True)

                # File counts for shipment/evidence attachments
                dist_ids = [d.get("dist_id") for d in distributions_out if d.get("dist_id")]
                if dist_ids:
                    count_rows = query_db(
                        """
                        SELECT dist_id, COUNT(*) AS cnt
                        FROM device_distribution_records
                        WHERE dist_id = ANY(%s)
                          AND fields_json ILIKE %s
                        GROUP BY dist_id
                        """,
                        (dist_ids, '%"Record Type": "shipment_record"%'),
                    ) or []
                    counts = {r.get("dist_id"): int(r.get("cnt") or 0) for r in count_rows}
                    
                    # Also get document details for focused distribution
                    doc_rows = query_db(
                        """
                        SELECT dist_id, id, original_filename, stored_filename, uploaded_at, fields_json
                        FROM device_distribution_records
                        WHERE dist_id = ANY(%s)
                          AND fields_json ILIKE %s
                        ORDER BY uploaded_at DESC
                        """,
                        (dist_ids, '%"Record Type": "shipment_record"%'),
                    ) or []
                    
                    # Group documents by dist_id
                    docs_by_dist = defaultdict(list)
                    for doc_row in doc_rows:
                        did = doc_row.get("dist_id")
                        if did:
                            try:
                                fields = json.loads(doc_row.get("fields_json") or "{}")
                                file_type = fields.get("File Type") or fields.get("file_type") or ""
                            except Exception:
                                file_type = ""
                            docs_by_dist[did].append({
                                "id": doc_row.get("id"),
                                "original_filename": doc_row.get("original_filename") or "file",
                                "stored_filename": doc_row.get("stored_filename"),
                                "uploaded_at": doc_row.get("uploaded_at") or "",
                                "file_type": file_type,
                            })
                    
                    for d in distributions_out:
                        did = d.get("dist_id")
                        if did in counts:
                            d["files_count"] = counts[did]
                        if did in docs_by_dist:
                            d["documents"] = docs_by_dist[did]

                return distributions_out

            normalized_rows = fetch_distribution_records(exclude_sources=EXCLUDED_DIST_SOURCES)
            if not normalized_rows:
                return render_template(
                    "admin_distribution_log.html",
                    distributions=[],
                    all_reps=query_db("SELECT * FROM reps WHERE active = 1 ORDER BY name") or [],
                    all_customers=query_db("SELECT id, facility_name, city, state FROM customers ORDER BY facility_name") or [],
                    today=today.strftime('%Y-%m-%d'),
                    date_from=date_from,
                    date_to=date_to,
                    shipstation_count=0,
                    customer_focus=None,
                    customer_focus_distributions=[],
                    focus_dist_id=focus_dist_id,
                )

            filtered_rows = []
            for r in normalized_rows:
                ship_date = r.get("ship_date") or ""
                if dist_id_filter and int(r.get("dist_id") or 0) != dist_id_filter:
                    continue
                if order_group_filter and (r.get("order_group") or "") != order_group_filter:
                    continue
                if ship_date and date_from and ship_date < date_from:
                    continue
                if ship_date and date_to and ship_date > date_to:
                    continue
                if not ship_date and date_from:
                    continue
                filtered_rows.append(r)

            distributions = _build_distributions(filtered_rows)

            # Customer focus card (same distribution list, but restricted to that customer)
            customer_focus = None
            customer_focus_distributions: list[dict] = []
            if customer_id_filter:
                customer_focus = normalize_row(
                    query_db(
                        "SELECT id, facility_name, company_key, city, state FROM customers WHERE id = %s",
                        (int(customer_id_filter),),
                        one=True,
                    )
                )
                focus_rows = fetch_distribution_records(
                    exclude_sources=EXCLUDED_DIST_SOURCES,
                    customer_ids=[int(customer_id_filter)],
                    date_from=date_from,
                    date_to=date_to,
                ) or []
                if order_group_filter:
                    focus_rows = [r for r in focus_rows if (r.get("order_group") or "") == order_group_filter]
                customer_focus_distributions = _build_distributions(focus_rows)

            shipstation_count = sum(1 for d in distributions if (d.get("source") or "").lower() == "shipstation")

            return render_template(
                "admin_distribution_log.html",
                distributions=distributions,
                all_reps=query_db("SELECT * FROM reps WHERE active = 1 ORDER BY name") or [],
                all_customers=query_db("SELECT id, facility_name, city, state FROM customers ORDER BY facility_name") or [],
                today=today.strftime('%Y-%m-%d'),
                date_from=date_from,
                date_to=date_to,
                shipstation_count=shipstation_count,
                customer_focus=customer_focus,
                customer_focus_distributions=customer_focus_distributions,
                focus_dist_id=focus_dist_id,
                order_group_filter=order_group_filter,
            )
        
        except Exception as e:
            import traceback
            traceback.print_exc()
            return f"Error loading distribution log: {e}", 500
    
    @app.route("/admin/customer-database")
    def admin_customer_database():
        """Legacy route. Redirect to CRM-style customer database."""
        return redirect(url_for("admin_customers"))

    @app.route("/admin/customers")
    def admin_customers():
        """CRM-style customer list (editable customers + rep assignments + notes)."""
        auth_token = (request.args.get("token") or "").strip()
        iframe_auth_token = "silq2412"  # Simple token for iframe access (GET-only)

        if not is_admin() and auth_token != iframe_auth_token:
            return redirect(url_for("admin_login"))
        try:
            q = (request.args.get("q") or "").strip()
            state = (request.args.get("state") or "").strip().upper()
            rep_id_raw = (request.args.get("rep_id") or "").strip()
            try:
                rep_id = int(rep_id_raw) if rep_id_raw else None
            except Exception:
                rep_id = None

            where = []
            params = []
            if q:
                where.append("(c.facility_name ILIKE %s OR c.company_key ILIKE %s)")
                params.extend([f"%{q}%", f"%{q}%"])
            if state:
                where.append("(c.state = %s)")
                params.append(state)
            if rep_id:
                where.append(
                    "(c.primary_rep_id = %s OR EXISTS (SELECT 1 FROM customer_rep_assignments cra WHERE cra.customer_id = c.id AND cra.rep_id = %s))"
                )
                params.extend([rep_id, rep_id])
            where_sql = ("WHERE " + " AND ".join(where)) if where else ""

            customers = query_db(
                f"""
                SELECT c.*, r.name AS primary_rep_name
                FROM customers c
                LEFT JOIN reps r ON r.id = c.primary_rep_id
                {where_sql}
                ORDER BY c.updated_at DESC NULLS LAST, c.id DESC
                """,
                tuple(params),
            ) or []

            assign_rows = query_db(
                """
                SELECT cra.customer_id, cra.rep_id, cra.is_primary, r.name AS rep_name
                FROM customer_rep_assignments cra
                JOIN reps r ON r.id = cra.rep_id
                ORDER BY r.name
                """
            ) or []
            reps_by_customer = {}
            for ar in assign_rows:
                cid = ar.get("customer_id")
                if cid not in reps_by_customer:
                    reps_by_customer[cid] = []
                reps_by_customer[cid].append({
                    "id": ar.get("rep_id"),
                    "name": ar.get("rep_name"),
                    "is_primary": bool(ar.get("is_primary")),
                })

            note_rows = query_db(
                """
                SELECT customer_id, COUNT(*) AS note_count, MAX(created_at) AS last_note_at
                FROM customer_notes
                GROUP BY customer_id
                """
            ) or []
            notes_by_customer = {nr.get("customer_id"): nr for nr in note_rows}

            all_reps = query_db("SELECT id, name FROM reps WHERE active = 1 ORDER BY name") or []

            customers_out = []
            for c in customers:
                cdict = normalize_row(c)
                cid = cdict.get("id") if isinstance(cdict, dict) else None
                meta = notes_by_customer.get(cid) or {}
                customers_out.append({
                    **(cdict or {}),
                    "primary_rep_name": (cdict or {}).get("primary_rep_name"),
                    "assigned_reps": reps_by_customer.get(cid, []),
                    "note_count": meta.get("note_count", 0),
                    "last_note_at": meta.get("last_note_at"),
                })

            # Order history (compact): list all distributions/orders per customer (single query)
            customer_ids = [c.get("id") for c in customers_out if c.get("id")]
            # Rich recent orders (canonical order_group rollups)
            customer_ids = [c.get("id") for c in customers_out if c.get("id")]
            recent_orders_by_customer = {}
            if customer_ids:
                rows = fetch_distribution_records(
                    exclude_sources=EXCLUDED_DIST_SOURCES,
                    customer_ids=customer_ids,
                )
                summaries = build_order_group_summaries(rows)
                for (cid, _og), rec in summaries.items():
                    if cid not in recent_orders_by_customer:
                        recent_orders_by_customer[cid] = []
                    recent_orders_by_customer[cid].append(rec)

                for cid, lst in recent_orders_by_customer.items():
                    lst.sort(key=lambda r: r.get("last_ship_date") or r.get("first_ship_date") or "", reverse=True)
                    recent_orders_by_customer[cid] = lst[:5]

            for c in customers_out:
                cid = c.get("id")
                c["recent_orders"] = recent_orders_by_customer.get(cid, [])

            # Format current date for template (avoid Jinja2 date filter which doesn't exist)
            current_date = datetime.now().strftime("%Y-%m-%d")
            
            response = make_response(
                render_template(
                "admin_customers.html",
                customers=customers_out,
                all_reps=all_reps,
                filters={"q": q, "state": state, "rep_id": rep_id},
                current_date=current_date,
                )
            )

            if auth_token == iframe_auth_token and not is_admin():
                # Allow iframe embedding from silq.tech
                response.headers['X-Frame-Options'] = 'ALLOW-FROM https://www.silq.tech'
                response.headers['Content-Security-Policy'] = "frame-ancestors 'self' https://www.silq.tech https://silq.tech"

            return response
        except Exception as e:
            import traceback
            traceback.print_exc()
            return f"Error loading customers: {e}", 500

    @app.route("/admin/customers/<int:customer_id>", methods=["GET", "POST"])
    def admin_customer_crm_profile(customer_id):
        auth_token = (request.args.get("token") or "").strip()
        iframe_auth_token = "silq2412"  # Simple token for iframe access (GET-only)

        if not is_admin():
            if request.method == "GET" and auth_token == iframe_auth_token:
                pass
            else:
                return redirect(url_for("admin_login"))
        try:
            customer = normalize_row(query_db("SELECT * FROM customers WHERE id = %s", (customer_id,), one=True))
            if not customer:
                return "Customer not found", 404

            all_reps = query_db("SELECT id, name FROM reps WHERE active = 1 ORDER BY name") or []
            assigned = query_db(
                "SELECT rep_id, is_primary FROM customer_rep_assignments WHERE customer_id = %s",
                (customer_id,),
            ) or []
            assigned_ids = {a.get("rep_id") for a in assigned}
            is_primary_map = {a.get("rep_id"): bool(a.get("is_primary")) for a in assigned}

            if request.method == "POST":
                facility_name = (request.form.get("facility_name") or "").strip()
                address1 = (request.form.get("address1") or "").strip()
                address2 = (request.form.get("address2") or "").strip()
                city = (request.form.get("city") or "").strip()
                state = (request.form.get("state") or "").strip().upper()
                zip_code = (request.form.get("zip") or "").strip()
                contact_name = (request.form.get("contact_name") or "").strip()
                contact_phone = (request.form.get("contact_phone") or "").strip()
                contact_email = (request.form.get("contact_email") or "").strip()
                primary_rep_id = request.form.get("primary_rep_id")
                try:
                    primary_rep_id_int = int(primary_rep_id) if primary_rep_id else None
                except Exception:
                    primary_rep_id_int = None

                execute_db(
                    """
                    UPDATE customers
                    SET facility_name=%s, address1=%s, address2=%s, city=%s, state=%s, zip=%s,
                        contact_name=%s, contact_phone=%s, contact_email=%s, primary_rep_id=%s, updated_at=%s
                    WHERE id=%s
                    """,
                    (
                        facility_name or customer.get("facility_name"),
                        address1 or None,
                        address2 or None,
                        city or None,
                        state or None,
                        zip_code or None,
                        contact_name or None,
                        contact_phone or None,
                        contact_email or None,
                        primary_rep_id_int,
                        datetime.utcnow(),
                        customer_id,
                    ),
                )

                rep_ids = request.form.getlist("rep_ids")
                rep_ids_int = []
                for rid in rep_ids:
                    try:
                        rep_ids_int.append(int(rid))
                    except Exception:
                        pass
                if primary_rep_id_int and primary_rep_id_int not in rep_ids_int:
                    rep_ids_int.append(primary_rep_id_int)

                execute_db("DELETE FROM customer_rep_assignments WHERE customer_id = %s", (customer_id,))
                for rid in sorted(set(rep_ids_int)):
                    execute_db(
                        """
                        INSERT INTO customer_rep_assignments (customer_id, rep_id, is_primary)
                        VALUES (%s, %s, %s)
                        ON CONFLICT (customer_id, rep_id) DO UPDATE SET is_primary = EXCLUDED.is_primary
                        """,
                        (customer_id, rid, bool(primary_rep_id_int and rid == primary_rep_id_int)),
                    )

                note_text = (request.form.get("new_note") or "").strip()
                if note_text:
                    note_date_str = request.form.get("note_date") or datetime.utcnow().date().isoformat()
                    try:
                        note_date = datetime.fromisoformat(note_date_str).date() if note_date_str else datetime.utcnow().date()
                    except Exception:
                        note_date = datetime.utcnow().date()
                    execute_db(
                        "INSERT INTO customer_notes (customer_id, note_text, note_date, created_at, updated_at, author) VALUES (%s, %s, %s, %s, %s, %s)",
                        (customer_id, note_text, note_date, datetime.utcnow(), datetime.utcnow(), "admin"),
                    )

                flash("Customer updated.", "success")
                return redirect(url_for("admin_customer_crm_profile", customer_id=customer_id))

            notes = query_db(
                "SELECT id, customer_id, note_text, note_date, created_at, updated_at, author FROM customer_notes WHERE customer_id = %s ORDER BY created_at DESC",
                (customer_id,),
            ) or []

            # High-level stats and shipment-level order history (do NOT collapse by order_number).
            rows = fetch_distribution_records(
                exclude_sources=EXCLUDED_DIST_SOURCES,
                customer_ids=[customer_id],
            )
            shipments = {}
            for r in rows or []:
                did = r.get("dist_id")
                if not did:
                    continue
                rec = shipments.get(did)
                if not rec:
                    rec = {
                        "dist_id": did,
                        "ship_date": (r.get("ship_date") or "")[:10],
                        "order_number": r.get("order_number") or r.get("distribution_number") or f"DIST-{did}",
                        "source": r.get("source") or "manual",
                        "sku_units": {},
                        "total_units": 0,
                    }
                    shipments[did] = rec

                sku = (r.get("sku") or "Unknown").strip() or "Unknown"
                try:
                    qty_int = int(r.get("qty") or 0)
                except Exception:
                    qty_int = 0
                rec["sku_units"][sku] = int(rec["sku_units"].get(sku, 0)) + qty_int
                rec["total_units"] += qty_int

            shipments_all = list(shipments.values())
            shipments_all.sort(key=lambda x: x.get("ship_date") or "", reverse=True)

            ship_dates = [s.get("ship_date") for s in shipments_all if s.get("ship_date")]
            first_ship_date = min(ship_dates) if ship_dates else None
            last_ship_date = max(ship_dates) if ship_dates else None
            total_units = sum(int(s.get("total_units") or 0) for s in shipments_all)
            customer_stats = {
                "orders": len(shipments_all),
                "units": total_units,
                "first_ship_date": first_ship_date,
                "last_ship_date": last_ship_date,
            }

            # Pagination for shipment history table
            try:
                page = int(request.args.get("page") or "1")
            except Exception:
                page = 1
            if page < 1:
                page = 1
            per_page = 25
            total_orders = len(shipments_all)
            start = (page - 1) * per_page
            end = start + per_page
            orders_page_raw = shipments_all[start:end]
            has_prev = page > 1
            has_next = end < total_orders

            # Attach computed display helpers expected by the template.
            orders_page = []
            for s in orders_page_raw:
                sku_breakdown = [
                    {"sku": sku, "units": units}
                    for sku, units in sorted((s.get("sku_units") or {}).items(), key=lambda kv: kv[1], reverse=True)
                ]
                sku_summary = ", ".join(f"{x['sku']} ({x['units']})" for x in sku_breakdown[:4]) if sku_breakdown else ""
                if sku_breakdown and len(sku_breakdown) > 4:
                    sku_summary += ", …"
                orders_page.append(
                    {
                        "dist_id": s.get("dist_id"),
                        "first_ship_date": s.get("ship_date") or "",
                        "order_number": s.get("order_number") or "",
                        "total_units": s.get("total_units") or 0,
                        "sku_count": len(sku_breakdown),
                        "source_label": _source_label({s.get("source") or ""}),
                        "sku_summary": sku_summary or "-",
                    }
                )

            # Distribution Record attachments for this customer (admin-only download)
            dist_record_rows = []
            if is_admin():
                dist_record_rows = query_db(
                    """
                    SELECT ddr.id AS record_id,
                           ddr.dist_id,
                           ddr.original_filename,
                           ddr.uploaded_at,
                           dd.ship_date,
                                                     dd.order_number,
                                                     dd.distribution_number
                    FROM device_distribution_records ddr
                    JOIN devices_distributed dd ON dd.id = ddr.dist_id
                    WHERE COALESCE(dd.customer_id, ddr.customer_id) = %s
                      AND (
                        ddr.fields_json ILIKE %s
                        OR ddr.fields_json ILIKE %s
                      )
                    ORDER BY dd.ship_date DESC NULLS LAST, ddr.uploaded_at DESC
                    """,
                    (customer_id, '%"Record Type": "shipment_record"%', '%"Source": "shipment_record"%'),
                ) or []
            dist_records = [normalize_row(r) for r in (dist_record_rows or [])]

            response = make_response(
                render_template(
                "admin_customer_crm_profile.html",
                customer=customer,
                all_reps=all_reps,
                assigned_ids=assigned_ids,
                is_primary_map=is_primary_map,
                notes=notes,
                stats=customer_stats,
                order_history=orders_page,
                order_history_total=total_orders,
                order_history_page=page,
                order_history_has_prev=has_prev,
                order_history_has_next=has_next,
                dist_records=dist_records,
                )
            )

            if request.method == "GET" and auth_token == iframe_auth_token and not is_admin():
                response.headers['X-Frame-Options'] = 'ALLOW-FROM https://www.silq.tech'
                response.headers['Content-Security-Policy'] = "frame-ancestors 'self' https://www.silq.tech https://silq.tech"

            return response
        except Exception as e:
            import traceback
            traceback.print_exc()
            return f"Error loading customer profile: {e}", 500
    
    @app.route("/admin/customer/<int:customer_id>/notes/json", methods=["GET"])
    def admin_customer_notes_json(customer_id):
        """Return customer notes as JSON for modal display"""
        if not is_admin():
            return jsonify({"error": "Unauthorized"}), 403
        
        try:
            notes = query_db(
                "SELECT id, customer_id, note_text, note_date, created_at, updated_at, author FROM customer_notes WHERE customer_id = %s ORDER BY created_at DESC",
                (customer_id,),
            ) or []
            
            notes_list = []
            for note in notes:
                notes_list.append({
                    "id": note.get("id"),
                    "note_text": note.get("note_text", ""),
                    "note_date": str(note.get("note_date", "")) if note.get("note_date") else None,
                    "created_at": str(note.get("created_at", "")) if note.get("created_at") else None,
                    "updated_at": str(note.get("updated_at", "")) if note.get("updated_at") else None,
                    "author": note.get("author", "")
                })
            
            return jsonify({"notes": notes_list})
        except Exception as e:
            import traceback
            traceback.print_exc()
            return jsonify({"error": str(e)}), 500
    
    @app.route("/admin/manual-distribution-entry", methods=["POST"])
    def admin_manual_distribution_entry():
        """
        Admin manually enters distribution when rep emails form.
        Supports file attachment (scanned form image/PDF) as evidence.
        Uses canonical field normalization for consistency with ShipStation data.
        """
        if not is_admin():
            return redirect(url_for("admin_login"))

        wants_json = (request.form.get("ajax") == "1") or (
            (request.headers.get("X-Requested-With") or "").lower() == "xmlhttprequest"
        )
        
        try:
            # Get form data
            rep_id = (request.form.get("rep_id") or "").strip()
            ship_date = request.form.get("ship_date")
            order_number = request.form.get("order_number") or f"MANUAL-{int(datetime.now().timestamp())}"
            facility_name = (request.form.get("facility_name") or "").strip()
            
            # Customer selection (new)
            customer_id = request.form.get("customer_id", "").strip()
            customer_row = None
            customer_data = {}
            
            if customer_id:
                try:
                    customer_id_int = int(customer_id)
                    customer_row = normalize_row(query_db(
                        "SELECT * FROM customers WHERE id = %s",
                        (customer_id_int,),
                        one=True
                    ))
                    if customer_row:
                        customer_data = {
                            "Facility Name": customer_row.get("facility_name", ""),
                            "Address1": customer_row.get("address1", ""),
                            "Address2": customer_row.get("address2", ""),
                            "City": customer_row.get("city", ""),
                            "State": customer_row.get("state", ""),
                            "Zip": customer_row.get("postal", ""),
                            "Contact Name": customer_row.get("contact_name", ""),
                            "Phone": customer_row.get("phone", ""),
                            "Email": customer_row.get("email", ""),
                        }
                        # Use customer's facility name if provided
                        if customer_data["Facility Name"]:
                            facility_name = customer_data["Facility Name"]
                except Exception as e:
                    print(f"Error loading customer {customer_id}: {e}")
                    customer_row = None

            sku_list = request.form.getlist("sku[]") or request.form.getlist("sku")
            lot_list = request.form.getlist("lot_number[]") or request.form.getlist("lot_number")
            qty_list = request.form.getlist("quantity[]") or request.form.getlist("quantity")
            sku_rows = []
            for idx, sku_val in enumerate(sku_list):
                sku_val = (sku_val or "").strip()
                lot_val = (lot_list[idx] if idx < len(lot_list) else "").strip()
                qty_val = qty_list[idx] if idx < len(qty_list) else ""
                try:
                    qty_int = int(qty_val)
                except Exception:
                    qty_int = 0
                if sku_val and lot_val and qty_int > 0:
                    sku_rows.append({"sku": sku_val, "lot": lot_val, "quantity": qty_int})
            
            # Handle optional file attachment (scanned form image)
            file = request.files.get("evidence_file")
            stored_filename = "MANUAL_ENTRY"
            original_filename = "MANUAL_ENTRY"
            evidence_stored_filename = None
            evidence_original_filename = None
            
            # Validate required fields
            if not all([rep_id, ship_date, facility_name]) or not sku_rows:
                msg = "Missing required fields: rep, ship date, facility name, and at least one SKU row"
                if wants_json:
                    return make_response(json.dumps({"ok": False, "error": msg}), 400, {"Content-Type": "application/json"})
                flash(msg, "danger")
                return redirect(url_for("admin_distribution_log"))

            try:
                rep_id_int = int(rep_id)
            except Exception:
                msg = "Invalid rep selection"
                if wants_json:
                    return make_response(json.dumps({"ok": False, "error": msg}), 400, {"Content-Type": "application/json"})
                flash(msg, "danger")
                return redirect(url_for("admin_distribution_log"))

            # Customer linkage and rep assignment
            if not customer_row:
                # Fall back to existing behavior
                customer_row = normalize_row(find_or_create_customer(facility_name=facility_name, primary_rep_id=rep_id_int))
                customer_id = customer_row.get("id") if customer_row else None
            
            ensure_rep_assignment(customer_id, rep_id_int, make_primary_if_none=True)
            
            # Create distribution header
            dist_id = execute_db(
                """INSERT INTO devices_distributed (rep_id, order_number, ship_date, source, created_at, customer_id) 
                   VALUES (%s, %s, %s, 'manual', %s, %s) RETURNING id""",
                (rep_id_int, order_number, ship_date, datetime.now().isoformat(), customer_id),
                returning_id=True
            )
            
            if not dist_id:
                msg = "Failed to create distribution record"
                if wants_json:
                    return make_response(json.dumps({"ok": False, "error": msg}), 500, {"Content-Type": "application/json"})
                flash(msg, "danger")
                return redirect(url_for("admin_distribution_log"))
            
            # Handle file upload if provided (with security validation)
            if file and file.filename:
                # Validate file upload (security hardening)
                is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions=app.config["ALLOWED_DISTRIBUTION_EXTENSIONS"])
                if not is_valid:
                    msg = error_msg or "Invalid file type"
                    if wants_json:
                        return make_response(json.dumps({"ok": False, "error": msg}), 400, {"Content-Type": "application/json"})
                    flash(msg, "danger")
                    return redirect(url_for("admin_distribution_log"))
                
                saved = save_distribution_record_file(int(dist_id), file)
                evidence_stored_filename = saved.get("stored_filename")
                evidence_original_filename = saved.get("original_filename")

                # Store a separate shipment-record attachment row (do not mix with SKU line items)
                evidence_fields = {
                    "Record Type": "shipment_record",
                    "Source": "shipment_record",
                    "Origin": "manual_entry_evidence",
                    "Facility Name": facility_name,
                    "Order Number": order_number,
                    "Ship Date": ship_date,
                }
                # Merge customer data into evidence file fields_json
                if customer_data:
                    evidence_fields.update(customer_data)
                execute_db(
                    """INSERT INTO device_distribution_records
                       (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, customer_id)
                       VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                    (
                        rep_id_int,
                        dist_id,
                        evidence_stored_filename,
                        evidence_original_filename,
                        datetime.now().isoformat(),
                        json.dumps(evidence_fields),
                        customer_id,
                    ),
                )
            
            total_units = 0
            for row in sku_rows:
                total_units += row.get("quantity", 0)
                raw_fields = {
                    "Facility Name": facility_name,
                    "SKU": row.get("sku"),
                    "Lot": row.get("lot"),
                    "Quantity": str(row.get("quantity", 0)),
                    "Distribution Date": ship_date,
                    "Order Number": order_number
                }

                normalized_fields = normalize_fields_json(raw_fields, source="manual")
                # Merge customer data if available
                if customer_data:
                    normalized_fields.update(customer_data)
                if customer_id:
                    normalized_fields["Company Key"] = canonical_customer_key(facility_name)

                execute_db(
                    """INSERT INTO device_distribution_records 
                       (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, customer_id) 
                       VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                    (rep_id_int, dist_id, stored_filename, original_filename, 
                     datetime.now().isoformat(), json.dumps(normalized_fields), customer_id)
                )

            company_key = canonical_customer_key(facility_name)
            execute_db(
                     """INSERT INTO new_customer_records 
                         (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, company_key, customer_id)
                         VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
                (
                    rep_id_int,
                    dist_id,
                    stored_filename,
                    original_filename,
                    datetime.now().isoformat(),
                    json.dumps({
                        "Facility Name": facility_name,
                        "Order Number": order_number,
                        "Ship Date": ship_date,
                        "Total Units": total_units,
                    }),
                    company_key,
                    customer_id,
                ),
            )
            
            flash(f"✅ Distribution entry saved successfully (Order: {order_number})", "success")
            if evidence_original_filename:
                flash(f"📎 Evidence file attached: {evidence_original_filename}", "info")

            if wants_json:
                return make_response(
                    json.dumps({"ok": True, "dist_id": dist_id, "order_number": order_number}),
                    200,
                    {"Content-Type": "application/json"},
                )
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            if wants_json:
                return make_response(
                    json.dumps({"ok": False, "error": f"Error saving distribution: {str(e)}"}),
                    500,
                    {"Content-Type": "application/json"},
                )
            flash(f"Error saving distribution: {str(e)}", "danger")
        
        return redirect(url_for("admin_distribution_log"))

    @app.route("/admin/distributions/<int:dist_id>/edit", methods=["GET", "POST"])
    def admin_edit_distribution(dist_id: int):
        if not is_admin():
            return redirect(url_for("admin_login"))

        dist = normalize_row(
            query_db(
                """
                SELECT dd.*, c.facility_name AS customer_facility
                FROM devices_distributed dd
                LEFT JOIN customers c ON c.id = dd.customer_id
                WHERE dd.id = %s
                """,
                (dist_id,),
                one=True,
            )
        )
        if not dist:
            return "Distribution not found", 404

        source = (dist.get("source") or "").lower()
        # Allow editing of ShipStation distributions (show warning in form instead)

        all_reps = query_db("SELECT id, name FROM reps WHERE active = 1 ORDER BY name") or []

        # Load existing line items (exclude shipment_record attachments)
        rows = query_db(
            "SELECT id, fields_json FROM device_distribution_records WHERE dist_id = %s ORDER BY id",
            (dist_id,),
        ) or []
        items = []
        non_attachment_ids = []
        for r in rows:
            try:
                f = json.loads(r.get("fields_json") or "{}")
            except Exception:
                f = {}
            rt = (f.get("Record Type") or f.get("record_type") or "").strip().lower()
            src = (f.get("Source") or f.get("source") or "").strip().lower()
            if rt == "shipment_record" or src == "shipment_record":
                continue
            sku = (f.get("SKU") or f.get("sku") or "").strip()
            lot = (f.get("Lot") or f.get("lot") or "").strip()
            qty = f.get("Quantity") or f.get("qty") or ""
            try:
                qty_int = int(float(qty)) if str(qty).strip() else 0
            except Exception:
                qty_int = 0
            if sku or lot or qty_int:
                items.append({"sku": sku, "lot": lot, "quantity": qty_int})
                non_attachment_ids.append(r.get("id"))

        if request.method == "POST":
            rep_id_raw = (request.form.get("rep_id") or "").strip()
            ship_date = (request.form.get("ship_date") or "").strip()
            order_number = (request.form.get("order_number") or "").strip()
            facility_name = (request.form.get("facility_name") or "").strip()

            sku_list = request.form.getlist("sku[]") or request.form.getlist("sku")
            lot_list = request.form.getlist("lot_number[]") or request.form.getlist("lot_number")
            qty_list = request.form.getlist("quantity[]") or request.form.getlist("quantity")
            sku_rows = []
            for idx, sku_val in enumerate(sku_list):
                sku_val = (sku_val or "").strip()
                lot_val = (lot_list[idx] if idx < len(lot_list) else "").strip()
                qty_val = qty_list[idx] if idx < len(qty_list) else ""
                try:
                    qty_int = int(qty_val)
                except Exception:
                    qty_int = 0
                if sku_val and lot_val and qty_int > 0:
                    sku_rows.append({"sku": sku_val, "lot": lot_val, "quantity": qty_int})

            if not all([rep_id_raw, ship_date, facility_name]) or not sku_rows:
                flash("Missing required fields.", "danger")
                return redirect(url_for("admin_edit_distribution", dist_id=dist_id))

            try:
                rep_id_int = int(rep_id_raw)
            except Exception:
                flash("Invalid rep selection.", "danger")
                return redirect(url_for("admin_edit_distribution", dist_id=dist_id))

            if not order_number:
                order_number = dist.get("order_number") or f"DIST-{dist_id}"

            customer_row = normalize_row(find_or_create_customer(facility_name=facility_name, primary_rep_id=rep_id_int))
            customer_id = customer_row.get("id") if customer_row else None
            ensure_rep_assignment(customer_id, rep_id_int, make_primary_if_none=True)

            execute_db(
                "UPDATE devices_distributed SET rep_id=%s, order_number=%s, ship_date=%s, customer_id=%s WHERE id=%s",
                (rep_id_int, order_number, ship_date, customer_id, dist_id),
            )

            # Remove previous line items (keep shipment_record attachments)
            if non_attachment_ids:
                execute_db("DELETE FROM device_distribution_records WHERE id = ANY(%s)", (non_attachment_ids,))

            total_units = 0
            for row in sku_rows:
                total_units += row.get("quantity", 0)
                raw_fields = {
                    "Facility Name": facility_name,
                    "SKU": row.get("sku"),
                    "Lot": row.get("lot"),
                    "Quantity": str(row.get("quantity", 0)),
                    "Distribution Date": ship_date,
                    "Order Number": order_number,
                }
                normalized_fields = normalize_fields_json(raw_fields, source=source or "manual")
                if customer_id:
                    normalized_fields["Company Key"] = canonical_customer_key(facility_name)
                execute_db(
                    """INSERT INTO device_distribution_records
                       (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, customer_id)
                       VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                    (
                        rep_id_int,
                        dist_id,
                        "EDITED",
                        "EDITED",
                        datetime.utcnow().isoformat(timespec="seconds") + "Z",
                        json.dumps(normalized_fields),
                        customer_id,
                    ),
                )

            execute_db("DELETE FROM new_customer_records WHERE dist_id = %s", (dist_id,))
            company_key = canonical_customer_key(facility_name)
            execute_db(
                """INSERT INTO new_customer_records
                   (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, company_key, customer_id)
                   VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
                (
                    rep_id_int,
                    dist_id,
                    "EDITED",
                    "EDITED",
                    datetime.utcnow().isoformat(timespec="seconds") + "Z",
                    json.dumps(
                        {
                            "Facility Name": facility_name,
                            "Order Number": order_number,
                            "Ship Date": ship_date,
                            "Total Units": total_units,
                        }
                    ),
                    company_key,
                    customer_id,
                ),
            )

            flash("Distribution updated.", "success")
            return redirect(url_for("admin_distribution_log"))

        return render_template(
            "admin_distribution_edit.html",
            dist=dist,
            items=items or [{"sku": "", "lot": "", "quantity": 0}],
            all_reps=all_reps,
        )

    @app.route("/admin/distributions/<int:dist_id>/delete", methods=["POST"])
    def admin_delete_distribution(dist_id: int):
        if not is_admin():
            return redirect(url_for("admin_login"))

        def _delete_distribution_by_id(did: int) -> dict:
            """Delete a distribution header and its related records + attachment files."""
            dist_row = normalize_row(query_db("SELECT * FROM devices_distributed WHERE id = %s", (did,), one=True))
            if not dist_row:
                return {"ok": False, "error": "Distribution not found"}

            # Delete any attachment files in DIST_RECORDS_DIR
            recs = query_db(
                "SELECT id, stored_filename, fields_json FROM device_distribution_records WHERE dist_id = %s",
                (did,),
            ) or []
            deleted_files = 0
            for r in recs:
                try:
                    f = json.loads(r.get("fields_json") or "{}")
                except Exception:
                    f = {}
                rt = (f.get("Record Type") or f.get("record_type") or "").strip().lower()
                src = (f.get("Source") or f.get("source") or "").strip().lower()
                if rt == "shipment_record" or src == "shipment_record":
                    stored = r.get("stored_filename")
                    if stored:
                        delete_distribution_record_blob(stored)
                        deleted_files += 1

            execute_db("DELETE FROM device_distribution_records WHERE dist_id = %s", (did,))
            execute_db("DELETE FROM new_customer_records WHERE dist_id = %s", (did,))
            execute_db("DELETE FROM devices_distributed WHERE id = %s", (did,))
            return {"ok": True, "deleted_files": deleted_files}

        result = _delete_distribution_by_id(dist_id)
        if not result.get("ok"):
            flash(result.get("error") or "Delete failed.", "danger")
        else:
            flash("Distribution deleted.", "success")
        return redirect(url_for("admin_distribution_log"))

    @app.route("/admin/system/delete-manual-distributions", methods=["POST"])
    def admin_delete_manual_distributions():
        if not is_admin():
            return redirect(url_for("admin_login"))

        confirm = (request.form.get("confirm") or "").strip()
        dry_run = (request.form.get("dry_run") or "").strip().lower() in {"1", "true", "yes", "y"}
        date_from = (request.form.get("date_from") or "").strip()
        date_to = (request.form.get("date_to") or "").strip()

        if confirm != "DELETE_MANUAL":
            flash("Type DELETE_MANUAL to confirm.", "warning")
            return redirect(url_for("admin_distribution_log"))

        # psycopg2 uses %s placeholders; do not embed literal % in SQL.
        where = ["coalesce(source,'') ILIKE %s"]
        params = ["manual%"]
        if date_from:
            where.append("ship_date >= %s")
            params.append(date_from)
        if date_to:
            where.append("ship_date <= %s")
            params.append(date_to)
        where_sql = " AND ".join(where)

        rows = query_db(
            f"SELECT id FROM devices_distributed WHERE {where_sql} ORDER BY ship_date NULLS LAST, id ASC",
            tuple(params),
        ) or []
        dist_ids = [int(r.get("id")) for r in rows if r.get("id")]

        if dry_run:
            flash(f"Dry-run: would delete {len(dist_ids)} manual distributions.", "info")
            return redirect(url_for("admin_distribution_log"))

        # Reuse the same deletion behavior as the single-delete path.
        deleted = 0
        errors = 0
        deleted_files = 0
        for did in dist_ids:
            try:
                # Inline the helper from admin_delete_distribution scope (dup is acceptable here to keep refactor minimal)
                dist_row = normalize_row(query_db("SELECT * FROM devices_distributed WHERE id = %s", (did,), one=True))
                if not dist_row:
                    continue
                recs = query_db(
                    "SELECT id, stored_filename, fields_json FROM device_distribution_records WHERE dist_id = %s",
                    (did,),
                ) or []
                for r in recs:
                    try:
                        f = json.loads(r.get("fields_json") or "{}")
                    except Exception:
                        f = {}
                    rt = (f.get("Record Type") or f.get("record_type") or "").strip().lower()
                    src = (f.get("Source") or f.get("source") or "").strip().lower()
                    if rt == "shipment_record" or src == "shipment_record":
                        stored = r.get("stored_filename")
                        if stored:
                            delete_distribution_record_blob(stored)
                            deleted_files += 1

                execute_db("DELETE FROM device_distribution_records WHERE dist_id = %s", (did,))
                execute_db("DELETE FROM new_customer_records WHERE dist_id = %s", (did,))
                execute_db("DELETE FROM devices_distributed WHERE id = %s", (did,))
                deleted += 1
            except Exception:
                errors += 1

        flash(
            f"Deleted manual distributions: deleted={deleted} | errors={errors} | deleted_files={deleted_files}",
            "success" if errors == 0 else "warning",
        )
        return redirect(url_for("admin_distribution_log"))

    @app.route("/admin/distribution-log/export", methods=["GET"])
    def admin_distribution_log_export():
        """Export the Distribution Log as CSV (excludes shipment_record attachments)."""
        if not is_admin():
            return redirect(url_for("admin_login"))

        try:
            import io
            import csv as _csv
            from datetime import datetime

            today = datetime.now()
            date_from = request.args.get('date_from', '2025-01-01')
            date_to = request.args.get('date_to', today.strftime('%Y-%m-%d'))
            if date_from and date_to and date_from > date_to:
                date_from, date_to = date_to, date_from

            dist_id_filter = None
            dist_id_raw = (request.args.get("dist_id") or "").strip()
            if dist_id_raw:
                try:
                    dist_id_filter = int(dist_id_raw)
                except Exception:
                    dist_id_filter = None

            normalized_rows = fetch_distribution_records(exclude_sources=EXCLUDED_DIST_SOURCES)
            filtered_rows = []
            for r in normalized_rows or []:
                ship_date = r.get("ship_date") or ""
                if dist_id_filter and int(r.get("dist_id") or 0) != dist_id_filter:
                    continue
                if ship_date and date_from and ship_date < date_from:
                    continue
                if ship_date and date_to and ship_date > date_to:
                    continue
                if not ship_date and date_from:
                    continue
                filtered_rows.append(r)

            shipments = {}
            for r in filtered_rows:
                ship_date = (r.get("ship_date") or "Unknown")
                shipment_key = r.get("shipment_key") or f"{r.get('order_group')}::{ship_date}::{r.get('dist_id')}"
                order_num = r.get("order_number") or f"DIST-{r.get('dist_id')}"
                if shipment_key not in shipments:
                    shipments[shipment_key] = {
                        "dist_id": r.get("dist_id"),
                        "order_number": order_num,
                        "ship_date": ship_date,
                        "facility": r.get("facility_label") or r.get("facility_key"),
                        "rep_name": r.get("rep_name") or "Unassigned",
                        "source": r.get("source") or "manual",
                        "tracking_number": r.get("tracking_number") or "",
                        "total_qty": 0,
                        "item_count": 0,
                        "items": [],
                    }
                shipments[shipment_key]["items"].append(
                    {
                        "sku": r.get("sku", "Unknown"),
                        "lot": r.get("lot", "UNKNOWN"),
                        "quantity": r.get("qty", 0),
                    }
                )
                shipments[shipment_key]["total_qty"] += r.get("qty", 0)
                shipments[shipment_key]["item_count"] = len(shipments[shipment_key]["items"])

            distributions = list(shipments.values())
            distributions.sort(key=lambda x: x.get("ship_date") or "", reverse=True)

            buf = io.StringIO()
            w = _csv.writer(buf)
            w.writerow([
                "Ship Date",
                "Order #",
                "Facility",
                "Rep",
                "Items",
                "Total Qty",
                "Source",
                "Tracking",
                "Dist ID",
            ])
            for d in distributions:
                w.writerow([
                    (d.get("ship_date") or "")[:10],
                    d.get("order_number") or "",
                    d.get("facility") or "",
                    d.get("rep_name") or "",
                    d.get("item_count") or 0,
                    d.get("total_qty") or 0,
                    d.get("source") or "",
                    d.get("tracking_number") or "",
                    d.get("dist_id") or "",
                ])

            out = make_response(buf.getvalue())
            out.headers["Content-Type"] = "text/csv; charset=utf-8"
            out.headers["Content-Disposition"] = f"attachment; filename=distribution_log_{date_from}_to_{date_to}.csv"
            return out
        except Exception as e:
            import traceback
            traceback.print_exc()
            return f"Error exporting distribution log: {e}", 500

    @app.route("/admin/distribution-records", methods=["GET"])
    def admin_distribution_records_list():
        if not is_admin():
            abort(403)
        dist_id_raw = (request.args.get("dist_id") or "").strip()
        try:
            dist_id_int = int(dist_id_raw)
        except Exception:
            return make_response(json.dumps({"ok": False, "error": "Invalid dist_id"}), 400, {"Content-Type": "application/json"})

        rows = query_db(
            "SELECT id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, file_type FROM device_distribution_records WHERE dist_id = %s ORDER BY uploaded_at DESC",
            (dist_id_int,),
        ) or []

        out = []
        for r in rows:
            try:
                f = json.loads(r.get("fields_json") or "{}")
            except Exception:
                f = {}
            rt = (f.get("Record Type") or f.get("record_type") or "").strip().lower()
            src = (f.get("Source") or f.get("source") or "").strip().lower()
            if rt != "shipment_record" and src != "shipment_record":
                continue
            out.append(
                {
                    "id": r.get("id"),
                    "dist_id": r.get("dist_id"),
                    "original_filename": r.get("original_filename"),
                    "uploaded_at": r.get("uploaded_at"),
                    "file_type": r.get("file_type"),
                }
            )

        return make_response(json.dumps({"ok": True, "records": out}), 200, {"Content-Type": "application/json"})

    @app.route("/admin/distribution-records/resolve", methods=["GET"])
    def admin_distribution_records_resolve():
        """Resolve a filename-derived order number to candidate distributions.

        Intended for one-off migrations where attachments live on an admin workstation.
        Requires admin session OR DIST_RECORDS_IMPORT_TOKEN.
        """
        if not is_admin() and not has_dist_records_import_token():
            abort(403)

        order_raw = (request.args.get("order") or request.args.get("order_number") or "").strip()
        if not order_raw:
            return make_response(json.dumps({"ok": False, "error": "Missing order"}), 400, {"Content-Type": "application/json"})

        # Normalize to alnum-only (SO0000123 etc)
        order_norm = re.sub(r"[^A-Z0-9]", "", order_raw.upper())
        if not order_norm:
            return make_response(json.dumps({"ok": False, "error": "Invalid order"}), 400, {"Content-Type": "application/json"})

        # Also try numeric-only matching (ignore leading zeros) to support filenames that
        # only contain the numeric portion, e.g. "Sales Order 200" matching "SO0000200".
        order_digits = re.sub(r"[^0-9]", "", order_norm)

        rows = query_db(
            """
            SELECT
                dd.id,
                dd.order_number,
                dd.ship_date,
                dd.rep_id,
                dd.source,
                dd.customer_id,
                c.facility_name AS customer_name,
                c.company_key AS customer_company_key,
                c.state AS customer_state
            FROM devices_distributed dd
            LEFT JOIN customers c ON c.id = dd.customer_id
            WHERE regexp_replace(upper(coalesce(dd.order_number, '')), '[^A-Z0-9]', '', 'g') = %s
               OR (
                    %s <> ''
                    AND NULLIF(ltrim(regexp_replace(upper(coalesce(dd.order_number, '')), '[^0-9]', '', 'g'), '0'), '')
                        = NULLIF(ltrim(%s, '0'), '')
                  )
            ORDER BY dd.ship_date NULLS LAST, dd.id ASC
            """,
            (order_norm, order_digits, order_digits),
        ) or []

        out = []
        for r in rows:
            out.append(
                {
                    "dist_id": r.get("id"),
                    "order_number": r.get("order_number"),
                    "ship_date": r.get("ship_date"),
                    "rep_id": r.get("rep_id"),
                    "source": r.get("source"),
                    "customer_id": r.get("customer_id"),
                    "customer_name": r.get("customer_name"),
                    "customer_company_key": r.get("customer_company_key"),
                    "customer_state": r.get("customer_state"),
                }
            )

        return make_response(
            json.dumps({"ok": True, "order_norm": order_norm, "order_digits": order_digits, "candidates": out}),
            200,
            {"Content-Type": "application/json"},
        )

    @app.route("/admin/distribution-records/upload", methods=["POST"])
    def admin_distribution_records_upload():
        if not is_admin() and not has_dist_records_import_token():
            abort(403)
        dist_id_raw = (request.form.get("dist_id") or "").strip()
        file_type = (request.form.get("file_type") or "").strip()
        # Validate file_type if provided
        if file_type and file_type not in ("sales_order", "packing_slip", "device_distribution_form"):
            return make_response(json.dumps({"ok": False, "error": "Invalid file_type"}), 400, {"Content-Type": "application/json"})
        try:
            dist_id_int = int(dist_id_raw)
        except Exception:
            return make_response(json.dumps({"ok": False, "error": "Invalid dist_id"}), 400, {"Content-Type": "application/json"})

        file = request.files.get("file")
        if not file or not file.filename:
            return make_response(json.dumps({"ok": False, "error": "No file"}), 400, {"Content-Type": "application/json"})

        def _compute_sha1_import_key(file_storage) -> str:
            h = hashlib.sha1()
            try:
                try:
                    file_storage.stream.seek(0)
                except Exception:
                    pass
                while True:
                    chunk = file_storage.stream.read(1024 * 1024)
                    if not chunk:
                        break
                    h.update(chunk)
            finally:
                try:
                    file_storage.stream.seek(0)
                except Exception:
                    pass
            return f"fs_sha1::{h.hexdigest()}"

        # If caller didn't send an import_key, compute one server-side.
        import_key = (request.form.get("import_key") or "").strip() or _compute_sha1_import_key(file)

        # Duplicate / repair behavior by import_key (idempotent + heals missing blobs)
        existing = normalize_row(
            query_db(
                """
                SELECT id, stored_filename
                FROM device_distribution_records
                WHERE dist_id = %s
                  AND fields_json ILIKE %s
                ORDER BY id ASC
                LIMIT 1
                """,
                (dist_id_int, f"%{import_key}%"),
                one=True,
            )
        )
        if existing:
            stored_existing = (existing.get("stored_filename") or "").strip()
            if stored_existing and dist_record_blob_exists(stored_existing):
                return make_response(json.dumps({"ok": True, "skipped": True}), 200, {"Content-Type": "application/json"})

            if stored_existing:
                try:
                    save_distribution_record_file_to(stored_existing, file)
                    update_fields = [secure_filename(file.filename) or (file.filename or "upload"), datetime.utcnow().isoformat(timespec="seconds") + "Z"]
                    if file_type:
                        update_fields.append(file_type)
                    update_fields.append(int(existing.get("id")))
                    if file_type:
                        execute_db(
                            "UPDATE device_distribution_records SET original_filename=%s, uploaded_at=%s, file_type=%s WHERE id=%s",
                            tuple(update_fields),
                        )
                    else:
                        execute_db(
                            "UPDATE device_distribution_records SET original_filename=%s, uploaded_at=%s WHERE id=%s",
                            tuple(update_fields),
                        )
                    return make_response(
                        json.dumps({"ok": True, "repaired": True, "record_id": existing.get("id")}),
                        200,
                        {"Content-Type": "application/json"},
                    )
                except Exception as e:
                    return make_response(
                        json.dumps({"ok": False, "error": f"Repair failed: {e}"}),
                        500,
                        {"Content-Type": "application/json"},
                    )

        file = request.files.get("file")
        if not file or not file.filename:
            return make_response(json.dumps({"ok": False, "error": "No file"}), 400, {"Content-Type": "application/json"})
        
        # Validate file upload (security hardening)
        is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions=app.config["ALLOWED_DISTRIBUTION_EXTENSIONS"])
        if not is_valid:
            return make_response(json.dumps({"ok": False, "error": error_msg or "Invalid file"}), 400, {"Content-Type": "application/json"})

        dist = normalize_row(query_db("SELECT * FROM devices_distributed WHERE id = %s", (dist_id_int,), one=True))
        if not dist:
            return make_response(json.dumps({"ok": False, "error": "Distribution not found"}), 404, {"Content-Type": "application/json"})

        saved = save_distribution_record_file(dist_id_int, file)
        stored_filename = saved.get("stored_filename")
        original_filename = saved.get("original_filename")
        uploaded_at = datetime.utcnow().isoformat(timespec="seconds") + "Z"
        fields = {
            "Record Type": "shipment_record",
            "Source": "shipment_record",
            "Origin": "admin_upload" if is_admin() else "token_import",
            "Import Key": import_key or None,
            "dist_id": dist_id_int,
            "Order Number": dist.get("order_number"),
            "Ship Date": dist.get("ship_date"),
        }

        # Remove null-ish keys for cleaner JSON
        if not fields.get("Import Key"):
            fields.pop("Import Key", None)

        execute_db(
            """INSERT INTO device_distribution_records
               (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, customer_id, file_type)
               VALUES (%s, %s, %s, %s, %s, %s, %s, %s)""",
            (
                dist.get("rep_id"),
                dist_id_int,
                stored_filename,
                original_filename,
                uploaded_at,
                json.dumps(fields),
                dist.get("customer_id"),
                file_type if file_type else None,
            ),
        )

        return make_response(json.dumps({"ok": True}), 200, {"Content-Type": "application/json"})

    @app.route("/admin/distribution-records/import-master", methods=["POST"])
    def admin_distribution_records_import_master():
        """Enqueue a master Sales Order PDF import.

        Dry run has been removed: every import is a real import.

        Auth: admin session OR DIST_RECORDS_IMPORT_TOKEN.
        Request: multipart/form-data
          - file: PDF
          - attach_all: optional (default 1)
          - limit: optional int (max pages to process)
        """
        if not is_admin() and not has_dist_records_import_token():
            abort(403)

        file = request.files.get("file")
        if not file or not getattr(file, "filename", None):
            return make_response(json.dumps({"ok": False, "error": "No file"}), 400, {"Content-Type": "application/json"})
        
        # Validate file upload (security hardening) - PDF only for master salesorder
        is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions={".pdf"})
        if not is_valid:
            return make_response(json.dumps({"ok": False, "error": error_msg or "Invalid file. Only PDF files allowed."}), 400, {"Content-Type": "application/json"})

        attach_all = (request.form.get("attach_all") or request.args.get("attach_all") or "1").strip().lower() not in {"0", "false", "no", "n"}
        limit_raw = (request.form.get("limit") or request.args.get("limit") or "").strip()
        limit = None
        if limit_raw:
            try:
                limit = max(0, int(limit_raw))
            except Exception:
                limit = None

        ok, msg, http_status = _start_master_salesorder_import_job_from_filestorage(file, attach_all=attach_all, limit=limit)
        return make_response(json.dumps({"ok": ok, "message": msg}), http_status, {"Content-Type": "application/json"})

    MASTER_IMPORT_JOB_KEY = "master_salesorder_pdf"

    def _now_iso_utc() -> str:
        return datetime.utcnow().isoformat(timespec="seconds") + "Z"

    def _master_job_default_progress() -> dict:
        return {
            "status": "idle",
            "message": "",
            "started_at": None,
            "updated_at": None,
            "error": None,
            "cancel_requested": False,
            "error_samples": [],
            "pages_total": 0,
            "pages_processed": 0,
            "attached": 0,
            "skipped_duplicate": 0,
            "repaired": 0,
            "unmatched": 0,
            "no_order_number": 0,
            "errors": 0,
            "master_sha1": None,
            "attach_all": True,
        }

    def _ensure_master_import_job_row() -> None:
        try:
            execute_db(
                """
                INSERT INTO master_import_jobs (job_key, status, started_at, updated_at, cancel_requested, error, progress_json)
                VALUES (%s, %s, %s, %s, 0, NULL, %s)
                ON CONFLICT (job_key) DO NOTHING
                """,
                (
                    MASTER_IMPORT_JOB_KEY,
                    "idle",
                    None,
                    _now_iso_utc(),
                    json.dumps(_master_job_default_progress()),
                ),
            )
        except Exception:
            pass

    def _get_master_import_job_row() -> dict:
        _ensure_master_import_job_row()
        row = normalize_row(
            query_db(
                "SELECT job_key, status, started_at, updated_at, cancel_requested, error, progress_json FROM master_import_jobs WHERE job_key = %s",
                (MASTER_IMPORT_JOB_KEY,),
                one=True,
            )
        )
        return row or {}

    def _get_master_import_progress() -> dict:
        row = _get_master_import_job_row()
        try:
            progress = json.loads(row.get("progress_json") or "{}")
        except Exception:
            progress = {}
        base = _master_job_default_progress()
        base.update(progress or {})
        # Mirror DB columns (authoritative status/error/cancel flags)
        base["status"] = row.get("status") or base.get("status")
        base["started_at"] = row.get("started_at") or base.get("started_at")
        base["updated_at"] = row.get("updated_at") or base.get("updated_at")
        base["error"] = row.get("error")
        base["cancel_requested"] = bool(int(row.get("cancel_requested") or 0))
        return base

    def _set_master_import_state(*, status: str | None = None, message: str | None = None, error: str | None = None, cancel_requested: bool | None = None, progress_updates: dict | None = None) -> None:
        row = _get_master_import_job_row()
        try:
            progress = json.loads(row.get("progress_json") or "{}")
        except Exception:
            progress = {}
        merged = _master_job_default_progress()
        merged.update(progress or {})
        if progress_updates:
            for k, v in progress_updates.items():
                merged[k] = v
        if message is not None:
            merged["message"] = message
        if status is not None:
            merged["status"] = status
        merged["updated_at"] = _now_iso_utc()

        new_status = status if status is not None else (row.get("status") or merged.get("status") or "idle")
        new_started_at = row.get("started_at")
        if new_status == "running" and not new_started_at:
            new_started_at = _now_iso_utc()

        new_cancel = row.get("cancel_requested")
        if cancel_requested is not None:
            new_cancel = 1 if cancel_requested else 0
            merged["cancel_requested"] = bool(cancel_requested)

        execute_db(
            """
            UPDATE master_import_jobs
               SET status=%s,
                   started_at=%s,
                   updated_at=%s,
                   cancel_requested=%s,
                   error=%s,
                   progress_json=%s
             WHERE job_key=%s
            """,
            (
                new_status,
                new_started_at,
                merged.get("updated_at"),
                int(new_cancel or 0),
                error,
                json.dumps(merged),
                MASTER_IMPORT_JOB_KEY,
            ),
        )

    def _is_master_import_running() -> bool:
        row = _get_master_import_job_row()
        status = (row.get("status") or "").strip().lower()
        if status != "running":
            return False

        # If the job hasn't updated recently, treat it as stale.
        try:
            updated_at = (row.get("updated_at") or "").strip()
            if updated_at:
                # stored as ISO UTC with trailing 'Z'
                ts = updated_at[:-1] if updated_at.endswith("Z") else updated_at
                last = datetime.fromisoformat(ts)
                age = (datetime.utcnow() - last).total_seconds()
                if age > 30 * 60:
                    _set_master_import_state(status="stale", message="Import became unresponsive and was marked stale", error=row.get("error"))
                    return False
        except Exception:
            pass

        return True

    def _is_master_import_cancel_requested() -> bool:
        row = _get_master_import_job_row()
        try:
            return bool(int(row.get("cancel_requested") or 0))
        except Exception:
            return False

    def _request_master_import_cancel() -> None:
        _set_master_import_state(cancel_requested=True, message="Cancel requested...", status="running")

    def _start_master_salesorder_import_job_from_filestorage(file_storage, *, attach_all: bool, limit: int | None) -> tuple[bool, str, int]:
        if _is_master_import_running():
            return False, "Import already running", 409

        import tempfile
        import os as _os

        try:
            tmp = tempfile.NamedTemporaryFile(prefix="master_so_", suffix=".pdf", delete=False)
            tmp_path = tmp.name
            with tmp:
                while True:
                    chunk = file_storage.stream.read(1024 * 1024)
                    if not chunk:
                        break
                    tmp.write(chunk)
        except Exception as e:
            return False, f"Failed to save temp PDF: {e}", 500

        # Initialize job state in DB immediately.
        _set_master_import_state(
            status="running",
            message="Queued",
            error=None,
            cancel_requested=False,
            progress_updates={
                "status": "running",
                "message": "Queued",
                "started_at": _now_iso_utc(),
                "pages_total": 0,
                "pages_processed": 0,
                "attached": 0,
                "skipped_duplicate": 0,
                "repaired": 0,
                "unmatched": 0,
                "no_order_number": 0,
                "errors": 0,
                "error_samples": [],
                "master_sha1": None,
                "attach_all": bool(attach_all),
            },
        )

        def _runner():
            try:
                with app.app_context():
                    _run_master_salesorder_import_from_pdf_path(tmp_path, limit=limit, attach_all=attach_all)
            finally:
                try:
                    _os.unlink(tmp_path)
                except Exception:
                    pass

        t = Thread(target=_runner, daemon=True)
        t.start()
        return True, "Master import started", 200

    def _sha1_file_helper(path: str) -> str:
        """Helper function to compute SHA1 hash of a file."""
        h = hashlib.sha1()
        with open(path, "rb") as fp:
            while True:
                chunk = fp.read(1024 * 1024)
                if not chunk:
                    break
                h.update(chunk)
        return h.hexdigest()

    def _run_master_salesorder_import_from_pdf_path(
        pdf_path: str,
        *,
        limit: int | None,
        attach_all: bool,
    ) -> None:
        """Background job to import a master Sales Order PDF.

        Runs in a daemon thread so the UI can poll progress instead of holding an HTTP request
        open for many minutes (prevents timeouts / OOM from giant responses).
        """
        import gc
        import io
        import os as _os
        import traceback
        from pypdf import PdfReader, PdfWriter

        try:
            master_sha1 = _sha1_file_helper(pdf_path)

            # Preflight S3/Spaces connectivity early.
            if _dist_records_s3_enabled():
                try:
                    bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
                    if bucket:
                        _dist_records_s3_client().head_bucket(Bucket=bucket)
                except Exception as e:
                    _set_master_import_state(status="error", error=f"S3 preflight failed: {e}", message="Failed")
                    return

            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            max_pages = total_pages if limit is None else min(total_pages, limit)

            _set_master_import_state(
                status="running",
                message="Running",
                error=None,
                progress_updates={
                    "pages_total": int(total_pages),
                    "master_sha1": master_sha1,
                    "attach_all": bool(attach_all),
                },
            )

            attached = 0
            skipped_duplicate = 0
            repaired = 0
            unmatched = 0
            no_order_number = 0
            errors = 0
            error_samples: list[str] = []

            for page_index in range(max_pages):
                if _is_master_import_cancel_requested():
                    _set_master_import_state(status="completed", message="Import canceled", error=None)
                    return

                _set_master_import_state(
                    status="running",
                    message=f"Processing page {page_index + 1} of {max_pages}...",
                    error=None,
                    progress_updates={
                        "pages_processed": page_index + 1,
                        "attached": attached,
                        "skipped_duplicate": skipped_duplicate,
                        "repaired": repaired,
                        "unmatched": unmatched,
                        "no_order_number": no_order_number,
                        "errors": errors,
                        "error_samples": error_samples,
                    },
                )

                try:
                    page = reader.pages[page_index]
                    text = page.extract_text() or ""

                    # Extract ORDER NUMBER (digits) from the header region.
                    upper = text.upper()
                    start = upper.find("ORDER NUMBER")
                    if start == -1:
                        start = 0
                    end = upper.find("SALESPERSON", start)
                    if end == -1:
                        end = min(len(text), start + 2500)
                    chunk = text[start:end]
                    lines = [ln.strip() for ln in chunk.splitlines() if ln.strip()]
                    filtered = [ln for ln in lines if ln.upper() not in {"ORDER NUMBER:", "ORDER DATE:", "CUSTOMER NUMBER:"}]
                    extracted = None
                    for ln in filtered:
                        if re.fullmatch(r"\d{5,}", ln):
                            extracted = ln
                            break
                    if not extracted:
                        m = re.search(r"\b\d{5,}\b", chunk)
                        extracted = m.group(0) if m else None

                    if not extracted:
                        no_order_number += 1
                        continue

                    # Resolve candidates (prefer ShipStation).
                    order_norm = re.sub(r"[^A-Z0-9]", "", (extracted or "").upper())
                    order_digits = re.sub(r"[^0-9]", "", order_norm)
                    cand_rows = query_db(
                        """
                        SELECT id, order_number, ship_date, rep_id, source, customer_id
                        FROM devices_distributed
                        WHERE regexp_replace(upper(coalesce(order_number, '')), '[^A-Z0-9]', '', 'g') = %s
                           OR (
                                %s <> ''
                                AND NULLIF(ltrim(regexp_replace(upper(coalesce(order_number, '')), '[^0-9]', '', 'g'), '0'), '')
                                    = NULLIF(ltrim(%s, '0'), '')
                              )
                        ORDER BY ship_date NULLS LAST, id ASC
                        """,
                        (order_norm, order_digits, order_digits),
                    ) or []
                    candidates = [
                        {
                            "dist_id": r.get("id"),
                            "source": r.get("source"),
                        }
                        for r in cand_rows
                    ]
                    shipstation_candidates = [c for c in candidates if (c.get("source") or "").strip().lower() == "shipstation"]
                    if shipstation_candidates:
                        candidates = shipstation_candidates

                    if not candidates:
                        unmatched += 1
                        continue

                    if (not attach_all) and len(candidates) != 1:
                        unmatched += 1
                        continue

                    targets = candidates if attach_all else candidates[:1]
                    # Always generate the 1-page PDF bytes so we can hash it for a stable import key.
                    w = PdfWriter()
                    w.add_page(page)
                    b = io.BytesIO()
                    w.write(b)
                    page_bytes = b.getvalue()
                    page_sha1 = hashlib.sha1(page_bytes).hexdigest()
                    # IMPORTANT: import key is based on page content, not the master PDF file.
                    # This avoids duplicating records when you upload a newer master PDF that contains
                    # the same pages plus additional pages.
                    import_key = f"so_page_sha1::{page_sha1}::order::{extracted}"

                    for c in targets:
                        dist_id_int = int(c.get("dist_id"))
                        dist = normalize_row(query_db("SELECT * FROM devices_distributed WHERE id = %s", (dist_id_int,), one=True))
                        if not dist:
                            continue

                        existing = normalize_row(
                            query_db(
                                """
                                SELECT id, stored_filename
                                FROM device_distribution_records
                                WHERE dist_id = %s
                                  AND fields_json ILIKE %s
                                ORDER BY id ASC
                                LIMIT 1
                                """,
                                (dist_id_int, f"%{import_key}%"),
                                one=True,
                            )
                        )

                        if existing:
                            stored_existing = (existing.get("stored_filename") or "").strip()
                            if stored_existing and dist_record_blob_exists(stored_existing):
                                skipped_duplicate += 1
                                continue
                            if stored_existing:
                                save_distribution_record_bytes_to(stored_existing, page_bytes, content_type="application/pdf")
                                execute_db(
                                    "UPDATE device_distribution_records SET original_filename=%s, uploaded_at=%s WHERE id=%s",
                                    (
                                        f"SO_{extracted}.pdf",
                                        _now_iso_utc(),
                                        int(existing.get("id")),
                                    ),
                                )
                                repaired += 1
                                continue

                        saved = save_distribution_record_bytes(dist_id_int, f"SO_{extracted}.pdf", page_bytes, content_type="application/pdf")
                        uploaded_at = _now_iso_utc()
                        fields = {
                            "Record Type": "shipment_record",
                            "Source": "shipment_record",
                            "Origin": "master_pdf_import_bg",
                            "Import Key": import_key,
                            "dist_id": dist_id_int,
                            "Order Number": dist.get("order_number"),
                            "Ship Date": dist.get("ship_date"),
                            "Master PDF SHA1": master_sha1,
                            "Master Page": page_index + 1,
                            "Master Extracted Order": extracted,
                            "Page SHA1": page_sha1,
                        }
                        execute_db(
                            """INSERT INTO device_distribution_records
                               (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, customer_id)
                               VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                            (
                                dist.get("rep_id"),
                                dist_id_int,
                                saved.get("stored_filename"),
                                saved.get("original_filename"),
                                uploaded_at,
                                json.dumps(fields),
                                dist.get("customer_id"),
                            ),
                        )
                        attached += 1

                except Exception as e:
                    errors += 1
                    if len(error_samples) < 5:
                        try:
                            error_samples.append(f"Page {page_index + 1}: {type(e).__name__}: {e}")
                        except Exception:
                            error_samples.append(f"Page {page_index + 1}: error")

                    _set_master_import_state(progress_updates={"errors": errors, "error_samples": error_samples})

                if (page_index + 1) % 10 == 0:
                    gc.collect()

            _set_master_import_state(
                status="completed",
                message="Completed",
                error=None,
                progress_updates={
                    "pages_processed": max_pages,
                    "attached": attached,
                    "skipped_duplicate": skipped_duplicate,
                    "repaired": repaired,
                    "unmatched": unmatched,
                    "no_order_number": no_order_number,
                    "errors": errors,
                    "error_samples": error_samples,
                },
            )
        except Exception as e:
            traceback.print_exc()
            _set_master_import_state(status="error", message="Failed", error=str(e))

    @app.route("/admin/distribution-records/import-master/start", methods=["POST"])
    def admin_distribution_records_import_master_start():
        if not is_admin():
            abort(403)

        file = request.files.get("file")
        if not file or not getattr(file, "filename", None):
            return make_response(json.dumps({"ok": False, "error": "No file"}), 400, {"Content-Type": "application/json"})
        
        # Validate file upload (security hardening) - PDF only for master salesorder
        is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions={".pdf"})
        if not is_valid:
            return make_response(json.dumps({"ok": False, "error": error_msg or "Invalid file. Only PDF files allowed."}), 400, {"Content-Type": "application/json"})

        attach_all = (request.form.get("attach_all") or "1").strip().lower() not in {"0", "false", "no", "n"}
        limit_raw = (request.form.get("limit") or "").strip()
        limit = None
        if limit_raw:
            try:
                limit = max(0, int(limit_raw))
            except Exception:
                limit = None

        ok, msg, http_status = _start_master_salesorder_import_job_from_filestorage(file, attach_all=attach_all, limit=limit)
        return make_response(json.dumps({"ok": ok, "message": msg}), http_status, {"Content-Type": "application/json"})

    @app.route("/admin/distribution-records/import-master/progress", methods=["GET"])
    def admin_distribution_records_import_master_progress():
        if not is_admin():
            abort(403)
        return make_response(
            json.dumps({"ok": True, "running": _is_master_import_running(), "progress": _get_master_import_progress()}),
            200,
            {"Content-Type": "application/json"},
        )

    @app.route("/admin/distribution-records/import-master/cancel", methods=["POST"])
    def admin_distribution_records_import_master_cancel():
        if not is_admin():
            abort(403)
        _request_master_import_cancel()
        return make_response(json.dumps({"ok": True}), 200, {"Content-Type": "application/json"})

    @app.route("/admin/distribution-records/import-master-ui", methods=["GET"])
    def admin_distribution_records_import_master_ui():
        """Browser UI for importing a master Sales Order PDF.

        This avoids needing a local script/env vars: admins can upload the PDF and run the import.
        """
        if not is_admin():
            return redirect(url_for("admin_login"))

        return render_template("admin_distribution_records_import_master.html")

    # ------------------------------------------------------------------ Label_Bulk PDF import (1 page per distribution)

    LABEL_BULK_IMPORT_JOB_KEY = "label_bulk_pdf"

    def _label_bulk_job_default_progress() -> dict:
        return {
            "status": "idle",
            "message": "",
            "started_at": None,
            "updated_at": None,
            "error": None,
            "cancel_requested": False,
            "error_samples": [],
            "pages_total": 0,
            "pages_processed": 0,
            "attached": 0,
            "skipped_duplicate": 0,
            "repaired": 0,
            "unmatched": 0,
            "no_order_number": 0,
            "no_tracking": 0,
            "errors": 0,
            "master_sha1": None,
        }

    def _ensure_label_bulk_import_job_row() -> None:
        try:
            execute_db(
                """
                INSERT INTO master_import_jobs (job_key, status, started_at, updated_at, cancel_requested, error, progress_json)
                VALUES (%s, %s, %s, %s, 0, NULL, %s)
                ON CONFLICT (job_key) DO NOTHING
                """,
                (
                    LABEL_BULK_IMPORT_JOB_KEY,
                    "idle",
                    None,
                    _now_iso_utc(),
                    json.dumps(_label_bulk_job_default_progress()),
                ),
            )
        except Exception:
            pass

    def _get_label_bulk_import_job_row() -> dict:
        _ensure_label_bulk_import_job_row()
        row = normalize_row(
            query_db(
                "SELECT job_key, status, started_at, updated_at, cancel_requested, error, progress_json FROM master_import_jobs WHERE job_key = %s",
                (LABEL_BULK_IMPORT_JOB_KEY,),
                one=True,
            )
        )
        return row or {}

    def _get_label_bulk_import_progress() -> dict:
        row = _get_label_bulk_import_job_row()
        try:
            progress = json.loads(row.get("progress_json") or "{}")
        except Exception:
            progress = {}
        base = _label_bulk_job_default_progress()
        base.update(progress or {})
        base["status"] = row.get("status") or base.get("status")
        base["started_at"] = row.get("started_at") or base.get("started_at")
        base["updated_at"] = row.get("updated_at") or base.get("updated_at")
        base["error"] = row.get("error")
        base["cancel_requested"] = bool(int(row.get("cancel_requested") or 0))
        return base

    def _set_label_bulk_import_state(*, status: str | None = None, message: str | None = None, error: str | None = None, cancel_requested: bool | None = None, progress_updates: dict | None = None) -> None:
        row = _get_label_bulk_import_job_row()
        try:
            progress = json.loads(row.get("progress_json") or "{}")
        except Exception:
            progress = {}
        merged = _label_bulk_job_default_progress()
        merged.update(progress or {})
        if progress_updates:
            for k, v in progress_updates.items():
                merged[k] = v
        if message is not None:
            merged["message"] = message
        if status is not None:
            merged["status"] = status
        merged["updated_at"] = _now_iso_utc()

        new_status = status if status is not None else (row.get("status") or merged.get("status") or "idle")
        new_started_at = row.get("started_at")
        if new_status == "running" and not new_started_at:
            new_started_at = _now_iso_utc()

        new_cancel = row.get("cancel_requested")
        if cancel_requested is not None:
            new_cancel = 1 if cancel_requested else 0
            merged["cancel_requested"] = bool(cancel_requested)

        execute_db(
            """
            UPDATE master_import_jobs
               SET status=%s,
                   started_at=%s,
                   updated_at=%s,
                   cancel_requested=%s,
                   error=%s,
                   progress_json=%s
             WHERE job_key=%s
            """,
            (
                new_status,
                new_started_at,
                merged.get("updated_at"),
                int(new_cancel or 0),
                error,
                json.dumps(merged),
                LABEL_BULK_IMPORT_JOB_KEY,
            ),
        )

    def _is_label_bulk_import_running() -> bool:
        row = _get_label_bulk_import_job_row()
        status = (row.get("status") or "").strip().lower()
        if status != "running":
            return False

        try:
            updated_at = (row.get("updated_at") or "").strip()
            if updated_at:
                ts = updated_at[:-1] if updated_at.endswith("Z") else updated_at
                last = datetime.fromisoformat(ts)
                age = (datetime.utcnow() - last).total_seconds()
                if age > 30 * 60:
                    _set_label_bulk_import_state(status="stale", message="Import became unresponsive and was marked stale", error=row.get("error"))
                    return False
        except Exception:
            pass
        return True

    def _is_label_bulk_import_cancel_requested() -> bool:
        row = _get_label_bulk_import_job_row()
        try:
            return bool(int(row.get("cancel_requested") or 0))
        except Exception:
            return False

    def _request_label_bulk_import_cancel() -> None:
        _set_label_bulk_import_state(cancel_requested=True, message="Cancel requested...", status="running")

    def _canonical_tracking_digits(value: str | None) -> str:
        return re.sub(r"[^0-9]", "", str(value or ""))

    def _extract_label_bulk_keys(text: str) -> tuple[str | None, str | None]:
        """Return (order_number, tracking_digits) extracted from a label page."""
        t = text or ""
        # Order: typically "Order # SO 00000245".
        m = re.search(r"Order\s*#\s*SO\s*([A-Za-z0-9\-]+)", t, flags=re.IGNORECASE)
        order_raw = None
        if m:
            order_raw = m.group(1)
        else:
            m2 = re.search(r"\bSO\s*0*([0-9]{3,})\b", t, flags=re.IGNORECASE)
            if m2:
                # Keep the canonical SO prefix for consistency.
                order_raw = f"SO{m2.group(1)}"

        order_norm = canonical_order_number(order_raw) if order_raw else ""
        order_digits = re.sub(r"[^0-9]", "", order_norm)

        # Tracking: find long digit sequences (often spaced), keep the longest plausible.
        candidates = re.findall(r"\b\d[\d ]{8,}\d\b", t)
        best = ""
        for c in candidates:
            d = _canonical_tracking_digits(c)
            if not d:
                continue
            # Ignore obvious placeholders.
            if set(d) == {"0"}:
                continue

            # Some PDFs contain padded/duplicated variants like 000<real_tracking>.
            d = d.lstrip("0")
            if not d:
                continue
            if order_digits and d == order_digits:
                continue
            if len(d) < 10:
                continue
            if len(d) > len(best):
                best = d

        return (order_norm or None, best or None)

    def _resolve_single_distribution_for_label_page(order_norm: str | None, tracking_digits: str | None) -> int | None:
        """Resolve to exactly one devices_distributed.id, or None."""
        # Prefer tracking match when present.
        if tracking_digits:
            rows = query_db(
                """
                SELECT id, source
                FROM devices_distributed
                WHERE regexp_replace(coalesce(tracking_number, ''), '[^0-9]', '', 'g') = %s
                ORDER BY ship_date NULLS LAST, id ASC
                """,
                (tracking_digits,),
            ) or []
            if len(rows) == 1:
                return int(rows[0].get("id"))
            if rows:
                return None

        if not order_norm:
            return None

        order_digits = re.sub(r"[^0-9]", "", order_norm)
        cand_rows = query_db(
            """
            SELECT id, source
            FROM devices_distributed
            WHERE regexp_replace(upper(coalesce(order_number, '')), '[^A-Z0-9]', '', 'g') = %s
               OR (
                    %s <> ''
                    AND NULLIF(ltrim(regexp_replace(upper(coalesce(order_number, '')), '[^0-9]', '', 'g'), '0'), '')
                        = NULLIF(ltrim(%s, '0'), '')
                  )
            ORDER BY ship_date NULLS LAST, id ASC
            """,
            (order_norm, order_digits, order_digits),
        ) or []

        # Prefer ShipStation if present.
        shipstation = [r for r in cand_rows if (r.get("source") or "").strip().lower() == "shipstation"]
        if shipstation:
            cand_rows = shipstation

        if len(cand_rows) == 1:
            return int(cand_rows[0].get("id"))
        return None

    def _start_label_bulk_import_job_from_filestorage(file_storage, *, limit: int | None) -> tuple[bool, str, int]:
        if _is_label_bulk_import_running():
            return False, "Import already running", 409

        import tempfile
        import os as _os

        try:
            tmp = tempfile.NamedTemporaryFile(prefix="label_bulk_", suffix=".pdf", delete=False)
            tmp_path = tmp.name
            with tmp:
                while True:
                    chunk = file_storage.stream.read(1024 * 1024)
                    if not chunk:
                        break
                    tmp.write(chunk)
        except Exception as e:
            return False, f"Failed to save temp PDF: {e}", 500

        _set_label_bulk_import_state(
            status="running",
            message="Queued",
            error=None,
            cancel_requested=False,
            progress_updates={
                "status": "running",
                "message": "Queued",
                "started_at": _now_iso_utc(),
                "pages_total": 0,
                "pages_processed": 0,
                "attached": 0,
                "skipped_duplicate": 0,
                "repaired": 0,
                "unmatched": 0,
                "no_order_number": 0,
                "no_tracking": 0,
                "errors": 0,
                "error_samples": [],
                "master_sha1": None,
            },
        )

        def _runner():
            try:
                with app.app_context():
                    _run_label_bulk_import_from_pdf_path(tmp_path, limit=limit)
            finally:
                try:
                    _os.unlink(tmp_path)
                except Exception:
                    pass

        t = Thread(target=_runner, daemon=True)
        t.start()
        return True, "Label_Bulk import started", 200

    def _run_label_bulk_import_from_pdf_path(pdf_path: str, *, limit: int | None) -> None:
        import gc
        import io
        import traceback
        from pypdf import PdfReader, PdfWriter

        try:
            master_sha1 = _sha1_file_helper(pdf_path)

            if _dist_records_s3_enabled():
                try:
                    bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
                    if bucket:
                        _dist_records_s3_client().head_bucket(Bucket=bucket)
                except Exception as e:
                    _set_label_bulk_import_state(status="error", error=f"S3 preflight failed: {e}", message="Failed")
                    return

            reader = PdfReader(pdf_path)
            total_pages = len(reader.pages)
            max_pages = total_pages if limit is None else min(total_pages, limit)

            _set_label_bulk_import_state(
                status="running",
                message="Running",
                error=None,
                progress_updates={
                    "pages_total": int(total_pages),
                    "master_sha1": master_sha1,
                },
            )

            attached = 0
            skipped_duplicate = 0
            repaired = 0
            unmatched = 0
            no_order_number = 0
            no_tracking = 0
            errors = 0
            error_samples: list[str] = []

            for page_index in range(max_pages):
                if _is_label_bulk_import_cancel_requested():
                    _set_label_bulk_import_state(status="completed", message="Import canceled", error=None)
                    return

                _set_label_bulk_import_state(
                    status="running",
                    message=f"Processing page {page_index + 1} of {max_pages}...",
                    error=None,
                    progress_updates={
                        "pages_processed": page_index + 1,
                        "attached": attached,
                        "skipped_duplicate": skipped_duplicate,
                        "repaired": repaired,
                        "unmatched": unmatched,
                        "no_order_number": no_order_number,
                        "no_tracking": no_tracking,
                        "errors": errors,
                        "error_samples": error_samples,
                    },
                )

                try:
                    page = reader.pages[page_index]
                    text = page.extract_text() or ""
                    order_norm, tracking_digits = _extract_label_bulk_keys(text)
                    if not order_norm:
                        no_order_number += 1
                    if not tracking_digits:
                        no_tracking += 1

                    dist_id_int = _resolve_single_distribution_for_label_page(order_norm, tracking_digits)
                    if not dist_id_int:
                        unmatched += 1
                        continue

                    dist = normalize_row(query_db("SELECT * FROM devices_distributed WHERE id = %s", (int(dist_id_int),), one=True))
                    if not dist:
                        unmatched += 1
                        continue

                    w = PdfWriter()
                    w.add_page(page)
                    b = io.BytesIO()
                    w.write(b)
                    page_bytes = b.getvalue()
                    page_sha1 = hashlib.sha1(page_bytes).hexdigest()
                    import_key = f"label_page_sha1::{page_sha1}::order::{order_norm or ''}::trk::{tracking_digits or ''}"

                    existing_label_rows = query_db(
                        """
                        SELECT id, stored_filename
                        FROM device_distribution_records
                        WHERE dist_id = %s
                          AND fields_json ILIKE %s
                        ORDER BY id ASC
                        LIMIT 2
                        """,
                        (int(dist_id_int), '%"Doc Type": "label_bulk"%'),
                    ) or []
                    existing_label = normalize_row(existing_label_rows[0]) if existing_label_rows else None

                    # Enforce 1 label per distribution. If one exists, update/repair it.
                    if existing_label:
                        stored_existing = (existing_label.get("stored_filename") or "").strip()
                        if stored_existing:
                            save_distribution_record_bytes_to(stored_existing, page_bytes, content_type="application/pdf")
                            fields = {
                                "Record Type": "shipment_record",
                                "Source": "shipment_record",
                                "Origin": "label_bulk_import_bg",
                                "Doc Type": "label_bulk",
                                "Import Key": import_key,
                                "dist_id": int(dist_id_int),
                                "Order Number": dist.get("order_number"),
                                "Ship Date": dist.get("ship_date"),
                                "Master PDF SHA1": master_sha1,
                                "Master Page": page_index + 1,
                                "Page SHA1": page_sha1,
                                "Label Extracted Order": order_norm,
                                "Label Tracking Digits": tracking_digits,
                            }
                            execute_db(
                                "UPDATE device_distribution_records SET original_filename=%s, uploaded_at=%s, fields_json=%s WHERE id=%s",
                                (
                                    f"LABEL_{order_norm or 'UNKNOWN'}_{tracking_digits or 'NO_TRACK'}.pdf",
                                    _now_iso_utc(),
                                    json.dumps(fields),
                                    int(existing_label.get("id")),
                                ),
                            )
                            repaired += 1
                            continue
                        skipped_duplicate += 1
                        continue

                    # Idempotency by import key (page content hash)
                    existing_by_key = normalize_row(
                        query_db(
                            """
                            SELECT id, stored_filename
                            FROM device_distribution_records
                            WHERE dist_id = %s
                              AND fields_json ILIKE %s
                            ORDER BY id ASC
                            LIMIT 1
                            """,
                            (int(dist_id_int), f"%{import_key}%"),
                            one=True,
                        )
                    )
                    if existing_by_key:
                        stored_existing = (existing_by_key.get("stored_filename") or "").strip()
                        if stored_existing and dist_record_blob_exists(stored_existing):
                            skipped_duplicate += 1
                            continue
                        if stored_existing:
                            save_distribution_record_bytes_to(stored_existing, page_bytes, content_type="application/pdf")
                            repaired += 1
                            continue

                    saved = save_distribution_record_bytes(
                        int(dist_id_int),
                        f"LABEL_{order_norm or 'UNKNOWN'}_{tracking_digits or 'NO_TRACK'}.pdf",
                        page_bytes,
                        content_type="application/pdf",
                    )
                    fields = {
                        "Record Type": "shipment_record",
                        "Source": "shipment_record",
                        "Origin": "label_bulk_import_bg",
                        "Doc Type": "label_bulk",
                        "Import Key": import_key,
                        "dist_id": int(dist_id_int),
                        "Order Number": dist.get("order_number"),
                        "Ship Date": dist.get("ship_date"),
                        "Master PDF SHA1": master_sha1,
                        "Master Page": page_index + 1,
                        "Page SHA1": page_sha1,
                        "Label Extracted Order": order_norm,
                        "Label Tracking Digits": tracking_digits,
                    }
                    execute_db(
                        """INSERT INTO device_distribution_records
                           (rep_id, dist_id, stored_filename, original_filename, uploaded_at, fields_json, customer_id)
                           VALUES (%s, %s, %s, %s, %s, %s, %s)""",
                        (
                            dist.get("rep_id"),
                            int(dist_id_int),
                            saved.get("stored_filename"),
                            saved.get("original_filename"),
                            _now_iso_utc(),
                            json.dumps(fields),
                            dist.get("customer_id"),
                        ),
                    )
                    attached += 1

                except Exception as e:
                    errors += 1
                    if len(error_samples) < 5:
                        try:
                            error_samples.append(f"Page {page_index + 1}: {type(e).__name__}: {e}")
                        except Exception:
                            error_samples.append(f"Page {page_index + 1}: error")
                    _set_label_bulk_import_state(progress_updates={"errors": errors, "error_samples": error_samples})

                if (page_index + 1) % 10 == 0:
                    gc.collect()

            _set_label_bulk_import_state(
                status="completed",
                message="Completed",
                error=None,
                progress_updates={
                    "pages_processed": max_pages,
                    "attached": attached,
                    "skipped_duplicate": skipped_duplicate,
                    "repaired": repaired,
                    "unmatched": unmatched,
                    "no_order_number": no_order_number,
                    "no_tracking": no_tracking,
                    "errors": errors,
                    "error_samples": error_samples,
                },
            )
        except Exception as e:
            traceback.print_exc()
            _set_label_bulk_import_state(status="error", message="Failed", error=str(e))

    @app.route("/admin/distribution-records/import-label-bulk/start", methods=["POST"])
    def admin_distribution_records_import_label_bulk_start():
        if not is_admin():
            abort(403)
        file = request.files.get("file")
        if not file or not getattr(file, "filename", None):
            return make_response(json.dumps({"ok": False, "error": "No file"}), 400, {"Content-Type": "application/json"})
        limit_raw = (request.form.get("limit") or "").strip()
        limit = None
        if limit_raw:
            try:
                limit = max(0, int(limit_raw))
            except Exception:
                limit = None
        ok, msg, http_status = _start_label_bulk_import_job_from_filestorage(file, limit=limit)
        return make_response(json.dumps({"ok": ok, "message": msg}), http_status, {"Content-Type": "application/json"})

    @app.route("/admin/distribution-records/import-label-bulk/progress", methods=["GET"])
    def admin_distribution_records_import_label_bulk_progress():
        if not is_admin():
            abort(403)
        return make_response(
            json.dumps({"ok": True, "running": _is_label_bulk_import_running(), "progress": _get_label_bulk_import_progress()}),
            200,
            {"Content-Type": "application/json"},
        )

    @app.route("/admin/distribution-records/import-label-bulk/cancel", methods=["POST"])
    def admin_distribution_records_import_label_bulk_cancel():
        if not is_admin():
            abort(403)
        _request_label_bulk_import_cancel()
        return make_response(json.dumps({"ok": True}), 200, {"Content-Type": "application/json"})

    @app.route("/admin/distribution-records/import-label-bulk-ui", methods=["GET"])
    def admin_distribution_records_import_label_bulk_ui():
        if not is_admin():
            return redirect(url_for("admin_login"))
        return render_template("admin_distribution_records_import_label_bulk.html")

    @app.route("/admin/lotlog/reconcile-distributions", methods=["POST"])
    def admin_lotlog_reconcile_distributions():
        if not is_admin():
            return redirect(url_for("admin_login"))

        only_unknown = (request.form.get("only_unknown") or "1").strip() not in {"0", "false", "no"}
        force = (request.form.get("force") or "0").strip() in {"1", "true", "yes"}
        limit_raw = (request.form.get("limit") or "").strip()
        limit = None
        if limit_raw:
            try:
                limit = int(limit_raw)
            except Exception:
                limit = None

        stats = reconcile_distribution_records_from_lotlog(only_unknown=only_unknown, force=force, limit=limit)
        flash(
            " | ".join(
                [
                    f"Reconciled distribution records from Lot Log: updated={stats.get('updated')}",
                    f"skipped={stats.get('skipped')}",
                    f"no_lot={stats.get('no_lot')}",
                    f"no_mapping={stats.get('no_mapping')}",
                    f"errors={stats.get('errors')}",
                ]
            ),
            "success" if (stats.get("errors") or 0) == 0 else "warning",
        )
        return redirect(url_for("admin_lotlog"))

    @app.route("/admin/distribution-records/<int:record_id>/delete", methods=["POST"])
    def admin_distribution_records_delete(record_id: int):
        if not is_admin():
            abort(403)
        rec = normalize_row(query_db("SELECT * FROM device_distribution_records WHERE id = %s", (record_id,), one=True))
        if not rec:
            return make_response(json.dumps({"ok": False, "error": "Not found"}), 404, {"Content-Type": "application/json"})
        try:
            f = json.loads(rec.get("fields_json") or "{}")
        except Exception:
            f = {}
        rt = (f.get("Record Type") or f.get("record_type") or "").strip().lower()
        src = (f.get("Source") or f.get("source") or "").strip().lower()
        if rt != "shipment_record" and src != "shipment_record":
            return make_response(json.dumps({"ok": False, "error": "Not a shipment record"}), 400, {"Content-Type": "application/json"})

        stored_filename = rec.get("stored_filename")
        if stored_filename:
            delete_distribution_record_blob(stored_filename)

        execute_db("DELETE FROM device_distribution_records WHERE id = %s", (record_id,))
        return make_response(json.dumps({"ok": True}), 200, {"Content-Type": "application/json"})

    @app.route("/admin/distribution-records/<int:record_id>/download", methods=["GET"])
    def admin_distribution_records_download(record_id: int):
        if not is_admin():
            abort(403)
        rec = normalize_row(query_db("SELECT * FROM device_distribution_records WHERE id = %s", (record_id,), one=True))
        if not rec:
            abort(404)
        try:
            f = json.loads(rec.get("fields_json") or "{}")
        except Exception:
            f = {}
        rt = (f.get("Record Type") or f.get("record_type") or "").strip().lower()
        src = (f.get("Source") or f.get("source") or "").strip().lower()
        if rt != "shipment_record" and src != "shipment_record":
            abort(404)
        stored_filename = rec.get("stored_filename")
        if not stored_filename:
            abort(404)

        # If S3 is configured, redirect to a short-lived presigned download URL.
        if _dist_records_s3_enabled():
            try:
                bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
                if not bucket:
                    abort(404)
                key = _dist_records_s3_key(stored_filename)
                client = _dist_records_s3_client()
                # Fail fast if the object is missing.
                client.head_object(Bucket=bucket, Key=key)

                original = (rec.get("original_filename") or "download").strip() or "download"
                # Tell browsers to download with a friendly name.
                url = client.generate_presigned_url(
                    "get_object",
                    Params={
                        "Bucket": bucket,
                        "Key": key,
                        "ResponseContentDisposition": f'attachment; filename="{original}"',
                    },
                    ExpiresIn=60 * 10,
                )
                return redirect(url)
            except Exception:
                abort(404)

        # Filesystem fallback
        try:
            file_path = _safe_dist_record_path(stored_filename)
        except Exception:
            abort(400)
        if not file_path.exists():
            abort(404)
        return send_from_directory(file_path.parent, file_path.name, as_attachment=True)

    @app.route("/admin/new-customers")
    def admin_new_customer_records():
        """Customer database with SS4-style grouping."""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            import json
            customers = query_db("SELECT * FROM new_customer_records ORDER BY id DESC")
            if not customers:
                return render_template("admin_new_customer_records.html", records=[])

            # Group by company_key (SS4-style)
            grouped = {}
            for c in customers:
                try:
                    fields = json.loads(c["fields_json"])
                except Exception as e:
                    import traceback
                    print(f"[ERROR] admin_customer_merge (JSON parse): {e}")
                    traceback.print_exc()
                    fields = {}

                facility = fields.get("Facility Name", "").strip()
                key = normalize_company_key(facility)
                if not key:
                    key = f"UNKNOWN_{c['id']}"

                if key not in grouped:
                    orders = query_db(
                        """
                        SELECT DISTINCT dd.id, dd.distribution_number, dd.order_number, dd.ship_date, dd.source
                        FROM devices_distributed dd
                        JOIN device_distribution_records ddr ON ddr.dist_id = dd.id
                        WHERE ddr.fields_json::text ILIKE %s
                        ORDER BY dd.ship_date DESC
                        """,
                        (f"%{facility}%",),
                    )
                    grouped[key] = {
                        "id": c["id"],
                        "fields": fields,
                        "orders": orders or [],
                        "original_filename": c["original_filename"],
                        "uploaded_at": c["uploaded_at"],
                        "duplicate_count": 1,
                        "duplicate_ids": [c["id"]]
                    }
                else:
                    grouped[key]["duplicate_count"] += 1
                    grouped[key]["duplicate_ids"].append(c["id"])
                    if c["uploaded_at"] > grouped[key]["uploaded_at"]:
                        grouped[key]["id"] = c["id"]
                        grouped[key]["uploaded_at"] = c["uploaded_at"]
                        grouped[key]["original_filename"] = c["original_filename"]

            records = list(grouped.values())
            return render_template("admin_new_customer_records.html", records=records)
        except Exception as e:
            print(f"Error in new customers: {e}")
            import traceback
            traceback.print_exc()
            return f"Error loading customer records: {e}", 500

    @app.route("/admin/customer/<int:cust_id>", methods=["GET", "POST"])
    def admin_customer_profile(cust_id):
        """Customer detail view."""
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            import json
            from datetime import datetime
            
            cust = normalize_row(query_db("SELECT * FROM new_customer_records WHERE id = %s", (cust_id,), one=True))
            if not cust:
                return "Customer not found", 404

            try:
                fields = json.loads(cust["fields_json"])
            except Exception as e:
                import traceback
                print(f"[ERROR] admin_customer_profile (JSON parse): {e}")
                traceback.print_exc()
                fields = {}

            customer = {
                "id": cust["id"],
                "fields": fields,
                "original_filename": cust["original_filename"],
                "uploaded_at": cust["uploaded_at"],
                "company_key": cust.get("company_key"),
            }

            # Find matching customer record by company_key to link notes
            matching_customer_id = None
            if customer.get("company_key"):
                matching_customer = query_db(
                    "SELECT id FROM customers WHERE company_key = %s LIMIT 1",
                    (customer["company_key"],),
                    one=True
                )
                if matching_customer:
                    matching_customer_id = matching_customer.get("id")

            # Handle POST (add note)
            if request.method == "POST":
                if not matching_customer_id:
                    flash("Cannot add notes: No matching customer record found. Notes require a customer record in the customers table.", "warning")
                    return redirect(url_for("admin_customer_profile", cust_id=cust_id))
                
                note_text = (request.form.get("note_text") or "").strip()
                note_date_str = request.form.get("note_date") or datetime.utcnow().date().isoformat()
                
                if not note_text:
                    flash("Note text is required.", "danger")
                    return redirect(url_for("admin_customer_profile", cust_id=cust_id))
                
                try:
                    note_date = datetime.fromisoformat(note_date_str).date() if note_date_str else datetime.utcnow().date()
                except Exception:
                    note_date = datetime.utcnow().date()
                
                execute_db(
                    "INSERT INTO customer_notes (customer_id, note_text, note_date, created_at, updated_at, author) VALUES (%s, %s, %s, %s, %s, %s)",
                    (matching_customer_id, note_text, note_date, datetime.utcnow(), datetime.utcnow(), "admin"),
                )
                
                flash("Note added successfully.", "success")
                return redirect(url_for("admin_customer_profile", cust_id=cust_id))

            facility = fields.get("Facility Name", "").strip()
            orders = []
            if facility:
                raw_rows = query_db(
                    """
                    SELECT dd.id, dd.distribution_number, dd.order_number, dd.ship_date, dd.source, ddr.fields_json
                    FROM devices_distributed dd
                    JOIN device_distribution_records ddr ON ddr.dist_id = dd.id
                    WHERE ddr.fields_json::text ILIKE %s
                    ORDER BY dd.ship_date DESC
                    """,
                    (f"%{facility}%",),
                )

                grouped = {}
                if raw_rows:
                    for r in raw_rows:
                        onum = r["order_number"] or r.get("distribution_number") or f"DIST-{r['id']}"
                        if onum not in grouped:
                            grouped[onum] = {
                                "order_number": onum,
                                "ship_date": r["ship_date"] or "Unknown",
                                "source": r["source"],
                                "items": [],
                                "total_qty": 0,
                            }
                        try:
                            fi = json.loads(r["fields_json"])
                        except Exception as e:
                            import traceback
                            print(f"[ERROR] admin_customer_profile (fields_json parse): {e}")
                            traceback.print_exc()
                            fi = {}
                        qty = int(float(fi.get("Quantity", 0) or 0))
                        grouped[onum]["items"].append({
                            "sku": fi.get("SKU", "-"),
                            "quantity": qty,
                            "lot": fi.get("Lot", "-")
                        })
                        grouped[onum]["total_qty"] += qty
                orders = sorted(grouped.values(), key=lambda o: o["ship_date"] or "", reverse=True)
            
            # Normalize all dates to YYYY-MM-DD strings before passing to template
            for order in orders:
                if order.get("ship_date"):
                    order["ship_date"] = _normalize_ship_date_ymd(order["ship_date"]) or ""
            
            # Normalize customer uploaded_at if used
            if customer.get("uploaded_at"):
                customer["uploaded_at"] = _normalize_ship_date_ymd(customer["uploaded_at"]) or ""

            # Load notes if customer record exists
            notes = []
            if matching_customer_id:
                notes = query_db(
                    "SELECT id, customer_id, note_text, note_date, created_at, updated_at, author FROM customer_notes WHERE customer_id = %s ORDER BY created_at DESC",
                    (matching_customer_id,),
                ) or []
            
            # Normalize note dates
            for note in notes:
                if note.get("note_date"):
                    note["note_date"] = _normalize_ship_date_ymd(note["note_date"]) or ""
                if note.get("created_at"):
                    note["created_at"] = _normalize_ship_date_ymd(note["created_at"]) or ""
                if note.get("updated_at"):
                    note["updated_at"] = _normalize_ship_date_ymd(note["updated_at"]) or ""

            return render_template("admin_customer_profile.html", customer=customer, orders=orders, notes=notes, matching_customer_id=matching_customer_id, now=datetime.now())
        except Exception as e:
            print(f"Error in customer profile {cust_id}: {e}")
            import traceback
            traceback.print_exc()
            return f"Error loading customer profile: {e}", 500

    # ------------------------------------------------------------------ Admin Management Routes
    
    @app.route("/admin/reps", methods=["GET", "POST"])
    def admin_reps():
        if not is_admin():
            return redirect(url_for("admin_login"))
        if request.method == "POST":
            name = request.form.get("name", "").strip()
            slug = request.form.get("slug", "").strip().lower()
            email = request.form.get("email", "").strip()
            password = request.form.get("password", "").strip()
            # Address fields
            address_line1 = request.form.get("address_line1", "").strip()
            address_line2 = request.form.get("address_line2", "").strip()
            city = request.form.get("city", "").strip()
            state = request.form.get("state", "").strip().upper()
            zip_code = request.form.get("zip", "").strip()
            # Active field: default to 1 (active) if checkbox is checked, else 0
            active = 1 if request.form.get("active") else 0
            
            if not name or not slug or not password:
                flash("Missing required fields", "danger")
            else:
                # Validate state (2 letters A-Z)
                if state and not (len(state) == 2 and state.isalpha()):
                    flash("State must be exactly 2 letters (e.g., CA, NY)", "danger")
                    return redirect(url_for("admin_reps"))
                
                # Validate zip (5 digits or ZIP+4 format: 12345-6789)
                if zip_code and not (len(zip_code) == 5 and zip_code.isdigit() or 
                                    (len(zip_code) == 10 and zip_code[5] == '-' and zip_code[:5].isdigit() and zip_code[6:].isdigit())):
                    flash("ZIP code must be 5 digits (12345) or ZIP+4 format (12345-6789)", "danger")
                    return redirect(url_for("admin_reps"))
                
                pw_hash = generate_password_hash(password)
                try:
                    email_norm = (email or "").strip().lower()
                    if email_norm:
                        # Prevent collisions among active reps
                        existing_active = query_db(
                            "SELECT id FROM reps WHERE active = 1 AND LOWER(email) = %s LIMIT 1",
                            (email_norm,),
                            one=True,
                        )
                        if existing_active:
                            flash("That email is already used by an active rep. Deactivate the old rep or use a different email.", "danger")
                            return redirect(url_for("admin_reps"))

                        # If reusing an email from an inactive rep, clear it from inactive records
                        execute_db(
                            "UPDATE reps SET email = NULL WHERE active = 0 AND LOWER(email) = %s",
                            (email_norm,),
                        )

                    # Use convert_active_to_db_type to ensure integer value (1/0)
                    active_value = convert_active_to_db_type(active)
                    rep_id = execute_db(
                        "INSERT INTO reps (name, slug, password_hash, email, active, address_line1, address_line2, city, state, zip) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING id",
                        (name, slug, pw_hash, email or None, active_value, address_line1 or None, address_line2 or None, city or None, state or None, zip_code or None),
                        returning_id=True,
                    )
                    flash(f"Rep '{name}' created successfully.", "success")
                    if email:
                        login_url = url_for("rep_login", _external=True)
                        ok, msg = send_email(
                            email,
                            "Your Rep Portal account is ready",
                            f"Hi {name},\n\nYour Rep Portal is ready. Sign in here: {login_url}\n\n"
                            "If you need to set or reset your password, use the Forgot password link on the login page.",
                        )
                        if not ok:
                            flash(f"Rep created, but welcome email was not sent: {msg}", "warning")
                except Exception as e:
                    flash(f"Error creating rep: {e}", "danger")

        reps = query_db("SELECT * FROM reps ORDER BY active DESC, name") or []
        reps_norm = [normalize_row(r) or {} for r in reps]
        active_reps = [r for r in reps_norm if int(r.get("active") or 0) == 1]
        inactive_reps = [r for r in reps_norm if int(r.get("active") or 0) != 1]
        return render_template(
            "admin_reps.html",
            reps=reps_norm,
            active_reps=active_reps,
            inactive_reps=inactive_reps,
        )

    @app.route("/admin/email-test", methods=["GET", "POST"])
    def admin_email_test():
        if not is_admin():
            return redirect(url_for("admin_login"))

        default_to = (request.args.get("to") or app.config.get("SMTP_USERNAME") or "").strip()
        smtp_config = {
            "server": app.config.get("SMTP_SERVER") or "",
            "port": str(app.config.get("SMTP_PORT") or ""),
            "use_tls": bool(app.config.get("SMTP_USE_TLS")),
            "username": app.config.get("SMTP_USERNAME") or "",
            "email_from": app.config.get("EMAIL_FROM") or "",
        }

        if request.method == "POST":
            to_address = (request.form.get("to") or "").strip()
            subject = (request.form.get("subject") or "").strip() or "eQMS SMTP test"
            body = (request.form.get("body") or "").strip()
            if not body:
                body = f"This is a test email from eQMS.\n\nSent at: {datetime.utcnow().isoformat()}Z"

            # region agent log
            _agent_log(
                "EMAIL_SEND_MISSING_FN",
                "Proto1.py:admin_email_test",
                "about to call send_email",
                {
                    "has_smtp_server": bool(app.config.get("SMTP_SERVER")),
                    "smtp_port": str(app.config.get("SMTP_PORT") or ""),
                    "send_email_defined": bool("send_email" in globals()),
                    "to_len": len(to_address or ""),
                    "subject_len": len(subject or ""),
                },
            )
            # endregion
            try:
                ok, msg = send_email(to_address, subject, body)
            except Exception as e:
                # region agent log
                _agent_log(
                    "EMAIL_SEND_MISSING_FN",
                    "Proto1.py:admin_email_test",
                    "send_email raised",
                    {"exc_type": type(e).__name__},
                )
                # endregion
                raise
            if ok:
                flash(f"Test email sent to {to_address} from {smtp_config['email_from']}.", "success")
            else:
                flash(f"Test email failed: {msg}", "danger")
            return redirect(url_for("admin_email_test", to=to_address))

        return render_template("admin_email_test.html", default_to=default_to, smtp_config=smtp_config)

    @app.route("/admin/reps/edit/<int:rep_id>", methods=["GET", "POST"])
    def admin_rep_edit(rep_id):
        if not is_admin():
            return redirect(url_for("admin_login"))
        rep = query_db("SELECT * FROM reps WHERE id = %s", (rep_id,), one=True)
        if not rep:
            flash("Rep not found", "danger")
            return redirect(url_for("admin_reps"))
        
        if request.method == "POST":
            name = request.form.get("name", "").strip()
            slug = request.form.get("slug", "").strip().lower()
            email = request.form.get("email", "").strip()
            password = request.form.get("password", "").strip()
            active = 1 if request.form.get("active") else 0
            # Address fields
            address_line1 = request.form.get("address_line1", "").strip()
            address_line2 = request.form.get("address_line2", "").strip()
            city = request.form.get("city", "").strip()
            state = request.form.get("state", "").strip().upper()
            zip_code = request.form.get("zip", "").strip()
            
            # Validate state (2 letters A-Z)
            if state and not (len(state) == 2 and state.isalpha()):
                flash("State must be exactly 2 letters (e.g., CA, NY)", "danger")
                return redirect(url_for("admin_rep_edit", rep_id=rep_id))
            
            # Validate zip (5 digits or ZIP+4 format: 12345-6789)
            if zip_code and not (len(zip_code) == 5 and zip_code.isdigit() or 
                                (len(zip_code) == 10 and zip_code[5] == '-' and zip_code[:5].isdigit() and zip_code[6:].isdigit())):
                flash("ZIP code must be 5 digits (12345) or ZIP+4 format (12345-6789)", "danger")
                return redirect(url_for("admin_rep_edit", rep_id=rep_id))
            
            try:
                email_norm = (email or "").strip().lower()
                if email_norm and active == 1:
                    existing_active = query_db(
                        "SELECT id FROM reps WHERE active = 1 AND LOWER(email) = %s AND id <> %s LIMIT 1",
                        (email_norm, rep_id),
                        one=True,
                    )
                    if existing_active:
                        flash("That email is already used by an active rep. Deactivate the old rep or use a different email.", "danger")
                        return redirect(url_for("admin_rep_edit", rep_id=rep_id))

                    # If reusing an email from an inactive rep, clear it from inactive records
                    execute_db(
                        "UPDATE reps SET email = NULL WHERE active = 0 AND LOWER(email) = %s AND id <> %s",
                        (email_norm, rep_id),
                    )

                if password:
                    pw_hash = generate_password_hash(password)
                    # Use convert_active_to_db_type to ensure integer value (1/0)
                    active_value = convert_active_to_db_type(active)
                    execute_db("UPDATE reps SET name=%s, slug=%s, email=%s, active=%s, password_hash=%s, address_line1=%s, address_line2=%s, city=%s, state=%s, zip=%s WHERE id=%s",
                             (name, slug, email or None, active_value, pw_hash, address_line1 or None, address_line2 or None, city or None, state or None, zip_code or None, rep_id))
                    if email:
                        login_url = url_for("rep_login", _external=True)
                        ok, msg = send_email(
                            email,
                            "Your rep portal password was changed",
                            f"Hi {name},\n\nAn administrator updated your password. You can sign in at {login_url}."
                            " If you did not request this change, please reset your password from the login page.",
                        )
                        if not ok:
                            flash(f"Password updated, but email was not sent: {msg}", "warning")
                else:
                    # Use convert_active_to_db_type to ensure integer value (1/0)
                    active_value = convert_active_to_db_type(active)
                    execute_db("UPDATE reps SET name=%s, slug=%s, email=%s, active=%s, address_line1=%s, address_line2=%s, city=%s, state=%s, zip=%s WHERE id=%s",
                             (name, slug, email or None, active_value, address_line1 or None, address_line2 or None, city or None, state or None, zip_code or None, rep_id))
                flash("Rep updated successfully.", "success")
                return redirect(url_for("admin_reps"))
            except Exception as e:
                flash(f"Error updating rep: {e}", "danger")
        
        return render_template("admin_rep_edit.html", rep=rep)

    @app.route("/admin/reps/<int:rep_id>/delete", methods=["POST"])
    def admin_rep_delete(rep_id):
        """Delete an inactive rep (active reps cannot be deleted)"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        rep = query_db("SELECT * FROM reps WHERE id = %s", (rep_id,), one=True)
        if not rep:
            flash("Rep not found", "danger")
            return redirect(url_for("admin_reps"))
        
        # Check if rep is active
        rep_normalized = normalize_row(rep)
        is_active = int(rep_normalized.get("active") or 0) == 1
        
        if is_active:
            flash("Cannot delete active reps. Deactivate the rep first, then delete.", "danger")
            return redirect(url_for("admin_reps"))
        
        try:
            # Delete rep (cascade will handle related records if foreign keys are set up)
            # If there are foreign key constraints, we may need to handle them explicitly
            execute_db("DELETE FROM reps WHERE id = %s", (rep_id,))
            flash(f"Rep '{rep_normalized.get('name', 'Unknown')}' deleted successfully.", "success")
        except Exception as e:
            # Check if it's a foreign key constraint error
            error_msg = str(e).lower()
            if "foreign key" in error_msg or "constraint" in error_msg:
                flash(f"Cannot delete rep: There are related records (distributions, shipments, etc.) that reference this rep. Please reassign or remove those records first.", "danger")
            else:
                flash(f"Error deleting rep: {e}", "danger")
        
        return redirect(url_for("admin_reps"))

    @app.route("/admin/shipments/unassigned")
    def admin_unassigned_shipments():
        """List all unassigned ShipStation shipments for manual rep assignment"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        try:
            # Get unassigned ShipStation shipments with items
            unassigned = query_db(
                """
                SELECT dd.id, dd.order_number, dd.ship_date, dd.tracking_number, 
                       dd.ss_shipment_id, dd.created_at, dd.customer_id,
                       (SELECT json_agg(json_build_object(
                           'sku', fields_json->>'SKU',
                           'lot', fields_json->>'Lot',
                           'quantity', fields_json->>'Quantity',
                           'facility', fields_json->>'Facility Name'
                       ))
                       FROM device_distribution_records 
                       WHERE dist_id = dd.id) as items
                FROM devices_distributed dd
                WHERE dd.rep_id IS NULL 
                  AND dd.source = 'shipstation'
                ORDER BY dd.ship_date DESC, dd.created_at DESC
                LIMIT 500
                """
            ) or []
            
            # Check for potential conflicts (same facility, same date, different source)
            for ship in unassigned:
                ship_date = (ship.get("ship_date") or "")[:10] if ship.get("ship_date") else None
                items = ship.get("items") or []
                facility = items[0].get("facility") if items else None
                customer_id = ship.get("customer_id")
                
                if facility and ship_date:
                    # Check for manual entries to same customer on same date
                    if customer_id:
                        conflicts = query_db(
                            """
                            SELECT dd.id, dd.order_number, dd.source, r.name as rep_name, dd.rep_id
                            FROM devices_distributed dd
                            LEFT JOIN reps r ON r.id = dd.rep_id
                            WHERE dd.customer_id = %s
                              AND dd.ship_date LIKE %s
                              AND dd.source != 'shipstation'
                              AND dd.rep_id IS NOT NULL
                            LIMIT 5
                            """,
                            (customer_id, f"{ship_date}%")
                        ) or []
                    else:
                        # Fallback: check by facility name and date
                        conflicts = query_db(
                            """
                            SELECT dd.id, dd.order_number, dd.source, r.name as rep_name, dd.rep_id
                            FROM devices_distributed dd
                            LEFT JOIN reps r ON r.id = dd.rep_id
                            JOIN device_distribution_records ddr ON ddr.dist_id = dd.id
                            WHERE dd.source != 'shipstation'
                              AND ddr.fields_json->>'Facility Name' ILIKE %s
                              AND dd.ship_date LIKE %s
                              AND dd.rep_id IS NOT NULL
                            LIMIT 5
                            """,
                            (f"%{facility}%", f"{ship_date}%")
                        ) or []
                    ship["conflicts"] = conflicts
            
            all_reps = query_db("SELECT id, name, slug FROM reps WHERE active = 1 ORDER BY name") or []
            
            return render_template("admin_unassigned_shipments.html", 
                                 shipments=unassigned, 
                                 all_reps=all_reps)
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f"Error loading unassigned shipments: {e}", "danger")
            return redirect(url_for("admin_dashboard"))
    
    @app.route("/admin/shipments/<int:dist_id>/assign", methods=["POST"])
    def admin_assign_shipment(dist_id):
        """Assign a ShipStation shipment to a rep"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        rep_id = request.form.get("rep_id", "").strip()
        if not rep_id:
            flash("Please select a rep", "danger")
            return redirect(url_for("admin_unassigned_shipments"))
        
        try:
            rep_id_int = int(rep_id)
            
            # Verify shipment exists and is unassigned
            shipment = query_db(
                "SELECT id, order_number, source, customer_id, ship_date FROM devices_distributed WHERE id = %s AND rep_id IS NULL AND source = 'shipstation'",
                (dist_id,),
                one=True
            )
            
            if not shipment:
                flash("Shipment not found or already assigned", "danger")
                return redirect(url_for("admin_unassigned_shipments"))
            
            # Check for potential conflicts before assigning
            customer_id = shipment.get("customer_id")
            ship_date = (shipment.get("ship_date") or "")[:10]
            conflicts = []
            
            if customer_id and ship_date:
                # Check for manual entries to same customer on same date with same rep
                conflicts = query_db(
                    """
                    SELECT dd.id, dd.order_number, dd.source, r.name as rep_name
                    FROM devices_distributed dd
                    LEFT JOIN reps r ON r.id = dd.rep_id
                    WHERE dd.customer_id = %s
                      AND dd.ship_date LIKE %s
                      AND dd.source != 'shipstation'
                      AND dd.rep_id = %s
                    LIMIT 5
                    """,
                    (customer_id, f"{ship_date}%", rep_id_int)
                ) or []
            
            if conflicts:
                conflict_msg = f"Warning: {len(conflicts)} manual entr{'y' if len(conflicts) == 1 else 'ies'} exist for same facility/date with this rep. Order{'s' if len(conflicts) != 1 else ''}: {', '.join([c.get('order_number', 'N/A') for c in conflicts])}. Both entries will appear in tracing report."
                flash(conflict_msg, "warning")
            
            # Assign rep to shipment and all related records
            execute_db(
                "UPDATE devices_distributed SET rep_id = %s WHERE id = %s",
                (rep_id_int, dist_id)
            )
            
            execute_db(
                "UPDATE device_distribution_records SET rep_id = %s WHERE dist_id = %s",
                (rep_id_int, dist_id)
            )
            
            execute_db(
                "UPDATE new_customer_records SET rep_id = %s WHERE dist_id = %s",
                (rep_id_int, dist_id)
            )
            
            rep = query_db("SELECT name FROM reps WHERE id = %s", (rep_id_int,), one=True)
            rep_name = rep.get("name") if rep else "Unknown"
            
            flash(f"Shipment {shipment.get('order_number', dist_id)} assigned to {rep_name}", "success")
            return redirect(url_for("admin_unassigned_shipments"))
            
        except ValueError:
            flash("Invalid rep ID", "danger")
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f"Error assigning shipment: {e}", "danger")
        
        return redirect(url_for("admin_unassigned_shipments"))
    
    @app.route("/admin/shipments/bulk-assign", methods=["POST"])
    def admin_bulk_assign_shipments():
        """Bulk assign multiple shipments to a rep"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        rep_id = request.form.get("rep_id", "").strip()
        shipment_ids = request.form.getlist("shipment_ids")
        
        if not rep_id or not shipment_ids:
            flash("Please select a rep and at least one shipment", "danger")
            return redirect(url_for("admin_unassigned_shipments"))
        
        try:
            rep_id_int = int(rep_id)
            assigned_count = 0
            
            for dist_id_str in shipment_ids:
                try:
                    dist_id = int(dist_id_str)
                    
                    # Verify shipment exists and is unassigned
                    shipment = query_db(
                        "SELECT id FROM devices_distributed WHERE id = %s AND rep_id IS NULL AND source = 'shipstation'",
                        (dist_id,),
                        one=True
                    )
                    
                    if shipment:
                        execute_db(
                            "UPDATE devices_distributed SET rep_id = %s WHERE id = %s",
                            (rep_id_int, dist_id)
                        )
                        execute_db(
                            "UPDATE device_distribution_records SET rep_id = %s WHERE dist_id = %s",
                            (rep_id_int, dist_id)
                        )
                        execute_db(
                            "UPDATE new_customer_records SET rep_id = %s WHERE dist_id = %s",
                            (rep_id_int, dist_id)
                        )
                        assigned_count += 1
                except Exception:
                    continue
            
            rep = query_db("SELECT name FROM reps WHERE id = %s", (rep_id_int,), one=True)
            rep_name = rep.get("name") if rep else "Unknown"
            
            flash(f"Assigned {assigned_count} shipment(s) to {rep_name}", "success")
            
        except ValueError:
            flash("Invalid rep ID", "danger")
        except Exception as e:
            import traceback
            traceback.print_exc()
            flash(f"Error bulk assigning shipments: {e}", "danger")
        
        return redirect(url_for("admin_unassigned_shipments"))

    @app.route("/admin/lotlog", methods=["GET", "POST"])
    def admin_lotlog():
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        if request.method == "POST":
            file = request.files.get("file")
            filename = file.filename if file else ""
            # Validate file upload (security hardening)
            if file and filename:
                is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions={".xlsx", ".xls"})
                if not is_valid:
                    flash(error_msg or "Invalid file. Only Excel files (.xlsx, .xls) allowed.", "danger")
                    return redirect(url_for("admin_lot_log"))
            
            if file and filename and filename.lower().endswith((".xlsx", ".xls")):
                try:
                    import openpyxl
                    from datetime import datetime
                    
                    file.stream.seek(0)
                    wb = openpyxl.load_workbook(file.stream)
                    ws = wb.active if wb else None
                    count = 0
                    
                    # Structure: Lot, SKU, Correct Lot Name, Manufacturing Date, Expiration Date, Total Units
                    if ws:
                        for row in ws.iter_rows(min_row=2, values_only=True):
                            if not row or not row[0]:
                                continue
                            
                            lot = str(row[0]).strip()
                            sku = str(row[1]).strip() if len(row) > 1 and row[1] else ""
                            correct = str(row[2]).strip() if len(row) > 2 and row[2] else lot
                            
                            # Handle dates
                            mfg = None
                            exp = None
                            total_units = None
                            
                            if len(row) > 3 and row[3]:
                                try:
                                    from datetime import date
                                    if isinstance(row[3], (datetime, date)):
                                        mfg = row[3].strftime('%Y-%m-%d')
                                    else:
                                        mfg = str(row[3])[:10]
                                except Exception as e:
                                    import traceback
                                    print(f"[ERROR] admin_lot_log_import (mfg date parse): {e}")
                                    traceback.print_exc()
                            
                            if len(row) > 4 and row[4]:
                                try:
                                    from datetime import date
                                    if isinstance(row[4], (datetime, date)):
                                        exp = row[4].strftime('%Y-%m-%d')
                                    else:
                                        exp = str(row[4])[:10]
                                except Exception as e:
                                    import traceback
                                    print(f"[ERROR] admin_lot_log_import (exp date parse): {e}")
                                    traceback.print_exc()
                        
                        if len(row) > 5 and row[5]:
                            try:
                                total_units = int(float(str(row[5])))
                            except Exception as e:
                                import traceback
                                print(f"[ERROR] admin_lot_log_import (total_units parse): {e}")
                                traceback.print_exc()
                        
                        # Upsert
                        exists = query_db("SELECT id FROM lot_log WHERE lot_number=%s", (lot,), one=True)
                        now = datetime.now().isoformat()
                        
                        if exists:
                            execute_db("""
                                UPDATE lot_log 
                                SET sku=%s, correct_lot=%s, mfg_date=%s, exp_date=%s, total_units=%s, updated_at=%s 
                                WHERE lot_number=%s
                            """, (sku, correct, mfg, exp, total_units, now, lot))
                        else:
                            execute_db("""
                                INSERT INTO lot_log (lot_number, sku, correct_lot, mfg_date, exp_date, total_units, updated_at) 
                                VALUES (%s, %s, %s, %s, %s, %s, %s)
                            """, (lot, sku, correct, mfg, exp, total_units, now))
                        count += 1
                    
                    flash(f"Processed {count} lot entries.", "success")
                except Exception as e:
                    flash(f"Error processing file: {e}", "danger")
            else:
                flash("Please upload a valid Excel file", "danger")
        
        recent = query_db("SELECT * FROM lot_log ORDER BY updated_at DESC LIMIT 50") or []
        total_row = query_db("SELECT COUNT(*) as cnt FROM lot_log", one=True)
        total_lots = total_row[0] if total_row and len(total_row) > 0 else 0
        
        return render_template("admin_lotlog.html", recent_lots=recent, total_lots=total_lots)

    @app.route("/admin/lotlog/edit/<int:log_id>", methods=["GET", "POST"])
    def admin_lotlog_edit(log_id):
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        entry = query_db("SELECT * FROM lot_log WHERE id = %s", (log_id,), one=True)
        if not entry:
            flash("Lot entry not found", "danger")
            return redirect(url_for("admin_lotlog"))
        
        if request.method == "POST":
            from datetime import datetime
            lot_number = request.form.get("lot_number", "").strip()
            sku = request.form.get("sku", "").strip()
            correct_lot = request.form.get("correct_lot", "").strip()
            mfg_date = request.form.get("mfg_date", "").strip() or None
            exp_date = request.form.get("exp_date", "").strip() or None
            total_units = request.form.get("total_units", "").strip()
            
            try:
                units = int(total_units) if total_units else None
                execute_db("""
                    UPDATE lot_log 
                    SET lot_number=%s, sku=%s, correct_lot=%s, mfg_date=%s, exp_date=%s, total_units=%s, updated_at=%s 
                    WHERE id=%s
                """, (lot_number, sku, correct_lot, mfg_date, exp_date, units, datetime.now().isoformat(), log_id))
                flash("Lot updated successfully.", "success")
                return redirect(url_for("admin_lotlog"))
            except Exception as e:
                flash(f"Error updating lot: {e}", "danger")
        
        return render_template("admin_lotlog_edit.html", entry=entry)

    @app.route("/admin/lotlog/delete/<int:log_id>", methods=["POST"])
    def admin_lotlog_delete(log_id):
        if not is_admin():
            return redirect(url_for("admin_login"))
        try:
            execute_db("DELETE FROM lot_log WHERE id = %s", (log_id,))
            flash("Lot deleted successfully.", "success")
        except Exception as e:
            flash(f"Error deleting lot: {e}", "danger")
        return redirect(url_for("admin_lotlog"))
    
    # ------------------------------------------------------------------ Admin Received Route
    
    @app.route("/admin/received", methods=["GET", "POST"])
    def admin_received():
        """Manage rep shipments received: packing slip + receiving inspection + lot selection + units"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        if request.method == "POST":
            rep_id = request.form.get("rep_id")
            packing_file = request.files.get("packing_slip")
            receiving_inspection_file = request.files.get("receiving_inspection")
            # Get multiple line items (arrays from form)
            lot_numbers = request.form.getlist("lot_number[]")
            units_received_list = request.form.getlist("units_received[]")
            notes = request.form.get("notes", "").strip()
            
            # Validation
            if not rep_id:
                flash("Rep is required.", "danger")
            elif not packing_file or not packing_file.filename:
                flash("Packing slip is required.", "danger")
            elif not receiving_inspection_file or not receiving_inspection_file.filename:
                flash("Receiving inspection form is required.", "danger")
            elif not lot_numbers or len(lot_numbers) == 0:
                flash("At least one lot selection is required.", "danger")
            elif len(lot_numbers) != len(units_received_list):
                flash("Mismatch between lot numbers and units received.", "danger")
            else:
                # Validate all line items
                line_items = []
                for i, lot_number in enumerate(lot_numbers):
                    lot_number = lot_number.strip()
                    units_str = units_received_list[i].strip() if i < len(units_received_list) else ""
                    
                    if not lot_number:
                        flash(f"Line item {i+1}: Lot is required.", "danger")
                        break
                    if not units_str:
                        flash(f"Line item {i+1}: Units received is required.", "danger")
                        break
                    
                    try:
                        units_received = int(units_str)
                        if units_received <= 0:
                            flash(f"Line item {i+1}: Units received must be greater than 0.", "danger")
                            break
                        
                        # Validate lot exists
                        lot_exists = query_db(
                            "SELECT id, correct_lot FROM lot_log WHERE correct_lot = %s OR lot_number = %s LIMIT 1",
                            (lot_number, lot_number),
                            one=True,
                        )
                        if not lot_exists:
                            flash(f"Line item {i+1}: Lot {lot_number} not found in LotLog. Please add it first.", "danger")
                            break
                        
                        # Prefer canonical correct_lot
                        canonical_lot = (lot_exists.get("correct_lot") or lot_number) if isinstance(lot_exists, dict) else lot_number
                        line_items.append((canonical_lot, units_received))
                    except ValueError:
                        flash(f"Line item {i+1}: Units received must be a valid integer.", "danger")
                        break
                
                # If validation passed, create shipment
                if len(line_items) == len(lot_numbers):
                    # Get rep info
                    rep_row = query_db("SELECT * FROM reps WHERE id = %s", (int(rep_id),), one=True)
                    if not rep_row:
                        flash("Invalid rep.", "danger")
                    else:
                        rep = normalize_row(rep_row)
                        # Create shipment in transaction
                        conn = get_db()
                        if conn:
                            try:
                                with conn.cursor() as cur:
                                    # Create shipment header
                                    created_at = datetime.utcnow().isoformat(timespec="seconds") + "Z"
                                    # Create title from first lot
                                    first_lot = line_items[0][0] if line_items else "Unknown"
                                    title = f"Shipment - {first_lot}" if len(line_items) == 1 else f"Shipment - {len(line_items)} lots"
                                    cur.execute("""
                                        INSERT INTO devices_received (rep_id, title, description, created_at)
                                        VALUES (%s, %s, %s, %s)
                                        RETURNING id
                                    """, (rep_id, title, notes, created_at))
                                    shipment_id = cur.fetchone()[0]
                                    
                                    # Save files to stable folder structure: uploads/received_shipments/<shipment_id>/
                                    base_dir = Path(app.config["UPLOAD_ROOT"]) / "received_shipments" / str(shipment_id)
                                    base_dir.mkdir(parents=True, exist_ok=True)
                                    
                                    # Validate file extensions
                                    allowed_extensions = {'.pdf', '.png', '.jpg', '.jpeg', '.webp'}
                                    packing_ext = Path(packing_file.filename).suffix.lower()
                                    inspection_ext = Path(receiving_inspection_file.filename).suffix.lower()
                                    
                                    if packing_ext not in allowed_extensions:
                                        raise ValueError(f"Packing slip must be PDF, PNG, JPG, JPEG, or WEBP. Got: {packing_ext}")
                                    if inspection_ext not in allowed_extensions:
                                        raise ValueError(f"Receiving inspection must be PDF, PNG, JPG, JPEG, or WEBP. Got: {inspection_ext}")
                                    
                                    # Save packing slip with standardized name
                                    packing_path = base_dir / f"packing_slip{packing_ext}"
                                    packing_file.save(packing_path)
                                    
                                    # Save receiving inspection with standardized name
                                    inspection_path = base_dir / f"receiving_inspection{inspection_ext}"
                                    receiving_inspection_file.save(inspection_path)
                                    
                                    # Store relative paths in database
                                    packing_rel_path = f"received_shipments/{shipment_id}/packing_slip{packing_ext}"
                                    inspection_rel_path = f"received_shipments/{shipment_id}/receiving_inspection{inspection_ext}"
                                    
                                    # Update shipment with file paths
                                    cur.execute("""
                                        UPDATE devices_received
                                        SET packing_list_filename = %s,
                                            packing_list_original = %s,
                                            recv_inspection_filename = %s,
                                            recv_inspection_original = %s
                                        WHERE id = %s
                                    """, (
                                        packing_rel_path,
                                        packing_file.filename,
                                        inspection_rel_path,
                                        receiving_inspection_file.filename,
                                        shipment_id
                                    ))
                                    
                                    # Create line items for all lots
                                    total_units = 0
                                    for lot_number, units_received in line_items:
                                        cur.execute("""
                                            INSERT INTO shipment_line_items (shipment_id, lot_number, units_received, created_at)
                                            VALUES (%s, %s, %s, %s)
                                        """, (shipment_id, lot_number, units_received, created_at))
                                        total_units += units_received
                                    
                                    conn.commit()
                                    if len(line_items) == 1:
                                        flash(f"Shipment created successfully: {total_units} units of lot {line_items[0][0]}.", "success")
                                    else:
                                        flash(f"Shipment created successfully: {total_units} total units across {len(line_items)} lots.", "success")
                            except Exception as e:
                                conn.rollback()
                                flash(f"Error creating shipment: {e}", "danger")
                        else:
                            flash("Database connection error.", "danger")
        
        # GET: Show form and existing shipments
        # Get shipments with their line items (lot, units)
        shipments_raw = query_db("""
            SELECT dr.id, dr.created_at, dr.packing_list_filename, dr.recv_inspection_filename,
                   r.name AS rep_name, r.slug AS rep_slug
            FROM devices_received dr
            JOIN reps r ON dr.rep_id = r.id
            ORDER BY dr.created_at DESC
        """) or []
        
        # Defensive default: ensure shipments_raw is a list
        if shipments_raw is None:
            shipments_raw = []
        
        # Enrich with line item data (show all line items)
        shipments = []
        for s in shipments_raw:
            line_items = query_db("""
                SELECT lot_number, units_received
                FROM shipment_line_items
                WHERE shipment_id = %s
                ORDER BY created_at ASC
            """, (s["id"],)) or []
            
            shipment_dict = dict(s)
            # Store all line items
            shipment_dict["line_items"] = line_items
            # For backward compatibility, also store first item
            if line_items:
                shipment_dict["lot_number"] = line_items[0].get("lot_number")
                shipment_dict["units_received"] = sum(int(item.get("units_received") or 0) for item in line_items)
            else:
                shipment_dict["lot_number"] = None
                shipment_dict["units_received"] = None
            shipments.append(shipment_dict)
        
        reps = query_db("SELECT id, name, slug FROM reps WHERE active = 1 ORDER BY name") or []
        
        # Lot dropdown should show ONLY canonical "correct_lot" values (unique list).
        # We still display a SKU hint when available.
        lots = query_db(
            """
            SELECT
                correct_lot AS lot_number,
                MAX(NULLIF(sku, '')) AS sku
            FROM lot_log
            WHERE correct_lot IS NOT NULL AND btrim(correct_lot) <> ''
            GROUP BY correct_lot
            ORDER BY correct_lot
            """
        ) or []
        
        return render_template("admin_received.html", shipments=shipments, reps=reps, lots=lots)
    
    # ------------------------------------------------------------------ Admin Tracing Routes
    
    @app.route("/admin/tracing/generate/<month>")
    def admin_tracing_generate(month):
        if not is_admin():
            return redirect(url_for("admin_login"))
        results = generate_tracing_reports_for_month(month)
        count = sum(1 for _, ok in results if ok)
        if count:
            flash(f"Generated {count} tracing reports for {month}.", "success")
        else:
            flash("No tracing data found. Tracing report generation requires master distribution logs.", "warning")
        return redirect(url_for("admin_tracing_overview"))
    
    @app.route("/admin/tracing/<int:report_id>/regenerate", methods=["POST"])
    def admin_tracing_regenerate(report_id):
        """Regenerate a single tracing report"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        report = normalize_row(query_db(
            """SELECT tr.*, r.name as rep_name, r.slug
               FROM tracing_reports tr
               JOIN reps r ON r.id = tr.rep_id
               WHERE tr.id = %s""",
            (report_id,),
            one=True
        ))
        
        if not report:
            flash("Tracing report not found.", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        rep = normalize_row(query_db("SELECT * FROM reps WHERE id = %s", (report.get("rep_id"),), one=True))
        if not rep:
            flash("Rep not found.", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        month = report.get("month")
        csv_path = generate_tracing_report_for_rep(rep, month)
        
        if csv_path is not None:
            flash(f"Tracing report regenerated successfully for {report.get('rep_name')} - {month}.", "success")
        else:
            flash(f"No distribution data found for {report.get('rep_name')} - {month}.", "warning")
        
        return redirect(url_for("admin_tracing_view", report_id=report_id))
    
    # Email functionality removed - use CSV downloads instead
    # @app.route("/admin/tracing/send/<month>")
    # def admin_tracing_send(month):
    #     ... removed ...

    @app.route("/admin/tracing/global/<month>/download")
    def admin_tracing_global_download(month: str):
        """Download global monthly tracing report as CSV (admin-only)"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        # Generate report if it doesn't exist
        csv_path = generate_global_tracing_report(month)
        if csv_path is None or not csv_path.exists():
            flash(f"Failed to generate global tracing report for {month}", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        # Read CSV file and serve as download
        try:
            with open(csv_path, 'r', encoding='utf-8') as f:
                csv_content = f.read()
            
            filename = f"Global_Tracing_Report_{month}.csv"
            response = make_response(csv_content)
            response.headers['Content-Type'] = 'text/csv; charset=utf-8'
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        except Exception as e:
            flash(f"Error reading global tracing report: {str(e)}", "danger")
            return redirect(url_for("admin_tracing_overview"))

    @app.route("/admin/tracing")
    def admin_tracing_overview():
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        smtp_configured = bool((app.config.get("SMTP_SERVER") or "").strip())
        
        try:
            reports_raw = query_db(
                """
                SELECT tr.*, r.name as rep_name, r.slug
                FROM tracing_reports tr
                JOIN reps r ON r.id = tr.rep_id
                ORDER BY tr.month DESC, r.name ASC
                """
            ) or []
        except Exception as e:
            print(f"ERROR in admin_tracing_overview: Database query failed: {e}")
            import traceback
            traceback.print_exc()
            flash("Error loading tracing reports. Please check server logs.", "danger")
            reports_raw = []
        
        # Helper function to safely convert datetime to string
        def safe_datetime_to_str(dt_value):
            """Safely convert datetime to string, return empty string on any error"""
            if dt_value is None:
                return ""
            try:
                if isinstance(dt_value, datetime):
                    return dt_value.isoformat()[:19]  # YYYY-MM-DD HH:MM:SS
                elif isinstance(dt_value, str):
                    return dt_value[:19] if len(dt_value) >= 19 else dt_value
                else:
                    return str(dt_value)[:19] if dt_value else ""
            except Exception:
                return ""
        
        # Helper function to safely extract filename from path
        def safe_extract_filename(path_value):
            """Safely extract filename from path string"""
            if not path_value:
                return ""
            try:
                path_str = str(path_value).strip()
                if not path_str:
                    return ""
                try:
                    return Path(path_str).name
                except Exception:
                    # Fallback: string manipulation
                    if "/" in path_str:
                        return path_str.split("/")[-1]
                    elif "\\" in path_str:
                        return path_str.split("\\")[-1]
                    else:
                        return path_str
            except Exception:
                return ""
        
        # Process each report with comprehensive error handling
        # CRITICAL: Convert DictRow to plain dict BEFORE adding computed fields
        # psycopg2.extras.DictRow does not allow adding new keys (KeyError on assignment)
        reports = []
        for r in reports_raw:
            try:
                # Convert DictRow to plain dict immediately to allow adding computed fields
                # This prevents KeyError when trying to set r["report_filename"] on DictRow
                row = dict(r) if not isinstance(r, dict) else r.copy()
                
                # Check for missing expected fields and log (server-side only, no secrets)
                missing_fields = []
                if not row.get("id"):
                    missing_fields.append("id")
                if not row.get("rep_id"):
                    missing_fields.append("rep_id")
                if not row.get("month"):
                    missing_fields.append("month")
                if not row.get("rep_name"):
                    missing_fields.append("rep_name")
                if not row.get("slug"):
                    missing_fields.append("slug")
                if missing_fields:
                    report_id = row.get("id", "unknown")
                    month = row.get("month", "unknown")
                    rep_name = row.get("rep_name", "unknown")
                    print(f"WARNING: Tracing report missing fields - id={report_id}, month={month}, rep={rep_name}, missing={missing_fields}")
                
                # Ensure ALL required keys exist with safe defaults
                report = {
                    "id": row.get("id") or 0,
                    "rep_id": row.get("rep_id") or 0,
                    "month": str(row.get("month") or "").strip() or "Unknown",
                    "rep_name": str(row.get("rep_name") or "").strip() or "Unknown Rep",
                    "slug": str(row.get("slug") or "").strip() or "",
                    "status": str(row.get("status") or "").strip() or "draft",
                    "report_path": str(row.get("report_path") or "").strip() if row.get("report_path") else "",
                    "approval_file_path": str(row.get("approval_file_path") or "").strip() if row.get("approval_file_path") else "",
                    "email_to": str(row.get("email_to") or "").strip() if row.get("email_to") else "",
                    "data_issue": False,  # Flag for malformed records
                }
                
                # Extract filenames safely (now safe because report is a plain dict)
                report["approval_filename"] = safe_extract_filename(report["approval_file_path"])
                report["report_filename"] = safe_extract_filename(report["report_path"])
                
                # Handle missing report_path: show "Missing file" indicator
                if not report["report_path"]:
                    report["report_filename"] = ""
                    report["missing_file"] = True
                else:
                    report["missing_file"] = False
                
                # Convert datetime objects to strings safely
                report["generated_at"] = safe_datetime_to_str(row.get("generated_at"))
                report["email_sent_at"] = safe_datetime_to_str(row.get("email_sent_at"))
                report["approval_uploaded_at"] = safe_datetime_to_str(row.get("approval_uploaded_at"))
                
                reports.append(report)
                
            except Exception as e:
                # If a record is malformed, create a safe placeholder
                report_id = r.get("id", "unknown") if hasattr(r, "get") else "unknown"
                month = r.get("month", "unknown") if hasattr(r, "get") else "unknown"
                rep_name = r.get("rep_name", "unknown") if hasattr(r, "get") else "unknown"
                print(f"ERROR processing report row - id={report_id}, month={month}, rep={rep_name}: {e}")
                import traceback
                traceback.print_exc()
                
                # Create a safe placeholder record with "Data Issue" badge
                reports.append({
                    "id": report_id if isinstance(report_id, int) else 0,
                    "rep_id": 0,
                    "month": "Unknown",
                    "rep_name": "Data Issue",
                    "slug": "",
                    "status": "error",
                    "report_path": "",
                    "approval_file_path": "",
                    "approval_filename": "",
                    "report_filename": "",
                    "email_to": "",
                    "generated_at": "",
                    "email_sent_at": "",
                    "approval_uploaded_at": "",
                    "missing_file": False,
                    "data_issue": True,
                })
        
        try:
            return render_template("admin_tracing_overview.html", reports=reports, smtp_configured=smtp_configured)
        except Exception as e:
            print(f"CRITICAL ERROR rendering admin_tracing_overview template: {e}")
            import traceback
            traceback.print_exc()
            # Return a simple error page with correlation ID
            error_id = secrets.token_hex(8)
            return f"""<html><head><title>Error</title></head><body>
                <h1>Error Loading Tracing Overview</h1>
                <p>Error ID: {error_id}</p>
                <p>Please check server logs for details.</p>
                <p><a href="/admin/dashboard">Return to Dashboard</a></p>
            </body></html>""", 500

    @app.route("/admin/tracing/<int:report_id>/approve", methods=["POST"])
    def admin_tracing_approve(report_id):
        if not is_admin():
            return redirect(url_for("admin_login"))
        execute_db(
            "UPDATE tracing_reports SET status='approved', approval_uploaded_at=%s WHERE id=%s",
            (datetime.now(), report_id),
        )
        flash("Tracing report marked approved.", "success")
        return redirect(url_for("admin_tracing_overview"))

    @app.route("/admin/tracing/<int:report_id>/dispute", methods=["POST"])
    def admin_tracing_dispute(report_id):
        if not is_admin():
            return redirect(url_for("admin_login"))
        note = request.form.get("note", "")
        execute_db(
            "UPDATE tracing_reports SET status='disputed', dispute_note=%s WHERE id=%s",
            (note, report_id),
        )
        flash("Tracing report marked disputed.", "warning")
        return redirect(url_for("admin_tracing_overview"))

    @app.route("/admin/tracing/<int:report_id>/edit", methods=["GET", "POST"])
    def admin_tracing_edit(report_id):
        """Edit a draft tracing report CSV"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        report = normalize_row(query_db(
            "SELECT * FROM tracing_reports WHERE id = %s",
            (report_id,),
            one=True
        ))
        
        if not report:
            abort(404)
        
        # Only allow editing draft reports
        if report.get("status") != "draft":
            flash("Only draft reports can be edited.", "warning")
            return redirect(url_for("admin_tracing_overview"))
        
        if request.method == "POST":
            # Get edited rows from form
            rows = []
            row_count = int(request.form.get("row_count", 0))
            
            for i in range(row_count):
                if request.form.get(f"delete_{i}") == "1":
                    continue  # Skip deleted rows
                
                rows.append({
                    "Ship Date": request.form.get(f"ship_date_{i}", ""),
                    "Order #": request.form.get(f"order_number_{i}", ""),
                    "Facility": request.form.get(f"facility_{i}", ""),
                    "City": request.form.get(f"city_{i}", ""),
                    "State": request.form.get(f"state_{i}", ""),
                    "SKU": request.form.get(f"sku_{i}", ""),
                    "Lot": request.form.get(f"lot_{i}", ""),
                    "Quantity": request.form.get(f"quantity_{i}", ""),
                })
            
            # Write updated CSV to DB (persistent storage)
            fieldnames = ["Ship Date", "Order #", "Facility", "City", "State", "SKU", "Lot", "Quantity"]
            csv_buffer = io.StringIO()
            writer = csv.DictWriter(csv_buffer, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(rows)
            csv_content = csv_buffer.getvalue()
            csv_buffer.close()
            
            # Update DB with new CSV content
            edited_at = datetime.now()
            performed_by = session.get("admin_username", "admin") or "admin"
            execute_db(
                "UPDATE tracing_reports SET report_csv_content = %s, edited_at = %s WHERE id = %s",
                (csv_content, edited_at, report_id)
            )
            
            # Also write to filesystem for backward compatibility (if directory exists)
            report_path_str = report.get("report_path", "")
            if report_path_str:
                report_path = Path(report_path_str)
                if not report_path.is_absolute():
                    report_path = BASE_DIR / report_path
                try:
                    report_path.parent.mkdir(parents=True, exist_ok=True)
                    with open(report_path, "w", newline="", encoding="utf-8") as f:
                        f.write(csv_content)
                except Exception as e:
                    print(f"WARNING: Could not write CSV to filesystem: {e}. Report updated in DB only.")
            
            # Record history event
            try:
                execute_db(
                    """INSERT INTO tracing_report_history (report_id, event_type, event_at, performed_by, note)
                       VALUES (%s, 'edited', %s, %s, 'Report edited via web interface')""",
                    (report_id, edited_at, performed_by)
                )
            except Exception:
                pass  # History recording is optional
            
            flash("Tracing report updated successfully.", "success")
            
            return redirect(url_for("admin_tracing_overview"))
        
        # GET: Show edit form
        rows = []
        # Prefer DB storage
        csv_content = report.get("report_csv_content")
        if csv_content:
            reader = csv.DictReader(io.StringIO(csv_content))
            rows = list(reader)
        else:
            # Fallback to filesystem
            report_path_str = report.get("report_path", "")
            if report_path_str:
                report_path = Path(report_path_str)
                if not report_path.is_absolute():
                    report_path = BASE_DIR / report_path
                if report_path.exists():
                    with open(report_path, "r", encoding="utf-8") as f:
                        reader = csv.DictReader(f)
                        rows = list(reader)
        
        return render_template("admin_tracing_edit.html", report=report, rows=rows)

    @app.route("/admin/tracing/<int:report_id>/view")
    def admin_tracing_view(report_id):
        """View tracing report CSV in browser with pagination"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        report = normalize_row(query_db(
            """SELECT tr.*, r.name as rep_name, r.slug
               FROM tracing_reports tr
               JOIN reps r ON r.id = tr.rep_id
               WHERE tr.id = %s""",
            (report_id,),
            one=True
        ))
        
        if not report:
            abort(404)
        
        # Get DB distribution count for comparison (same query as generate_tracing_report_for_rep)
        db_dist_count = 0
        db_dist_query_error = None
        rep_id = report.get("rep_id")
        month = report.get("month")
        if rep_id and month:
            try:
                year = int(month[:4])
                mon = int(month[5:7])
                start_date = f"{year}-{mon:02d}-01"
                end_date = f"{year + 1}-01-01" if mon == 12 else f"{year}-{mon + 1:02d}-01"
                
                db_rows = query_db(
                    """
                    SELECT dd.id, ddr.fields_json
                    FROM devices_distributed dd
                    JOIN device_distribution_records ddr ON ddr.dist_id = dd.id
                    WHERE dd.rep_id = %s 
                      AND dd.ship_date >= %s 
                      AND dd.ship_date < %s
                    """,
                    (rep_id, start_date, end_date),
                ) or []
                
                # Count valid rows (same filter as generate_tracing_report_for_rep)
                for r in db_rows:
                    try:
                        fields = json.loads(r["fields_json"])
                    except Exception:
                        fields = {}
                    sku = fields.get("SKU", "").strip() if fields.get("SKU") else ""
                    lot = str(fields.get("Lot", "") or "").strip()
                    quantity_str = fields.get("Quantity", 0) or 0
                    try:
                        quantity = int(float(quantity_str))
                    except (ValueError, TypeError):
                        quantity = 0
                    # Only count valid rows (same filter as generation)
                    if not ((not sku and quantity <= 0) or (not sku and not lot and quantity <= 0)):
                        db_dist_count += 1
            except Exception as e:
                db_dist_query_error = str(e)
                print(f"[TRACING VIEW] Error querying DB distributions: {e}")
        
        # Read CSV file
        report_path_str = report.get("report_path", "")
        report_path = Path(report_path_str)
        if not report_path.is_absolute():
            report_path = BASE_DIR / report_path

        # Auto-heal: if DB path is missing/stale but canonical path exists, use it and update DB.
        try:
            canonical = get_tracing_report_path(
                Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")),
                report.get("slug") or "",
                report.get("month") or "",
                ext="csv",
            )
            if (not report_path_str or not report_path.exists()) and canonical.exists():
                new_rel = str(canonical.relative_to(BASE_DIR))
                execute_db("UPDATE tracing_reports SET report_path = %s WHERE id = %s", (new_rel, report_id))
                report["report_path"] = new_rel
                report_path_str = new_rel
                report_path = canonical
                print(f"[TRACING VIEW] report_id={report_id} auto_healed_path='{new_rel}'")
        except Exception as e:
            # Never block view on path healing
            print(f"[TRACING VIEW] report_id={report_id} auto_heal_failed error='{e}'")
        
        rows = []
        total_rows = 0
        page = int(request.args.get("page", 1))
        per_page = int(request.args.get("per_page", 50))
        csv_error = None
        csv_error_type = None  # 'missing', 'empty', 'unreadable'
        
        # Prefer DB storage over filesystem (persistent across redeploys)
        csv_content = report.get("report_csv_content")
        if csv_content:
            # Read from DB storage
            try:
                reader = csv.DictReader(io.StringIO(csv_content))
                all_rows = list(reader)
                total_rows = len(all_rows)
                if total_rows == 0:
                    csv_error = "Report exists but contains no distribution records. The CSV may only have headers."
                    csv_error_type = "empty"
                else:
                    start = (page - 1) * per_page
                    end = start + per_page
                    rows = all_rows[start:end]
            except Exception as e:
                csv_error = f"Error parsing CSV from database: {str(e)}"
                csv_error_type = "unreadable"
                print(f"[TRACING VIEW] report_id={report_id} reason='db_unreadable' error='{str(e)}'")
        elif report_path.exists():
            # Fallback to filesystem (backward compatibility)
            try:
                with open(report_path, "r", encoding="utf-8") as f:
                    reader = csv.DictReader(f)
                    all_rows = list(reader)
                    total_rows = len(all_rows)
                    if total_rows == 0:
                        csv_error = "Report file exists but contains no distribution records."
                        csv_error_type = "empty"
                    else:
                        start = (page - 1) * per_page
                        end = start + per_page
                        rows = all_rows[start:end]
            except Exception as e:
                csv_error = f"Error reading CSV file: {str(e)}"
                csv_error_type = "unreadable"
        else:
            csv_error = f"Report file not found and no DB content available. Path: {report_path_str}"
            csv_error_type = "missing"
            print(f"[TRACING VIEW] report_id={report_id} path='{report_path_str}' reason='file_not_found_no_db'")
        
        total_pages = (total_rows + per_page - 1) // per_page if total_rows > 0 else 1
        
        # Calculate summary stats from all rows (for summary box)
        summary_stats = {"total_units": 0, "unique_orders": set(), "unique_facilities": set()}
        if csv_content:
            # Use DB content for summary
            try:
                reader = csv.DictReader(io.StringIO(csv_content))
                all_summary_rows = list(reader)
                for row in all_summary_rows:
                    try:
                        qty = int(float(row.get('Quantity', 0) or 0))
                        summary_stats["total_units"] += qty
                    except (ValueError, TypeError):
                        pass
                    order_num = (row.get('Order #') or '').strip()
                    if order_num:
                        summary_stats["unique_orders"].add(order_num)
                    facility = (row.get('Facility') or '').strip()
                    if facility:
                        summary_stats["unique_facilities"].add(facility)
            except Exception:
                pass
        elif report_path.exists() and not csv_error:
            # Fallback to filesystem for summary
            try:
                with open(report_path, "r", encoding="utf-8") as f:
                    reader = csv.DictReader(f)
                    all_summary_rows = list(reader)
                    for row in all_summary_rows:
                        try:
                            qty = int(float(row.get('Quantity', 0) or 0))
                            summary_stats["total_units"] += qty
                        except (ValueError, TypeError):
                            pass
                        order_num = (row.get('Order #') or '').strip()
                        if order_num:
                            summary_stats["unique_orders"].add(order_num)
                        facility = (row.get('Facility') or '').strip()
                        if facility:
                            summary_stats["unique_facilities"].add(facility)
            except Exception:
                pass
        
        summary_stats["unique_orders"] = len(summary_stats["unique_orders"])
        summary_stats["unique_facilities"] = len(summary_stats["unique_facilities"])
        
        return render_template("admin_tracing_view.html", 
                             report=report, 
                             rows=rows, 
                             page=page, 
                             per_page=per_page,
                             total_rows=total_rows,
                             total_pages=total_pages,
                             summary_stats=summary_stats,
                             csv_error=csv_error,
                             csv_error_type=csv_error_type,
                             db_dist_count=db_dist_count,
                             db_dist_query_error=db_dist_query_error)

    @app.route("/admin/tracing/<int:report_id>/replace", methods=["GET", "POST"])
    def admin_tracing_replace(report_id):
        """Upload replacement CSV with versioning"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        report = normalize_row(query_db(
            """SELECT tr.*, r.name as rep_name, r.slug
               FROM tracing_reports tr
               JOIN reps r ON r.id = tr.rep_id
               WHERE tr.id = %s""",
            (report_id,),
            one=True
        ))
        
        if not report:
            abort(404)
        
        if request.method == "POST":
            file = request.files.get("file")
            note = request.form.get("note", "").strip()
            
            if not file or not file.filename:
                flash("Please select a CSV file to upload.", "warning")
                return redirect(url_for("admin_tracing_replace", report_id=report_id))
            
            # Validate file upload (security hardening)
            is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions=app.config["ALLOWED_CSV_EXTENSIONS"])
            if not is_valid:
                flash(error_msg or "Invalid file. Only CSV files allowed.", "danger")
                return redirect(url_for("admin_tracing_replace", report_id=report_id))
            
            try:
                # Read uploaded file to validate
                uploaded_content = file.read()
                file.seek(0)  # Reset for saving
                
                # Try to parse CSV
                csv_str = uploaded_content.decode('utf-8')
                reader = csv.DictReader(csv_str.splitlines())
                fieldnames = reader.fieldnames or []
                required_fields = ["Ship Date", "Order #", "Facility", "City", "State", "SKU", "Lot", "Quantity"]
                
                missing_fields = [f for f in required_fields if f not in fieldnames]
                if missing_fields:
                    flash(f"CSV missing required columns: {', '.join(missing_fields)}", "danger")
                    return redirect(url_for("admin_tracing_replace", report_id=report_id))
                
                # Get current report path
                current_report_path_str = report.get("report_path", "")
                current_report_path = Path(current_report_path_str)
                if not current_report_path.is_absolute():
                    current_report_path = BASE_DIR / current_report_path
                
                # Create versioned directory
                report_dir = current_report_path.parent
                versions_dir = report_dir / "versions"
                versions_dir.mkdir(exist_ok=True)
                
                # Save old file to versions directory
                old_file_path = None
                if current_report_path.exists():
                    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    old_version_path = versions_dir / f"{current_report_path.stem}_v{timestamp}{current_report_path.suffix}"
                    import shutil
                    shutil.copy2(current_report_path, old_version_path)
                    old_file_path = str(old_version_path.relative_to(BASE_DIR))
                
                # Save new file (overwrite current) - ensure path is safe
                safe_path = ensure_safe_path(current_report_path.parent, current_report_path.name)
                file.save(safe_path)
                new_file_path = str(current_report_path.relative_to(BASE_DIR))
                
                # Update database
                replaced_at = datetime.now()
                replaced_by = session.get("admin_username", "admin") or "admin"
                
                execute_db(
                    """UPDATE tracing_reports 
                       SET report_path = %s, replaced_at = %s, replaced_by = %s, replacement_note = %s, edited_at = %s
                       WHERE id = %s""",
                    (new_file_path, replaced_at, replaced_by, note, replaced_at, report_id)
                )
                
                # Record history
                execute_db(
                    """INSERT INTO tracing_report_history (report_id, event_type, event_at, performed_by, file_path, note, old_file_path)
                       VALUES (%s, 'replaced', %s, %s, %s, %s, %s)""",
                    (report_id, replaced_at, replaced_by, new_file_path, note, old_file_path)
                )
                
                flash(f"Tracing report replaced successfully. Old version saved to versions/ directory.", "success")
                return redirect(url_for("admin_tracing_view", report_id=report_id))
                
            except Exception as e:
                flash(f"Error replacing report: {e}", "danger")
                import traceback
                traceback.print_exc()
        
        return render_template("admin_tracing_replace.html", report=report)

    @app.route("/admin/tracing/<int:report_id>/history")
    def admin_tracing_history(report_id):
        """View document history for a tracing report"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        report = normalize_row(query_db(
            """SELECT tr.*, r.name as rep_name, r.slug
               FROM tracing_reports tr
               JOIN reps r ON r.id = tr.rep_id
               WHERE tr.id = %s""",
            (report_id,),
            one=True
        ))
        
        if not report:
            abort(404)
        
        # Get history events
        history = query_db(
            """SELECT * FROM tracing_report_history 
               WHERE report_id = %s 
               ORDER BY event_at DESC""",
            (report_id,)
        ) or []
        
        # Build full history including DB events
        full_history = []
        
        # Report generated
        if report.get("generated_at"):
            full_history.append({
                "event_type": "generated",
                "event_at": report.get("generated_at"),
                "performed_by": "system",
                "note": "Report generated",
                "file_path": report.get("report_path")
            })
        
        # Report replaced (from history table)
        for h in history:
            full_history.append(h)
        
        # Report emailed
        if report.get("email_sent_at"):
            full_history.append({
                "event_type": "emailed",
                "event_at": report.get("email_sent_at"),
                "performed_by": "system",
                "note": f"Email sent to {report.get('email_to', 'rep')}"
            })
        
        # Report approved
        if report.get("approval_uploaded_at"):
            full_history.append({
                "event_type": "approved",
                "event_at": report.get("approval_uploaded_at"),
                "performed_by": "rep",
                "note": "Approval recorded",
                "file_path": report.get("approval_file_path")
            })
        
        # Report edited
        if report.get("edited_at"):
            full_history.append({
                "event_type": "edited",
                "event_at": report.get("edited_at"),
                "performed_by": report.get("replaced_by", "admin"),
                "note": "Report edited via web interface"
            })
        
        # Sort by date
        full_history.sort(key=lambda x: (
            x["event_at"] if isinstance(x["event_at"], datetime) 
            else datetime.fromisoformat(str(x["event_at"]).replace('Z', '+00:00')) 
            if isinstance(x["event_at"], str) 
            else datetime.now()
        ), reverse=True)
        
        return render_template("admin_tracing_history.html", report=report, history=full_history)

    # ------------------------------------------------------------------ Distribution Log Approval Routes
    
    @app.route("/admin/distribution-log-approvals")
    def admin_distribution_log_approvals():
        """List all monthly distribution log approvals"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        try:
            approvals_raw = query_db(
                """SELECT * FROM distribution_log_approvals 
                   ORDER BY month DESC"""
            ) or []
        except Exception as e:
            print(f"ERROR in admin_distribution_log_approvals: {e}")
            import traceback
            traceback.print_exc()
            flash("Error loading distribution log approvals.", "danger")
            approvals_raw = []
        
        approvals = []
        for a in approvals_raw:
            approval = normalize_row(a)
            approvals.append({
                "month": approval.get("month", ""),
                "status": approval.get("status", "draft"),
                "generated_at": approval.get("generated_at"),
                "email_sent_at": approval.get("email_sent_at"),
                "approval_uploaded_at": approval.get("approval_uploaded_at"),
                "approval_sender_email": approval.get("approval_sender_email"),
            })
        
        return render_template("admin_distribution_log_approvals.html", approvals=approvals)
    
    @app.route("/admin/distribution-log-approvals/<month>")
    def admin_distribution_log_approval_detail(month):
        """View/upload approval for specific month"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        month_normalized = month.strip()
        if len(month_normalized) != 7 or month_normalized[4] != '-':
            flash("Invalid month format. Expected YYYY-MM", "danger")
            return redirect(url_for("admin_distribution_log_approvals"))
        
        try:
            approval = normalize_row(
                query_db(
                    "SELECT * FROM distribution_log_approvals WHERE month = %s",
                    (month_normalized,),
                    one=True
                )
            )
        except Exception as e:
            print(f"ERROR: {e}")
            approval = None
        
        # Check if report file exists
        report_exists = False
        report_path = None
        if approval and approval.get("report_path"):
            report_path = BASE_DIR / approval["report_path"]
            report_exists = report_path.exists()
        
        # Check if approval .eml file exists
        eml_exists = False
        eml_path = None
        if approval and approval.get("approval_eml_path"):
            eml_path = BASE_DIR / approval["approval_eml_path"]
            eml_exists = eml_path.exists()
        
        return render_template(
            "admin_distribution_log_approval_detail.html",
            approval=approval,
            month=month_normalized,
            report_exists=report_exists,
            report_path=report_path,
            eml_exists=eml_exists,
            eml_path=eml_path
        )
    
    @app.route("/admin/distribution-log-approvals/<month>/generate", methods=["POST"])
    def admin_distribution_log_generate(month):
        """Generate distribution log CSV for month"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        path = generate_distribution_log_for_month(month)
        if path:
            flash(f"Distribution log generated for {month}", "success")
        else:
            flash(f"Failed to generate distribution log for {month}", "danger")
        return redirect(url_for("admin_distribution_log_approval_detail", month=month))
    
    @app.route("/admin/distribution-log-approvals/<month>/send", methods=["POST"])
    def admin_distribution_log_send(month):
        """Send approval request email"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        ok, msg = send_distribution_log_email(month)
        if ok:
            flash(f"Approval request email sent for {month}", "success")
        else:
            flash(f"Failed to send email: {msg}", "danger")
        return redirect(url_for("admin_distribution_log_approval_detail", month=month))
    
    @app.route("/admin/distribution-log-approvals/<month>/upload", methods=["POST"])
    def admin_distribution_log_upload(month):
        """Upload approval .eml file"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        month_normalized = month.strip()
        if len(month_normalized) != 7 or month_normalized[4] != '-':
            flash("Invalid month format", "danger")
            return redirect(url_for("admin_distribution_log_approvals"))
        
        file = request.files.get("eml_file")
        if not file or not file.filename:
            flash("Please select a .eml file to upload", "warning")
            return redirect(url_for("admin_distribution_log_approval_detail", month=month))
        
        # Validate file upload (security hardening)
        is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions={".eml"})
        if not is_valid:
            flash(error_msg or "Please upload a .eml file", "warning")
            return redirect(url_for("admin_distribution_log_approval_detail", month=month))
        
        try:
            # Ensure month is safe (no path traversal)
            if ".." in month_normalized or "/" in month_normalized or "\\" in month_normalized:
                flash("Invalid month format", "danger")
                return redirect(url_for("admin_distribution_log_approval_detail", month=month))
            
            # Save uploaded file temporarily
            month_dir = Path(app.config.get("DIST_LOG_APPROVALS_DIR", BASE_DIR / "distribution_log_approvals")) / month_normalized
            month_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate secure filename with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            stored_filename = f"Approval_{month_normalized}_{timestamp}.eml"
            
            # Ensure path is safe
            safe_path = ensure_safe_path(month_dir, stored_filename)
            file.save(safe_path)
            eml_path = safe_path
            
            # Parse and validate .eml file
            eml_data = parse_approval_eml(eml_path)
            if not eml_data:
                eml_path.unlink()  # Delete invalid file
                flash("Failed to parse .eml file", "danger")
                return redirect(url_for("admin_distribution_log_approval_detail", month=month))
            
            is_valid, error_msg = validate_approval_eml(eml_data)
            if not is_valid:
                eml_path.unlink()  # Delete invalid file
                flash(f"Approval validation failed: {error_msg}", "danger")
                return redirect(url_for("admin_distribution_log_approval_detail", month=month))
            
            # Store approval
            eml_relative_path = str(eml_path.relative_to(BASE_DIR))
            approval_uploaded_at = datetime.now()
            approved_by = session.get("admin_username") or "admin"
            
            execute_db(
                """UPDATE distribution_log_approvals
                   SET status = 'approved',
                       approval_eml_path = %s,
                       approval_sender_email = %s,
                       approval_subject = %s,
                       approval_date = %s,
                       approval_uploaded_at = %s,
                       approved_by = %s
                   WHERE month = %s""",
                (
                    eml_relative_path,
                    eml_data.get("from", ""),
                    eml_data.get("subject", ""),
                    eml_data.get("date"),
                    approval_uploaded_at,
                    approved_by,
                    month_normalized
                )
            )
            
            flash(f"Approval uploaded and validated successfully for {month}", "success")
            return redirect(url_for("admin_distribution_log_approval_detail", month=month))
            
        except Exception as e:
            print(f"ERROR uploading approval .eml: {e}")
            import traceback
            traceback.print_exc()
            flash(f"Error uploading approval: {str(e)}", "danger")
            return redirect(url_for("admin_distribution_log_approval_detail", month=month))
    
    @app.route("/admin/distribution-log-approvals/<month>/download/<filename>")
    def admin_distribution_log_download(month, filename):
        """Download report or approval file"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        try:
            month_normalized = month.strip()
            month_dir = Path(app.config.get("DIST_LOG_APPROVALS_DIR", BASE_DIR / "distribution_log_approvals")) / month_normalized
            file_path = month_dir / secure_filename(filename)
            
            # Security: ensure file is within the month directory
            try:
                file_path.resolve().relative_to(month_dir.resolve())
            except ValueError:
                print(f"[DIST LOG DOWNLOAD] Security check failed: month={month}, filename={filename}")
                return render_template_string("""
                <!DOCTYPE html>
                <html>
                <head><title>File Not Found</title></head>
                <body style="font-family: Arial, sans-serif; padding: 40px; text-align: center;">
                    <h1>File Not Found</h1>
                    <p>The requested file could not be found for {{ month }}.</p>
                    <p><small>Security validation failed. Please contact an administrator.</small></p>
                    <p><a href="{{ url_for('admin_distribution_log_approvals') }}">← Back to Distribution Log Approvals</a></p>
                </body>
                </html>
                """, month=month_normalized), 404
            
            if not file_path.exists() or not file_path.is_file():
                print(f"[DIST LOG DOWNLOAD] File not found: month={month}, filename={filename}, path={file_path}")
                return render_template_string("""
                <!DOCTYPE html>
                <html>
                <head><title>File Not Found</title></head>
                <body style="font-family: Arial, sans-serif; padding: 40px; text-align: center;">
                    <h1>File Not Found</h1>
                    <p>The file "{{ filename }}" could not be found for {{ month }}.</p>
                    <p><small>The file may have been deleted or moved. Please regenerate the report or re-upload the approval.</small></p>
                    <p><a href="{{ url_for('admin_distribution_log_approvals') }}">← Back to Distribution Log Approvals</a></p>
                </body>
                </html>
                """, filename=secure_filename(filename), month=month_normalized), 404
            
            return send_from_directory(str(month_dir), secure_filename(filename))
        except Exception as e:
            print(f"[DIST LOG DOWNLOAD] Error: month={month}, filename={filename}, error={str(e)}")
            return render_template_string("""
            <!DOCTYPE html>
            <html>
            <head><title>Download Error</title></head>
            <body style="font-family: Arial, sans-serif; padding: 40px; text-align: center;">
                <h1>Download Error</h1>
                <p>An error occurred while downloading the file. Please try again or contact an administrator.</p>
                <p><a href="{{ url_for('admin_distribution_log_approvals') }}">← Back to Distribution Log Approvals</a></p>
            </body>
            </html>
            """), 500

    @app.route("/admin/documents")
    def admin_documents():
        """Central admin document manager"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        # Get filters
        rep_filter = request.args.get("rep", "")
        month_filter = request.args.get("month", "")
        doc_type_filter = request.args.get("type", "")
        
        # Build query for tracing reports
        query = """
            SELECT tr.*, r.name as rep_name, r.slug
            FROM tracing_reports tr
            JOIN reps r ON r.id = tr.rep_id
            WHERE 1=1
        """
        params = []
        
        if rep_filter:
            query += " AND r.slug = %s"
            params.append(rep_filter)
        
        if month_filter:
            query += " AND tr.month = %s"
            params.append(month_filter)
        
        query += " ORDER BY tr.month DESC, r.name ASC"
        
        reports = query_db(query, tuple(params)) or []
        
        # Get unlinked documents
        ensure_rep_documents_table()
        unlinked_query = """
            SELECT rd.*, r.name as rep_name, r.slug
            FROM rep_documents rd
            JOIN reps r ON rd.rep_id = r.id
            WHERE rd.distribution_id IS NULL
        """
        unlinked_params = []
        
        if rep_filter:
            unlinked_query += " AND r.slug = %s"
            unlinked_params.append(rep_filter)
        
        unlinked_query += " ORDER BY rd.uploaded_at DESC LIMIT 50"
        
        unlinked_docs = query_db(unlinked_query, tuple(unlinked_params)) or []
        
        # Get all reps for filter dropdown
        all_reps = query_db("SELECT slug, name FROM reps WHERE active = 1 ORDER BY name") or []
        
        return render_template("admin_documents.html",
                             reports=reports,
                             unlinked_docs=unlinked_docs,
                             all_reps=all_reps,
                             rep_filter=rep_filter,
                             month_filter=month_filter,
                             doc_type_filter=doc_type_filter)

    @app.route("/tracing/approve/<token>")
    def tracing_approve(token):
        """Public route for reps to approve their tracing reports via tokenized link"""
        try:
            # Validate token
            token_record = normalize_row(
                query_db(
                    """SELECT tat.*, tr.rep_id, tr.month, tr.report_path, r.name as rep_name, r.slug
                       FROM tracing_approval_tokens tat
                       JOIN tracing_reports tr ON tr.id = tat.report_id
                       JOIN reps r ON r.id = tr.rep_id
                       WHERE tat.token = %s AND tat.used = FALSE AND tat.expires_at > %s""",
                    (token, datetime.now()),
                    one=True
                )
            )
            
            if not token_record:
                return render_template("tracing_approve.html", 
                                     success=False, 
                                     message="Invalid or expired approval link. Please contact support."), 400
            
            report_id = token_record["report_id"]
            rep_id = token_record["rep_id"]
            month = token_record["month"]
            rep_slug = token_record["slug"]
            
            # Get report path to determine directory
            report_path_str = token_record["report_path"]
            report_path = Path(report_path_str)
            if not report_path.is_absolute():
                report_path = BASE_DIR / report_path
            # Ensure the directory exists
            report_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Directory where evidence file will be written
            evidence_dir = report_path.parent
            
            # Generate evidence file
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            evidence_filename = f"Tracing_Approval_{month}_{timestamp}.txt"
            evidence_path = evidence_dir / evidence_filename
            
            approval_timestamp = datetime.now()
            ip_address = request.remote_addr if request else None
            user_agent = request.headers.get("User-Agent") if request else None
            
            # Write evidence file
            evidence_content = f"""Tracing Report Approval Evidence
=====================================

Report Month: {month}
Rep Name: {token_record["rep_name"]}
Rep ID: {rep_id}
Approval Timestamp: {approval_timestamp.isoformat()}
IP Address: {ip_address or "N/A"}
User Agent: {user_agent or "N/A"}
Token: {token[:16]}...{token[-8:]} (truncated for security)

Approval Status: APPROVED

This file serves as evidence that the rep has approved their monthly tracing report via the tokenized approval link.

"""
            
            evidence_dir.mkdir(parents=True, exist_ok=True)
            with open(evidence_path, "w", encoding="utf-8") as f:
                f.write(evidence_content)
            
            # Calculate relative path for database storage
            rel_evidence_path = str(evidence_path.relative_to(BASE_DIR))
            
            # Update token as used
            execute_db(
                """UPDATE tracing_approval_tokens 
                   SET used = TRUE, used_at = %s, ip_address = %s, user_agent = %s
                   WHERE token = %s""",
                (approval_timestamp, ip_address, user_agent, token)
            )
            
            # Update tracing_reports record
            execute_db(
                """UPDATE tracing_reports 
                   SET status = 'approved', 
                       approval_uploaded_at = %s, 
                       approval_file_path = %s
                   WHERE id = %s""",
                (approval_timestamp, rel_evidence_path, report_id)
            )
            
            # Record history event
            try:
                execute_db(
                    """INSERT INTO tracing_report_history (report_id, event_type, event_at, performed_by, file_path, note)
                       VALUES (%s, 'approved', %s, 'rep', %s, 'Approval recorded via token link')""",
                    (report_id, approval_timestamp, rel_evidence_path)
                )
            except Exception:
                pass  # History recording is optional
            
            return render_template("tracing_approve.html", 
                                 success=True, 
                                 message=f"Thank you! Your approval for the {month} tracing report has been recorded.",
                                 rep_name=token_record["rep_name"],
                                 month=month)
        
        except Exception as e:
            import traceback
            traceback.print_exc()
            return render_template("tracing_approve.html", 
                                 success=False, 
                                 message=f"An error occurred processing your approval: {str(e)}"), 500
    
    # Email functionality removed - use CSV downloads instead
    # @app.route("/admin/tracing/<int:report_id>/resend", methods=["POST"])
    # def admin_tracing_resend(report_id):
    #     ... removed ...
    
    @app.route("/admin/tracing/<int:report_id>/upload-approval", methods=["POST"])
    def admin_tracing_upload_approval(report_id):
        """Admin uploads historical approval file (.eml, PDF, DOCX) for a tracing report"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        report = normalize_row(query_db(
            """SELECT tr.*, r.slug, r.name as rep_name
               FROM tracing_reports tr
               JOIN reps r ON r.id = tr.rep_id
               WHERE tr.id = %s""",
            (report_id,),
            one=True
        ))
        
        if not report:
            flash("Tracing report not found", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        file = request.files.get("approval_file")
        if not file or not file.filename:
            flash("No file selected", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        # Validate file upload (security hardening)
        is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions=app.config["ALLOWED_TRACING_EXTENSIONS"])
        if not is_valid:
            flash(error_msg or "Invalid file", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        slug = report.get("slug")
        month = report.get("month")
        if not slug or not month:
            flash("Invalid report data", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        # Normalize slug and validate month format
        slug_normalized = slug.lower().strip()
        month_normalized = month.strip()
        
        import re
        if not re.match(r'^\d{4}-\d{2}$', month_normalized):
            flash(f"Invalid month format: {month_normalized} (expected YYYY-MM)", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        # Save file with path traversal prevention
        tracing_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")) / slug_normalized / month_normalized
        tracing_dir.mkdir(parents=True, exist_ok=True)
        
        # Ensure slug and month are safe (no path traversal)
        if ".." in str(slug) or ".." in str(month) or "/" in str(slug) or "\\" in str(slug):
            flash("Invalid path", "danger")
            return redirect(url_for("admin_tracing_overview"))
        
        # Generate filename with timestamp for .eml files, preserve original for others
        ext = Path(secure_name).suffix.lower()
        if ext == ".eml":
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            saved_filename = f"Tracing_Approval_{month}_{timestamp}{ext}"
        else:
            saved_filename = f"Tracing_Approval_{month}{ext}"
        
        # Ensure path is safe
        safe_path = ensure_safe_path(tracing_dir, saved_filename)
        file.save(safe_path)
        
        # Update database record
        rel_path = str((tracing_dir / saved_filename).relative_to(BASE_DIR))
        execute_db(
            "UPDATE tracing_reports SET approval_file_path = %s, approval_uploaded_at = %s, status = 'approved' WHERE id = %s",
            (rel_path, datetime.now(), report_id)
        )
        
        flash(f"Approval file uploaded successfully for {report.get('rep_name', 'rep')} - {month}", "success")
        return redirect(url_for("admin_tracing_overview"))
    
    @app.route("/admin/tracing/<int:report_id>/download-csv")
    def admin_tracing_download_csv(report_id):
        """Download tracing report CSV by report_id (simplified route for CSV downloads)"""
        return admin_tracing_download(report_id, "report")
    
    @app.route("/admin/tracing/download/<int:report_id>/<file_type>")
    def admin_tracing_download(report_id, file_type):
        """Download tracing report or approval file by report_id"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        try:
            # Get report from database
            report = normalize_row(query_db(
                """SELECT tr.*, r.slug
                   FROM tracing_reports tr
                   JOIN reps r ON r.id = tr.rep_id
                   WHERE tr.id = %s""",
                (report_id,),
                one=True
            ))
            
            if not report:
                abort(404, "Tracing report not found")
            
            slug = report.get("slug")
            month = report.get("month")
            
            if not slug or not month:
                abort(404, "Tracing report missing required fields (slug or month)")
            
            # Normalize slug and validate month format
            slug_normalized = slug.lower().strip()
            month_normalized = month.strip()
            
            import re
            if not re.match(r'^\d{4}-\d{2}$', month_normalized):
                abort(400, f"Invalid month format: {month_normalized} (expected YYYY-MM)")
            
            # Determine which file to download
            if file_type == "report":
                # Prefer DB storage (persistent across redeploys)
                csv_content = report.get("report_csv_content")
                if csv_content:
                    # Serve from DB storage
                    filename = f"Tracing_Report_{month_normalized}.csv"
                    response = make_response(csv_content)
                    response.headers['Content-Type'] = 'text/csv; charset=utf-8'
                    response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
                    return response
                
                # Fallback to filesystem (backward compatibility)
                tracing_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")) / slug_normalized / month_normalized
                file_path = None
                report_path_str = report.get("report_path", "")
                if report_path_str:
                    report_path = Path(report_path_str)
                    if not report_path.is_absolute():
                        report_path = BASE_DIR / report_path
                    if report_path.exists():
                        file_path = report_path
                    elif tracing_dir.exists():
                        filename = Path(report_path_str).name
                        candidate = tracing_dir / filename
                        if candidate.exists():
                            file_path = candidate
                if not file_path and tracing_dir.exists():
                    files = list(tracing_dir.glob("Tracing_Report_*"))
                    if files:
                        file_path = files[0]
                
                if file_path and file_path.exists():
                    return send_file(str(file_path), as_attachment=True, download_name=file_path.name)
                else:
                    print(f"[TRACING DOWNLOAD] report_id={report_id} file_type={file_type} reason='file_not_found_no_db'")
                    return render_template_string("""
                    <!DOCTYPE html>
                    <html>
                    <head><title>File Not Found</title></head>
                    <body style="font-family: Arial, sans-serif; padding: 40px; text-align: center;">
                        <h1>File Not Found</h1>
                        <p>The tracing report file could not be found for {{ slug }}/{{ month }}.</p>
                        <p><small>The file may have been deleted or moved. Please regenerate the report.</small></p>
                        <p><a href="{{ url_for('admin_tracing_overview') }}">← Back to Tracing Overview</a></p>
                    </body>
                    </html>
                    """, slug=slug_normalized, month=month_normalized), 404
            
            elif file_type == "approval" or file_type == "evidence":
                file_path = None
                approval_path_str = report.get("approval_file_path", "")
                if approval_path_str:
                    approval_path = Path(approval_path_str)
                    if not approval_path.is_absolute():
                        approval_path = BASE_DIR / approval_path
                    if approval_path.exists():
                        file_path = approval_path
                    else:
                        # Try to find file in tracing directory
                        filename = Path(approval_path_str).name
                        candidate = tracing_dir / filename
                        if candidate.exists():
                            file_path = candidate
                if not file_path:
                    # Fallback: search for any file starting with Tracing_Approval
                    files = list(tracing_dir.glob("Tracing_Approval_*"))
                    if files:
                        file_path = files[0]
                    else:
                        print(f"[TRACING DOWNLOAD] report_id={report_id} file_type={file_type} path='{approval_path_str}' resolved='{tracing_dir}' reason='approval_file_not_found'")
                        # Return custom 404 page instead of abort
                        return render_template_string("""
                        <!DOCTYPE html>
                        <html>
                        <head><title>File Not Found</title></head>
                        <body style="font-family: Arial, sans-serif; padding: 40px; text-align: center;">
                            <h1>File Not Found</h1>
                            <p>The approval file could not be found for {{ slug }}/{{ month }}.</p>
                            <p><small>The file may have been deleted or moved. Please contact an administrator.</small></p>
                            <p><a href="{{ url_for('admin_tracing_overview') }}">← Back to Tracing Overview</a></p>
                        </body>
                        </html>
                        """, slug=slug, month=month), 404
                return send_file(file_path, as_attachment=True)
            
            else:
                # Return custom 404 page for invalid file type
                return render_template_string("""
                <!DOCTYPE html>
                <html>
                <head><title>Invalid File Type</title></head>
                <body style="font-family: Arial, sans-serif; padding: 40px; text-align: center;">
                    <h1>Invalid File Type</h1>
                    <p>The file type "{{ file_type }}" is not valid. Valid types are: "report" or "approval".</p>
                    <p><a href="{{ url_for('admin_tracing_overview') }}">← Back to Tracing Overview</a></p>
                </body>
                </html>
                """, file_type=file_type), 404
                
        except Exception as e:
            # Only log unexpected errors, not 404s (which are already handled above)
            if hasattr(e, 'code') and e.code == 404:
                raise
            print(f"[TRACING DOWNLOAD] report_id={report_id} file_type={file_type} error='{str(e)}'")
            # Return custom error page
            return render_template_string("""
            <!DOCTYPE html>
            <html>
            <head><title>Download Error</title></head>
            <body style="font-family: Arial, sans-serif; padding: 40px; text-align: center;">
                <h1>Download Error</h1>
                <p>An error occurred while downloading the file. Please try again or contact an administrator.</p>
                <p><a href="{{ url_for('admin_tracing_overview') }}">← Back to Tracing Overview</a></p>
            </body>
            </html>
            """), 500
    
    @app.route("/admin/rep/<slug>/tracing/<month>/download/<filename>")
    def admin_download_tracing_file(slug, month, filename):
        """Download tracing report or approval file (legacy route)"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        try:
            # Sanitize filename (prevent directory traversal)
            filename = Path(filename).name
            
            tracing_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")) / slug / month
            if not tracing_dir.exists():
                print(f"Tracing directory not found: {tracing_dir}")
                abort(404)
            
            file_path = tracing_dir / filename
            if not file_path.exists():
                # Try case-insensitive search
                files = list(tracing_dir.glob(f"*{filename}*"))
                if files:
                    file_path = files[0]
                    filename = file_path.name
                else:
                    print(f"File not found: {file_path}")
                    abort(404)
            
            return send_from_directory(tracing_dir, filename, as_attachment=True)
        except Exception as e:
            print(f"Error downloading tracing file: {e}")
            import traceback
            traceback.print_exc()
            abort(404)
    
    # ------------------------------------------------------------------ Rep Tracing Management (Admin Only)
    
    @app.route("/admin/rep/<slug>/tracing")
    def admin_rep_tracing(slug):
        """View tracing reports and approvals for a specific rep"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        rep = get_rep_by_slug(slug)
        if not rep:
            abort(404)
        
        # Get tracing reports from database
        db_reports = query_db(
            """SELECT month, status, approval_uploaded_at, approval_file_path, report_path
               FROM tracing_reports 
               WHERE rep_id = %s 
               ORDER BY month DESC""",
            (rep["id"],)
        ) or []
        db_reports_by_month = {r["month"]: r for r in db_reports}
        
        # Get tracing directory for this rep
        tracing_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")) / slug
        tracing_months = []
        
        if tracing_dir.exists():
            for month_dir in sorted(tracing_dir.iterdir(), reverse=True):
                if month_dir.is_dir():
                    month = month_dir.name
                    files = {"report": None, "approval": None}
                    for f in month_dir.iterdir():
                        if f.is_file():
                            if f.stem.startswith("Tracing_Report"):
                                files["report"] = f.name
                            elif f.stem.startswith("Tracing_Approval"):
                                files["approval"] = f.name
                    
                    # Get database info for this month
                    db_info = db_reports_by_month.get(month, {})
                    
                    # Extract filename from approval_file_path if it exists
                    approval_filename = None
                    if db_info.get("approval_file_path"):
                        approval_filename = Path(db_info["approval_file_path"]).name
                    
                    tracing_months.append({
                        "month": month,
                        "report_file": files["report"],
                        "approval_file": files["approval"] or approval_filename,
                        "status": db_info.get("status", "draft"),
                        "approval_uploaded_at": db_info.get("approval_uploaded_at"),
                        "approval_file_path": db_info.get("approval_file_path")
                    })
        
        return render_template("admin_rep_tracing.html", rep=rep, tracing_months=tracing_months)
    
    @app.route("/admin/rep/<slug>/tracing/<month>/upload", methods=["POST"])
    def admin_upload_tracing_approval(slug, month):
        """Admin uploads tracing approval file (.eml, PDF, DOCX) for a rep"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        rep = get_rep_by_slug(slug)
        if not rep:
            abort(404)
        
        file = request.files.get("approval_file")
        if not file or not file.filename:
            flash("No file selected", "danger")
            return redirect(url_for("admin_rep_tracing", slug=slug))
        
        # Validate file upload (security hardening)
        is_valid, error_msg, secure_name = validate_upload_file(file, allowed_extensions=app.config["ALLOWED_TRACING_EXTENSIONS"])
        if not is_valid:
            flash(error_msg or "Invalid file", "danger")
            return redirect(url_for("admin_rep_tracing", slug=slug))
        
        # Ensure slug and month are safe (no path traversal)
        if ".." in str(slug) or ".." in str(month) or "/" in str(slug) or "\\" in str(slug):
            flash("Invalid path", "danger")
            return redirect(url_for("admin_rep_tracing", slug=slug))
        
        tracing_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")) / slug / month
        tracing_dir.mkdir(parents=True, exist_ok=True)
        
        # Generate filename with timestamp for .eml files, preserve original for others
        ext = Path(secure_name).suffix.lower()
        if ext == ".eml":
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            saved_filename = f"Tracing_Approval_{month}_{timestamp}{ext}"
        else:
            saved_filename = f"Tracing_Approval_{month}{ext}"
        
        # Ensure path is safe
        safe_path = ensure_safe_path(tracing_dir, saved_filename)
        file.save(safe_path)
        
        # Update database record
        report = query_db(
            "SELECT id FROM tracing_reports WHERE rep_id = %s AND month = %s",
            (rep["id"], month),
            one=True
        )
        if report:
            rel_path = str((tracing_dir / saved_filename).relative_to(BASE_DIR))
            execute_db(
                "UPDATE tracing_reports SET approval_file_path = %s, approval_uploaded_at = %s, status = 'approved' WHERE id = %s",
                (rel_path, datetime.now(), report["id"])
            )
        
        flash(f"Approval file uploaded for {month}", "success")
        return redirect(url_for("admin_rep_tracing", slug=slug))

    # ------------------------------------------------------------------ Rep Routes

    @app.route("/rep/login", methods=["GET", "POST"])
    @app.route("/rep/<slug>/login", methods=["GET", "POST"])
    def rep_login(slug=None):
        identifier_default = (slug or request.args.get("slug") or "").lower()
        next_url = request.args.get("next")
        rep_prefill = get_rep_by_slug(identifier_default) if identifier_default else None

        if request.method == "POST":
            identifier = request.form.get("identifier") or identifier_default
            password = request.form.get("password", "")
            rep_row = find_rep_by_identifier(identifier)
            rep = normalize_row(rep_row)
            if rep and check_password_hash(rep.get("password_hash", ""), password):
                session.clear()
                session["rep_slug"] = rep.get("slug")
                session["rep_id"] = rep.get("id")
                session.permanent = True
                flash("Logged in.", "success")
                return redirect(next_url or request.form.get("next") or url_for("rep_dashboard", slug=rep.get("slug")))
            flash("Invalid email/slug or password.", "danger")

        return render_template("rep_login.html", rep=rep_prefill, identifier=identifier_default, next_url=next_url)

    @app.route("/rep/forgot-password", methods=["GET", "POST"])
    @app.route("/rep/<slug>/forgot-password", methods=["GET", "POST"])
    def rep_forgot_password(slug=None):
        identifier_default = (slug or request.args.get("slug") or "").lower()
        if request.method == "POST":
            email_in = (request.form.get("email") or "").strip().lower()
            rep_row = find_rep_by_identifier(email_in)
            rep = normalize_row(rep_row)
            if rep and rep.get("email"):
                # Invalidate any prior unused tokens for this rep
                execute_db("UPDATE rep_password_resets SET used = TRUE WHERE rep_id = %s AND used = FALSE", (rep.get("id"),))

                token = secrets.token_urlsafe(32)
                expires_at = datetime.utcnow() + timedelta(hours=1)
                execute_db(
                    "INSERT INTO rep_password_resets (rep_id, token, expires_at, used) VALUES (%s, %s, %s, FALSE)",
                    (rep.get("id"), token, expires_at),
                )

                reset_url = url_for("rep_reset_password", token=token, _external=True)
                # region agent log
                _agent_log(
                    "EMAIL_SEND_MISSING_FN",
                    "Proto1.py:rep_forgot_password",
                    "about to call send_email",
                    {
                        "send_email_defined": bool("send_email" in globals()),
                        "has_smtp_server": bool(app.config.get("SMTP_SERVER")),
                        "has_rep_email": True,
                    },
                )
                # endregion
                try:
                    ok, msg = send_email(
                        rep.get("email"),
                        "Reset your Rep Portal password",
                        f"Hi {rep.get('name')},\n\n"
                        "We received a request to reset your Rep Portal password. Use the secure link below to set a new password. "
                        "This link expires in 60 minutes.\n\n"
                        f"{reset_url}\n\n"
                        "If you did not request this, you can ignore this email.",
                    )
                except Exception as e:
                    # region agent log
                    _agent_log(
                        "EMAIL_SEND_MISSING_FN",
                        "Proto1.py:rep_forgot_password",
                        "send_email raised",
                        {"exc_type": type(e).__name__},
                    )
                    # endregion
                    ok, msg = False, f"exception:{type(e).__name__}"
                if not ok:
                    # Security: do not reveal whether email exists; just log for ops.
                    print(f"[PASSWORD_RESET_ERROR] Failed to send reset email: {msg}")
            flash("If that email is in our system, we've sent reset instructions.", "info")
            return redirect(url_for("rep_login", slug=identifier_default))
        return render_template("rep_forgot_password.html", slug=identifier_default)

    @app.route("/rep/reset-password", methods=["GET", "POST"])
    def rep_reset_password():
        token = request.args.get("token") or request.form.get("token")
        if not token:
            flash("Reset token missing.", "danger")
            return redirect(url_for("rep_login"))

        reset_row = normalize_row(query_db("SELECT * FROM rep_password_resets WHERE token = %s", (token,), one=True))
        now = datetime.utcnow()
        if not reset_row or reset_row.get("used") or (reset_row.get("expires_at") and reset_row.get("expires_at") < now):
            flash("Reset link is invalid or has expired.", "danger")
            return redirect(url_for("rep_login"))

        if request.method == "POST":
            pw1 = request.form.get("password", "")
            pw2 = request.form.get("password_confirm", "")
            if not pw1 or len(pw1) < 8:
                flash("Password must be at least 8 characters.", "danger")
                return redirect(url_for("rep_reset_password", token=token))
            if pw1 != pw2:
                flash("Passwords do not match.", "danger")
                return redirect(url_for("rep_reset_password", token=token))

            rep_id = reset_row.get("rep_id")
            pw_hash = generate_password_hash(pw1)
            execute_db("UPDATE reps SET password_hash = %s WHERE id = %s", (pw_hash, rep_id))
            execute_db("UPDATE rep_password_resets SET used = TRUE, used_at = %s WHERE token = %s", (datetime.utcnow(), token))

            rep_row = normalize_row(query_db("SELECT slug FROM reps WHERE id = %s", (rep_id,), one=True))
            slug_target = rep_row.get("slug") if isinstance(rep_row, dict) else None
            flash("Password updated. Please log in.", "success")
            return redirect(url_for("rep_login", slug=slug_target))

        return render_template("rep_reset_password.html", token=token)

    @app.route("/rep/<slug>/logout")
    def rep_logout(slug):
        """Log out the current rep and clear all session data"""
        if get_logged_in_rep_slug() == slug.lower():
            session.pop("rep_slug", None)
            session.pop("rep_id", None)
        flash("Logged out.", "info")
        return redirect(url_for("index"))

    @app.route("/rep/<slug>/documents")
    @require_rep
    def rep_documents(slug):
        rep = normalize_row(g.get("current_rep") or get_rep_by_slug(slug.lower()))
        if rep is None:
            abort(404)
        ensure_rep_documents_table()
        docs = query_db(
            """
            SELECT * FROM rep_documents
            WHERE rep_id = %s
            ORDER BY uploaded_at DESC
            """,
            (rep["id"],),
        ) or []
        doc_type_labels = {
            "receiving_inspection": "Receiving Inspection Form",
            "device_distribution": "Device Distribution Form",
        }
        return render_template(
            "rep_documents.html",
            rep=rep,
            documents=docs,
            doc_type_labels=doc_type_labels,
            uploader_role="admin" if is_admin() else "rep",
        )

    @app.route("/files/rep-docs/<slug>/<int:doc_id>")
    @require_rep
    def download_rep_document(slug, doc_id):
        rep = normalize_row(g.get("current_rep") or get_rep_by_slug(slug.lower()))
        if rep is None:
            abort(404)
        ensure_rep_documents_table()
        doc = query_db(
            "SELECT * FROM rep_documents WHERE id = %s AND rep_id = %s",
            (doc_id, rep["id"]),
            one=True,
        )
        doc = normalize_row(doc)
        if not doc:
            abort(404)
        file_path = BASE_DIR / doc["stored_filename"]
        if not file_path.exists():
            abort(404)
        return send_from_directory(file_path.parent, file_path.name, as_attachment=True)

    @app.route("/viewer")
    def viewer():
        if not (is_admin() or get_logged_in_rep_slug()):
            flash("Please log in.", "warning")
            return redirect(url_for("index"))
        file_url = request.args.get("url")
        if not file_url:
            flash("No doc", "danger")
            return redirect(request.referrer or url_for("index"))
        return render_template("view_docx.html", file_url=file_url)

    def get_rep_inventory(rep_id: int) -> dict:
        """
        Get inventory summary for a rep: received, distributed, and current inventory.
        
        Returns:
            dict with keys:
                - total_received: sum of units_received from shipment_line_items
                - total_distributed: sum of quantities from device_distribution_records
                - current_inventory: total_received - total_distributed
                - recent_received: list of recent received entries (date, lot, qty, links)
        """
        # Calculate total received
        received_sum = query_db("""
            SELECT COALESCE(SUM(sli.units_received), 0) AS total
            FROM shipment_line_items sli
            JOIN devices_received dr ON sli.shipment_id = dr.id
            WHERE dr.rep_id = %s
        """, (rep_id,), one=True)
        total_received = int(received_sum.get("total", 0) if received_sum else 0)
        
        # Calculate total distributed (from device_distribution_records fields_json)
        # This is approximate - we sum "Quantity" or "Qty" fields from JSON
        distributed_records = query_db("""
            SELECT fields_json
            FROM device_distribution_records
            WHERE rep_id = %s
        """, (rep_id,)) or []
        
        total_distributed = 0
        for rec in distributed_records:
            try:
                fields = json.loads(rec.get("fields_json", "{}") or "{}")
                # Try common quantity field names
                qty = fields.get("Quantity") or fields.get("Qty") or fields.get("quantity") or fields.get("qty") or 0
                if isinstance(qty, (int, float)):
                    total_distributed += int(qty)
                elif isinstance(qty, str):
                    try:
                        total_distributed += int(float(qty))
                    except (ValueError, TypeError):
                        pass
            except (json.JSONDecodeError, TypeError):
                pass
        
        current_inventory = total_received - total_distributed
        
        # Get recent received entries with line items
        recent_received = query_db("""
            SELECT 
                dr.created_at,
                sli.lot_number,
                sli.units_received,
                dr.id AS shipment_id,
                dr.packing_list_filename,
                dr.recv_inspection_filename
            FROM shipment_line_items sli
            JOIN devices_received dr ON sli.shipment_id = dr.id
            WHERE dr.rep_id = %s
            ORDER BY dr.created_at DESC
            LIMIT 10
        """, (rep_id,)) or []
        
        # If no line items but shipments exist, get shipments directly
        if not recent_received:
            shipments_only = query_db("""
                SELECT 
                    dr.created_at,
                    dr.id AS shipment_id,
                    dr.packing_list_filename,
                    dr.recv_inspection_filename
                FROM devices_received dr
                WHERE dr.rep_id = %s
                ORDER BY dr.created_at DESC
                LIMIT 10
            """, (rep_id,)) or []
            for s in shipments_only:
                s["lot_number"] = None
                s["units_received"] = None
            recent_received = shipments_only
        
        return {
            "total_received": total_received,
            "total_distributed": total_distributed,
            "current_inventory": current_inventory,
            "recent_received": recent_received
        }

    @app.route("/rep/<slug>")
    @require_rep
    def rep_dashboard(slug):
        try:
            rep_obj = g.get("current_rep") or get_rep_by_slug(slug.lower())
            rep = normalize_row(rep_obj)
            if rep is not None and not isinstance(rep, dict):
                rep = normalize_row(query_db("SELECT * FROM reps WHERE id = %s", (rep,), one=True))
            if rep is None and rep_obj:
                # Last attempt: look up by slug
                rep = normalize_row(get_rep_by_slug(slug.lower()))
            if rep is None:
                abort(404)
            
            rep_id = rep.get("id")
            if not rep_id:
                abort(404)
            
            # Get all distributions with error handling
            try:
                distributions_raw = query_db("""
                    SELECT DISTINCT
                        d.id,
                        d.rep_id,
                        d.order_number,
                        d.ship_date,
                        d.created_at,
                        d.source,
                        d.tracking_number,
                        d.distribution_number,
                        dr.title AS shipment_title
                    FROM devices_distributed d
                    LEFT JOIN devices_received dr ON d.shipment_id = dr.id
                    WHERE d.rep_id = %s
                    ORDER BY d.ship_date DESC NULLS LAST, d.created_at DESC
                """, (rep_id,)) or []
            except Exception as e:
                print(f"Error querying distributions: {e}")
                import traceback
                traceback.print_exc()
                distributions_raw = []
            
            distributions = []
            for dist in distributions_raw:
                try:
                    dist_id = dist.get("id")
                    if not dist_id:
                        continue
                    
                    # Get distribution records with error handling
                    try:
                        dist_records = query_db("""
                            SELECT ddr.*, ddr.fields_json, ddr.original_filename, ddr.stored_filename, ddr.uploaded_at
                            FROM device_distribution_records ddr
                            WHERE ddr.dist_id = %s
                            ORDER BY ddr.id
                        """, (dist_id,)) or []
                    except Exception as e:
                        print(f"Error querying distribution records for dist_id {dist_id}: {e}")
                        dist_records = []
                    
                    parsed_records = []
                    evidence_files = []
                    
                    for rec in dist_records:
                        try:
                            fields_json_str = rec.get("fields_json") or "{}"
                            if isinstance(fields_json_str, str):
                                fields = json.loads(fields_json_str)
                            else:
                                fields = fields_json_str if isinstance(fields_json_str, dict) else {}
                        except Exception as e:
                            print(f"Error parsing fields_json for record {rec.get('id')}: {e}")
                            fields = {}
                        
                        # Check if this is a SKU record
                        has_sku = bool(fields.get("SKU") or fields.get("sku"))
                        has_lot = bool(fields.get("Lot") or fields.get("lot"))
                        has_quantity = bool(fields.get("Quantity") or fields.get("quantity"))
                        
                        if has_sku or (has_lot and has_quantity):
                            try:
                                quantity_val = fields.get("Quantity") or fields.get("quantity", 0) or 0
                                quantity_int = int(float(str(quantity_val)))
                            except Exception as e:
                                import traceback
                                print(f"[ERROR] rep_dashboard (quantity parse): {e}")
                                traceback.print_exc()
                                quantity_int = 0
                            
                            parsed_records.append({
                                "id": rec.get("id"),
                                "sku": fields.get("SKU") or fields.get("sku", ""),
                                "lot": fields.get("Lot") or fields.get("lot", ""),
                                "quantity": quantity_int,
                                "facility_name": fields.get("Facility Name") or fields.get("facility_name", ""),
                                "city": fields.get("City") or fields.get("city", ""),
                                "state": fields.get("State") or fields.get("state", ""),
                            })
                        else:
                            # Evidence file
                            evidence_files.append({
                                "id": rec.get("id"),
                                "original_filename": rec.get("original_filename") or "",
                                "stored_filename": rec.get("stored_filename") or "",
                                "uploaded_at": rec.get("uploaded_at") or "",
                            })
                    
                    # Build distribution object
                    dist_obj = dict(dist)
                    dist_obj["records"] = parsed_records
                    dist_obj["evidence_files"] = evidence_files
                    dist_obj["total_units"] = sum(r.get("quantity", 0) for r in parsed_records)
                    
                    # Get facility name
                    facility_name = ""
                    if parsed_records:
                        facility_name = parsed_records[0].get("facility_name", "")
                    elif evidence_files:
                        # Try to get from evidence file fields_json if available
                        try:
                            first_evidence = dist_records[0] if dist_records else None
                            if first_evidence:
                                ev_fields = json.loads(first_evidence.get("fields_json") or "{}")
                                facility_name = ev_fields.get("Facility Name") or ""
                        except Exception as e:
                            import traceback
                            print(f"[ERROR] rep_dashboard (facility_name parse): {e}")
                            traceback.print_exc()
                    
                    dist_obj["facility_name"] = facility_name
                    distributions.append(dist_obj)
                    
                except Exception as e:
                    print(f"Error processing distribution {dist.get('id')}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            # Sort distributions
            try:
                distributions.sort(key=lambda d: (
                    d.get("ship_date") or d.get("created_at") or "",
                ), reverse=True)
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (distributions sort): {e}")
                traceback.print_exc()
            
            # Get other data with error handling
            try:
                forms = query_db("SELECT * FROM form_types WHERE active = 1 ORDER BY name") or []
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (forms query): {e}")
                traceback.print_exc()
                forms = []
            
            try:
                submissions = query_db("SELECT s.*, f.name AS form_name FROM submissions s JOIN form_types f ON s.form_type_id = f.id WHERE s.rep_id = %s ORDER BY s.uploaded_at DESC", (rep_id,)) or []
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (submissions query): {e}")
                traceback.print_exc()
                submissions = []
            
            try:
                training_docs = query_db("SELECT * FROM training_docs WHERE active = 1 ORDER BY title") or []
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (training_docs query): {e}")
                traceback.print_exc()
                training_docs = []
            
            try:
                # Get received shipments with line items (lot, units)
                received_raw = query_db("SELECT * FROM devices_received WHERE rep_id = %s ORDER BY created_at DESC", (rep_id,)) or []
                
                # Enrich with line item data
                received = []
                for r in received_raw:
                    line_items = query_db("""
                        SELECT lot_number, units_received
                        FROM shipment_line_items
                        WHERE shipment_id = %s
                        ORDER BY created_at DESC
                        LIMIT 1
                    """, (r["id"],)) or []
                    
                    rec_dict = dict(r)
                    if line_items:
                        rec_dict["lot_number"] = line_items[0].get("lot_number")
                        rec_dict["units_received"] = line_items[0].get("units_received")
                    else:
                        rec_dict["lot_number"] = None
                        rec_dict["units_received"] = None
                    received.append(rec_dict)
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (received query): {e}")
                traceback.print_exc()
                received = []
            
            try:
                master_log = query_db("SELECT * FROM rep_master_logs WHERE rep_id = %s AND log_type = 'distribution' ORDER BY updated_at DESC LIMIT 1", (rep_id,), one=True)
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (master_log query): {e}")
                traceback.print_exc()
                master_log = None
            
            try:
                customer_records = query_db("SELECT * FROM new_customer_records WHERE rep_id = %s ORDER BY uploaded_at DESC", (rep_id,)) or []
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (customer_records query): {e}")
                traceback.print_exc()
                customer_records = []
            
            customers_list = []
            for c in customer_records:
                try:
                    f = json.loads(c.get('fields_json', '{}') or '{}')
                except Exception as e:
                    import traceback
                    print(f"[ERROR] rep_dashboard (customer fields_json parse): {e}")
                    traceback.print_exc()
                    f = {}
                try:
                    # Convert uploaded_at datetime to string if needed
                    uploaded_at = c.get('uploaded_at')
                    if uploaded_at:
                        if isinstance(uploaded_at, datetime):
                            date_str = uploaded_at.isoformat()[:10]
                        elif isinstance(uploaded_at, str):
                            date_str = uploaded_at[:10] if len(uploaded_at) >= 10 else uploaded_at
                        else:
                            date_str = str(uploaded_at)[:10]
                    else:
                        date_str = ""
                    
                    customers_list.append({
                        "name": f.get("Facility Name", "Unknown"),
                        "city": f.get("City", ""),
                        "date": date_str,
                        "files": [{"name": "New Customer Form", "url": url_for('download_customer_file', slug=slug, folder=f"cust_{c.get('dist_id')}", filename=c.get('original_filename', '')) if c.get('dist_id') else "#"}]
                    })
                except Exception as e:
                    import traceback
                    print(f"[ERROR] rep_dashboard (customer record processing): {e}")
                    traceback.print_exc()

            templates_fs = []
            policies_fs = []
            complaint_docs = []
            try:
                form_templates_dir = Path(app.config.get("FORM_TEMPLATES_DIR", BASE_DIR / "form_templates"))
                if form_templates_dir.exists():
                    for f in sorted(form_templates_dir.iterdir()):
                        if f.is_file():
                            templates_fs.append({"filename": f.name, "label": f.stem.replace("_", " ")})
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (form_templates_dir): {e}")
                traceback.print_exc()
            
            try:
                quality_docs_dir = Path(app.config.get("QUALITY_DOCS_DIR", BASE_DIR / "quality_docs"))
                if quality_docs_dir.exists():
                    for f in sorted(quality_docs_dir.iterdir()):
                        if f.is_file():
                            policies_fs.append({"filename": f.name, "label": f.stem.replace("_", " ")})
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (quality_docs_dir): {e}")
                traceback.print_exc()
            
            try:
                product_complaints_dir = Path(app.config.get("PRODUCT_COMPLAINTS_DIR", BASE_DIR / "product_complaints"))
                if product_complaints_dir.exists():
                    for f in sorted(product_complaints_dir.iterdir()):
                        if f.is_file():
                            complaint_docs.append({"filename": f.name, "label": f.stem.replace("_", " ")})
            except Exception as e:
                import traceback
                print(f"[ERROR] rep_dashboard (product_complaints_dir): {e}")
                traceback.print_exc()

            # Calculate tracing months from ship_date (not created_at) to include manual entries correctly
            tracing_months = []
            months_seen = set()
            for d in distributions:
                try:
                    # Use ship_date if available, fallback to created_at
                    date_str = d.get("ship_date") or d.get("created_at") or ""
                    if not date_str or len(date_str) < 7:
                        continue
                    month = date_str[:7]  # Extract YYYY-MM
                    if month in months_seen:
                        continue
                    months_seen.add(month)
                    try:
                        label = datetime(int(month[:4]), int(month[5:7]), 1).strftime("%B %Y")
                    except Exception as e:
                        import traceback
                        print(f"[ERROR] rep_dashboard (month label format): {e}")
                        traceback.print_exc()
                        label = month
                    
                    # Check for tracing report status from database
                    try:
                        tracing_report = query_db("""
                            SELECT status, report_path, email_sent_at, approval_uploaded_at, approval_file_path
                            FROM tracing_reports
                            WHERE rep_id = %s AND month = %s
                        """, (rep_id, month), one=True)
                    except Exception as e:
                        import traceback
                        print(f"[ERROR] rep_dashboard (tracing_report query): {e}")
                        traceback.print_exc()
                        tracing_report = None
                    
                    month_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")) / slug / month
                    tracing_file = None
                    approval_file = None
                    status = "draft"
                    
                    if tracing_report:
                        status = tracing_report.get("status") or "draft"
                        report_path = tracing_report.get("report_path")
                        if report_path:
                            try:
                                path_obj = Path(report_path)
                                if path_obj.is_absolute():
                                    tracing_file = path_obj.name
                                else:
                                    tracing_file = Path(report_path).name
                            except Exception as e:
                                import traceback
                                print(f"[ERROR] rep_dashboard (tracing_file path): {e}")
                                traceback.print_exc()
                        approval_file_path = tracing_report.get("approval_file_path")
                        if approval_file_path:
                            try:
                                path_obj = Path(approval_file_path)
                                if path_obj.is_absolute():
                                    approval_file = path_obj.name
                                else:
                                    approval_file = Path(approval_file_path).name
                            except Exception as e:
                                import traceback
                                print(f"[ERROR] rep_dashboard (approval_file path): {e}")
                                traceback.print_exc()
                    elif month_dir.exists():
                        # Fallback to filesystem check
                        try:
                            for f in month_dir.iterdir():
                                if f.is_file():
                                    if f.stem.startswith(f"Tracing_Report_{month}"):
                                        tracing_file = f.name
                                    elif f.stem.startswith(f"Tracing_Approval_{month}"):
                                        approval_file = f.name
                        except Exception as e:
                            import traceback
                            print(f"[ERROR] rep_dashboard (filesystem tracing check): {e}")
                            traceback.print_exc()
                    
                    # Convert datetime objects to strings for template
                    email_sent_at_str = None
                    approval_uploaded_at_str = None
                    if tracing_report:
                        email_sent_at = tracing_report.get("email_sent_at")
                        if email_sent_at:
                            if isinstance(email_sent_at, datetime):
                                email_sent_at_str = email_sent_at.isoformat()[:10]
                            elif isinstance(email_sent_at, str):
                                email_sent_at_str = email_sent_at[:10] if len(email_sent_at) >= 10 else email_sent_at
                            else:
                                email_sent_at_str = str(email_sent_at)[:10]
                        
                        approval_uploaded_at = tracing_report.get("approval_uploaded_at")
                        if approval_uploaded_at:
                            if isinstance(approval_uploaded_at, datetime):
                                approval_uploaded_at_str = approval_uploaded_at.isoformat()[:10]
                            elif isinstance(approval_uploaded_at, str):
                                approval_uploaded_at_str = approval_uploaded_at[:10] if len(approval_uploaded_at) >= 10 else approval_uploaded_at
                            else:
                                approval_uploaded_at_str = str(approval_uploaded_at)[:10]
                    
                    tracing_months.append({
                        "month": month,
                        "label": label,
                        "tracing_file": tracing_file,
                        "approval_file": approval_file,
                        "status": status,
                        "email_sent_at": email_sent_at_str,
                        "approval_uploaded_at": approval_uploaded_at_str,
                    })
                except Exception as e:
                    print(f"Error processing tracing month: {e}")
                    continue
            
            tracing_months.sort(key=lambda m: m.get("month", ""), reverse=True)
            
            # Calculate statistics with error handling
            try:
                total_distributions = len(distributions)
                total_units = sum(d.get("total_units", 0) for d in distributions)
                unique_customers = len(set(
                    r.get("facility_name") for d in distributions 
                    for r in d.get("records", []) 
                    if r.get("facility_name")
                ))
                unique_orders = len(set(
                    d.get("order_number") for d in distributions 
                    if d.get("order_number")
                ))
                
                current_month = datetime.now().strftime("%Y-%m")
                current_month_distributions = [
                    d for d in distributions 
                    if (d.get("ship_date") or d.get("created_at") or "")[:7] == current_month
                ]
                current_month_units = sum(d.get("total_units", 0) for d in current_month_distributions)
            except Exception as e:
                print(f"Error calculating statistics: {e}")
                total_distributions = 0
                total_units = 0
                unique_customers = 0
                unique_orders = 0
                current_month_distributions = []
                current_month_units = 0
            
            # Get rep ZIP for initial target facilities load
            rep_zip = rep.get("zip") or ""
            rep_zip_clean = "".join(c for c in rep_zip if c.isdigit())[:5] if rep_zip else ""
            
            # Get inventory summary
            try:
                inventory = get_rep_inventory(rep_id)
            except Exception as e:
                print(f"Error calculating inventory: {e}")
                inventory = {
                    "total_received": 0,
                    "total_distributed": 0,
                    "current_inventory": 0,
                    "recent_received": []
                }
            
            return render_template("rep_dashboard.html", 
                rep=rep, 
                forms=forms or [],
                submissions=submissions or [],
                templates_fs=templates_fs or [],
                policies_fs=policies_fs or [],
                complaint_docs=complaint_docs or [],
                training_docs=training_docs or [],
                received=received or [],
                distributions=distributions,
                master_log=master_log,
                customer_info_entries=customers_list or [],
                tracing_months=tracing_months or [],
                total_distributions=total_distributions,
                total_units=total_units,
                unique_customers=unique_customers,
                rep_zip=rep_zip_clean,
                unique_orders=unique_orders,
                current_month_distributions=len(current_month_distributions),
                current_month_units=current_month_units,
                inventory=inventory
            )
        except Exception as e:
            print(f"CRITICAL ERROR in rep_dashboard: {e}")
            import traceback
            traceback.print_exc()
            # Don't try to render error.html - return simple error response
            error_msg = str(e)
            return f"""<html><head><title>Error</title></head><body>
                <h1>Error Loading Dashboard</h1>
                <p>{error_msg}</p>
                <p><a href="/">Return to Home</a></p>
            </body></html>""", 500

    @app.route("/files/distribution-records/<slug>/<int:record_id>/download", endpoint="download_distribution_record_file")
    @require_rep
    def download_distribution_record_file(slug, record_id):
        """Download evidence file from device_distribution_records"""
        try:
            rep = normalize_row(g.get("current_rep") or get_rep_by_slug(slug.lower()))
            if rep is None:
                abort(404)
            
            rep_id = rep.get("id")
            if not rep_id:
                abort(404)
            
            # Get record and verify rep access
            record = normalize_row(query_db(
                "SELECT * FROM device_distribution_records WHERE id = %s AND rep_id = %s",
                (record_id, rep_id),
                one=True
            ))
            if not record:
                abort(404)
            
            stored_filename = record.get("stored_filename")
            if not stored_filename:
                abort(404)
            
            # Get file path (handle both filesystem and S3)
            if _dist_records_s3_enabled():
                try:
                    bucket = (app.config.get("DIST_RECORDS_S3_BUCKET") or "").strip()
                    if not bucket:
                        abort(404)
                    key = _dist_records_s3_key(stored_filename)
                    client = _dist_records_s3_client()
                    client.head_object(Bucket=bucket, Key=key)  # Fail fast if object missing
                    
                    original = (record.get("original_filename") or "download").strip() or "download"
                    url = client.generate_presigned_url(
                        "get_object",
                        Params={
                            "Bucket": bucket,
                            "Key": key,
                            "ResponseContentDisposition": f'attachment; filename="{original}"',
                        },
                        ExpiresIn=60 * 10,  # 10 minutes
                    )
                    return redirect(url)
                except Exception as e:
                    print(f"Error downloading from S3: {e}")
                    abort(404)
            else:
                # Filesystem download
                try:
                    file_path = _safe_dist_record_path(stored_filename)
                except Exception as e:
                    print(f"Error getting file path: {e}")
                    abort(400)
                if not file_path.exists():
                    abort(404)
                original = (record.get("original_filename") or file_path.name).strip() or file_path.name
                return send_from_directory(file_path.parent, file_path.name, as_attachment=True, download_name=original)
        except Exception as e:
            print(f"Error downloading distribution record file: {e}")
            import traceback
            traceback.print_exc()
            abort(404)

    # ------------------------------------------------------------------ file downloads

    @app.route("/files/templates/<filename>", endpoint="download_template")
    def download_template(filename):
        if not (is_admin() or get_logged_in_rep_slug()):
            abort(403)
        base_dir = Path(app.config.get("FORM_TEMPLATES_DIR", BASE_DIR / "form_templates"))
        if not base_dir.exists():
            abort(404)
        return send_from_directory(base_dir, filename, as_attachment=True)

    @app.route("/files/policies/<filename>", endpoint="download_policy")
    def download_policy(filename):
        if not (is_admin() or get_logged_in_rep_slug()):
            abort(403)
        base_dir = Path(app.config.get("QUALITY_DOCS_DIR", BASE_DIR / "quality_docs"))
        if not base_dir.exists():
            abort(404)
        return send_from_directory(base_dir, filename, as_attachment=True)

    @app.route("/files/training/<filename>", endpoint="download_training")
    def download_training(filename):
        if not (is_admin() or get_logged_in_rep_slug()):
            abort(403)
        base_dir = Path(app.config.get("TRAINING_DOCS_DIR", BASE_DIR / "training_docs"))
        if not base_dir.exists():
            abort(404)
        return send_from_directory(base_dir, filename, as_attachment=True)

    @app.route("/files/complaints/<filename>", endpoint="download_complaint")
    def download_complaint(filename):
        if not (is_admin() or get_logged_in_rep_slug()):
            abort(403)
        base_dir = Path(app.config.get("PRODUCT_COMPLAINTS_DIR", BASE_DIR / "product_complaints"))
        if not base_dir.exists():
            abort(404)
        return send_from_directory(base_dir, filename, as_attachment=True)

    @app.route("/files/received/<int:shipment_id>/<kind>", endpoint="download_received_file")
    def download_received_file(shipment_id, kind):
        """Download received shipment files (admin or rep access)"""
        # Check auth: admin or rep who owns the shipment
        if is_admin():
            shipment = normalize_row(query_db("SELECT * FROM devices_received WHERE id = %s", (shipment_id,), one=True))
        else:
            rep_slug = get_logged_in_rep_slug()
            if not rep_slug:
                abort(403)
            rep = normalize_row(get_rep_by_slug(rep_slug))
            if not rep:
                abort(403)
            shipment = normalize_row(query_db("SELECT * FROM devices_received WHERE id = %s AND rep_id = %s", (shipment_id, rep["id"]), one=True))
        
        if not shipment:
            abort(404)
        
        # Map kind to filename column
        col_map = {
            "packing_slip": "packing_list_filename",
            "receiving_inspection": "recv_inspection_filename",
            "packing": "packing_list_filename",  # Legacy support
            "inspection": "recv_inspection_filename",  # Legacy support
        }
        col = col_map.get(kind)
        fn = shipment.get(col) if col else None
        if not fn:
            abort(404)
        
        # New path structure: received_shipments/<shipment_id>/<filename>
        # Filename may be relative path or just filename
        if "/" in fn:
            # Already a relative path (new structure)
            file_path = Path(app.config["UPLOAD_ROOT"]) / fn
        else:
            # Legacy: try old structure first, then new
            rep_slug = shipment.get("rep_slug") or ""
            old_path = Path(app.config["UPLOAD_ROOT"]) / "received" / rep_slug / str(shipment_id) / fn
            new_path = Path(app.config["UPLOAD_ROOT"]) / "received_shipments" / str(shipment_id) / fn
            if old_path.exists():
                file_path = old_path
            else:
                file_path = new_path
        
        if not file_path.exists():
            abort(404)
        # Display inline for images, download for PDFs
        as_attachment = file_path.suffix.lower() == '.pdf'
        return send_from_directory(file_path.parent, file_path.name, as_attachment=as_attachment)

    @app.route("/files/distributed/<slug>/<int:dist_id>/<kind>", endpoint="download_distributed_file")
    @require_rep
    def download_distributed_file(slug, dist_id, kind):
        rep = normalize_row(g.get("current_rep") or get_rep_by_slug(slug.lower()))
        if rep is None:
            abort(404)
        dist = normalize_row(query_db("SELECT * FROM devices_distributed WHERE id = %s AND rep_id = %s", (dist_id, rep["id"]), one=True))
        if not dist:
            abort(404)
        if kind == "ack":
            fn = dist.get("ack_filename")
        elif kind == "cust":
            fn = dist.get("cust_info_filename")
        else:
            abort(404)
        if not fn:
            abort(404)
        file_path = Path(app.config["UPLOAD_ROOT"]) / "distributed" / slug / str(dist_id) / fn
        if not file_path.exists():
            abort(404)
        return send_from_directory(file_path.parent, file_path.name, as_attachment=True)

    @app.route("/files/distributed/master/<slug>", endpoint="download_dist_master")
    @require_rep
    def download_dist_master(slug):
        rep = normalize_row(g.get("current_rep") or get_rep_by_slug(slug.lower()))
        if rep is None:
            abort(404)
        master_log = normalize_row(query_db("SELECT * FROM rep_master_logs WHERE rep_id = %s AND log_type = 'distribution' ORDER BY updated_at DESC LIMIT 1", (rep["id"],), one=True))
        if not master_log or not master_log.get("stored_filename"):
            abort(404)
        file_path = BASE_DIR / master_log["stored_filename"]
        if not file_path.exists():
            abort(404)
        return send_from_directory(file_path.parent, file_path.name, as_attachment=True)

    @app.route("/files/submissions/<slug>/<form_type_id>/<filename>", endpoint="download_submission")
    @require_rep
    def download_submission(slug, form_type_id, filename):
        rep = normalize_row(g.get("current_rep") or get_rep_by_slug(slug.lower()))
        if rep is None:
            abort(404)
        sub = normalize_row(query_db("SELECT * FROM submissions WHERE rep_id = %s AND form_type_id = %s AND stored_filename LIKE %s", (rep["id"], form_type_id, f"%{filename}"), one=True))
        if not sub:
            abort(404)
        file_path = BASE_DIR / sub["stored_filename"]
        if not file_path.exists():
            abort(404)
        return send_from_directory(file_path.parent, file_path.name, as_attachment=True)

    @app.route("/files/customers/<slug>/<folder>/<filename>", endpoint="download_customer_file")
    @require_rep
    def download_customer_file(slug, folder, filename):
        rep = g.get("current_rep") or get_rep_by_slug(slug.lower())
        if rep is None:
            abort(404)
        file_path = Path(app.config["CUSTOMER_INFO_DIR"]) / slug / folder / filename
        if not file_path.exists():
            abort(404)
        return send_from_directory(file_path.parent, file_path.name, as_attachment=True)

    @app.route("/files/tracing/<slug>/<month>/<kind>", endpoint="download_tracing_file")
    @require_rep
    def download_tracing_file(slug, month, kind):
        rep = g.get("current_rep") or get_rep_by_slug(slug.lower())
        if rep is None:
            abort(404)
        base_dir = Path(app.config.get("TRACING_DIR", BASE_DIR / "tracing_reports")) / slug / month
        if not base_dir.exists():
            abort(404)
        
        # Try exact match first, then case-insensitive glob
        files = []
        kind_lower = kind.lower()
        
        if kind_lower == "tracing":
            # Look for Tracing_Report_*.csv (exact case first)
            files = list(base_dir.glob("Tracing_Report_*.csv"))
            if not files:
                # Try case-insensitive
                files = [f for f in base_dir.glob("*") if f.is_file() and "tracing" in f.name.lower() and f.suffix.lower() == ".csv"]
        elif kind_lower == "approval":
            # Look for *Approval*.pdf or *approval*.pdf
            files = list(base_dir.glob("*Approval*.pdf"))
            if not files:
                files = list(base_dir.glob("*approval*.pdf"))
            if not files:
                # Try case-insensitive
                files = [f for f in base_dir.glob("*") if f.is_file() and "approval" in f.name.lower() and f.suffix.lower() == ".pdf"]
        else:
            # Generic search (case-insensitive)
            files = [f for f in base_dir.glob("*") if f.is_file() and kind_lower in f.name.lower()]
        
        if not files:
            abort(404)
        
        return send_from_directory(files[0].parent, files[0].name, as_attachment=True)

    # ------------------------------------------------------------------ Rep Target Facilities API
    
    @app.route("/api/rep-targets", methods=["GET"])
    def api_rep_targets():
        """
        API endpoint for fetching target facilities near a ZIP code.
        
        Returns standardized facility objects with:
        - facility_id, name (facility_name), address, city, state, zip
        - distance_miles, catheter_days
        - cauti_flag (boolean), cauti_label (string)
        - Additional fields: latitude, longitude, latest_end_date, hospital_type, etc.
        """
        # Check rep authentication
        rep = get_current_rep()
        if not rep and not is_admin():
            return jsonify({"error": "Authentication required"}), 403
        
        try:
            center_zip = request.args.get("center_zip", "").strip()
            center_lat_raw = request.args.get("center_lat", "").strip()
            center_lon_raw = request.args.get("center_lon", "").strip()
            radius_raw = request.args.get("radius", "75").strip()
            limit_raw = request.args.get("limit", "75").strip()
            
            # Support both center_zip and center_lat/center_lon (prefer lat/lon if provided)
            center_lat = None
            center_lon = None
            
            if center_lat_raw and center_lon_raw:
                # Use provided coordinates
                try:
                    center_lat = float(center_lat_raw)
                    center_lon = float(center_lon_raw)
                    # Validate coordinate ranges
                    if not (-90 <= center_lat <= 90) or not (-180 <= center_lon <= 180):
                        return jsonify({"error": "Invalid coordinates: lat must be -90 to 90, lon must be -180 to 180"}), 400
                except (ValueError, TypeError):
                    return jsonify({"error": "Invalid coordinate format: center_lat and center_lon must be numbers"}), 400
            elif center_zip:
                # Use ZIP code (existing behavior)
                # If no ZIP provided, try to use rep's ZIP
                if not center_zip:
                    rep_zip = rep.get("zip") or "" if rep else ""
                    center_zip = "".join(c for c in rep_zip if c.isdigit())[:5] if rep_zip else ""
                
                # Validate ZIP code format (5 digits)
                if not center_zip or len(center_zip) != 5 or not center_zip.isdigit():
                    print(f"[API] Invalid ZIP code provided: {center_zip}")
                    return jsonify({"error": "Please provide a valid 5-digit ZIP code"}), 400
            else:
                # Try to use rep's ZIP as fallback
                rep_zip = rep.get("zip") or "" if rep else ""
                center_zip = "".join(c for c in rep_zip if c.isdigit())[:5] if rep_zip else ""
                if not center_zip or len(center_zip) != 5 or not center_zip.isdigit():
                    return jsonify({"error": "Please provide center_zip or center_lat/center_lon"}), 400
            
            # Parse and validate radius (0-300 miles)
            try:
                radius = float(radius_raw)
                if radius < 0 or radius > 300:
                    print(f"[API] Radius out of range: {radius}, clamping to 0-300")
                    radius = min(max(radius, 0), 300)
            except (ValueError, TypeError):
                print(f"[API] Invalid radius value: {radius_raw}, using default 75.0")
                radius = 75.0
            
            # Parse and validate limit (1-200)
            try:
                limit = int(limit_raw)
                if limit < 1 or limit > 200:
                    print(f"[API] Limit out of range: {limit}, clamping to 1-200")
                    limit = min(max(limit, 1), 200)
            except (ValueError, TypeError):
                print(f"[API] Invalid limit value: {limit_raw}, using default 75")
                limit = 75
            
            # Verify get_targets function is available
            if not get_targets:
                print("[API] ERROR: get_targets function not available")
                return jsonify({"error": "Target facilities module not available"}), 503
            
            # Check if cache exists (using same priority as hospital_targets.py)
            # Priority: facility_targets_cache.parquet > .csv > legacy facility_catheter_days.csv
            hospital_cache_exists = False
            hospital_cache_filename = None
            try:
                from pathlib import Path
                from hospital_targets import FACILITY_TARGETS_CACHE_PARQUET, FACILITY_TARGETS_CACHE_CSV, FACILITY_CACHE_CSV
                
                # Check in priority order (same as load_facility_cache)
                if FACILITY_TARGETS_CACHE_PARQUET.exists():
                    hospital_cache_exists = True
                    hospital_cache_filename = FACILITY_TARGETS_CACHE_PARQUET.name
                elif FACILITY_TARGETS_CACHE_CSV.exists():
                    hospital_cache_exists = True
                    hospital_cache_filename = FACILITY_TARGETS_CACHE_CSV.name
                elif FACILITY_CACHE_CSV.exists():
                    hospital_cache_exists = True
                    hospital_cache_filename = FACILITY_CACHE_CSV.name
            except Exception as cache_check_error:
                # If we can't check cache, log but continue (might be dev mode)
                print(f"[API] Warning: Could not verify cache existence: {cache_check_error}")
                import traceback
                traceback.print_exc()
                hospital_cache_exists = False
                hospital_cache_filename = None
            
            # Always check cache existence and return 503 if missing (regardless of cache-only mode)
            # This ensures we distinguish "cache missing" from "valid empty result"
            if not hospital_cache_exists:
                import os
                cache_only = os.environ.get("CACHE_ONLY", "").lower() in ("1", "true", "yes") or \
                            os.environ.get("CACHE_ONLY_MODE", "").lower() in ("1", "true", "yes")
                
                # Return 503 with cache_missing status - this signals to frontend that it's a system issue
                # not a geographic "no facilities in radius" issue
                print(f"[API] ERROR: Hospital cache not found - returning 503 (cache_missing)")
                return jsonify({
                    "error": "Target facilities data is temporarily unavailable",
                    "status": "cache_missing",
                    "ok": False,
                    "message": "Hospital cache file not found. Expected one of: facility_targets_cache.parquet, facility_targets_cache.csv, or facility_catheter_days.csv",
                    "cache_only_mode": cache_only
                }), 503
            
            # Cache exists - proceed with query
            print(f"[API] Hospital cache found: {hospital_cache_filename}")
            if center_lat is not None and center_lon is not None:
                print(f"[API] Request: center_lat={center_lat}, center_lon={center_lon}, radius={radius}mi, limit={limit}")
                facilities = get_targets(center_lat=center_lat, center_lon=center_lon, radius_miles=radius, limit=limit)
            else:
                print(f"[API] Request: ZIP={center_zip}, radius={radius}mi, limit={limit}")
                facilities = get_targets(center_zip=center_zip, radius_miles=radius, limit=limit)
            
            # If cache exists but query returned empty, that's a valid result (no facilities in radius)
            # This is different from cache_missing - it means cache loaded but no facilities match the search criteria
            
            # Get center coordinates for response
            center_coords = None
            geocode_error = None
            if center_lat is not None and center_lon is not None:
                # Use provided coordinates
                center_coords = (center_lat, center_lon)
            else:
                # Geocode ZIP code
                try:
                    if zip_to_latlon:
                        center_coords = zip_to_latlon(center_zip)
                        if not center_coords:
                            geocode_error = f"ZIP code {center_zip} could not be geocoded"
                            print(f"[API] Warning: {geocode_error}")
                except Exception as e:
                    geocode_error = f"Geocoding error: {str(e)}"
                    print(f"[API] Warning: {geocode_error}")
            
            # Validate facility coordinates
            facilities_with_coords = sum(1 for f in facilities if f.get("latitude") is not None and f.get("longitude") is not None)
            facilities_without_coords = len(facilities) - facilities_with_coords
            
            # Warn if many facilities are missing coordinates
            if len(facilities) > 0 and facilities_without_coords > len(facilities) * 0.05:
                print(f"[API] Warning: {facilities_without_coords}/{len(facilities)} facilities missing coordinates (>{5}%)")
            
            # Ensure output schema is standardized (get_targets already standardizes, but verify)
            standardized_results = []
            for fac in facilities:
                standardized = fac.copy()
                # Ensure required fields exist
                standardized.setdefault("name", standardized.get("facility_name", ""))
                standardized.setdefault("cauti_flag", bool(standardized.get("cauti_is_worse", False)))
                standardized.setdefault("cauti_label", standardized.get("cauti_compared_to_national", "") or "")
                standardized_results.append(standardized)
            
            # Count CAUTI-worse facilities for response metadata
            cauti_worse_count = sum(1 for f in standardized_results if f.get("cauti_flag", False))
            
            # Admin debug option (only for admin users)
            debug_info = None
            if is_admin() and request.args.get("debug") == "1":
                # Load cache to get statistics (already loaded, but we can get counts)
                from hospital_targets import load_facility_cache
                try:
                    all_facilities = load_facility_cache(dev_fallback=False)
                    total_facilities = len(all_facilities) if all_facilities else 0
                    all_facilities_with_coords = sum(1 for f in all_facilities if f.get("latitude") is not None and f.get("longitude") is not None) if all_facilities else 0
                    debug_info = {
                        "request_zip": center_zip if center_zip else None,
                        "request_center_lat": center_lat if center_lat is not None else None,
                        "request_center_lon": center_lon if center_lon is not None else None,
                        "request_radius_miles": radius,
                        "request_limit": limit,
                        "center_coords": {"lat": center_coords[0], "lon": center_coords[1]} if center_coords else None,
                        "geocode_error": geocode_error,
                        "total_facilities_in_cache": total_facilities,
                        "facilities_with_coords_in_cache": all_facilities_with_coords,
                        "returned_count": len(standardized_results),
                        "returned_with_coords": facilities_with_coords,
                        "returned_without_coords": facilities_without_coords,
                        "cache_file": hospital_cache_filename
                    }
                except Exception as debug_error:
                    debug_info = {"error": str(debug_error)}
            
            # At this point, cache exists (we returned 503 if missing)
            # Empty results here means valid "no facilities in radius" result
            # Determine if empty result is legitimate (geocoding succeeded but no facilities in radius)
            # vs. a system issue (geocoding failed or cache problem)
            is_legitimate_empty = (
                center_coords is not None and 
                len(standardized_results) == 0 and
                hospital_cache_exists
            )
            
            response_data = {
                "ok": True,
                "center_zip": center_zip if center_zip else None,
                "center_lat": center_lat if center_lat is not None else None,
                "center_lon": center_lon if center_lon is not None else None,
                "center_coords": {"lat": center_coords[0], "lon": center_coords[1]} if center_coords else None,
                "geocode_error": geocode_error,
                "geocode_success": center_coords is not None,
                "radius": radius,
                "limit": limit,
                "count": len(standardized_results),
                "cauti_worse_count": cauti_worse_count,
                "facilities_with_coords": facilities_with_coords,
                "facilities_without_coords": facilities_without_coords,
                "is_legitimate_empty": is_legitimate_empty,
                "results": standardized_results,
                # Status: "ok" means cache loaded successfully (even if 0 results)
                # "cache_missing" would have returned 503 above
                "status": "ok"
            }
            
            # Add debug info if requested (admin only)
            if debug_info:
                response_data["debug"] = debug_info
            
            # Log appropriately - empty results with cache present = valid geographic result
            if len(standardized_results) == 0:
                if center_lat is not None and center_lon is not None:
                    print(f"[API] Response: 0 facilities found in radius {radius}mi from coordinates ({center_lat}, {center_lon}) (valid empty result - cache loaded: {hospital_cache_filename})")
                else:
                    print(f"[API] Response: 0 facilities found in radius {radius}mi from ZIP {center_zip} (valid empty result - cache loaded: {hospital_cache_filename})")
            else:
                print(f"[API] Response: {len(standardized_results)} facilities ({cauti_worse_count} CAUTI-worse) from cache: {hospital_cache_filename}")
            # region agent log
            _agent_log(
                "MAP_DOCTORS_API",
                "Proto1.py:api_rep_targets",
                "rep-targets response",
                {
                    "used_zip": bool(center_zip),
                    "zip_len": len(center_zip or ""),
                    "used_coords": center_lat is not None and center_lon is not None,
                    "radius": float(radius) if isinstance(radius, (int, float)) else None,
                    "limit": int(limit) if isinstance(limit, int) else None,
                    "cache_ok": bool(hospital_cache_exists),
                    "status": "ok",
                    "count": int(len(standardized_results)),
                    "coords_present": int(facilities_with_coords),
                },
            )
            # endregion
            return jsonify(response_data)
        
        except ValueError as e:
            # Parameter validation errors should return 400
            print(f"[API] Validation error: {e}")
            return jsonify({"error": f"Invalid parameter: {str(e)}"}), 400
        except Exception as e:
            # Unexpected errors: in cache-only mode, return 503 if it's a cache-related error
            import traceback
            import os
            cache_only = os.environ.get("CACHE_ONLY", "").lower() in ("1", "true", "yes") or \
                        os.environ.get("CACHE_ONLY_MODE", "").lower() in ("1", "true", "yes")
            
            error_str = str(e).lower()
            if cache_only and ("cache" in error_str or "file not found" in error_str or "does not exist" in error_str):
                print(f"[API] Cache-related error in cache-only mode: {e}")
                traceback.print_exc()
                return jsonify({
                    "error": "Target facilities data is temporarily unavailable",
                    "status": "cache_missing",
                    "ok": False
                }), 503
            else:
                print(f"[API] Unexpected error: {e}")
                traceback.print_exc()
                return jsonify({"error": f"Internal server error: {str(e)}"}), 500
    
    @app.route("/api/rep-targets/export", methods=["GET"])
    def api_rep_targets_export():
        """
        Export current facility list as CSV.
        Uses same parameters as /api/rep-targets to get matching results.
        """
        # Check rep authentication
        rep = get_current_rep()
        if not rep and not is_admin():
            return jsonify({"error": "Authentication required"}), 403
        
        try:
            # Get same parameters as api_rep_targets
            center_zip = request.args.get("center_zip", "").strip()
            radius_raw = request.args.get("radius", "75").strip()
            limit_raw = request.args.get("limit", "75").strip()
            
            # Use rep's ZIP if not provided
            if not center_zip and rep:
                rep_zip = rep.get("zip") or ""
                center_zip = "".join(c for c in rep_zip if c.isdigit())[:5] if rep_zip else ""
            
            # Validate ZIP
            if not center_zip or len(center_zip) != 5 or not center_zip.isdigit():
                return jsonify({"error": "Please provide a valid 5-digit ZIP code"}), 400
            
            # Parse radius and limit (same validation as api_rep_targets)
            try:
                radius = min(max(float(radius_raw), 0), 300)
                limit = min(max(int(limit_raw), 1), 200)
            except (ValueError, TypeError):
                radius = 75.0
                limit = 75
            
            # Get facilities (reuse api_rep_targets logic)
            if not get_targets:
                return jsonify({"error": "Target facilities module not available"}), 503
            
            if center_lat is not None and center_lon is not None:
                facilities = get_targets(center_lat=center_lat, center_lon=center_lon, radius_miles=radius, limit=limit)
            else:
                facilities = get_targets(center_zip=center_zip, radius_miles=radius, limit=limit)
            
            # Standardize results
            standardized_results = []
            for fac in facilities:
                standardized = fac.copy()
                standardized.setdefault("name", standardized.get("facility_name", ""))
                standardized.setdefault("cauti_flag", bool(standardized.get("cauti_is_worse", False)))
                standardized_results.append(standardized)
            
            # Generate CSV
            import io
            from flask import Response
            
            output = io.StringIO()
            writer = csv.DictWriter(output, fieldnames=[
                "facility_id", "facility_name", "address", "city", "state", "zip",
                "latitude", "longitude", "distance_miles", "catheter_days",
                "cauti_flag", "cauti_label", "hospital_type"
            ])
            writer.writeheader()
            
            for fac in standardized_results:
                writer.writerow({
                    "facility_id": fac.get("facility_id", ""),
                    "facility_name": fac.get("name") or fac.get("facility_name", ""),
                    "address": fac.get("address", ""),
                    "city": fac.get("city", ""),
                    "state": fac.get("state", ""),
                    "zip": fac.get("zip", ""),
                    "latitude": fac.get("latitude", ""),
                    "longitude": fac.get("longitude", ""),
                    "distance_miles": fac.get("distance_miles", ""),
                    "catheter_days": fac.get("catheter_days", 0),
                    "cauti_flag": "Yes" if fac.get("cauti_flag") else "No",
                    "cauti_label": fac.get("cauti_label", ""),
                    "hospital_type": fac.get("hospital_type", "")
                })
            
            filename = f"facilities_{center_zip}_{radius}mi_{datetime.now().strftime('%Y%m%d')}.csv"
            return Response(
                output.getvalue(),
                mimetype="text/csv",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
            
        except ValueError as e:
            # Parameter validation errors should return 400, not 500
            print(f"[API Export] Validation error: {e}")
            return jsonify({"error": f"Invalid parameter: {str(e)}"}), 400
        except Exception as e:
            # Unexpected errors
            print(f"[API Export] Error: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({"error": f"Error exporting CSV: {str(e)}"}), 500
    
    @app.route("/api/facility-doctors", methods=["GET"])
    def api_facility_doctors():
        """
        API endpoint for fetching doctors affiliated with a facility.
        
        Reads from cache/facility_doctors_cache.json (cache-only mode, no raw datasets required).
        Returns empty list if no doctors found (not an error).
        """
        # Check rep authentication
        rep = get_current_rep()
        if not rep and not is_admin():
            return jsonify({"error": "Authentication required"}), 403
        
        try:
            facility_id = request.args.get("facility_id", "").strip()
            
            if not facility_id:
                print("[API] Missing facility_id parameter")
                return jsonify({"error": "facility_id parameter required"}), 400
            
            # Validate facility_id format (should be 6 digits after normalization)
            from doctor_targets import normalize_facility_id
            normalized_id = normalize_facility_id(facility_id)
            if not normalized_id or len(normalized_id) != 6 or not normalized_id.isdigit():
                print(f"[API] Invalid facility_id format: {facility_id} (normalized: {normalized_id})")
                return jsonify({"error": "facility_id must be a valid 6-digit facility ID"}), 400
            
            if not get_doctors_for_facility:
                print("[API] ERROR: get_doctors_for_facility function not available")
                return jsonify({"error": "Doctor targeting module not available"}), 503
            
            # Check if doctor cache exists (for graceful failure in production)
            doctor_cache_exists = False
            import os
            try:
                from pathlib import Path
                
                cache_dir = Path(__file__).resolve().parent / "cache"
                doctor_cache_exists = (cache_dir / "facility_doctors_cache.json").exists()
                
                cache_only = os.environ.get("CACHE_ONLY", "").lower() in ("1", "true", "yes") or \
                            os.environ.get("CACHE_ONLY_MODE", "").lower() in ("1", "true", "yes")
                
                if not doctor_cache_exists and cache_only:
                    print("[API] ERROR: Doctor cache not found in cache-only mode - returning 503")
                    return jsonify({
                        "error": "Doctor data is temporarily unavailable",
                        "status": "cache_missing"
                    }), 503
            except Exception as cache_check_error:
                print(f"[API] Warning: Could not verify doctor cache existence: {cache_check_error}")
            
            # Get doctors for facility using cache-only mode (no raw datasets needed in production)
            # Cache-only mode prevents attempting to build cache if missing (returns empty list instead)
            cache_only_mode = os.environ.get("CACHE_ONLY", "").lower() in ("1", "true", "yes") or \
                            os.environ.get("CACHE_ONLY_MODE", "").lower() in ("1", "true", "yes")
            doctors = get_doctors_for_facility(facility_id, cache_only_mode=cache_only_mode)
            
            # If cache was missing and get_doctors returned empty, check if we should return 503
            if not doctors and cache_only_mode and not doctor_cache_exists:
                return jsonify({
                    "error": "Doctor data is temporarily unavailable",
                    "status": "cache_missing"
                }), 503
            
            print(f"[API] Request: facility_id={facility_id} (normalized: {normalized_id}), returned {len(doctors)} doctors")
            # region agent log
            _agent_log(
                "MAP_DOCTORS_API",
                "Proto1.py:api_facility_doctors",
                "facility-doctors response",
                {
                    "facility_id_len": len(facility_id or ""),
                    "normalized_ok": bool(normalized_id and len(normalized_id) == 6),
                    "cache_only_mode": bool(cache_only_mode),
                    "doctor_cache_exists": bool(doctor_cache_exists),
                    "count": int(len(doctors)),
                },
            )
            # endregion
            return jsonify({
                "ok": True,
                "facility_id": normalized_id,  # Return normalized ID
                "count": len(doctors),
                "doctors": doctors
            })
        except ValueError as e:
            # Parameter validation errors should return 400, not 500
            print(f"[API] Validation error: {e}")
            return jsonify({"error": f"Invalid parameter: {str(e)}"}), 400
        except Exception as e:
            # Unexpected errors: in cache-only mode, return 503 if it's a cache-related error
            import traceback
            import os
            cache_only = os.environ.get("CACHE_ONLY", "").lower() in ("1", "true", "yes") or \
                        os.environ.get("CACHE_ONLY_MODE", "").lower() in ("1", "true", "yes")
            
            error_str = str(e).lower()
            if cache_only and ("cache" in error_str or "file not found" in error_str or "does not exist" in error_str):
                print(f"[API] Cache-related error in cache-only mode: {e}")
                return jsonify({
                    "error": "Doctor data is temporarily unavailable",
                    "status": "cache_error"
                }), 503
            
            print(f"[API] ERROR: {e}")
            traceback.print_exc()
            return jsonify({"error": f"Error fetching doctors: {str(e)}"}), 500

    # ------------------------------------------------------------------ Admin Cache Status
    
    @app.route("/admin/cache-status")
    def admin_cache_status():
        """Admin-only endpoint to show cache status (build dates, dataset versions)"""
        if not is_admin():
            return redirect(url_for("admin_login"))
        
        try:
            from data_bootstrap import get_cache_metadata
            from pathlib import Path
            
            BASE_DIR = Path(__file__).resolve().parent
            CACHE_DIR = BASE_DIR / "cache"
            try:
                metadata = get_cache_metadata()
            except Exception as metadata_error:
                # If metadata unavailable, use empty dict (not critical)
                print(f"[Cache Status] Could not load metadata: {metadata_error}")
                metadata = {}
            
            # Check actual cache files (new production filenames, with legacy fallback)
            # Hospital cache: prefer facility_targets_cache.parquet > .csv > legacy facility_catheter_days.csv
            facility_cache_parquet = CACHE_DIR / "facility_targets_cache.parquet"
            facility_cache_csv = CACHE_DIR / "facility_targets_cache.csv"
            facility_cache_legacy = CACHE_DIR / "facility_catheter_days.csv"
            
            # Find which hospital cache file exists (priority order)
            facility_cache = None
            facility_cache_type = None
            if facility_cache_parquet.exists():
                facility_cache = facility_cache_parquet
                facility_cache_type = "parquet"
            elif facility_cache_csv.exists():
                facility_cache = facility_cache_csv
                facility_cache_type = "csv"
            elif facility_cache_legacy.exists():
                facility_cache = facility_cache_legacy
                facility_cache_type = "legacy"
            
            # Doctors cache: facility_doctors_cache.json (standardized name)
            doctors_cache = CACHE_DIR / "facility_doctors_cache.json"
            doctors_cache_legacy = CACHE_DIR / "facility_doctors.json"  # Legacy alias (backward compat)
            
            # Check legacy doctors cache if new one doesn't exist
            if not doctors_cache.exists() and doctors_cache_legacy.exists():
                doctors_cache = doctors_cache_legacy
            
            zip_cache = CACHE_DIR / "zip_geocodes.db"
            
            # Determine which cache file is actually being used (priority order)
            active_hospital_cache = None
            active_hospital_type = None
            if facility_cache_parquet.exists():
                active_hospital_cache = facility_cache_parquet
                active_hospital_type = "parquet"
            elif facility_cache_csv.exists():
                active_hospital_cache = facility_cache_csv
                active_hospital_type = "csv"
            elif facility_cache_legacy.exists():
                active_hospital_cache = facility_cache_legacy
                active_hospital_type = "legacy"
            
            active_doctor_cache = None
            active_doctor_type = None
            if doctors_cache.exists():
                active_doctor_cache = doctors_cache
                active_doctor_type = "standard"
            elif doctors_cache_legacy.exists():
                active_doctor_cache = doctors_cache_legacy
                active_doctor_type = "legacy"
            
            # Calculate coordinate statistics for facility cache
            coord_stats = {
                "total_facilities": 0,
                "coords_present": 0,
                "coords_missing": 0,
                "coords_present_pct": 0.0,
                "coords_missing_pct": 0.0
            }
            
            if active_hospital_cache and active_hospital_cache.exists():
                try:
                    if active_hospital_type == "parquet":
                        try:
                            import pandas as pd
                            try:
                                df = pd.read_parquet(active_hospital_cache)
                            except (ImportError, ModuleNotFoundError) as import_err:
                                # pyarrow/fastparquet not available - treat as missing
                                print(f"[Cache Status] Warning: Parquet engine not available (pyarrow/fastparquet), skipping parquet coord stats: {import_err}")
                                coord_stats["parquet_error"] = "Parquet engine not available (install pyarrow or fastparquet)"
                                raise
                            coord_stats["total_facilities"] = len(df)
                            coord_stats["coords_missing"] = int(df['lat'].isna().sum() if 'lat' in df.columns else 0)
                            coord_stats["coords_present"] = coord_stats["total_facilities"] - coord_stats["coords_missing"]
                            if coord_stats["total_facilities"] > 0:
                                coord_stats["coords_present_pct"] = round((coord_stats["coords_present"] / coord_stats["total_facilities"]) * 100, 1)
                                coord_stats["coords_missing_pct"] = round((coord_stats["coords_missing"] / coord_stats["total_facilities"]) * 100, 1)
                        except Exception as e:
                            print(f"[Cache Status] Error reading parquet for coord stats: {e}")
                    elif active_hospital_type in ("csv", "legacy"):
                        try:
                            import csv
                            import io
                            with open(active_hospital_cache, 'r', encoding='utf-8') as f:
                                # Skip comment lines (starting with #) and blank lines until header
                                lines = []
                                for line in f:
                                    stripped = line.strip()
                                    if stripped and not stripped.startswith('#'):
                                        lines.append(line)
                                
                                if lines:
                                    reader = csv.DictReader(io.StringIO(''.join(lines)))
                                    rows = list(reader)
                                    coord_stats["total_facilities"] = len(rows)
                                    coord_stats["coords_missing"] = sum(1 for r in rows if not r.get('lat') or r.get('lat') == '' or r.get('lat') == 'None')
                                    coord_stats["coords_present"] = coord_stats["total_facilities"] - coord_stats["coords_missing"]
                                    if coord_stats["total_facilities"] > 0:
                                        coord_stats["coords_present_pct"] = round((coord_stats["coords_present"] / coord_stats["total_facilities"]) * 100, 1)
                                        coord_stats["coords_missing_pct"] = round((coord_stats["coords_missing"] / coord_stats["total_facilities"]) * 100, 1)
                        except Exception as e:
                            print(f"[Cache Status] Error reading CSV for coord stats: {e}")
                except Exception as e:
                    print(f"[Cache Status] Error calculating coordinate statistics: {e}")
            
            cache_status = {
                "facility_cache": {
                    "exists": active_hospital_cache is not None and active_hospital_cache.exists(),
                    "filename": active_hospital_cache.name if active_hospital_cache else None,
                    "type": active_hospital_type,
                    "size_bytes": active_hospital_cache.stat().st_size if active_hospital_cache and active_hospital_cache.exists() else 0,
                    "size_mb": round(active_hospital_cache.stat().st_size / (1024 * 1024), 2) if active_hospital_cache and active_hospital_cache.exists() else 0,
                    "last_modified": datetime.fromtimestamp(active_hospital_cache.stat().st_mtime).isoformat() if active_hospital_cache and active_hospital_cache.exists() else None,
                    "built_at": metadata.get("facility_catheter_days", {}).get("built_at") or metadata.get("facility_targets_cache", {}).get("built_at"),
                    "dataset_version": metadata.get("facility_catheter_days", {}).get("dataset_version") or metadata.get("facility_targets_cache", {}).get("dataset_version"),
                    # Coordinate statistics
                    "coordinate_stats": coord_stats,
                    # Also check alternative filenames for debugging
                    "files_checked": {
                        "facility_targets_cache.parquet": facility_cache_parquet.exists(),
                        "facility_targets_cache.csv": facility_cache_csv.exists(),
                        "facility_catheter_days.csv": facility_cache_legacy.exists()
                    }
                },
                "doctors_cache": {
                    "exists": active_doctor_cache is not None and active_doctor_cache.exists(),
                    "filename": active_doctor_cache.name if active_doctor_cache else None,
                    "type": active_doctor_type,
                    "size_bytes": active_doctor_cache.stat().st_size if active_doctor_cache and active_doctor_cache.exists() else 0,
                    "size_mb": round(active_doctor_cache.stat().st_size / (1024 * 1024), 2) if active_doctor_cache and active_doctor_cache.exists() else 0,
                    "last_modified": datetime.fromtimestamp(active_doctor_cache.stat().st_mtime).isoformat() if active_doctor_cache and active_doctor_cache.exists() else None,
                    "built_at": metadata.get("facility_doctors", {}).get("built_at"),
                    "dataset_version": metadata.get("facility_doctors", {}).get("dataset_version"),
                    # Also check legacy filename
                    "files_checked": {
                        "facility_doctors_cache.json": doctors_cache.exists(),
                        "facility_doctors.json": doctors_cache_legacy.exists()
                    }
                },
                "zip_cache": {
                    "exists": zip_cache.exists(),
                    "size_bytes": zip_cache.stat().st_size if zip_cache.exists() else 0,
                    "size_kb": round(zip_cache.stat().st_size / 1024, 2) if zip_cache.exists() else 0
                },
                "bootstrap_config": {
                    "spaces_configured": bool(
                        (os.environ.get("SPACES_KEY") or os.environ.get("DO_SPACES_KEY")) and
                        (os.environ.get("SPACES_SECRET") or os.environ.get("DO_SPACES_SECRET")) and
                        (os.environ.get("SPACES_ENDPOINT") or os.environ.get("DO_SPACES_ENDPOINT")) and
                        (os.environ.get("SPACES_BUCKET") or os.environ.get("DO_SPACES_BUCKET"))
                    ),
                    "bucket": os.environ.get("SPACES_BUCKET", os.environ.get("DO_SPACES_BUCKET", "not set")),
                    "endpoint": os.environ.get("SPACES_ENDPOINT", os.environ.get("DO_SPACES_ENDPOINT", "not set")),
                    "cache_only_mode": (
                        os.environ.get("CACHE_ONLY", "").lower() in ("1", "true", "yes") or
                        os.environ.get("CACHE_ONLY_MODE", "").lower() in ("1", "true", "yes")
                    ),
                    "cache_prefix": os.environ.get("CACHE_PREFIX", "caches/"),
                    "hospital_cache_key_parquet": os.environ.get(
                        "HOSPITAL_CACHE_KEY_PARQUET",
                        f"{os.environ.get('CACHE_PREFIX', 'caches/').strip().rstrip('/')}/facility_targets_cache.parquet"
                    ),
                    "hospital_cache_key_csv": os.environ.get(
                        "HOSPITAL_CACHE_KEY_CSV",
                        f"{os.environ.get('CACHE_PREFIX', 'caches/').strip().rstrip('/')}/facility_targets_cache.csv"
                    ),
                    "hospital_cache_key_legacy": os.environ.get(
                        "HOSPITAL_CACHE_KEY_LEGACY",
                        f"{os.environ.get('CACHE_PREFIX', 'caches/').strip().rstrip('/')}/facility_catheter_days.csv"
                    ),
                    "doctor_cache_key": os.environ.get(
                        "DOCTOR_CACHE_KEY",
                        f"{os.environ.get('CACHE_PREFIX', 'caches/').strip().rstrip('/')}/facility_doctors_cache.json"
                    ),
                    "allow_raw_dataset_bootstrap": os.environ.get("ALLOW_RAW_DATASET_BOOTSTRAP", "").lower() in ("1", "true", "yes")
                }
            }
            
            # Render template with cache status data
            return render_template("admin_cache_status.html", cache_status=cache_status)
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return render_template("admin_cache_status.html", cache_status={"error": f"Error checking cache status: {str(e)}"}), 500

    with app.app_context():
        init_db()
        ensure_rep_documents_table()
    
    # Inject canonical field normalization into shipstation_sync module
    # This avoids circular import issues while allowing shipstation_sync to use normalize_fields_json
    set_normalize_fields_json_fn(normalize_fields_json)
    set_customer_helper_fns(find_or_create_customer, ensure_rep_assignment, pick_rep_for_customer)

    # Expose DB helpers for scripts/tests
    app.query_db = query_db  # type: ignore[attr-defined]
    app.execute_db = execute_db  # type: ignore[attr-defined]
    app.fetch_distribution_records = fetch_distribution_records  # type: ignore[attr-defined]
    app.canonical_customer_key = canonical_customer_key  # type: ignore[attr-defined]

    return app


if __name__ == "__main__":
    app = create_app()
    app.run(host='0.0.0.0', debug=True)