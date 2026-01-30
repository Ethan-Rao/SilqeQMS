#!/usr/bin/env python3
"""
Import script for Equipment and Suppliers from Excel/Word files.

Usage:
    python scripts/import_equipment_and_suppliers.py

Files expected (in repo root):
    - "Silq Equipment Master List.xlsx" (equipment data)
    - "SILQ Approved Supplier List Feb 2025.docx" (supplier data)

Idempotent: Safe to re-run without creating duplicates.
"""
from __future__ import annotations

import os
import re
import sys
from datetime import date, datetime
from pathlib import Path

# Ensure repo root is on sys.path
ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from sqlalchemy.orm import Session

from app.eqms.models import Permission, Role, User
from app.eqms.modules.equipment.models import Equipment, EquipmentSupplier
from app.eqms.modules.suppliers.models import Supplier
from scripts._db_utils import script_session


def _normalize_text(s: str | None) -> str:
    """Strip and normalize text."""
    return (s or "").strip()


def _parse_date(val) -> date | None:
    """Parse Excel date (datetime) or string date."""
    if val is None:
        return None
    if isinstance(val, datetime):
        return val.date()
    if isinstance(val, date):
        return val
    s = str(val).strip()
    if not s:
        return None
    # Try common formats
    for fmt in ("%Y-%m-%d", "%m/%d/%Y", "%m/%d/%y", "%Y/%m/%d"):
        try:
            return datetime.strptime(s, fmt).date()
        except ValueError:
            continue
    return None


def _parse_int(val) -> int | None:
    """Parse integer from various types."""
    if val is None:
        return None
    if isinstance(val, int):
        return val
    if isinstance(val, float):
        return int(val)
    s = str(val).strip()
    if not s:
        return None
    try:
        return int(float(s))
    except ValueError:
        return None


def _normalize_supplier_name(name: str) -> str:
    """Normalize supplier name for matching."""
    # Strip, title case, remove common suffixes
    name = name.strip()
    # Remove suffixes for matching
    suffixes = [", Inc.", " Inc.", ", Inc", " Inc", ", LLC", " LLC", ", Corp.", " Corp.", 
                ", Corporation", " Corporation", ", Ltd.", " Ltd.", ", Ltd", " Ltd",
                ", Limited", " Limited"]
    for suffix in suffixes:
        if name.endswith(suffix):
            name = name[:-len(suffix)]
    return name.strip().title()


def import_equipment_from_excel(filepath: str, s: Session, user: User) -> dict:
    """Import equipment from Excel file."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    if not os.path.exists(filepath):
        return {"error": f"File not found: {filepath}"}

    wb = load_workbook(filepath, data_only=True)
    ws = wb.active

    created = 0
    skipped = 0
    errors = []
    mfg_values = {}  # equip_code -> mfg for later linking

    # Assume first row is header
    headers = [cell.value for cell in ws[1]]
    
    # Map header names to columns (flexible)
    col_map = {}
    header_mappings = {
        "equip_code": ["equip code", "equipment code", "code", "equip_code", "id", "equipment id"],
        "status": ["status"],
        "description": ["description", "desc", "name"],
        "mfg": ["mfg", "manufacturer", "make"],
        "model_no": ["model no", "model", "model_no", "model number"],
        "serial_no": ["serial no", "serial", "serial_no", "serial number", "sn"],
        "date_in_service": ["date in service", "in service", "date_in_service", "service date"],
        "location": ["location", "loc"],
        "cal_interval": ["cal interval", "calibration interval", "cal_interval"],
        "last_cal_date": ["last cal", "last calibration", "last_cal_date", "cal date"],
        "cal_due_date": ["cal due", "calibration due", "cal_due_date"],
        "pm_interval": ["pm interval", "pm_interval"],
        "last_pm_date": ["last pm", "last_pm_date", "pm date"],
        "pm_due_date": ["pm due", "pm_due_date"],
        "comments": ["comments", "notes", "remarks"],
    }

    for i, h in enumerate(headers):
        if h is None:
            continue
        h_lower = str(h).strip().lower()
        for field, options in header_mappings.items():
            if h_lower in options:
                col_map[field] = i
                break

    if "equip_code" not in col_map:
        return {"error": "Could not find 'equip_code' column in Excel header"}

    for row in ws.iter_rows(min_row=2):
        vals = [cell.value for cell in row]
        
        equip_code = _normalize_text(str(vals[col_map["equip_code"]]) if col_map.get("equip_code") is not None else "")
        if not equip_code:
            continue

        # Check if exists
        existing = s.query(Equipment).filter(Equipment.equip_code == equip_code).one_or_none()
        if existing:
            skipped += 1
            continue

        def get_val(field):
            idx = col_map.get(field)
            return vals[idx] if idx is not None and idx < len(vals) else None

        mfg = _normalize_text(str(get_val("mfg") or ""))
        
        try:
            eq = Equipment(
                equip_code=equip_code,
                status=_normalize_text(str(get_val("status") or "")) or "Active",
                description=_normalize_text(str(get_val("description") or "")) or None,
                mfg=mfg or None,
                model_no=_normalize_text(str(get_val("model_no") or "")) or None,
                serial_no=_normalize_text(str(get_val("serial_no") or "")) or None,
                date_in_service=_parse_date(get_val("date_in_service")),
                location=_normalize_text(str(get_val("location") or "")) or None,
                cal_interval=_parse_int(get_val("cal_interval")),
                last_cal_date=_parse_date(get_val("last_cal_date")),
                cal_due_date=_parse_date(get_val("cal_due_date")),
                pm_interval=_parse_int(get_val("pm_interval")),
                last_pm_date=_parse_date(get_val("last_pm_date")),
                pm_due_date=_parse_date(get_val("pm_due_date")),
                comments=_normalize_text(str(get_val("comments") or "")) or None,
                created_at=datetime.utcnow(),
                updated_at=datetime.utcnow(),
                created_by_user_id=user.id,
                updated_by_user_id=user.id,
            )
            s.add(eq)
            created += 1
            if mfg:
                mfg_values[equip_code] = mfg
        except Exception as e:
            errors.append(f"Row {row[0].row}: {e}")

    s.flush()
    return {"created": created, "skipped": skipped, "errors": errors, "mfg_values": mfg_values}


def import_suppliers_from_docx(filepath: str, s: Session, user: User) -> dict:
    """Import suppliers from Word document."""
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    if not os.path.exists(filepath):
        return {"error": f"File not found: {filepath}"}

    doc = Document(filepath)
    created = 0
    skipped = 0
    errors = []

    # Try to find a table in the document
    if doc.tables:
        # Parse from table
        for table in doc.tables:
            rows = table.rows
            if len(rows) < 2:
                continue
            
            # Get header row
            headers = [cell.text.strip().lower() for cell in rows[0].cells]
            
            # Map columns
            col_map = {}
            header_mappings = {
                "name": ["name", "supplier", "company", "vendor"],
                "status": ["status", "approval status"],
                "category": ["category", "type"],
                "product_service_provided": ["products", "services", "products/services", "product_service_provided"],
                "address": ["address", "location"],
                "initial_listing_date": ["initial listing", "initial_listing_date", "date added", "listing date"],
                "certification_expiration": ["certification expiration", "cert expiration", "expiration"],
                "notes": ["notes", "comments", "remarks"],
            }
            
            for i, h in enumerate(headers):
                for field, options in header_mappings.items():
                    if h in options:
                        col_map[field] = i
                        break

            if "name" not in col_map:
                continue

            for row in rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                
                name = cells[col_map["name"]] if col_map.get("name") is not None else ""
                if not name:
                    continue

                # Check if exists (by normalized name)
                normalized_name = _normalize_supplier_name(name)
                existing = s.query(Supplier).filter(Supplier.name.ilike(f"%{name}%")).first()
                if existing:
                    skipped += 1
                    continue

                def get_cell(field):
                    idx = col_map.get(field)
                    return cells[idx] if idx is not None and idx < len(cells) else ""

                try:
                    sup = Supplier(
                        name=name,
                        status=get_cell("status") or "Pending",
                        category=get_cell("category") or None,
                        product_service_provided=get_cell("product_service_provided") or None,
                        address=get_cell("address") or None,
                        initial_listing_date=_parse_date(get_cell("initial_listing_date")),
                        certification_expiration=_parse_date(get_cell("certification_expiration")),
                        notes=get_cell("notes") or None,
                        created_at=datetime.utcnow(),
                        updated_at=datetime.utcnow(),
                        created_by_user_id=user.id,
                        updated_by_user_id=user.id,
                    )
                    s.add(sup)
                    created += 1
                except Exception as e:
                    errors.append(f"Supplier '{name}': {e}")
    else:
        # Try to parse from paragraphs (simpler format)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text or len(text) < 3:
                continue
            
            # Skip obvious non-supplier lines
            if text.lower().startswith(("approved", "supplier", "date", "note", "version")):
                continue

            # Assume each non-empty line could be a supplier name
            name = text.split("\t")[0].strip()  # Take first part if tab-separated
            name = name.split("  ")[0].strip()  # Take first part if space-separated
            
            if not name or len(name) < 2:
                continue

            # Check if exists
            existing = s.query(Supplier).filter(Supplier.name.ilike(f"%{name}%")).first()
            if existing:
                skipped += 1
                continue

            try:
                sup = Supplier(
                    name=name,
                    status="Pending",
                    created_at=datetime.utcnow(),
                    updated_at=datetime.utcnow(),
                    created_by_user_id=user.id,
                    updated_by_user_id=user.id,
                )
                s.add(sup)
                created += 1
            except Exception as e:
                errors.append(f"Supplier '{name}': {e}")

    s.flush()
    return {"created": created, "skipped": skipped, "errors": errors}


def link_equipment_suppliers(s: Session, user: User, mfg_values: dict | None = None) -> dict:
    """Link equipment to suppliers based on manufacturer name matching."""
    linked = 0
    skipped = 0

    # Get all equipment with mfg field populated
    equipment_list = s.query(Equipment).filter(Equipment.mfg.isnot(None)).all()
    suppliers = {_normalize_supplier_name(sup.name): sup for sup in s.query(Supplier).all()}

    for eq in equipment_list:
        mfg = eq.mfg
        if not mfg:
            continue

        normalized_mfg = _normalize_supplier_name(mfg)
        
        # Try exact match first
        supplier = suppliers.get(normalized_mfg)
        
        # If no exact match, try partial match
        if not supplier:
            for sup_name, sup in suppliers.items():
                if normalized_mfg.lower() in sup_name.lower() or sup_name.lower() in normalized_mfg.lower():
                    supplier = sup
                    break

        if not supplier:
            skipped += 1
            continue

        # Check if association already exists
        existing = (
            s.query(EquipmentSupplier)
            .filter(EquipmentSupplier.equipment_id == eq.id)
            .filter(EquipmentSupplier.supplier_id == supplier.id)
            .one_or_none()
        )
        if existing:
            skipped += 1
            continue

        # Create association
        assoc = EquipmentSupplier(
            equipment_id=eq.id,
            supplier_id=supplier.id,
            relationship_type="Manufacturer",
            created_at=datetime.utcnow(),
            created_by_user_id=user.id,
        )
        s.add(assoc)
        linked += 1

    s.flush()
    return {"linked": linked, "skipped": skipped}


def main():
    database_url = os.environ.get("DATABASE_URL") or "sqlite:///eqms.db"
    admin_email = (os.environ.get("ADMIN_EMAIL") or "admin@silqeqms.com").strip().lower()

    try:
        with script_session(database_url) as s:
            # Get admin user
            admin_user = s.query(User).filter(User.email == admin_email).one_or_none()
            if not admin_user:
                print(f"ERROR: Admin user '{admin_email}' not found. Run scripts/init_db.py first.")
                return

        # Default file paths (can be overridden via args)
        equipment_file = sys.argv[1] if len(sys.argv) > 1 else "Silq Equipment Master List.xlsx"
        suppliers_file = sys.argv[2] if len(sys.argv) > 2 else "SILQ Approved Supplier List Feb 2025.docx"

        print(f"Importing equipment from: {equipment_file}")
        eq_result = import_equipment_from_excel(equipment_file, s, admin_user)
        if "error" in eq_result:
            print(f"  ERROR: {eq_result['error']}")
        else:
            print(f"  Equipment: created={eq_result['created']}, skipped={eq_result['skipped']}")
            if eq_result.get("errors"):
                for err in eq_result["errors"][:5]:
                    print(f"    {err}")

        print(f"\nImporting suppliers from: {suppliers_file}")
        sup_result = import_suppliers_from_docx(suppliers_file, s, admin_user)
        if "error" in sup_result:
            print(f"  ERROR: {sup_result['error']}")
        else:
            print(f"  Suppliers: created={sup_result['created']}, skipped={sup_result['skipped']}")
            if sup_result.get("errors"):
                for err in sup_result["errors"][:5]:
                    print(f"    {err}")

        print("\nLinking equipment to suppliers (by manufacturer)...")
        link_result = link_equipment_suppliers(s, admin_user, eq_result.get("mfg_values"))
        print(f"  Associations: linked={link_result['linked']}, skipped={link_result['skipped']}")

            s.commit()
            print("\nImport complete. Changes committed.")
    except Exception as e:
        print(f"ERROR: {e}")
        raise


if __name__ == "__main__":
    main()
