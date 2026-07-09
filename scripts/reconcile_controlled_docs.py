"""
Phase 3 — Controlled-Document Housecleaning & Staging Assembly (Dev Prompt 4).

REPORT-ONLY. This script performs NO database writes, NO import, and NO deploy.
It reconciles active-vs-obsolete controlled documents from the authoritative
inputs on disk and assembles a clean, human-reviewable staging tree plus text
artifacts for coordinator + Ethan review. The real import runs later.

Deliverables (see the prompt):
  A. Reconciliation engine -> eQMS_Upload_Staging/reconciliation/
       - document_register_v2.csv
       - manifests/<doc_number>.json   (importer-ready; consumed later by
         scripts/import_document_control.py)
       - discrepancies.md
  B. Staging assembly -> eQMS_Upload_Staging/document_control/<DOC>/{current,superseded}/
     (records staged/cataloged separately). Binaries are NOT committed.
  C. Rebuilt consolidated DCO log -> eQMS_Upload_Staging/reconciliation/DCO_Log_v2.csv

Precedence for the current released revision (highest wins), in priority order:
  1. QMSInProcess/DCO0xx/ package masters (newest completed DCOs; released even
     if not yet in the legacy DCO log)
  2. DCO log new-rev entries (authoritative through DCO091)
  3. The "clean" folders (QM Documents/, Forms, Templates, and Travelers/)

Usage:
    python scripts/reconcile_controlled_docs.py [--no-copy] [--staging-root DIR]
"""
from __future__ import annotations

import argparse
import csv
import json
import re
import shutil
import sys
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))


# ---------------------------------------------------------------------------
# Scan roots (relative to repo root) with a source priority. Lower priority
# number wins when the same revision is found in more than one place.
# ---------------------------------------------------------------------------
SCAN_ROOTS: list[tuple[str, int, str]] = [
    ("QMSInProcess", 1, "QMSInProcess package"),
    ("QM Documents", 2, "QM Documents (clean)"),
    ("Forms, Templates, and Travelers", 2, "Forms/Templates/Travelers (clean)"),
    ("DCOs/Previous Revisions", 3, "DCOs/Previous Revisions"),
    ("Equipment", 4, "Equipment (records)"),
    ("DMR", 4, "DMR"),
    ("DHF", 4, "DHF"),
    ("SLQ-DHF", 4, "SLQ-DHF"),
    ("RiskManagement", 4, "RiskManagement"),
    ("Administration", 4, "Administration"),
]

DCO_LOG_XLSX = "DCOs/SILQ DCO Log.xlsx"
DOCNUM_LOG_XLSX = "DCOs/SILQ Document Number Log.xlsx"
COMPLETED_DCOS_DIR = "DCOs/CompletedDCOs"
QMS_INPROCESS_DIR = "QMSInProcess"

# Redline / working-artifact markers (case-insensitive). These are NOT releases.
REDLINE_RE = re.compile(r"(_redline(?:v\d+)?|_redlined|_redlinde|_rl)\b|redline", re.IGNORECASE)

# A controlled/records document number token, after normalization. Examples:
#   QM.SLQ001, FM1-QM.SLQ001, TMP3-QM.SLQ013, SP-C.SLQ001, SW.SLQ007, VV.SLQ012,
#   L.SLQ011, DC.SLQ001, AD.SLQ001. Also numeric records: SLQ-4000, RM-0018, PS-0006.
# Trailing "-\d+" handles web-content sub-items like WC.SLQ001-0.
DOTTED_DOCNUM_RE = re.compile(r"^[A-Z0-9]+(?:-[A-Z0-9]+)*\.(?:HX|SLQ)\d+(?:-\d+)?$", re.IGNORECASE)
NUMERIC_DOCNUM_RE = re.compile(r"^(?:SLQ|RM|PS)-\d+$", re.IGNORECASE)
REV_RE = re.compile(r"^[A-Z]{1,3}$")

# §4 corrections (baked in + surfaced in discrepancies.md).
# DCO085 line 1 records "QM.SLQ024" which is a typo for QM.SLQ034 (Org Chart).
DCO085_TYPO_TOKEN = "QM.SLQ024"
DCO085_TYPO_FIX = "QM.SLQ034"

# QM.SLQ004 through QM.SLQ010 and their forms/templates were marked OBSOLETE in
# their entirety by DCO095 (superseded by the QM.SLQ052 family).
OBSOLETE_BASE_NUMBERS = {4, 5, 6, 7, 8, 9, 10}
OBSOLETE_REASON = "Superseded in entirety by QM.SLQ052 Rev A (DCO095)."

# Artifacts confirmed by Ethan as misplaced/non-existent — do not treat as a real
# revision (and do not raise a "released master missing" flag for them). FM1-QM.SLQ018
# stays at Rev A; a "B" was mistakenly placed in the DCO093 folder (now removed).
MISPLACED_ARTIFACTS = {("FM1-QM.SLQ018", "B")}


# ---------------------------------------------------------------------------
# §4 coordinator's hand-reconciliation — verify programmatically, report drift.
# ---------------------------------------------------------------------------
def _expected_current() -> dict[str, str]:
    exp: dict[str, str] = {
        "QM.SLQ001": "B", "QM.SLQ002": "B", "QM.SLQ003": "C", "QM.SLQ004": "B",
        "QM.SLQ005": "B", "QM.SLQ011": "B", "QM.SLQ012": "C", "QM.SLQ013": "C",
        "QM.SLQ014": "C", "QM.SLQ015": "C", "QM.SLQ016": "D", "QM.SLQ017": "B",
        "QM.SLQ018": "B", "QM.SLQ019": "C", "QM.SLQ020": "E", "QM.SLQ021": "E",
        "QM.SLQ022": "C", "QM.SLQ023": "B", "QM.SLQ025": "A", "QM.SLQ026": "D",
        "QM.SLQ027": "F", "QM.SLQ028": "B", "QM.SLQ029": "B", "QM.SLQ030": "B",
        "QM.SLQ032": "A", "QM.SLQ033": "B", "QM.SLQ034": "G", "QM.SLQ035": "D",
        "QM.SLQ036": "F", "QM.SLQ037": "B", "QM.SLQ038": "C", "QM.SLQ039": "C",
        "QM.SLQ040": "C", "QM.SLQ043": "B", "QM.SLQ045": "B", "QM.SLQ046": "B",
        "QM.SLQ047": "B", "QM.SLQ048": "B", "QM.SLQ049": "B", "QM.SLQ050": "B",
        "QM.SLQ051": "B", "QM.SLQ052": "A",
    }
    for n in range(6, 11):  # QM.SLQ006–010 = A
        exp[f"QM.SLQ{n:03d}"] = "A"
    # Forms / templates newer than the clean folder (from §4).
    exp.update({
        "FM1-QM.SLQ001": "B", "FM2-QM.SLQ001": "C", "FM1-QM.SLQ014": "B",
        "FM1-QM.SLQ015": "B", "FM2-QM.SLQ015": "B", "FM7-QM.SLQ015": "B",
        "FM2-QM.SLQ017": "B", "FM1-QM.SLQ016": "B",
        # FM1-QM.SLQ018 stays at Rev A (Ethan's correction #3; the DCO093 "B" was
        # misplaced and removed), overriding the original §4 "B".
        "FM1-QM.SLQ018": "A",
        "FM1-QM.SLQ040": "B", "FM4-QM.SLQ050": "B",
        "FM1-QM.SLQ052": "A", "FM2-QM.SLQ052": "A", "FM3-QM.SLQ052": "A",
        "TMP1-QM.SLQ052": "A", "TMP2-QM.SLQ052": "A",
    })
    return exp


EXPECTED_CURRENT = _expected_current()


# ===========================================================================
# Pure helpers (unit-tested)
# ===========================================================================
def normalize_doc_number(token: str) -> str:
    """
    Normalize a legacy/loose document token to current SILQ numbering.

    - HX -> SLQ (the 2021 DCO007 re-baseline), e.g. QM.HX001 -> QM.SLQ001.
    - Insert the missing 'QM.' for bare form/template tokens, e.g.
      FM1-SLQ020 -> FM1-QM.SLQ020, TMP1-SLQ005 -> TMP1-QM.SLQ005.
    - Trim and collapse case of the family portion; keep as uppercased canonical.
    """
    t = (token or "").strip()
    if not t:
        return ""
    # Drop trailing punctuation/noise.
    t = t.strip().strip(",")
    # HX -> SLQ
    t = re.sub(r"\.HX(\d+)", r".SLQ\1", t, flags=re.IGNORECASE)
    t = re.sub(r"-HX(\d+)", r"-SLQ\1", t, flags=re.IGNORECASE)
    # Bare "QM.###" shorthand (e.g. DCO005 "QM.027") -> "QM.SLQ###".
    m = re.match(r"^QM\.(\d+)$", t, flags=re.IGNORECASE)
    if m:
        t = f"QM.SLQ{m.group(1)}"
    # Bare FM#-SLQ### / TMP#-SLQ### -> insert QM.
    m = re.match(r"^((?:FM|TMP)\d+)-SLQ(\d+)(-\d+)?$", t, flags=re.IGNORECASE)
    if m:
        t = f"{m.group(1).upper()}-QM.SLQ{m.group(2)}{m.group(3) or ''}"
    # Uppercase the SLQ/HX marker consistently.
    t = re.sub(r"\.slq", ".SLQ", t, flags=re.IGNORECASE)
    t = re.sub(r"\.hx", ".HX", t, flags=re.IGNORECASE)
    return t


def is_document_number(token: str) -> bool:
    """True if the (normalized) token looks like a controlled/records doc number."""
    t = normalize_doc_number(token)
    if not t:
        return False
    if t.upper().startswith("DCO"):
        return False
    return bool(DOTTED_DOCNUM_RE.match(t) or NUMERIC_DOCNUM_RE.match(t))


def is_document_control(doc_number: str) -> bool:
    """
    Controlled documents that belong in the eQMS Document Control module:
    QM.SLQ* SOPs/WIs and their controlled FM#-QM.SLQ* forms and TMP#-QM.SLQ*
    templates. Everything else (SW/VV/SP/BOM/MP/QC/TR/L/DC/AD/RM records) is
    routed to admin_docs libraries later.
    """
    t = (doc_number or "").upper()
    return bool(re.match(r"^(FM\d+-|TMP\d+-)?QM\.SLQ\d+$", t))


def rev_to_int(rev: str | None) -> int:
    """Excel-style base-26 revision ordering: A=1, B=2, ... Z=26, AA=27. 0 if unknown."""
    if not rev:
        return 0
    r = rev.strip().upper()
    if not re.fullmatch(r"[A-Z]+", r):
        return 0
    n = 0
    for ch in r:
        n = n * 26 + (ord(ch) - ord("A") + 1)
    return n


def is_redline(name: str) -> bool:
    return bool(REDLINE_RE.search(name or ""))


@dataclass
class ParsedFile:
    doc_number: str
    revision: str | None
    title: str
    ext: str
    is_redline: bool
    raw_stem: str


def parse_filename(filename: str) -> ParsedFile | None:
    """
    Parse a controlled-document filename into (doc_number, revision, title, ext).

    Handles real-world messiness:
      - mixed-case extensions (.DOCX)
      - redline/working suffixes (_RedLine, _Redlined, _RL, _RedLineV1, _RedLinde)
      - 'Rev C'-style revision tokens
      - underscore-embedded revision (SLQ-4000_A.pdf)
      - parenthetical noise like '(DCO050)'
    Returns None if the leading token is not a recognizable document number
    (e.g. 'DCO092.docx').
    """
    if not filename:
        return None
    p = Path(filename)
    ext = p.suffix.lower().lstrip(".")
    stem = p.stem.strip()
    redline = is_redline(p.name)

    tokens = stem.split()
    if not tokens:
        return None
    doc_token = tokens[0]

    revision: str | None = None
    title_tokens: list[str] = tokens[1:]

    # Underscore-embedded revision, e.g. SLQ-4000_A -> doc SLQ-4000, rev A.
    if "_" in doc_token and not is_document_number(doc_token):
        head, _, tail = doc_token.rpartition("_")
        if REV_RE.match(tail.upper()) and is_document_number(head):
            doc_token = head
            revision = tail.upper()

    if revision is None and title_tokens:
        first = title_tokens[0]
        if first.lower() == "rev" and len(title_tokens) >= 2:
            # "Rev C", "Rev D_Risk...", "Rev H_signed" -> take the leading letters
            # (break on whitespace or underscore, since \b treats '_' as a word char).
            m = re.match(r"^([A-Za-z]{1,3})(?=[\s_]|$)", title_tokens[1])
            if m:
                revision = m.group(1).upper()
                title_tokens = title_tokens[2:]
        elif re.fullmatch(r"[A-Za-z]{1,3}", first) and first.isupper():
            # A bare 1–3 letter uppercase token is the revision.
            revision = first.upper()
            title_tokens = title_tokens[1:]

    if not is_document_number(doc_token):
        return None

    doc_number = normalize_doc_number(doc_token)
    title = " ".join(title_tokens).strip()
    return ParsedFile(
        doc_number=doc_number,
        revision=revision,
        title=title,
        ext=ext,
        is_redline=redline,
        raw_stem=stem,
    )


def normalize_title(title: str) -> str:
    """Normalize a title for matching: lowercase, alnum+spaces only, collapsed."""
    t = (title or "").lower()
    t = re.sub(r"[^a-z0-9]+", " ", t)
    return re.sub(r"\s+", " ", t).strip()


# A file named by title with a trailing revision, e.g.
#   "Authorized Approvers Form Rev C.pdf" -> ("Authorized Approvers Form", "C")
#   "Electronic Signature Acknowledgement Form Rev B.pdf" -> (..., "B")
TITLE_REV_RE = re.compile(r"^(?P<title>.+?)[\s_]+Rev\s+(?P<rev>[A-Za-z]{1,3})\b.*$", re.IGNORECASE)


def match_title_named_file(
    filename: str,
    title_index: dict[str, str],
    ambiguous_titles: set[str] | None = None,
) -> tuple[str, str] | None:
    """
    Resolve a title-named file (no leading doc number) to (doc_number, revision)
    using a title->doc_number index built from the register/logs. Requires a
    trailing "Rev X" token and an unambiguous exact title match; otherwise None
    (the caller catalogs it as unmatched rather than force-linking).
    """
    if not filename:
        return None
    stem = Path(filename).stem.strip()
    m = TITLE_REV_RE.match(stem)
    if not m:
        return None
    key = normalize_title(m.group("title"))
    if not key or (ambiguous_titles and key in ambiguous_titles):
        return None
    doc = title_index.get(key)
    if not doc:
        return None
    return doc, m.group("rev").upper()


def classify_doc_type(doc_number: str, title: str) -> str:
    t = (doc_number or "").upper()
    title_l = (title or "").lower()
    if t.startswith("FM"):
        return "Form"
    if t.startswith("TMP"):
        return "Template"
    if t.startswith("TR"):
        return "Traveler"
    if "quality manual" in title_l:
        return "Manual"
    if "quality policy" in title_l:
        return "Policy"
    if "work instruction" in title_l or re.search(r"\bWI\b", title or ""):
        return "WI"
    return "SOP"


def classify_category(doc_number: str, title: str) -> str:
    """
    Best-effort subsystem/category for the Document Control browse UI. This is an
    editable default (surfaced for review), not an authoritative taxonomy.
    """
    s = f"{doc_number} {title}".lower()
    rules = [
        ("Design Control", ["design", "traceability matrix", "v&v", "verification", "validation plan"]),
        ("Risk Management", ["risk", "hazard", "fmea"]),
        ("Supplier & Purchasing", ["supplier", "purchasing", "purchase order", "approved supplier"]),
        ("CAPA & Complaints", ["capa", "complaint", "nonconform", "ncmr", "advisory", "recall", "field safety", "fsca"]),
        ("Production & Process", ["work order", "shipping", "receiving", "process validation", "device master", "workstation", "traveler", "manufacturing"]),
        ("Equipment & Facilities", ["equipment", "calibration", "maintenance", "environmental", "tag out"]),
        ("Post-Market & Regulatory", ["medical device reporting", "emdr", "regulatory", "post-market", "surveillance"]),
        ("Training", ["training"]),
        ("Management & Quality System", ["management review", "internal audit", "quality objective", "organization chart", "quality manual", "quality policy", "quality planning"]),
        ("Document & Records Control", ["document control", "document change", "good documentation", "electronic doc", "electronic signature", "part number", "dhr", "confidential patient", "identification and traceability"]),
        ("Sales & Order Management", ["sales order"]),
    ]
    for cat, kws in rules:
        if any(k in s for k in kws):
            return cat
    return "Quality Management System"


def dco_sort_key(dco_number: str) -> tuple[int, str]:
    m = re.match(r"^(\d+)([A-Z]*)$", (dco_number or "").strip().upper())
    if not m:
        return (10_000, dco_number or "")
    return (int(m.group(1)), m.group(2))


def parse_flexible_date(value) -> date | None:
    if value is None or value == "":
        return None
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    s = str(value).strip()
    if not s or s.lower() in {"new", "n/a", "-", "pending"}:
        return None
    fmts = ["%Y-%m-%d", "%m/%d/%Y", "%b %d %Y", "%d%b%Y", "%B %d %Y", "%b %d, %Y", "%d %b %Y"]
    for f in fmts:
        try:
            return datetime.strptime(s, f).date()
        except ValueError:
            continue
    return None


# ===========================================================================
# I/O — logs, Word docs, file scan
# ===========================================================================
@dataclass
class DcoLine:
    dco_number: str
    line_no: int
    document_number: str
    from_rev: str
    to_rev: str
    document_title: str
    originator: str
    date_requested: date | None
    effective_date: date | None
    impact_assessments: list[str]
    change_description: str = ""
    source: str = "log"
    corrections: list[str] = field(default_factory=list)


IMPACT_COLS = [(11, "Training"), (12, "Risk"), (13, "Validation"), (14, "Regulatory"), (15, "Material Disposition")]


def _clean(v) -> str:
    if v is None:
        return ""
    return str(v).strip()


def parse_dco_log(path: Path) -> list[DcoLine]:
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb["Sheet1"]
    lines: list[DcoLine] = []
    last_dco = ""
    for row in ws.iter_rows(min_row=4, values_only=True):
        row = list(row) + [None] * (16 - len(row))
        dco = _clean(row[0]) or last_dco
        token = _clean(row[2])
        if not token:
            continue
        last_dco = dco
        corrections: list[str] = []
        norm = normalize_doc_number(token)
        if dco == "085" and norm == DCO085_TYPO_TOKEN:
            corrections.append(f"DCO085: '{DCO085_TYPO_TOKEN}' is a typo for {DCO085_TYPO_FIX} (Organization Chart); recorded as {DCO085_TYPO_FIX}.")
            norm = DCO085_TYPO_FIX
        impacts = [name for col, name in IMPACT_COLS if _clean(row[col]).upper().startswith("Y")]
        lines.append(DcoLine(
            dco_number=dco,
            line_no=int(_clean(row[1])) if _clean(row[1]).isdigit() else 0,
            document_number=norm,
            from_rev=_clean(row[3]),
            to_rev=_clean(row[4]),
            document_title=_clean(row[5]),
            originator=_clean(row[6]),
            date_requested=parse_flexible_date(row[7]),
            effective_date=parse_flexible_date(row[10]),
            impact_assessments=impacts,
            corrections=corrections,
        ))
    wb.close()
    return lines


def _docx_find_change_table(doc):
    """Return the table whose header has Document + Revision + Description columns."""
    for tbl in doc.tables:
        if not tbl.rows:
            continue
        header = " | ".join((c.text or "").strip().lower() for c in tbl.rows[0].cells)
        if "document" in header and "revision" in header and ("description" in header or "reason" in header):
            return tbl
    return None


def _docx_originator(doc) -> str:
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            for i, c in enumerate(cells):
                if c.lower().startswith("originator") and i + 1 < len(cells):
                    val = cells[i + 1].strip()
                    if val:
                        return val
    return ""


def extract_dco_docx(path: Path, dco_number: str) -> list[DcoLine]:
    """Reconstruct DCO line rows from a DCO Word doc (used for DCO092–095)."""
    import docx

    doc = docx.Document(str(path))
    tbl = _docx_find_change_table(doc)
    if tbl is None:
        return []
    originator = _docx_originator(doc)
    out: list[DcoLine] = []
    line_no = 0
    for row in tbl.rows[1:]:
        cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
        if len(cells) < 5:
            continue
        title, number, cur_rev, eff, desc = cells[0], cells[1], cells[2], cells[3], cells[4]
        if not number or not is_document_number(number):
            # e.g. a trailing "Note:" row
            continue
        line_no += 1
        norm = normalize_doc_number(number)
        from_rev = "" if cur_rev.strip().lower() in {"new", "n/a", "-", ""} else cur_rev.strip()
        impacts = _heuristic_impacts(doc)
        out.append(DcoLine(
            dco_number=dco_number,
            line_no=line_no,
            document_number=norm,
            from_rev=from_rev or "-",
            to_rev="",  # filled from package scan / next_revision later
            document_title=title,
            originator=originator,
            date_requested=None,
            effective_date=parse_flexible_date(eff),
            impact_assessments=impacts,
            change_description=" ".join(desc.split()),
            source="docx",
        ))
    return out


_NEGATIVE_PHRASES = (
    "no new hazards", "no additional", "no software", "no process validation",
    "not required", "no additional product risk", "no changes affect",
    "no. these", "not applicable", "no additional system validation",
)


def _text_is_negative(text: str) -> bool:
    t = (text or "").strip().lower()
    if not t:
        return True
    if t.startswith(("no", "none", "n/a", "not ")):
        return True
    return any(p in t for p in _NEGATIVE_PHRASES)


def _section_data_cell(doc, header_needle: str) -> str:
    """Return the narrative data cell of the section whose header contains needle."""
    for t in doc.tables:
        if not t.rows:
            continue
        head = " | ".join((c.text or "").strip().lower() for c in t.rows[0].cells)
        if header_needle in head:
            # Prefer the last data row's last non-empty cell.
            for row in reversed(t.rows):
                texts = [(c.text or "").strip() for c in row.cells]
                nonempty = [x for x in texts if x and header_needle not in x.lower()]
                if nonempty:
                    return nonempty[-1]
    return ""


def _heuristic_impacts(doc) -> list[str]:
    """
    Best-effort impact_assessments for DCO092–095 (the Yes/No checkbox glyphs are
    not reliably machine-readable). Derived from the substance of each section's
    narrative cell: an assessment applies only when its section text is present
    and NOT a negative statement (e.g. "No additional risk assessment required").
    Flagged for human review in discrepancies.md.
    """
    impacts: list[str] = []

    # Training: a notification table listing >=1 applicable document.
    for t in doc.tables:
        if t.rows:
            head = " ".join((c.text or "").lower() for c in t.rows[0].cells)
            if "notification type" in head and len(t.rows) > 1:
                impacts.append("Training")
                break

    risk = _section_data_cell(doc, "risk assessment activities")
    if risk and not _text_is_negative(risk):
        impacts.append("Risk")

    vv = _section_data_cell(doc, "v&v activities")
    if vv and not _text_is_negative(vv):
        impacts.append("Validation")

    reg = _section_data_cell(doc, "regulatory assessment")
    if reg and not _text_is_negative(reg):
        impacts.append("Regulatory")

    disp = _section_data_cell(doc, "disposition requirements")
    if disp and not _text_is_negative(disp) and re.search(r"(scrap|use-as-is|use as is|rework)", disp.lower()):
        impacts.append("Material Disposition")

    order = ["Training", "Risk", "Validation", "Regulatory", "Material Disposition"]
    return [x for x in order if x in impacts]


@dataclass
class ScannedFile:
    doc_number: str
    revision: str | None
    title: str
    ext: str
    is_redline: bool
    priority: int
    source_label: str
    path: Path


def build_title_index(lines: list[DcoLine]) -> tuple[dict[str, str], set[str]]:
    """Map normalized document title -> doc_number. Titles mapping to more than
    one distinct doc number are returned as ambiguous (excluded from matching)."""
    index: dict[str, str] = {}
    ambiguous: set[str] = set()
    for ln in lines:
        if not ln.document_title or not is_document_number(ln.document_number):
            continue
        key = normalize_title(ln.document_title)
        if not key:
            continue
        existing = index.get(key)
        if existing and existing != ln.document_number:
            ambiguous.add(key)
        else:
            index[key] = ln.document_number
    for k in ambiguous:
        index.pop(k, None)
    return index, ambiguous


def scan_files(
    root: Path,
    title_index: dict[str, str] | None = None,
    ambiguous_titles: set[str] | None = None,
) -> tuple[list[ScannedFile], list[str]]:
    """Scan controlled-document folders. Returns (files, unmatched_title_named).

    Files whose leading token is a doc number are parsed directly. Files named by
    title with a trailing "Rev X" are linked via the title index when unambiguous;
    otherwise they are reported as unmatched for review.
    """
    title_index = title_index or {}
    out: list[ScannedFile] = []
    unmatched: list[str] = []
    for rel, priority, label in SCAN_ROOTS:
        base = root / rel
        if not base.exists():
            continue
        for fp in sorted(base.rglob("*")):
            if not fp.is_file():
                continue
            parsed = parse_filename(fp.name)
            if parsed is not None:
                out.append(ScannedFile(
                    doc_number=parsed.doc_number,
                    revision=parsed.revision,
                    title=parsed.title,
                    ext=parsed.ext,
                    is_redline=parsed.is_redline,
                    priority=priority,
                    source_label=label,
                    path=fp,
                ))
                continue
            # Fallback: title-named file with a trailing "Rev X".
            matched = match_title_named_file(fp.name, title_index, ambiguous_titles)
            if matched:
                doc_number, rev = matched
                out.append(ScannedFile(
                    doc_number=doc_number,
                    revision=rev,
                    title=fp.stem,
                    ext=fp.suffix.lower().lstrip("."),
                    is_redline=is_redline(fp.name),
                    priority=priority,
                    source_label=f"{label} (title-linked)",
                    path=fp,
                ))
            elif TITLE_REV_RE.match(fp.stem):
                unmatched.append(_rel(fp))
    return out, unmatched


# ===========================================================================
# Reconciliation
# ===========================================================================
MASTER_EXTS = ["docx", "xlsx", "doc", "xls"]  # preferred released-file formats
EVIDENCE_EXTS = ["pdf"]


@dataclass
class RevisionRecord:
    revision: str
    file: ScannedFile | None
    effective_date: date | None = None
    change_summary: str = ""


@dataclass
class DocRecord:
    doc_number: str
    title: str
    doc_type: str
    category: str
    destination: str  # "document_control" | "records"
    current_rev: str | None
    current_source: str
    current_file: ScannedFile | None
    revisions: list[RevisionRecord]  # ascending, only revs with a physical file
    status: str  # Released | Obsolete | Draft
    obsolete_reason: str = ""
    flags: list[str] = field(default_factory=list)


def _pick_best_file(files: list[ScannedFile]) -> ScannedFile | None:
    """Prefer master formats, then lowest source priority, then stable path order."""
    if not files:
        return None
    def key(f: ScannedFile):
        try:
            ext_rank = MASTER_EXTS.index(f.ext)
        except ValueError:
            ext_rank = len(MASTER_EXTS) + (0 if f.ext in EVIDENCE_EXTS else 1)
        return (f.priority, ext_rank, str(f.path))
    return sorted(files, key=key)[0]


def fill_docx_to_rev(docx_lines: list[DcoLine], scanned: list[ScannedFile]) -> None:
    """Populate to_rev for DCO092–095 rows from the package master revision."""
    from app.eqms.modules.document_control.service import next_revision

    pkg_rev: dict[tuple[str, str], str] = {}
    for f in scanned:
        if f.is_redline or not f.revision:
            continue
        for part in f.path.parts:
            m = re.match(r"^DCO(\d{3})$", part)
            if m:
                key = (m.group(1), f.doc_number)
                if rev_to_int(f.revision) > rev_to_int(pkg_rev.get(key)):
                    pkg_rev[key] = f.revision.upper()
    for ln in docx_lines:
        if ln.to_rev:
            continue
        key = (ln.dco_number, ln.document_number)
        if key in pkg_rev:
            ln.to_rev = pkg_rev[key]
        elif ln.from_rev in {"-", ""}:
            ln.to_rev = "A"
        else:
            try:
                ln.to_rev = next_revision(ln.from_rev)
            except ValueError:
                ln.to_rev = ""


def reconcile(
    scanned: list[ScannedFile],
    dco_lines: list[DcoLine],
    docx_lines: list[DcoLine],
) -> list[DocRecord]:
    # Highest to_rev known from the change logs per doc number.
    log_rev: dict[str, str] = {}
    log_eff: dict[tuple[str, str], date | None] = {}
    log_desc: dict[tuple[str, str], str] = {}
    for ln in dco_lines + docx_lines:
        if not ln.document_number or not is_document_number(ln.document_number):
            continue
        if rev_to_int(ln.to_rev) > rev_to_int(log_rev.get(ln.document_number)):
            log_rev[ln.document_number] = ln.to_rev
        if ln.to_rev:
            log_eff[(ln.document_number, ln.to_rev.upper())] = ln.effective_date
            if ln.change_description:
                log_desc[(ln.document_number, ln.to_rev.upper())] = ln.change_description

    # Group non-redline files by doc number; track redline revisions separately.
    by_doc: dict[str, list[ScannedFile]] = {}
    redline_revs: dict[str, set[str]] = {}
    for f in scanned:
        if (f.doc_number, (f.revision or "").upper()) in MISPLACED_ARTIFACTS:
            continue  # misplaced artifact confirmed by Ethan; ignore entirely
        if f.is_redline:
            if f.revision:
                redline_revs.setdefault(f.doc_number, set()).add(f.revision.upper())
            continue
        by_doc.setdefault(f.doc_number, []).append(f)

    all_docs = set(by_doc) | set(log_rev)
    records: list[DocRecord] = []

    for doc_number in sorted(all_docs):
        files = by_doc.get(doc_number, [])
        flags: list[str] = []

        # QM.SLQ018 defensive correction: treat a stray 'C' master as 'B'.
        if doc_number == "QM.SLQ018":
            for f in files:
                if (f.revision or "").upper() == "C":
                    flags.append("QM.SLQ018: on-disk file labeled 'C' remapped to 'B' per DCO093 correction; rename source master to 'B'.")
                    f.revision = "B"

        file_revs = {(f.revision or "").upper() for f in files if f.revision}
        max_file_rev = max((rev_to_int(r) for r in file_revs), default=0)
        max_log_rev = rev_to_int(log_rev.get(doc_number))
        current_int = max(max_file_rev, max_log_rev)

        current_rev = None
        for r in file_revs | {log_rev.get(doc_number, "")}:
            if rev_to_int(r) == current_int and r:
                current_rev = r.upper()
                break

        # Choose current file (files at current rev).
        cur_files = [f for f in files if (f.revision or "").upper() == current_rev] if current_rev else []
        current_file = _pick_best_file(cur_files)

        # Inference: an unlabeled "New" form/spec whose only master lacks a rev in
        # its filename, but the log gives a current rev -> treat that file as current.
        unlabeled = [f for f in files if not f.revision]
        if current_rev and not current_file and unlabeled and max_file_rev == 0:
            current_file = _pick_best_file(unlabeled)
            flags.append(f"Filename omits a revision; inferred current rev {current_rev} from the DCO log.")
            unlabeled = [f for f in unlabeled if f is not current_file]

        current_source = current_file.source_label if current_file else "(none)"

        if current_rev and not current_file:
            flags.append(f"Current revision {current_rev} known from the DCO log but NO physical file located.")
        if not files and current_rev:
            flags.append("No physical file found in any scanned folder for this document.")
        for f in unlabeled:
            flags.append(f"File has no parseable revision: {f.path.name}")

        # Released master missing but a redline exists at a higher revision.
        rl = redline_revs.get(doc_number, set())
        higher_redline = [r for r in rl if rev_to_int(r) > current_int]
        for r in sorted(higher_redline, key=rev_to_int):
            flags.append(f"Rev {r} exists only as a redline (working artifact); released master is missing.")

        # Title / type / category.
        title = current_file.title if current_file and current_file.title else ""
        if not title:
            for ln in docx_lines + dco_lines:
                if ln.document_number == doc_number and ln.document_title:
                    title = ln.document_title
                    break
        doc_type = classify_doc_type(doc_number, title)
        category = classify_category(doc_number, title)
        destination = "document_control" if is_document_control(doc_number) else "records"

        # Build ascending revision records that have a physical (non-redline) file.
        revs_present = sorted(
            {(f.revision or "").upper() for f in files if f.revision},
            key=rev_to_int,
        )
        rev_records: list[RevisionRecord] = []
        for r in revs_present:
            rfiles = [f for f in files if (f.revision or "").upper() == r]
            best = _pick_best_file(rfiles)
            rev_records.append(RevisionRecord(
                revision=r,
                file=best,
                effective_date=log_eff.get((doc_number, r)),
                change_summary=log_desc.get((doc_number, r), ""),
            ))
        # Inferred current file (unlabeled filename) — attach to the current rev.
        if current_rev and current_file and current_rev not in {rr.revision for rr in rev_records}:
            if (current_file.revision or "").upper() != current_rev:
                rev_records.append(RevisionRecord(
                    revision=current_rev,
                    file=current_file,
                    effective_date=log_eff.get((doc_number, current_rev)),
                    change_summary=log_desc.get((doc_number, current_rev), ""),
                ))
                rev_records.sort(key=lambda rr: rev_to_int(rr.revision))

        # Status / obsolete.
        status = "Released" if current_rev else "Draft"
        obsolete_reason = ""
        base_num = _base_number(doc_number)
        if base_num in OBSOLETE_BASE_NUMBERS and is_document_control(doc_number):
            status = "Obsolete"
            obsolete_reason = OBSOLETE_REASON
            flags.append("Marked OBSOLETE by DCO095 (QM.SLQ004–010 family superseded by QM.SLQ052).")

        # Superseded revisions with no file are worth noting for history gaps.
        if current_rev and rev_to_int(current_rev) > 1:
            have = {rr.revision for rr in rev_records}
            for i in range(1, rev_to_int(current_rev)):
                lbl = _int_to_rev(i)
                if lbl not in have:
                    flags.append(f"No physical file for prior revision {lbl} (history gap).")

        records.append(DocRecord(
            doc_number=doc_number,
            title=title,
            doc_type=doc_type,
            category=category,
            destination=destination,
            current_rev=current_rev,
            current_source=current_source,
            current_file=current_file,
            revisions=rev_records,
            status=status,
            obsolete_reason=obsolete_reason,
            flags=flags,
        ))
    return records


def _base_number(doc_number: str) -> int | None:
    m = re.search(r"QM\.SLQ(\d+)$", (doc_number or "").upper())
    return int(m.group(1)) if m else None


def _int_to_rev(n: int) -> str:
    out = []
    while n > 0:
        n -= 1
        out.append(chr(ord("A") + n % 26))
        n //= 26
    return "".join(reversed(out))


def divergences_vs_expected(records: list[DocRecord]) -> list[str]:
    out: list[str] = []
    got = {r.doc_number: (r.current_rev or "-") for r in records}
    for doc, exp in sorted(EXPECTED_CURRENT.items()):
        actual = got.get(doc)
        if actual is None:
            out.append(f"{doc}: expected current {exp} but the document was not found on disk/logs.")
        elif actual != exp:
            out.append(f"{doc}: computed current {actual} != coordinator's §4 expectation {exp}.")
    return out


# ===========================================================================
# Emitters
# ===========================================================================
def _rel(p: Path) -> str:
    try:
        return str(p.relative_to(ROOT)).replace("\\", "/")
    except ValueError:
        return str(p).replace("\\", "/")


def write_register_csv(records: list[DocRecord], out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", newline="", encoding="utf-8") as fh:
        w = csv.writer(fh)
        w.writerow([
            "document_number", "title", "doc_type", "destination", "current_rev",
            "current_source_dco", "current_file_path", "superseded_revs",
            "status", "category", "flags",
        ])
        for r in sorted(records, key=lambda x: x.doc_number):
            superseded = [rr.revision for rr in r.revisions if rr.revision != r.current_rev]
            w.writerow([
                r.doc_number, r.title, r.doc_type, r.destination, r.current_rev or "",
                r.current_source, _rel(r.current_file.path) if r.current_file else "",
                ";".join(superseded), r.status, r.category, " | ".join(r.flags),
            ])


def write_manifests(records: list[DocRecord], out_dir: Path, staging_doc_root: Path) -> int:
    out_dir.mkdir(parents=True, exist_ok=True)
    count = 0
    for r in records:
        if r.destination != "document_control":
            continue
        if not r.revisions:
            continue
        revisions = []
        for rr in r.revisions:
            if not rr.file:
                continue
            staged_rel = _staged_relpath(r, rr, staging_doc_root)
            revisions.append({
                "revision": rr.revision,
                "file": staged_rel,
                "effective_date": rr.effective_date.isoformat() if rr.effective_date else None,
                "change_summary": rr.change_summary,
                "released": True,
            })
        if not revisions:
            continue
        manifest = {
            "doc_number": r.doc_number,
            "title": r.title,
            "doc_type": r.doc_type,
            "category": r.category,
            "status": r.status,
            "revisions": revisions,
        }
        if r.status == "Obsolete":
            manifest["obsolete_reason"] = r.obsolete_reason
        safe = r.doc_number.replace("/", "_")
        (out_dir / f"{safe}.json").write_text(
            json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8"
        )
        count += 1
    return count


def _staged_relpath(r: DocRecord, rr: RevisionRecord, staging_doc_root: Path) -> str:
    sub = "current" if rr.revision == r.current_rev else "superseded"
    fname = rr.file.path.name
    return f"{r.doc_number}/{sub}/{fname}"


def write_discrepancies(
    records: list[DocRecord],
    divergences: list[str],
    dco_corrections: list[str],
    out_path: Path,
    docx_dco_reconstructed: list[str],
    title_linked: list[str] | None = None,
    unmatched_title_files: list[str] | None = None,
) -> None:
    lines: list[str] = []
    lines.append("# Controlled-Document Reconciliation — Discrepancies & Flags")
    lines.append("")
    lines.append("_Report-only output. No documents were imported or modified. "
                 "Source files were copied (never moved/deleted) into `eQMS_Upload_Staging/`._")
    lines.append("")

    lines.append("## 1. Divergence from the coordinator's §4 reconciliation")
    if divergences:
        for d in divergences:
            lines.append(f"- {d}")
    else:
        lines.append("- None. Every §4 expected current revision matches the computed result.")
    lines.append("")

    lines.append("## 2. Baked-in corrections (§7/§8)")
    if dco_corrections:
        for c in sorted(set(dco_corrections)):
            lines.append(f"- {c}")
    else:
        lines.append("- None recorded.")
    lines.append("- QM.SLQ018: DCO093 master is expected to be Rev **B** (a stray 'C' label is a mistake). "
                 "If a 'C' file is found it is remapped to 'B'; otherwise no action needed.")
    lines.append("- FM1-QM.SLQ018 stays at **Rev A** (Ethan): the mistakenly-placed DCO093 'B' artifact has "
                 "been removed from disk and is ignored here; there is no Rev B.")
    lines.append("- FM1-QM.SLQ016 → **B** is real; the corrected released B master is now on disk and links.")
    lines.append("")

    lines.append("## 2b. Title-named files linked by title→doc-number matcher")
    if title_linked:
        for t in sorted(title_linked):
            lines.append(f"- {t}")
    else:
        lines.append("- None.")
    if unmatched_title_files:
        lines.append("")
        lines.append("**Unmatched title-named files (contain a 'Rev X' but no confident doc-number match — review, not force-linked):**")
        for t in sorted(unmatched_title_files):
            lines.append(f"- {t}")
    lines.append("")

    lines.append("## 3. DCO092–095 (not in the legacy DCO log)")
    lines.append("- Treated as **released** per Ethan and reconstructed into `DCO_Log_v2.csv` from the "
                 "`QMSInProcess\\DCO09x\\DCO09x.docx` change tables.")
    for d in docx_dco_reconstructed:
        lines.append(f"  - {d}")
    lines.append("- Impact-assessment columns for these DCOs were derived **heuristically** from the Word "
                 "narrative (checkbox glyphs are not reliably machine-readable); confirm before finalizing.")
    lines.append("- The legacy `SILQ DCO Log.xlsx` and `SILQ Document Number Log.xlsx` remain out of date "
                 "(Ethan's QA action, not automated here).")
    lines.append("")

    # Per-document flags.
    lines.append("## 4. Missing files / history gaps / per-document flags")
    any_flag = False
    for r in sorted(records, key=lambda x: x.doc_number):
        if r.flags:
            any_flag = True
            lines.append(f"- **{r.doc_number}** (current {r.current_rev or '—'}, {r.status}):")
            for f in r.flags:
                lines.append(f"  - {f}")
    if not any_flag:
        lines.append("- None.")
    lines.append("")

    # Documents with no locatable current file.
    lines.append("## 5. Documents with no locatable current file")
    nofile = [r for r in records if r.current_rev and not r.current_file]
    if nofile:
        for r in nofile:
            lines.append(f"- {r.doc_number} — current {r.current_rev} known from logs, no file on disk.")
    else:
        lines.append("- None.")
    lines.append("")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def build_dco_log_v2_rows(
    dco_lines: list[DcoLine],
    docx_lines: list[DcoLine],
    scanned: list[ScannedFile],
) -> list[dict]:
    # to_rev for DCO092–095 rows is filled earlier via fill_docx_to_rev().
    rows: list[dict] = []
    for ln in dco_lines + docx_lines:
        rows.append({
            "dco_number": ln.dco_number,
            "document_number": ln.document_number,
            "document_title": ln.document_title,
            "from_rev": ln.from_rev or "-",
            "to_rev": ln.to_rev,
            "change_description": ln.change_description,
            "originator": ln.originator,
            "date_requested": ln.date_requested.isoformat() if ln.date_requested else "",
            "effective_date": ln.effective_date.isoformat() if ln.effective_date else "",
            "impact_assessments": ";".join(ln.impact_assessments),
        })
    rows.sort(key=lambda r: (dco_sort_key(r["dco_number"]), r["document_number"]))
    return rows


def write_dco_log_v2(rows: list[dict], out_path: Path) -> None:
    cols = [
        "dco_number", "document_number", "document_title", "from_rev", "to_rev",
        "change_description", "originator", "date_requested", "effective_date",
        "impact_assessments",
    ]
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=cols)
        w.writeheader()
        for r in rows:
            w.writerow(r)


# ===========================================================================
# Staging assembly (copy only; never move/delete)
# ===========================================================================
def _force_remove(func, path, _exc):
    """rmtree onerror handler: clear read-only then retry; ignore if still locked
    (e.g. OneDrive sync locks). copy2 overwrites, so leftovers are harmless."""
    import os
    import stat
    try:
        os.chmod(path, stat.S_IWRITE)
        func(path)
    except Exception:
        pass


def assemble_staging(records: list[DocRecord], staging_root: Path) -> dict:
    doc_root = staging_root / "document_control"
    rec_root = staging_root / "records"
    # Best-effort clear of the binary subtrees so re-runs are deterministic (never
    # touches the committable reconciliation/ outputs).
    for d in (doc_root, rec_root):
        if d.exists():
            shutil.rmtree(d, onerror=_force_remove)
    stats = {"current": 0, "superseded": 0, "records": 0}
    for r in records:
        if r.destination == "document_control":
            base = doc_root / r.doc_number
            for rr in r.revisions:
                if not rr.file:
                    continue
                sub = "current" if rr.revision == r.current_rev else "superseded"
                dest = base / sub
                dest.mkdir(parents=True, exist_ok=True)
                shutil.copy2(rr.file.path, dest / rr.file.path.name)
                stats["current" if sub == "current" else "superseded"] += 1
        else:
            base = rec_root / r.doc_number
            for rr in r.revisions:
                if not rr.file:
                    continue
                base.mkdir(parents=True, exist_ok=True)
                shutil.copy2(rr.file.path, base / rr.file.path.name)
                stats["records"] += 1
    return stats


def ensure_gitignore(staging_root: Path) -> None:
    """Ignore staged binaries but keep the reconciliation/ text outputs committable."""
    gi = ROOT / ".gitignore"
    marker = "# eQMS document staging (binaries not committed)"
    block = "\n".join([
        marker,
        "eQMS_Upload_Staging/document_control/",
        "eQMS_Upload_Staging/records/",
        "!eQMS_Upload_Staging/reconciliation/",
        "",
    ])
    existing = gi.read_text(encoding="utf-8") if gi.exists() else ""
    if marker not in existing:
        with gi.open("a", encoding="utf-8") as fh:
            fh.write(("\n" if existing and not existing.endswith("\n") else "") + block)


# ===========================================================================
# Main
# ===========================================================================
def run(root: Path, staging_root: Path, do_copy: bool) -> dict:
    dco_lines = parse_dco_log(root / DCO_LOG_XLSX)

    # DCO092–095 from the QMSInProcess Word docs.
    docx_lines: list[DcoLine] = []
    reconstructed: list[str] = []
    for n in range(92, 96):
        dco = f"{n:03d}"
        p = root / QMS_INPROCESS_DIR / f"DCO{dco}" / f"DCO{dco}.docx"
        if p.exists():
            rows = extract_dco_docx(p, dco)
            docx_lines.extend(rows)
            reconstructed.append(f"DCO{dco}: {len(rows)} line item(s) reconstructed from {p.name}.")

    # Best-effort change descriptions for legacy DCOs from CompletedDCOs Word docs.
    completed_dir = root / COMPLETED_DCOS_DIR
    desc_by_key: dict[tuple[str, str], str] = {}
    if completed_dir.exists():
        for fp in sorted(completed_dir.glob("DCO*.docx")):
            m = re.match(r"^DCO(\d{3})", fp.name)
            if not m:
                continue
            dco = m.group(1)
            try:
                import docx as _docx
                d = _docx.Document(str(fp))
                tbl = _docx_find_change_table(d)
                if tbl is None:
                    continue
                for row in tbl.rows[1:]:
                    cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
                    if len(cells) < 5 or not is_document_number(cells[1]):
                        continue
                    desc_by_key[(dco, normalize_doc_number(cells[1]))] = " ".join(cells[4].split())
            except Exception:
                continue
    for ln in dco_lines:
        key = (ln.dco_number, ln.document_number)
        if not ln.change_description and key in desc_by_key:
            ln.change_description = desc_by_key[key]

    title_index, ambiguous_titles = build_title_index(dco_lines + docx_lines)
    scanned, unmatched_title_files = scan_files(root, title_index, ambiguous_titles)
    title_linked = [_rel(f.path) + f"  ->  {f.doc_number} Rev {f.revision}"
                    for f in scanned if "title-linked" in f.source_label]
    fill_docx_to_rev(docx_lines, scanned)
    records = reconcile(scanned, dco_lines, docx_lines)
    divergences = divergences_vs_expected(records)

    recon_dir = staging_root / "reconciliation"
    write_register_csv(records, recon_dir / "document_register_v2.csv")
    n_manifests = write_manifests(records, recon_dir / "manifests", staging_root / "document_control")

    corrections = [c for ln in dco_lines for c in ln.corrections]
    write_discrepancies(
        records, divergences, corrections, recon_dir / "discrepancies.md",
        reconstructed, title_linked, unmatched_title_files,
    )

    dco_rows = build_dco_log_v2_rows(dco_lines, docx_lines, scanned)
    write_dco_log_v2(dco_rows, recon_dir / "DCO_Log_v2.csv")

    copy_stats = {"current": 0, "superseded": 0, "records": 0}
    if do_copy:
        copy_stats = assemble_staging(records, staging_root)
        ensure_gitignore(staging_root)

    dc = [r for r in records if r.destination == "document_control"]
    rec = [r for r in records if r.destination == "records"]
    changed_vs_stale = _count_changed_vs_stale(dc, root)

    return {
        "records_total": len(records),
        "document_control": len(dc),
        "records_dest": len(rec),
        "manifests": n_manifests,
        "dco_rows": len(dco_rows),
        "dco_reconstructed": len(docx_lines),
        "divergences": len(divergences),
        "changed_vs_stale": changed_vs_stale,
        "copy_stats": copy_stats,
        "obsolete": len([r for r in records if r.status == "Obsolete"]),
        "no_current_file": len([r for r in records if r.current_rev and not r.current_file]),
        "title_linked": len(title_linked),
        "unmatched_title": len(unmatched_title_files),
    }


def _count_changed_vs_stale(dc_records: list[DocRecord], root: Path) -> int:
    """How many current revs differ from the stale QM Documents clean folder."""
    clean = {}
    base = root / "QM Documents"
    if base.exists():
        for fp in base.glob("*"):
            pf = parse_filename(fp.name)
            if pf and pf.revision:
                clean[pf.doc_number] = pf.revision.upper()
    changed = 0
    for r in dc_records:
        if r.doc_number in clean and r.current_rev and r.current_rev != clean[r.doc_number]:
            changed += 1
    return changed


def main(argv=None):
    parser = argparse.ArgumentParser(description="Reconcile controlled documents (report-only).")
    parser.add_argument("--staging-root", default="eQMS_Upload_Staging", help="Output staging directory")
    parser.add_argument("--no-copy", action="store_true", help="Emit text artifacts only; skip copying binaries")
    args = parser.parse_args(argv)

    staging_root = Path(args.staging_root)
    if not staging_root.is_absolute():
        staging_root = ROOT / staging_root

    stats = run(ROOT, staging_root, do_copy=not args.no_copy)

    print("Controlled-document reconciliation complete (report-only).")
    print(f"  Total documents reconciled : {stats['records_total']}")
    print(f"    -> Document Control      : {stats['document_control']}")
    print(f"    -> Records (admin_docs)   : {stats['records_dest']}")
    print(f"  Importer manifests written : {stats['manifests']}")
    print(f"  Current-rev changes vs stale QM Documents: {stats['changed_vs_stale']}")
    print(f"  Obsolete documents         : {stats['obsolete']}")
    print(f"  §4 divergences flagged     : {stats['divergences']}")
    print(f"  DCO_Log_v2 rows            : {stats['dco_rows']} "
          f"({stats['dco_reconstructed']} reconstructed from DCO092–095 Word docs)")
    print(f"  Docs with no current file  : {stats['no_current_file']}")
    print(f"  Title-linked files         : {stats['title_linked']} (unmatched title-named: {stats['unmatched_title']})")
    if not args.no_copy:
        cs = stats["copy_stats"]
        print(f"  Staged files (copied)      : current={cs['current']} superseded={cs['superseded']} records={cs['records']}")
    print(f"  Artifacts under            : {_rel(staging_root)}/reconciliation/")


if __name__ == "__main__":
    main()
