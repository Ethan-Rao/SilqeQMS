"""
Refresh readable-text (.md) extractions for the DC.SLQ002 deliverables in
QMSInProcess/DC.SLQ002/ into docs/QMS-Readable-Texts/12-DHF-Software/.

Also mirrors the project plan and executed scope form into
docs/QMS-Readable-Texts/20-QMSInProcess/DC.SLQ002/ so that folder tracks the
same QMSInProcess sources.

Phase 0 design-control records in the same folder (DCO090, design review minutes
and slides, FM1-QM.SLQ008 Phase 0) are extracted into
docs/QMS-Readable-Texts/20-QMSInProcess/DC.SLQ002/ only.

Docx outputs match the existing SilqQMS readable texts in that folder:
  - Plain paragraph text (no markdown heading prefixes)
  - Tables emitted as `[Table]` markers followed by markdown pipe tables
  - No source-attribution header

PDF output uses a short title/source header and per-page text (pdfplumber).

Usage:
    python scripts/refresh_dc_slq002_readable_texts.py
"""

from __future__ import annotations

import os
import shutil
import subprocess
import tempfile
from pathlib import Path

import pdfplumber
from docx import Document


def _copy_with_fallback(src: Path, dst: Path) -> None:
    """Copy ``src`` to ``dst``, falling back to PowerShell ``Copy-Item`` when
    Python's stdlib gets PermissionError. Some files held by OneDrive or
    Office processes deny direct ``open()`` but allow backup-style copies."""
    try:
        shutil.copyfile(str(src), str(dst))
        return
    except PermissionError:
        pass
    cmd = [
        "powershell",
        "-NoProfile",
        "-Command",
        f"Copy-Item -LiteralPath '{src}' -Destination '{dst}' -Force",
    ]
    subprocess.run(cmd, check=True)

ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "QMSInProcess" / "DC.SLQ002"
OUTPUT_DIR = ROOT / "docs" / "QMS-Readable-Texts" / "12-DHF-Software"
MIRROR_DIR = ROOT / "docs" / "QMS-Readable-Texts" / "20-QMSInProcess" / "DC.SLQ002"

# Filenames written under OUTPUT_DIR that are also copied to MIRROR_DIR.
MIRROR_NAMES: frozenset[str] = frozenset(
    {
        "DC.SLQ002 A Design Project Plan, SilqQMS EDMS Transition.md",
        "Executed Design Project Scope Form DC.SLQ002 BM_CT signed.md",
    }
)

# Map source docx file (basename) -> output md file (basename).
# Output names match the existing convention in the readable-texts folder.
FILE_MAP: dict[str, str] = {
    "DC.SLQ002 A Design Project Plan, SilqQMS EDMS Transition.docx":
        "DC.SLQ002 A Design Project Plan, SilqQMS EDMS Transition.md",
    "SW.SLQ007 A Software Validation Plan, SilqQMS.docx":
        "SW.SLQ007 A Software Validation Plan, SilqQMS.md",
    # Note: source docx has a literal '^J' (ASCII 0x0A) in the filename in place
    # of a comma+space; resolve at runtime by glob if exact-match fails.
    "SW.SLQ008 A Product Requirements Specification, SilqQMS.docx":
        "SW.SLQ008 A Product Requirements Specification, SilqQMS.md",
    "SW.SLQ009 A Software Verification Test Plan, SIlqQMS.docx":
        "SW.SLQ009 A Software Verification Test Plan, SIlqQMS.md",
    # Source docx filenames in QMSInProcess/DC.SLQ002/ use the product name
    # "Silq eQMS"; readable-text outputs keep the historical "SilqQMS" suffix
    # for stable links from prompts and editing guides.
    "SW.SLQ010 A Software Verification Test Procedure Silq eQMS.docx":
        "SW.SLQ010 A Software Verification Test Procedure, SilqQMS.md",
    "SW.SLQ011 A Software Validation Report Silq eQMS.docx":
        "SW.SLQ011 A Software Validation Report, SilqQMS.md",
    "SW.SLQ012 A Requirements Traceability Matrix Silq eQMS.docx":
        "SW.SLQ012 A Requirements Traceability Matrix, SilqQMS.md",
}

# Executed scope form is supplied as PDF (signed); extract to the same basename
# as the existing readable text in 12-DHF-Software.
PDF_MAP: dict[str, str] = {
    # Current signed PDF in QMSInProcess/DC.SLQ002/ (BM_CT variant may exist in
    # archives; readable output filename is kept for stable repo links).
    "Executed Design Project Scope Form DC.SLQ002.pdf":
        "Executed Design Project Scope Form DC.SLQ002 BM_CT signed.md",
}

# Phase 0 design-control records (same DC.SLQ002 folder) — readable texts live
# under 20-QMSInProcess/DC.SLQ002/ only (not duplicated into 12-DHF-Software).
PHASE0_DOCX: dict[str, str] = {
    "FM1-QM.SLQ008 A Design Review Meeting Minutes Form DC.SLQ002 Phase 0.docx":
        "FM1-QM.SLQ008 A Design Review Meeting Minutes Form DC.SLQ002 Phase 0.md",
}
PHASE0_PDF: dict[str, str] = {
    "Silq Design Review Meeting Minutes DC.SLQ002 - SilqQMS EDMS Transition Phase 0.pdf":
        "Silq Design Review Meeting Minutes DC.SLQ002 - SilqQMS EDMS Transition Phase 0.md",
    "Silq Design Review Meeting Slides DC.SLQ002 - SilqQMS EDMS Transition Phase 0.pdf":
        "Silq Design Review Meeting Slides DC.SLQ002 - SilqQMS EDMS Transition Phase 0.md",
    "DCO090.pdf": "DCO090.md",
}


def extract_docx_to_readable_text(docx_path: Path) -> str:
    """
    Walk the docx body in document order, emitting paragraphs as plain text
    (one per line) and tables as `[Table]\\n<markdown pipe table>`.

    Empty paragraphs are preserved (one blank line each) so that the cover-page
    layout from the source docx is roughly retained. Consecutive blank lines
    are collapsed to at most two for readability.

    The source file is first copied to a temp path before opening so that
    OneDrive cloud-only state or an existing Word lock does not cause a
    PackageNotFound / PermissionError.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = Path(tmp.name)
    try:
        _copy_with_fallback(docx_path, tmp_path)
        doc = Document(str(tmp_path))
        return _walk_doc(doc)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _walk_doc(doc: Document) -> str:
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)

    out_lines: list[str] = []
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            try:
                para = next(para_iter)
            except StopIteration:
                continue
            text = para.text.rstrip()
            out_lines.append(text)
        elif tag == "tbl":
            try:
                table = next(table_iter)
            except StopIteration:
                continue
            rows: list[list[str]] = []
            for row in table.rows:
                cells = [c.text.replace("\n", " ").strip() for c in row.cells]
                rows.append(cells)
            if not rows:
                continue
            mc = max(len(r) for r in rows)
            for r in rows:
                while len(r) < mc:
                    r.append("")
            md = ["[Table]"]
            md.append("| " + " | ".join(rows[0]) + " |")
            md.append("| " + " | ".join(["---"] * mc) + " |")
            for r in rows[1:]:
                md.append("| " + " | ".join(r) + " |")
            out_lines.append("")
            out_lines.append("\n".join(md))
            out_lines.append("")

    # Normalize whitespace: collapse 3+ consecutive blank lines to 2.
    normalized: list[str] = []
    blank_run = 0
    for line in out_lines:
        if line == "":
            blank_run += 1
            if blank_run > 2:
                continue
        else:
            blank_run = 0
        normalized.append(line)

    return "\n".join(normalized).rstrip() + "\n"


def find_source_file(filename: str) -> Path | None:
    direct = SOURCE_DIR / filename
    if direct.exists():
        return direct
    # Tolerate filename quirks (e.g. embedded ^J newline char that the shell
    # listed for SW.SLQ008): glob by simplified pattern.
    stem = filename.rsplit(".", 1)[0]
    needle = stem.split(",")[0].strip().lower()
    for cand in SOURCE_DIR.iterdir():
        if cand.is_file() and cand.suffix.lower() == ".docx":
            if needle in cand.name.lower():
                return cand
    return None


def find_pdf_file(filename: str) -> Path | None:
    direct = SOURCE_DIR / filename
    if direct.exists():
        return direct
    stem = filename.rsplit(".", 1)[0]
    needle = stem.split(",")[0].strip().lower()
    # Normalize hyphen variants (ASCII hyphen, en dash, em dash) in filenames.
    needle_norm = (
        needle.replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2212", "-")
    )
    exact_lower = stem.lower()
    matches: list[Path] = []
    for cand in SOURCE_DIR.iterdir():
        if cand.is_file() and cand.suffix.lower() == ".pdf":
            cand_stem = cand.stem.lower().replace("\u2013", "-").replace("\u2014", "-").replace("\u2212", "-")
            if cand_stem == exact_lower.replace("\u2013", "-").replace("\u2014", "-").replace("\u2212", "-"):
                return cand
            cn = cand.name.lower().replace("\u2013", "-").replace("\u2014", "-").replace("\u2212", "-")
            if needle_norm in cn or needle in cand.name.lower():
                matches.append(cand)
    if not matches:
        return None
    if len(matches) == 1:
        return matches[0]
    matches.sort(key=lambda p: (len(p.stem), p.name.lower()))
    return matches[0]


def extract_pdf_to_readable_text(pdf_path: Path, source_rel: str) -> str:
    """Readable markdown-style text with page markers (matches prior exports)."""
    title = pdf_path.stem
    lines: list[str] = [
        f"# {title}",
        "",
        f"**Source:** `{source_rel}`",
        "",
        "---",
        "",
    ]
    with pdfplumber.open(str(pdf_path)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            lines.append(f"--- Page {i} ---")
            raw = page.extract_text()
            if raw:
                lines.append(raw.strip())
            lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def write_output_md(out_path: Path, text: str) -> None:
    out_path.write_text(text, encoding="utf-8")
    if out_path.name in MIRROR_NAMES:
        MIRROR_DIR.mkdir(parents=True, exist_ok=True)
        mirror_path = MIRROR_DIR / out_path.name
        mirror_path.write_text(text, encoding="utf-8")
        print(f"  MR   {mirror_path.relative_to(ROOT)}")


def write_mirror_only(out_path: Path, text: str) -> None:
    """Write a readable text only under 20-QMSInProcess/DC.SLQ002/."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(text, encoding="utf-8")
    print(f"  MR   {out_path.relative_to(ROOT)}")


def main() -> int:
    if not SOURCE_DIR.exists():
        print(f"Source dir not found: {SOURCE_DIR}")
        return 1
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print(f"Refreshing readable texts in: {OUTPUT_DIR}")
    print(f"From source: {SOURCE_DIR}")
    print("-" * 60)

    failures = 0
    for src_name, out_name in FILE_MAP.items():
        src = find_source_file(src_name)
        if not src:
            print(f"  MISSING source: {src_name}")
            failures += 1
            continue
        out_path = OUTPUT_DIR / out_name
        try:
            text = extract_docx_to_readable_text(src)
            write_output_md(out_path, text)
            print(f"  OK   {src.name} -> {out_path.name}  ({len(text):,} chars)")
        except Exception as e:  # pragma: no cover - reported to caller
            print(f"  FAIL {src.name}: {e}")
            failures += 1

    for src_name, out_name in PDF_MAP.items():
        src = find_pdf_file(src_name)
        if not src:
            print(f"  MISSING source: {src_name}")
            failures += 1
            continue
        out_path = OUTPUT_DIR / out_name
        rel = src.relative_to(ROOT).as_posix()
        try:
            text = extract_pdf_to_readable_text(src, rel)
            write_output_md(out_path, text)
            print(f"  OK   {src.name} -> {out_path.name}  ({len(text):,} chars)")
        except Exception as e:  # pragma: no cover
            print(f"  FAIL {src.name}: {e}")
            failures += 1

    print("Phase 0 design-control records -> 20-QMSInProcess/DC.SLQ002/")
    for src_name, out_name in PHASE0_DOCX.items():
        src = find_source_file(src_name)
        if not src:
            print(f"  MISSING source: {src_name}")
            failures += 1
            continue
        out_path = MIRROR_DIR / out_name
        try:
            text = extract_docx_to_readable_text(src)
            write_mirror_only(out_path, text)
            print(f"  OK   {src.name} -> {out_path.name}  ({len(text):,} chars)")
        except Exception as e:  # pragma: no cover
            print(f"  FAIL {src.name}: {e}")
            failures += 1

    for src_name, out_name in PHASE0_PDF.items():
        src = find_pdf_file(src_name)
        if not src:
            print(f"  MISSING source: {src_name}")
            failures += 1
            continue
        out_path = MIRROR_DIR / out_name
        rel = src.relative_to(ROOT).as_posix()
        try:
            text = extract_pdf_to_readable_text(src, rel)
            write_mirror_only(out_path, text)
            print(f"  OK   {src.name} -> {out_path.name}  ({len(text):,} chars)")
        except Exception as e:  # pragma: no cover
            print(f"  FAIL {src.name}: {e}")
            failures += 1

    print("-" * 60)
    print("Done." if failures == 0 else f"Done with {failures} failure(s).")
    return 0 if failures == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main())
